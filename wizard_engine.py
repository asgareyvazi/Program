# ============================================================================
# UNIVERSAL PROGRAM & PROCEDURE WIZARD — ENGINE
# ============================================================================
# "I want a drilling program" -> the wizard asks structured inputs ->
# a complete, professional Word document is generated on the spot.
#
# Templates live in wizard_library.py (programs) and wizard_procedures.py
# (procedures). Content is based on worldwide industry practice:
# API RP 5C3/10B/13B/53/59, ISO 10400/10426, NORSOK D-010, IADC guidelines,
# Shell DEP / Saudi Aramco SAES / BP GP style master documents.
# ============================================================================

import re
import html
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

from PySide6.QtWidgets import (
    QDialog, QWizard, QWizardPage, QVBoxLayout, QHBoxLayout, QGridLayout,
    QLabel, QLineEdit, QComboBox, QDoubleSpinBox, QCheckBox, QTextEdit,
    QListWidget, QListWidgetItem, QPushButton, QFileDialog, QProgressBar,
    QScrollArea, QWidget, QFrame, QFormLayout, QMessageBox, QGroupBox,
    QAbstractItemView, QTableWidget, QTableWidgetItem, QHeaderView,
    QApplication, QSpinBox, QSplitter
)
from PySide6.QtCore import Qt, QSize, QUrl
from PySide6.QtGui import QFont, QDesktopServices, QColor, QBrush
from docx.shared import Pt, Cm, RGBColor

# Risk engine integration (lazy import inside methods; RiskLevel used in UI)
from drilling_risk_analyzer import RiskLevel

# ----------------------------------------------------------------------------
# INPUT SPEC + TEMPLATE DEFINITION
# ----------------------------------------------------------------------------

INPUT_TYPES = ("text", "number", "combo", "check", "textarea", "table")


class InputSpec:
    """Definition of one wizard input field."""

    def __init__(self, key: str, label: str, input_type: str = "text",
                 options: Optional[List[str]] = None, unit: str = "",
                 required: bool = False, default: str = "",
                 placeholder: str = "", group: str = "General",
                 columns: Optional[List[str]] = None):
        assert input_type in INPUT_TYPES, f"bad input type {input_type}"
        self.key = key
        self.label = label
        self.type = input_type
        self.options = options or []
        self.unit = unit
        self.required = required
        self.default = default
        self.placeholder = placeholder
        self.group = group
        self.columns = columns or []


class TemplateDef:
    """Definition of one generator template (program or procedure)."""

    def __init__(self, key: str, name: str, icon: str, kind: str,
                 description: str, inputs: Optional[List[InputSpec]] = None,
                 markdown: str = "", source_file: Optional[str] = None,
                 tokens: Optional[Dict[str, str]] = None,
                 meta: Optional[Dict[str, str]] = None):
        self.key = key
        self.name = name
        self.icon = icon
        self.kind = kind              # "Program" | "Procedure"
        self.description = description
        self.inputs = inputs or []
        self.markdown = markdown
        self.source_file = source_file      # path relative to programs dir
        self.tokens = tokens or {}          # token -> input key (for file templates)
        self.meta = meta or {}

    @property
    def full_markdown(self) -> str:
        if self.source_file:
            base = Path(__file__).resolve().parent / "programs"
            p = base / self.source_file
            if p.exists():
                return p.read_text(encoding="utf-8", errors="replace")
            return f"# Template source not found: {self.source_file}"
        return self.markdown


# ----------------------------------------------------------------------------
# FILLING
# ----------------------------------------------------------------------------

# Unified placeholder syntax — the whole system accepts BOTH {{key}} and
# {key} (legacy DB procedures use single braces).  `{` followed by a known
# key is replaced; anything else stays untouched.
PLACEHOLDER_RE = re.compile(r"\{\{([a-zA-Z0-9_]+)\}\}|\{([a-zA-Z0-9_]+)\}")


# ----------------------------------------------------------------------------
# SECTION EXTRACTION & SELECTION
# ----------------------------------------------------------------------------

def extract_sections(md_text: str):
    """Split markdown into (heading, content) blocks at '##' level."""
    lines = md_text.replace("\r\n", "\n").split("\n")
    heads = []
    for i, l in enumerate(lines):
        m = re.match(r"^##\s+(.+)$", l)
        if m:
            heads.append((m.group(1).strip(), i))
    sections = []
    for idx, (h, start) in enumerate(heads):
        end = heads[idx + 1][1] if idx + 1 < len(heads) else len(lines)
        sections.append((h, "\n".join(lines[start:end])))
    return sections


def render_selected(md_text: str, selected_heads):
    """Return markdown containing only the selected '##' sections.

    The document header (before the first '##') and any section whose
    heading contains 'APPROVAL' are always kept.
    """
    lines = md_text.replace("\r\n", "\n").split("\n")
    first_h2 = next((i for i, l in enumerate(lines) if l.startswith("## ")),
                    None)
    header = "\n".join(lines[:first_h2]) if first_h2 is not None else md_text

    parts = [header.rstrip()]
    for h, content in extract_sections(md_text):
        if h in selected_heads or "APPROVAL" in h.upper():
            parts.append(content.rstrip())
    return "\n\n".join(p for p in parts if p.strip())


# ----------------------------------------------------------------------------
# COMPANY NAME NEUTRALIZATION
# ----------------------------------------------------------------------------
# Company/brand names must never appear in generated output — knowledge
# stays, branding goes. Operator/contractor names are supplied by the user
# as inputs instead.

OPERATOR_NAMES = [
    r"\bNISOC\b", r"\bPECO\b", r"\bOEOC\b", r"\bCNPCI\b", r"\bPEDEC\b",
    r"\bPEDCO\b", r"\bGWDC\b", r"\bNIOC\b", r"\bNIDC\b", r"\bIOOC\b",
    r"\bMSA\b", r"\bNICO\b", r"\bKPE\b", r"\bSK/PECO\b",
    r"\bSaudi Aramco\b", r"\bAramco\b", r"\bADNOC\b", r"\bKOC\b",
    r"\bINOC\b", r"\bNPDC\b", r"\bIOEC\b", r"\bJOGPC\b",
    r"\bPEDEX\b", r"\bMND\b", r"\bNational Iranian Oil Company\b",
    r"\bIranian Offshore Oil Company\b", r"\bCentral Iranian Oil Fields\b",
]
# Well-code / reservoir patterns that must never appear in output.
# The catalog & templates are general; this is a defensive layer for any
# content that may come from internal knowledge documents.
WELL_PATTERNS = [
    (r"\bAZNS\s*[-–]?\s*[A-Z0-9]*\b", ""),
    (r"\bSIAH MAKAN\b", "the field"),
    (r"\bSI-?\d+\b", ""),
    (r"\bAZR[- ]?\d*\b", ""),
    (r"\bMB-W\w+\b", ""),
    (r"\bYAP1\b", ""),
    (r"\bF-?20\b", ""),
    # well/section codes (S372, S223, S19...) — but NOT API steel grades
    # S-135 / S135 / S-95 (drill-pipe & casing grades must be preserved)
    (r"\bS-?(?!(?:135|95)\b)\d{2,3}\b", ""),
    (r"\bW0\d{2,3}\w*\b", ""),
    (r"\bSarvak\b", "main reservoir"),
    (r"\bFahliyan\b", "HPHT reservoir"),
    (r"\bKazhdumi\b", "reservoir"),
    (r"\bGadvan\b", "reservoir"),
    # offshore well codes from the 25-doc & pp2 sets
    (r"\b(?:well\s+)?BL-?01P\b", "the well"),
    (r"\b(?:well\s+)?2S-?01\b", "the well"),
    (r"\b(?:well\s+)?D-?127\b", "the well"),
    (r"\b(?:well\s+)?NR-?3[56]\b", "the well"),
    # well/field codes reported by the user in enrichment output (Batch T):
    # MB-011 / MB-013 (Mansouri field wells), GS 4-2 / GS-5 / GS Mbr. 4-2
    # (Gachsaran), N 1-3-5. The optional "well " prefix avoids doubled
    # text like "well the offset well".
    (r"\b(?:well\s+)?MB-?\d{0,4}(?=$|[\s.,;:)])", "the offset well"),
    (r"\bGS[- ]?Mbr\.?[- ]?\d{1,2}(?:[-/ ]\d{1,2})?\b", "the interval"),
    (r"\bGS[- ]?\d{1,2}(?:[-/ ]\d{1,2})?\b", "the interval"),
    (r"\bN\s*[1-9](?:[-–][1-9]){1,2}\b", "the well"),
    # reservoir / formation names (must never appear — internal knowledge)
    (r"\bGachsaran\s+formation\b", "the formation"),
    (r"\bGachsaran\b", "the formation"),
    (r"\bAsmari\b", "the reservoir"),
    (r"\bPabdeh\b", "the formation"),
    (r"\bAghajari\b", "the formation"),
    (r"\bMishan\b", "the formation"),
    (r"\bGuri\b", "the formation"),
    (r"\bIlam\b", "the formation"),
    (r"\bBangestan\b", "the formation"),
    (r"\bKhami\b", "the formation"),
    (r"\bDalan\b", "the formation"),
    (r"\bDariyan\b", "the formation"),
    (r"\bSalman\b", "the field"),
    (r"\bBalal\b", "the field"),
    (r"\bDorood\b", "the field"),
    (r"\bForoozan\b", "the field"),
    (r"\bSiri\b", "the field"),
    (r"\bNaftshahr\b", "the field"),
    (r"\bKEPCO\b", ""),
    (r"\bYaran\b", ""),
    (r"\bKangan\b", ""),
    (r"\bNICO\b", ""),
    (r"\bIOOC\b", ""),
    (r"\bNDCO\b", ""),
    (r"\bP\.?P\.?Z\b", ""),
    (r"\bShah Deniz\b", "the field"),
    (r"\bSDX-?\d*\b", ""),
    (r"\bSPD\s*\d*\b", ""),
    (r"\bSD A-03\b", "the well"),
    (r"\bBP\b", ""),
    (r"\(Brown\)|\(BROWN\)", ""),
    (r"\bBrown\b(?=\s+(?:JM|CPH|HSR|HMC))", ""),
    (r"\bMI\b(?!-)", ""),
    (r"\bHALCO\b", ""),
    (r"\bAnadrill\b", ""),
    (r"\bWHE\b", ""),
    (r"\bDril-Quip\b", ""),
    (r"\bIngram-Cactus\b", ""),
    (r"\bTOTCO\b", ""),
    (r"\bWeco\b", ""),
    (r"\bVarco\b", ""),
    (r"\bKhazar\b", "the field"),
    (r"\bCK-?\d+\b", ""),
    (r"\bSR-?\d+\b", ""),
    (r"\bNSH-?\d+\b", ""),
    (r"\bDH-?\d+\b", ""),
    (r"\bMK[- ]?\d+\b", ""),
    (r"\bWDI-?\d+\b", ""),
    (r"\bAGH-?\d+\b", ""),
    (r"\bHE-E?7\b", ""),
    (r"\bD-?1(?:27|26|028)\b", "the well"),
    (r"\bADCO\b", ""),
    (r"\bSaudi Aramco\b", ""),
    (r"\bAramco\b", ""),
    (r"\bNimir\b", ""),
    (r"\bPETROM\b", ""),
    (r"\bExxonMobil\b", ""),
    (r"\bExxon\b", ""),
    (r"\bMobil\b", ""),
    # NOTE: IADC is a standards body (International Association of Drilling
    # Contractors) — citations like "per IADC guidelines" are legitimate and
    # must be preserved.  It is intentionally NOT blacklisted.
    (r"\bDevereux\b", ""),
    (r"\bWell Control School\b", ""),
    (r"\bAberdeen\b", ""),
    (r"\bChevron\b", ""),
    (r"\bChevronTexaco\b", ""),
    (r"\bStatoil\b", ""),
    (r"\bMaersk\b", ""),
    (r"\bTotalEnergies\b", ""),
    (r"\bTotal\b(?=\s+(?:E&P|EP|Energies|Exploration|S\.?A\.?))", ""),
    (r"\bShell\b", ""),
    (r"\bBP Exploration\b", ""),
]

SERVICE_NAMES = [
    r"\bSchlumberger\b", r"\bSLB\b", r"\bHalliburton\b",
    r"\bBaker Hughes\b", r"\bBaker\b", r"\bBakerlok\b",
    r"\bWeatherford\b",
    r"\bNOV\b", r"\bNational Oilwell\b", r"\bFMC\b", r"\bCameron\b",
    r"\bHydril\b", r"\bTenaris\b", r"\bVallourec\b", r"\bCoreLab\b",
    r"\bSGS\b", r"\bT\u00dcV\b", r"\bExpro\b", r"\bProserv\b",
    r"\bOdfjell\b", r"\bKCA Deutag\b", r"\bParker Drilling\b",
    r"\bEDC\b", r"\bEnsco\b", r"\bTransocean\b", r"\bMaersk Drilling\b",
    r"\bNabors\b", r"\bHelmerich\b", r"\bPatterson-UTI\b",
    r"\bIPM\b", r"\bDowell\b", r"\bSperry Sun\b", r"\bSperry\b",
    r"\bMartin Decker\b", r"\bVarco\b", r"\bVetco\b", r"\bTIW\b",
    r"\bReagan\b", r"\bElmagco\b", r"\bNormar\b", r"\bOmsco\b",
    r"\bBARA-WATE\b", r"\bAnadrill\b", r"\bGeoQuest\b",
    r"\bInTouch\b", r"\bT\.H\. Hill\b", r"\bTotco\b", r"\bWeco\b",
    r"\bModuspec\b", r"\bWestHou\b", r"\bCoilCADE\b", r"\bCoilCade\b",
]


def neutralize_text(md_text: str, operator_name: str = "",
                    contractor_name: str = "") -> str:
    """Remove hard-coded company/brand names from output text.

    The user's own operator / contractor names (entered as inputs) are
    substituted instead of generic placeholders. When not provided, the
    generic roles 'the Operator' / 'the Service Company' are used.

    Temporary sentinels are used so that a user-entered company name which
    itself contains a blacklisted token (e.g. 'PEDEC Oil') is never
    re-replaced.
    """
    op = (operator_name or "").strip() or "the Operator"
    con = (contractor_name or "").strip() or "the Service Company"
    OP_TMP = "\x00OP\x00"
    CON_TMP = "\x00CON\x00"

    out = md_text
    # Protect the user's own names first (whole-word, case-sensitive), so a
    # user name that contains a blacklisted token (e.g. 'PEDEC Oil Co') or
    # is very short is never re-replaced.
    if op.lower() != "the operator" and len(op) >= 2:
        out = re.sub(r"\b" + re.escape(op) + r"\b", OP_TMP, out)
    if con.lower() != "the service company" and len(con) >= 2:
        out = re.sub(r"\b" + re.escape(con) + r"\b", CON_TMP, out)

    # A blacklisted token that appears inside the user's operator name maps
    # to the operator; inside the contractor name maps to the contractor;
    # otherwise operator tokens -> operator, service tokens -> contractor.
    def _target(pat, op, con):
        """Decide which user name a blacklisted token maps to."""
        if re.search(pat, op, re.IGNORECASE):
            return OP_TMP
        if re.search(pat, con, re.IGNORECASE):
            return CON_TMP
        return OP_TMP  # operator tokens default to operator

    for pat in OPERATOR_NAMES:
        out = re.sub(pat, _target(pat, op, con), out, flags=re.IGNORECASE)
    for pat in SERVICE_NAMES:
        if re.search(pat, op, re.IGNORECASE):
            out = re.sub(pat, OP_TMP, out, flags=re.IGNORECASE)
        else:
            out = re.sub(pat, CON_TMP, out, flags=re.IGNORECASE)

    out = out.replace(OP_TMP, op)
    out = out.replace(CON_TMP, con)

    # Defensive: remove well codes / reservoir names (general output)
    for pat, repl in WELL_PATTERNS:
        out = re.sub(pat, repl, out, flags=re.IGNORECASE)
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\(\s*\)", "", out)
    out = re.sub(re.escape(op) + r"\s+" + re.escape(op), op, out)
    out = re.sub(re.escape(con) + r"\s+" + re.escape(con), con, out)
    return out


UNRESOLVED_RE = re.compile(r"\{\{([a-zA-Z0-9_]+)\}\}|\{([a-zA-Z0-9_]+)\}")


def scan_unresolved_placeholders(md_text: str) -> List[str]:
    """Final-output QA: every {{key}} / {key} left in the text after all
    enrichment stages is an unresolved parameter.  Returns the keys."""
    keys = set()
    for m in UNRESOLVED_RE.finditer(md_text or ""):
        keys.add(m.group(1) or m.group(2))
    return sorted(keys)


def fill_template(tdef: TemplateDef, values: Dict[str, str]) -> str:
    """Replace {{key}} / {key} placeholders (and token literals for file
    templates) with the entered values. Empty values become
    '[To Be Filled]'."""
    md = tdef.full_markdown

    def repl(m):
        key = m.group(1) or m.group(2)
        val = values.get(key, "")
        return val if str(val).strip() else "[To Be Filled]"

    md = PLACEHOLDER_RE.sub(repl, md)

    # Token-based replacement (used by file templates such as ESP master doc)
    if tdef.tokens:
        for token, key in tdef.tokens.items():
            val = values.get(key, "")
            if str(val).strip():
                md = md.replace(token, str(val))

    # Append real reference documents from the operations library
    try:
        from wizard_references import reference_markdown
        ref_section = reference_markdown(tdef.key)
        if ref_section and "REFERENCE DOCUMENTS" not in md:
            md = md.rstrip() + "\n\n---\n\n" + ref_section
    except Exception:
        pass

    return md


# ----------------------------------------------------------------------------
# MARKDOWN -> WORD (.docx)
# ----------------------------------------------------------------------------

ACCENT = RGBColor(0xE9, 0x45, 0x60)
DARK_BLUE = RGBColor(0x0F, 0x34, 0x60)
GOLD = RGBColor(0xC9, 0x9A, 0x2E)
GREY = RGBColor(0x80, 0x80, 0x90)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "EDF1F7"
HEAD_BG = "0F3460"

# Active output font settings (set by md_to_docx from user options)
_FONT_NAME = "Calibri"
_FONT_SIZE = 11.0


def _shade_cell(cell, color_hex: str):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'))


def _add_runs(paragraph, text: str, base_size: Optional[float] = None):
    if base_size is None:
        base_size = _FONT_SIZE
    """Add runs to a paragraph, honoring **bold**, *italic*, `code`."""
    # Split by tokens: **...**, *...*, `...`
    pattern = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    for part in pattern.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
        else:
            run.text = part
        run.font.size = Pt(base_size)


def _heading(doc, text: str, level: int):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 8)
    pf.space_after = Pt(6)

    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{HEAD_BG}" w:val="clear"/>'))
        r = p.add_run("  " + text.upper())
        r.bold = True
        r.font.size = Pt(_FONT_SIZE + 4)
        r.font.name = _FONT_NAME
        r.font.color.rgb = WHITE
    elif level == 2:
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(_FONT_SIZE + 1.5)
        r.font.name = _FONT_NAME
        r.font.color.rgb = ACCENT
    else:
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(_FONT_SIZE + 0.5)
        r.font.name = _FONT_NAME
        r.font.color.rgb = DARK_BLUE


def _add_table(doc, header: List[str], rows: List[List[str]],
               col_widths: Optional[List[float]] = None):
    from docx.shared import Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        _add_runs(p, str(h), base_size=_FONT_SIZE - 1.5)
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = WHITE
        _shade_cell(hdr[i], HEAD_BG)

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(len(header)):
            val = row[i] if i < len(row) else ""
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            _add_runs(p, str(val), base_size=_FONT_SIZE - 1.5)
            if ridx % 2 == 1:
                _shade_cell(cells[i], LIGHT_BG)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def md_to_docx(md_text: str, out_path: str,
               meta: Optional[Dict[str, str]] = None,
               options: Optional[Dict] = None) -> bool:
    """Convert the markdown template (already filled) into a Word document.

    options may contain: font, font_size, page (A4/Letter), orientation,
    margin_left/right/top/bottom (cm), cover, toc, header_text, footer_text.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml

    meta = meta or {}
    options = options or {}

    global _FONT_NAME, _FONT_SIZE
    font_name = options.get("font", "Calibri")
    font_size = float(options.get("font_size", 11))
    _FONT_NAME = font_name
    _FONT_SIZE = font_size
    page = options.get("page", "A4")
    orientation = options.get("orientation", "Portrait")
    ml = options.get("margin_left", 2.5)
    mr = options.get("margin_right", 2.0)
    mt = options.get("margin_top", 2.0)
    mb = options.get("margin_bottom", 2.0)

    doc = Document()
    # Base font on the Normal style (applies to everything by default)
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)

    sec = doc.sections[0]
    if page == "Letter":
        sec.page_width = Cm(21.59)
        sec.page_height = Cm(27.94)
    else:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
    if orientation == "Landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(mt)
    sec.bottom_margin = Cm(mb)
    sec.left_margin = Cm(ml)
    sec.right_margin = Cm(mr)

    # Header / footer
    if options.get("header_text"):
        hp = sec.header.paragraphs[0]
        hp.text = options["header_text"]
        for r in hp.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = GREY
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT "/>')
    run._r.addnext(fld)
    if options.get("footer_text"):
        r2 = fp.add_run("   " + options["footer_text"])
        r2.font.size = Pt(8)
        r2.font.color.rgb = GREY

    # Cover page
    if options.get("cover", True):
        title = ""
        for line in md_text.replace("\r\n", "\n").split("\n"):
            m = re.match(r"^#\s+(.+)$", line)
            if m:
                title = m.group(1).strip()
                break
        doc.add_paragraph("")
        for _ in range(3):
            doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title or meta.get("title", "DOCUMENT"))
        r.bold = True
        r.font.size = Pt(24)
        r.font.color.rgb = DARK_BLUE

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(meta.get("title", ""))
        r2.font.size = Pt(14)
        r2.font.color.rgb = ACCENT

        doc.add_paragraph("")
        for label in ("Operator", "Contractor", "Document Number",
                      "Revision", "Date", "Prepared By", "Reviewed By",
                      "Approved By"):
            val = meta.get({
                "Operator": "operator",
                "Contractor": "contractor",
                "Document Number": "document_number",
                "Revision": "revision",
                "Date": "date",
                "Prepared By": "prepared_by",
                "Reviewed By": "reviewed_by",
                "Approved By": "approved_by",
            }.get(label, ""), "")
            if val:
                p3 = doc.add_paragraph()
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r3 = p3.add_run(f"{label}: {val}")
                r3.font.size = Pt(11)
        doc.add_page_break()

    # Table of contents (real Word TOC field — update with F9)
    if options.get("toc", True):
        p = doc.add_paragraph()
        r = p.add_run("TABLE OF CONTENTS")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = DARK_BLUE
        doc.add_paragraph()
        fld_p = doc.add_paragraph()
        fld_run = fld_p.add_run()
        fld_xml = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" TOC \\o &quot;1-2&quot; \\h \\z \\u "/>')
        fld_run._r.addnext(fld_xml)
        doc.add_page_break()

    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip()

        # Fenced code block -> single monospace paragraph
        if line.startswith("```"):
            buf = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            for b in buf:
                r = p.add_run(b + "\n")
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            _heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$", line):
            doc.add_paragraph()
            i += 1
            continue

        # Table
        if line.lstrip().startswith("|") and i + 1 < n and \
                re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]) and \
                "-" in lines[i + 1]:
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            _add_table(doc, header, rows)
            continue

        # Blockquote
        if line.startswith(">"):
            buf = []
            while i < n and lines[i].startswith(">"):
                buf.append(lines[i].lstrip("> ").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run("▸ ")
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_BLUE
            _add_runs(p, " ".join(buf), base_size=10)
            for run in p.runs[1:]:
                run.italic = True
                run.font.color.rgb = DARK_BLUE
            continue

        # Checklist item
        if re.match(r"^\s*-\s*\[[ xX]\]\s+", line):
            while i < n and re.match(r"^\s*-\s*\[[ xX]\]\s+", lines[i]):
                text = re.sub(r"^\s*-\s*\[[ xX]\]\s+", "", lines[i])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                r = p.add_run("☐  ")
                r.font.size = Pt(11)
                _add_runs(p, text)
                i += 1
            continue

        # Unordered list
        if re.match(r"^\s*[-*+]\s+", line):
            while i < n and re.match(r"^\s*[-*+]\s+", lines[i]):
                text = re.sub(r"^\s*[-*+]\s+", "", lines[i])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                r = p.add_run("•  ")
                r.font.size = Pt(10)
                _add_runs(p, text)
                i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+[.)]\s+", line):
            while i < n and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                text = re.sub(r"^\s*\d+[.)]\s+", "", lines[i])
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                _add_runs(p, text)
                i += 1
            continue

        # Blank
        if not line.strip():
            i += 1
            continue

        # Paragraph — guaranteed progress: unmatched lines are consumed.
        buf = []
        while i < n:
            s = lines[i].strip()
            if not s or s.startswith("#") or s.startswith(">") or \
                    s.startswith("|") or s.startswith("```") or \
                    re.match(r"^\s*[-*+]\s+", lines[i]) or \
                    re.match(r"^\s*\d+[.)]\s+", lines[i]) or \
                    re.match(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$", lines[i]) or \
                    re.match(r"^\s*-\s*\[[ xX]\]\s+", lines[i]):
                break
            buf.append(s)
            i += 1
        if buf:
            p = doc.add_paragraph()
            _add_runs(p, " ".join(buf))
        else:
            p = doc.add_paragraph()
            _add_runs(p, lines[i].strip() if i < n else "")
            i += 1

    # Footer meta
    if meta:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            f"Generated by Drilling Program Generator Pro v3.1  •  "
            f"{meta.get('date', datetime.now().strftime('%d-%B-%Y'))}")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY

    doc.save(out_path)
    return True


# ----------------------------------------------------------------------------
# DYNAMIC INPUT WIDGETS
# ----------------------------------------------------------------------------

class TableInputWidget(QWidget):
    """Editable table input (add/remove rows) for table-type InputSpec."""

    def __init__(self, spec: InputSpec, parent=None):
        super().__init__(parent)
        self.columns = spec.columns or ["Item"]
        lay = QVBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)

        self.table = QTableWidget(0, len(self.columns))
        self.table.setHorizontalHeaderLabels(self.columns)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setMinimumHeight(140)
        lay.addWidget(self.table)

        btns = QHBoxLayout()
        btn_add = QPushButton("➕ Add Row")
        btn_add.setMaximumWidth(110)
        btn_add.clicked.connect(self.add_row)
        btn_rm = QPushButton("➖ Remove Row")
        btn_rm.setMaximumWidth(110)
        btn_rm.clicked.connect(self.remove_row)
        btns.addWidget(btn_add)
        btns.addWidget(btn_rm)
        btns.addStretch()
        lay.addLayout(btns)

    def add_row(self, data=None):
        row = self.table.rowCount()
        self.table.insertRow(row)
        for c in range(len(self.columns)):
            val = data[c] if data and c < len(data) else ""
            self.table.setItem(row, c, QTableWidgetItem(str(val)))

    def remove_row(self):
        rows = {i.row() for i in self.table.selectedItems()}
        for r in sorted(rows, reverse=True):
            self.table.removeRow(r)

    def values(self) -> List[List[str]]:
        out = []
        for r in range(self.table.rowCount()):
            row = []
            for c in range(self.table.columnCount()):
                item = self.table.item(r, c)
                row.append(item.text().strip() if item else "")
            if any(row):
                out.append(row)
        return out


def _build_field(spec: InputSpec) -> QWidget:
    if spec.type == "text":
        w = QLineEdit()
        w.setPlaceholderText(spec.placeholder or spec.label)
        if spec.default:
            w.setText(spec.default)
        return w
    if spec.type == "number":
        w = QDoubleSpinBox()
        w.setRange(0, 1e9)
        w.setDecimals(2)
        # "not entered" state: at the minimum the box shows a marker
        # instead of a fake 0 — a required numeric field left untouched is
        # reported as missing, never as 0 (engineering correctness).
        w.setSpecialValueText("[Not Entered]")
        if spec.default:
            try:
                w.setValue(float(spec.default))
            except (TypeError, ValueError):
                pass
        if spec.unit:
            w.setSuffix(f" {spec.unit}")
        return w
    if spec.type == "combo":
        w = QComboBox()
        w.addItems(spec.options or [])
        if spec.default and w.findText(spec.default) >= 0:
            w.setCurrentText(spec.default)
        # Batch Y — a combo with a defined option list is NOT free-text:
        # the user must pick one of the engineering options.  Only an
        # empty option list becomes an editable free-text field.
        w.setEditable(not bool(spec.options))
        return w
    if spec.type == "check":
        w = QCheckBox()
        if spec.default.lower() in ("1", "yes", "true"):
            w.setChecked(True)
        return w
    if spec.type == "textarea":
        w = QTextEdit()
        w.setMaximumHeight(90)
        if spec.default:
            w.setPlainText(spec.default)
        return w
    if spec.type == "table":
        return TableInputWidget(spec)
    raise ValueError(spec.type)


def _get_field_value(spec: InputSpec, w: QWidget) -> str:
    if isinstance(w, QLineEdit):
        return w.text().strip()
    if isinstance(w, QDoubleSpinBox):
        # untouched numeric field (showing its special "[Not Entered]"
        # marker at the minimum) reads as empty — unless the default is
        # literally zero, which is a deliberate value.
        if w.value() == w.minimum() and w.specialValueText():
            dflt = str(spec.default or "").strip()
            if dflt not in ("0", "0.0", "0.00"):
                return ""
        return str(w.value())
    if isinstance(w, QComboBox):
        return w.currentText().strip()
    if isinstance(w, QCheckBox):
        return "YES" if w.isChecked() else "NO"
    if isinstance(w, QTextEdit):
        return w.toPlainText().strip()
    if isinstance(w, TableInputWidget):
        rows = w.values()
        if not rows:
            return ""
        return " | ".join(" ; ".join(r) for r in rows)
    return ""


def table_to_md_rows(value: str, columns: List[str]) -> str:
    """Convert a table input value back into markdown table rows."""
    if not value:
        return ""
    out = ["| " + " | ".join(columns) + " |",
           "|" + "|".join(["---"] * len(columns)) + "|"]
    for row in value.split(" | "):
        cells = row.split(" ; ")
        out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


# ----------------------------------------------------------------------------
# WIZARD DIALOG
# ----------------------------------------------------------------------------

class _TemplatePage(QWizardPage):
    """Page 1: choose what kind of document to generate (friendly view)."""

    def __init__(self, templates: List[TemplateDef], parent=None):
        super().__init__(parent)
        self.setTitle("1. What would you like to generate?")
        self.setSubTitle(
            "Pick the document type. The software will ask which sections "
            "you want, collect the inputs, and produce an editable Word "
            "document.")

        lay = QVBoxLayout(self)
        splitter = QSplitter(Qt.Horizontal)

        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(0, 0, 0, 0)

        rowf = QHBoxLayout()
        self.search = QLineEdit()
        self.search.setPlaceholderText("🔍 Search document types...")
        self.search.textChanged.connect(lambda _t: self._apply_filter())
        rowf.addWidget(self.search, 1)
        self.filter = QComboBox()
        self.filter.addItems(["All types", "Programs", "Procedures"])
        self.filter.currentTextChanged.connect(self._apply_filter)
        rowf.addWidget(self.filter)
        ll.addLayout(rowf)

        self.listw = QListWidget()
        for t in templates:
            item = QListWidgetItem(f"{t.icon}  {t.name}")
            item.setData(Qt.UserRole, t.key)
            item.setToolTip(t.description)
            item.setSizeHint(QSize(0, 34))
            self.listw.addItem(item)
        self.listw.currentItemChanged.connect(self._on_select)
        ll.addWidget(self.listw)
        splitter.addWidget(left)

        right = QWidget()
        rl = QVBoxLayout(right)
        rl.setContentsMargins(10, 0, 0, 0)
        self.desc = QLabel("Select a document type to see details.")
        self.desc.setWordWrap(True)
        self.desc.setStyleSheet(
            "background:#1a1a2e;border:1px solid #0f3460;border-radius:8px;"
            "padding:14px;font-size:12px;color:#c0ccd8;")
        rl.addWidget(self.desc)
        rl.addStretch()
        splitter.addWidget(right)

        splitter.setSizes([420, 500])
        lay.addWidget(splitter)

        self._templates = templates
        if self.listw.count():
            self.listw.setCurrentRow(0)

    def selected_key(self) -> str:
        item = self.listw.currentItem()
        return item.data(Qt.UserRole) if item else ""

    def _apply_filter(self, text=None):
        text = text if text is not None else self.filter.currentText()
        q = self.search.text().strip().lower() if hasattr(self, "search") else ""
        for idx in range(self.listw.count()):
            item = self.listw.item(idx)
            t = self._by_key(item.data(Qt.UserRole))
            show = (text == "All types" or
                    (text == "Programs" and t.kind == "Program") or
                    (text == "Procedures" and t.kind == "Procedure"))
            if show and q and q not in item.text().lower():
                show = False
            item.setHidden(not show)

    def _by_key(self, key):
        for t in self._templates:
            if t.key == key:
                return t
        return None

    def _on_select(self, cur, _prev):
        if cur is None:
            return
        t = self._by_key(cur.data(Qt.UserRole))
        if t:
            n_sections = len(extract_sections(t.full_markdown))
            self.desc.setText(
                f"<b style='color:#e94560;'>{t.icon} {t.name}</b><br>"
                f"<span style='color:#4fc3f7;'>{t.kind}</span><br><br>"
                f"{t.description}<br><br>"
                f"<span style='color:#8a8a9a;'>Sections available: "
                f"{n_sections}  •  Fields to fill: {len(t.inputs)}</span>")


class _WellProfilePage(QWizardPage):
    """Page 2: well profile — the engineer's basis of design.

    The user (as a drilling engineer) specifies the well type, environment
    and operation. This (a) is written into the document as a BASIS OF
    DESIGN section, (b) filters which knowledge documents are used for
    enrichment, and (c) highlights recommended templates.
    """

    # operation -> recommended template keys (in priority order)
    OP_TEMPLATES = {
        "Drilling": ["drilling_program", "offshore_drilling_program",
                     "advanced_drilling_program"],
        "Workover": ["workover_program", "offshore_workover_program"],
        "Re-Entry": ["reentry_program"],
        "Sidetrack": ["reentry_program"],
        "Completion": ["esp_workover", "offshore_workover_program"],
        "P&A": ["abandonment_program"],
        "Well Testing": ["well_testing_program"],
        "Stimulation": ["stimulation_program"],
        "Fishing": ["fishing_program"],
    }

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("2. Well Profile (Basis of Design)")
        self.setSubTitle(
            "Define the well like a drilling engineer would: well type, "
            "environment and operation. The software uses this to filter "
            "the right knowledge documents and pre-fill the document.")

        lay = QVBoxLayout(self)

        grid = QGridLayout()
        grid.setHorizontalSpacing(14)
        grid.setVerticalSpacing(10)

        grid.addWidget(QLabel("Well Type:"), 0, 0)
        self.well_type = QComboBox()
        self.well_type.addItems(["Vertical", "Deviated", "Horizontal", "ERD",
                                 "HPHT", "Deepwater", "Multi-lateral"])
        grid.addWidget(self.well_type, 0, 1)

        grid.addWidget(QLabel("Environment:"), 1, 0)
        self.environment = QComboBox()
        self.environment.addItems(["Onshore", "Offshore Jack-up",
                                   "Semi-submersible", "Fixed Platform",
                                   "Caspian Sea"])
        grid.addWidget(self.environment, 1, 1)

        grid.addWidget(QLabel("Operation:"), 2, 0)
        self.operation = QComboBox()
        self.operation.addItems(list(self.OP_TEMPLATES.keys()))
        grid.addWidget(self.operation, 2, 1)

        grid.addWidget(QLabel("Hole Sections (optional):"), 3, 0)
        self.holes = QLineEdit()
        self.holes.setPlaceholderText("e.g. 36\", 26\", 17-1/2\", 12-1/4\", 8-1/2\"")
        grid.addWidget(self.holes, 3, 1)

        lay.addLayout(grid)

        # knowledge match label
        self.lbl_match = QLabel("")
        self.lbl_match.setWordWrap(True)
        self.lbl_match.setStyleSheet(
            "background:#0f3460;color:#e0e0e0;border-radius:6px;"
            "padding:8px;font-size:11px;")
        lay.addWidget(self.lbl_match)

        self.lbl_templates = QLabel("")
        self.lbl_templates.setWordWrap(True)
        self.lbl_templates.setStyleSheet("color:#8a8a9a;font-size:11px;")
        lay.addWidget(self.lbl_templates)

        for w in (self.well_type, self.environment, self.operation):
            w.currentTextChanged.connect(self._update_matches)

        self._catalog = None
        lay.addStretch(1)

    def _get_catalog(self):
        if self._catalog is None:
            try:
                from document_catalog import get_catalog
                self._catalog = get_catalog()
            except Exception:
                self._catalog = False
        return self._catalog or None

    def _update_matches(self):
        cat = self._get_catalog()
        if not cat:
            self.lbl_match.setText("Knowledge catalog not available.")
            return
        n = cat.matched_summary(
            well_type=self.well_type.currentText(),
            environment=self.environment.currentText(),
            operation=self.operation.currentText())
        self.lbl_match.setText(
            f"📚 {n} matching knowledge documents in the internal library "
            f"({self.well_type.currentText()} well / {self.environment.currentText()} "
            f"/ {self.operation.currentText()}). These will be used to enrich "
            f"your document.")
        # recommended templates
        keys = self.OP_TEMPLATES.get(self.operation.currentText(), [])
        names = []
        try:
            wiz = self.wizard()
            page0 = wiz.page(0)
            for k in keys:
                t = page0._by_key(k) if hasattr(page0, "_by_key") else None
                if t:
                    names.append(f"{t.icon} {t.name}")
        except Exception:
            pass
        self.lbl_templates.setText(
            "💡 Recommended document types: " + (" | ".join(names)
                                                 if names else "—"))

    def initializePage(self):
        self._update_matches()

    def profile(self) -> Dict:
        return {
            "well_type": self.well_type.currentText(),
            "environment": self.environment.currentText(),
            "operation": self.operation.currentText(),
            "holes": self.holes.text().strip(),
        }

    def profile_markdown(self) -> str:
        p = self.profile()
        holes = p["holes"] or "As per casing design"
        return ("## WELL PROFILE & BASIS OF DESIGN\n\n"
                "| Parameter | Value |\n|---|---|\n"
                f"| Well Type | {p['well_type']} |\n"
                f"| Environment | {p['environment']} |\n"
                f"| Operation | {p['operation']} |\n"
                f"| Hole Sections | {holes} |\n")


class _SectionsPage(QWizardPage):
    """Page 2: user picks which sections go into the document."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("3. Choose the sections")
        self.setSubTitle(
            "Tick the sections you want in the document. Unticked sections "
            "are omitted. The header and approval block are always kept.")

        lay = QVBoxLayout(self)
        btnrow = QHBoxLayout()
        btn_all = QPushButton("✓ Select All")
        btn_all.setMaximumWidth(130)
        btn_none = QPushButton("✗ Select None")
        btn_none.setMaximumWidth(130)
        btnrow.addWidget(btn_all)
        btnrow.addWidget(btn_none)
        btnrow.addStretch(1)
        lay.addLayout(btnrow)
        self.listw = QListWidget()
        self.listw.setSelectionMode(QAbstractItemView.SingleSelection)
        lay.addWidget(self.listw)

        self.lbl = QLabel("")
        self.lbl.setStyleSheet("color:#8a8a9a;font-size:11px;")
        lay.addWidget(self.lbl)

        self._sections: List[str] = []
        self._saved: Dict[str, List[str]] = {}

        # wire the select all / none buttons
        btn_all = self.findChildren(QPushButton)
        for b in btn_all:
            if b.text() == "✓ Select All":
                b.clicked.connect(lambda: self._set_all(True))
            elif b.text() == "✗ Select None":
                b.clicked.connect(lambda: self._set_all(False))

    def _set_all(self, on: bool):
        for i in range(self.listw.count()):
            item = self.listw.item(i)
            if item.flags() & Qt.ItemIsEnabled:
                item.setCheckState(Qt.Checked if on else Qt.Unchecked)

    def initializePage(self):
        self.listw.clear()
        wiz = self.wizard()
        page0 = wiz.page(0)
        key = page0.selected_key()
        tdef = page0._by_key(key) if hasattr(page0, "_by_key") else None
        if tdef is None:
            return
        sections = extract_sections(tdef.full_markdown)
        self._sections = [h for h, _ in sections]

        prev = self._saved.get(key)
        for h, _ in sections:
            item = QListWidgetItem(h)
            item.setData(Qt.UserRole, h)
            if "APPROVAL" in h.upper():
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                item.setCheckState(Qt.Checked)
                item.setFlags(item.flags() & ~Qt.ItemIsEnabled)
            else:
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                if prev is not None:
                    item.setCheckState(Qt.Checked if h in prev else Qt.Unchecked)
                else:
                    item.setCheckState(Qt.Checked)
            self.listw.addItem(item)
        self.lbl.setText(
            f"{len(sections)} sections — header + approval are always included")

    def selected_heads(self) -> List[str]:
        out = []
        for i in range(self.listw.count()):
            item = self.listw.item(i)
            if item.flags() & Qt.ItemIsEnabled and item.checkState() == Qt.Checked:
                out.append(item.data(Qt.UserRole))
        return out

    def validatePage(self) -> bool:
        wiz = self.wizard()
        key = wiz.page(0).selected_key()
        self._saved[key] = self.selected_heads()
        return True


class ROPCalibrationDialog(QDialog):
    """Enter offset-well ROP data and fit the Bourgoyne-Young style model.

    The fitted model (K, a, b, c, d) is stored on the inputs page and used
    by the Deep Engineering Verification section of the generated document.
    """

    SAMPLE = [
        # depth_ft, wob_klbf, rpm, mw_ppg, rop_actual_ft_hr
        (5000, 20, 90, 11.0, 35.0),
        (6500, 22, 95, 11.3, 31.5),
        (8000, 25, 100, 11.5, 28.0),
        (9500, 28, 105, 11.8, 23.5),
        (11000, 30, 110, 12.0, 20.0),
    ]

    def __init__(self, parent=None, current: Optional[dict] = None):
        super().__init__(parent)
        self.setWindowTitle("🧮 ROP Calibration — Offset-Well Data")
        self.setMinimumSize(760, 480)
        self.result: Optional[dict] = None

        lay = QVBoxLayout(self)

        head = QLabel(
            "Fit the ROP model (Bourgoyne-Young style) with offset-well data:\n"
            "ROP = K × WOB^a × RPM^b × e^(−c·D) × e^(d·(MW−MW_opt))\n"
            "Only K is fitted (exponents are industry defaults). "
            "All data stays local — no cloud.")
        head.setWordWrap(True)
        lay.addWidget(head)

        cols = ["Depth (ft)", "WOB (klbf)", "RPM", "MW (ppg)",
                "Actual ROP (ft/hr)"]
        self.table = QTableWidget(0, 5)
        self.table.setHorizontalHeaderLabels(cols)
        self.table.horizontalHeader().setSectionResizeMode(
            QHeaderView.Stretch)
        lay.addWidget(self.table)

        btns = QHBoxLayout()
        b_add = QPushButton("➕ Add row")
        b_add.clicked.connect(lambda: self._add_row())
        b_rem = QPushButton("➖ Remove row")
        b_rem.clicked.connect(self._remove_row)
        b_smp = QPushButton("📋 Load sample data")
        b_smp.clicked.connect(self._load_sample)
        b_fit = QPushButton("⚙️  Fit Model")
        b_fit.clicked.connect(self._fit)
        btns.addWidget(b_add); btns.addWidget(b_rem)
        btns.addWidget(b_smp); btns.addStretch(); btns.addWidget(b_fit)
        lay.addLayout(btns)

        self.fit_label = QLabel("No fit yet.")
        self.fit_label.setWordWrap(True)
        self.fit_label.setStyleSheet(
            "background:#1a1a2e;border:1px solid #0f3460;border-radius:6px;"
            "padding:8px;color:#4fc3f7;")
        lay.addWidget(self.fit_label)

        okb = QHBoxLayout()
        okb.addStretch()
        b_ok = QPushButton("✔ Accept Fit")
        b_ok.setEnabled(False)
        b_ok.clicked.connect(self._accept)
        self.btn_ok = b_ok
        b_cancel = QPushButton("Cancel")
        b_cancel.clicked.connect(self.reject)
        okb.addWidget(b_ok); okb.addWidget(b_cancel)
        lay.addLayout(okb)

        if current and current.get("points"):
            for p in current["points"]:
                self._add_row(p)
        else:
            self._add_row()

    # ------------------------------------------------------------------
    def _add_row(self, data=None):
        r = self.table.rowCount()
        self.table.insertRow(r)
        defaults = list(data) if data else ["", "", "", "", ""]
        for c, val in enumerate(defaults):
            item = QTableWidgetItem(str(val))
            item.setFlags(item.flags() | Qt.ItemIsEditable)
            self.table.setItem(r, c, item)

    def _remove_row(self):
        rows = sorted({i.row() for i in self.table.selectedItems()},
                      reverse=True)
        for r in rows:
            self.table.removeRow(r)

    def _load_sample(self):
        self.table.setRowCount(0)
        for p in self.SAMPLE:
            self._add_row(p)

    def _points(self) -> List[Dict]:
        pts = []
        for r in range(self.table.rowCount()):
            vals = []
            ok = True
            for c in range(5):
                txt = self.table.item(r, c).text().strip() if \
                    self.table.item(r, c) else ""
                if not txt:
                    ok = False
                    break
                try:
                    vals.append(float(txt))
                except ValueError:
                    ok = False
                    break
            if ok and vals[4] > 0:
                pts.append({"depth": vals[0], "wob": vals[1], "rpm": vals[2],
                            "mw": vals[3], "rop_actual": vals[4]})
        return pts

    def _fit(self):
        pts = self._points()
        if len(pts) < 3:
            QMessageBox.warning(
                self, "Not Enough Data",
                "Enter at least 3 complete offset data points "
                "(depth, WOB, RPM, MW, actual ROP).")
            return
        try:
            from engineering_deep import ROPCalibrator
            rc = ROPCalibrator()
            k = rc.calibrate(pts)
            # prediction quality: mean abs % error
            errs = []
            for p in pts:
                pred = rc.predict(p["wob"], p["rpm"], p["depth"], p["mw"])
                if p["rop_actual"] > 0:
                    errs.append(abs(pred - p["rop_actual"]) /
                                p["rop_actual"] * 100)
            mape = sum(errs) / len(errs) if errs else 0.0
            self.result = {
                "k": k, "a": 1.0, "b": 0.6, "c": 0.00005, "d": -0.05,
                "mw_opt_ppg": 10.0, "n_points": len(pts),
                "mape_pct": round(mape, 1),
                "points": pts,
            }
            self.fit_label.setText(
                f"✅ Fitted: K = {k:.4g} from {len(pts)} points | "
                f"mean abs. error ≈ {mape:.1f}% | "
                f"ROP(25 klbf, 100 rpm, 8,000 ft, 11.5 ppg) = "
                f"{rc.predict(25, 100, 8000, 11.5):.1f} ft/hr")
            self.btn_ok.setEnabled(True)
        except Exception as e:
            QMessageBox.critical(self, "Fit Failed", str(e))

    def _accept(self):
        if self.result:
            self.accept()


class _InputsPage(QWizardPage):
    """Page 2: dynamic input form for the selected template."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("4. Provide the inputs")
        self.setSubTitle("Fill the requested fields. Empty fields will appear "
                         "as [To Be Filled] in the generated document.")

        lay = QVBoxLayout(self)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        container = QWidget()
        self.form_lay = QVBoxLayout(container)
        self.form_lay.setSpacing(10)
        scroll.setWidget(container)
        lay.addWidget(scroll)

        self.widgets: Dict[str, QWidget] = {}
        self._specs: Dict[str, InputSpec] = {}
        self._groups: Dict[str, QFormLayout] = {}
        self._group_widgets: Dict[str, QGroupBox] = {}
        self.web_notes: str = ""

    def _ensure_group(self, name: str) -> QFormLayout:
        if name in self._groups:
            return self._groups[name]
        gb = QGroupBox(name)
        fl = QFormLayout(gb)
        fl.setSpacing(7)
        self._groups[name] = fl
        self._group_widgets[name] = gb
        self.form_lay.addWidget(gb)
        return fl

    def initializePage(self):
        # clear previous
        for w in list(self.widgets.values()):
            try:
                w.setParent(None)
                w.deleteLater()
            except Exception:
                pass
        self.widgets.clear()
        self._specs.clear()
        self._groups.clear()
        self._group_widgets.clear()
        for gb in list(self._group_widgets.values()):
            gb.deleteLater()
        self._group_widgets.clear()

        tdef = self._selected_template()
        if tdef is None:
            return
        for spec in tdef.inputs:
            fl = self._ensure_group(spec.group)
            w = _build_field(spec)
            label = spec.label
            if spec.required:
                label = "* " + label
            fl.addRow(label, w)
            self.widgets[spec.key] = w
            self._specs[spec.key] = spec

        # Engineering Basis — optional canonical inputs so that EVERY
        # document type can run validation, readiness, the calculation
        # register and the deep engineering checks (audit P1).  Only keys
        # the template does not already ask for are added.
        basis_specs = [
            InputSpec("mud_weight", "Mud Weight", "number", unit="ppg",
                      placeholder="e.g. 12.0", group="Engineering Basis"),
            InputSpec("formation_pressure", "Pore Pressure (EMW)",
                      "number", unit="ppg",
                      placeholder="e.g. 11.0", group="Engineering Basis"),
            InputSpec("fracture_gradient", "Fracture Gradient", "number",
                      unit="ppg", placeholder="e.g. 16.0",
                      group="Engineering Basis"),
            InputSpec("hole_size", "Hole Size", "number", unit="in",
                      placeholder="e.g. 12.25", group="Engineering Basis"),
            InputSpec("casing_size", "Casing Size", "number", unit="in",
                      placeholder="e.g. 9.625", group="Engineering Basis"),
            InputSpec("casing_depth", "Casing Depth", "number", unit="ft",
                      placeholder="e.g. 8000", group="Engineering Basis"),
            InputSpec("total_depth", "Total Depth", "number", unit="ft",
                      placeholder="e.g. 10000", group="Engineering Basis"),
            InputSpec("bop_wp", "BOP Working Pressure", "number",
                      unit="psi", placeholder="e.g. 10000",
                      group="Engineering Basis"),
            InputSpec("kick_tolerance", "Kick Tolerance", "number",
                      unit="ppg", group="Engineering Basis"),
            InputSpec("h2s_plan", "H2S Contingency Plan", "text",
                      placeholder="e.g. H2S detection, PPE, drills",
                      group="Engineering Basis"),
            InputSpec("acceptance_criteria", "Acceptance Criteria", "text",
                      placeholder="e.g. LOT/FIT, wellhead test criteria",
                      group="Engineering Basis"),
            InputSpec("risk_assessment", "Risk Assessment Status", "text",
                      placeholder="e.g. Completed — matrix attached",
                      group="Engineering Basis"),
            InputSpec("flow_rate", "Flow / Pump Rate", "number", unit="gpm",
                      group="Engineering Basis"),
            InputSpec("yield_point", "Mud Yield Point", "number",
                      unit="lb/100ft²", group="Engineering Basis"),
            InputSpec("plastic_viscosity", "Mud Plastic Viscosity",
                      "number", unit="cP", group="Engineering Basis"),
            InputSpec("trip_speed", "Trip Speed", "number", unit="ft/min",
                      group="Engineering Basis"),
            InputSpec("dp_id", "Drill Pipe ID", "number", unit="in",
                      placeholder="e.g. 4.276",
                      group="Engineering Basis"),
            InputSpec("tfa", "Bit Total Nozzle Area (TFA)", "number",
                      unit="in²", placeholder="e.g. 0.3312 (3×12/32)",
                      group="Engineering Basis"),
            InputSpec("surface_type", "Surface Equipment Type", "combo",
                      options=["Type 1 (simple)", "Type 2 (standard)",
                               "Type 3 (long)", "Type 4 (extended)"],
                      group="Engineering Basis"),
            InputSpec("bha_od", "BHA / Drill Collar OD", "number",
                      unit="in", group="Engineering Basis"),
            InputSpec("bha_length", "BHA / Drill Collar Length", "number",
                      unit="ft", group="Engineering Basis"),
            InputSpec("wob", "Weight on Bit", "number", unit="klbf",
                      group="Engineering Basis"),
            InputSpec("rpm", "Rotary Speed", "number", unit="rpm",
                      group="Engineering Basis"),
            InputSpec("sicp", "Shut-In Casing Pressure (SICP)", "number",
                      unit="psi", group="Engineering Basis"),
            InputSpec("slow_pump_pressure", "Slow Pump Rate Pressure (SPR)",
                      "number", unit="psi", group="Engineering Basis"),
            InputSpec("pump_output", "Pump Output", "number",
                      unit="bbl/stk", placeholder="e.g. 0.1",
                      group="Engineering Basis"),
            InputSpec("kill_method", "Preferred Kill Method", "combo",
                      options=["", "Wait and Weight", "Driller's Method",
                               "Bullheading"],
                      group="Engineering Basis"),
            InputSpec("sigma_v_grad", "Overburden Gradient (σv)",
                      "number", unit="psi/ft", placeholder="e.g. 1.0",
                      group="Engineering Basis"),
            InputSpec("sH_sv_ratio", "σH / σv ratio", "number",
                      placeholder="e.g. 0.95", group="Engineering Basis"),
            InputSpec("sh_sv_ratio", "σh / σv ratio", "number",
                      placeholder="e.g. 0.85", group="Engineering Basis"),
            InputSpec("ucs_psi", "Rock UCS (unconfined strength)",
                      "number", unit="psi", placeholder="e.g. 8000",
                      group="Engineering Basis"),
            InputSpec("friction_angle", "Rock Friction Angle", "number",
                      unit="°", placeholder="e.g. 30",
                      group="Engineering Basis"),
            InputSpec("tensile_strength", "Rock Tensile Strength",
                      "number", unit="psi", placeholder="e.g. 500",
                      group="Engineering Basis"),
            InputSpec("lot_pressure", "LOT / FIT Pressure", "number",
                      unit="psi", placeholder="e.g. 1400",
                      group="Engineering Basis"),
            InputSpec("cemented_length", "Cemented Interval Length",
                      "number", unit="ft", group="Engineering Basis"),
            InputSpec("lead_yield", "Slurry Yield", "number",
                      unit="ft³/sk", placeholder="e.g. 1.18",
                      group="Engineering Basis"),
            InputSpec("water_per_sack", "Mix Water per Sack", "number",
                      unit="gal/sk", placeholder="e.g. 5.2",
                      group="Engineering Basis"),
            InputSpec("fluid_loss", "Slurry Fluid Loss", "number",
                      unit="ml/30min", placeholder="e.g. 100",
                      group="Engineering Basis"),
            InputSpec("static_time", "Static Time Before Gelation",
                      "number", unit="h", placeholder="e.g. 4",
                      group="Engineering Basis"),
            InputSpec("max_temperature", "Max Exposure Temperature",
                      "number", unit="°F", placeholder="e.g. 350",
                      group="Engineering Basis"),
            InputSpec("co2_pct", "CO₂ Content", "number", unit="%",
                      group="Engineering Basis"),
            InputSpec("packer_set", "Completion — Packer Set & Tested?",
                      "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("tree_ok", "Completion — X-mas Tree Tested?",
                      "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("trsv_ok", "Completion — TRSV Function-Tested?",
                      "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("cement_verified", "Completion — Cement Verified "
                      "(CBL/VDL)?", "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("lot_type", "Test Type", "combo",
                      options=["LOT", "FIT"], group="Engineering Basis"),
            InputSpec("can_rotate", "Stuck-pipe symptoms — can rotate?",
                      "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("can_circulate",
                      "Stuck-pipe symptoms — can circulate?",
                      "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("can_move_pipe",
                      "Stuck-pipe symptoms — can move pipe?",
                      "combo", options=["", "Yes", "No"],
                      group="Engineering Basis"),
            InputSpec("casing_weight", "Casing Weight", "number",
                      unit="ppf", group="Engineering Basis"),
            InputSpec("temperature_change", "Temperature Change (ΔT)",
                      "number", unit="°F",
                      placeholder="e.g. 150 (thermal stress check)",
                      group="Engineering Basis"),
            InputSpec("wear_fraction", "Casing Wear", "number",
                      unit="fraction (0–1)", placeholder="e.g. 0.1",
                      group="Engineering Basis"),
            InputSpec("corrosion_allowance", "Corrosion Allowance",
                      "number", unit="in", placeholder="e.g. 0.02",
                      group="Engineering Basis"),
            InputSpec("trajectory_table", "Well Trajectory (MD | Inc | Az)",
                      "table", columns=["MD (ft)", "Inc (°)", "Az (°)"],
                      group="Engineering Basis"),
            InputSpec("offset_trajectory_table",
                      "Offset Well Trajectory + Surface Offset (for "
                      "anti-collision)", "table",
                      columns=["MD (ft)", "Inc (°)", "Az (°)",
                               "N0 (ft)", "E0 (ft)"],
                      group="Engineering Basis"),
            InputSpec("total_cost", "Estimated Total Cost", "number",
                      placeholder="e.g. 12000000", group="Engineering Basis"),
        ]
        missing = [s for s in basis_specs if s.key not in self.widgets]
        if missing:
            bfl = self._ensure_group(
                "⚙️  Engineering Basis (enables validation, calculation "
                "register & deep checks)")
            for spec in missing:
                w = _build_field(spec)
                bfl.addRow(spec.label, w)
                self.widgets[spec.key] = w
                self._specs[spec.key] = spec

        # Auto-prefill from the Well Profile page (roadmap: "Prefill
        # inputs from the well profile") — profile fields are mapped onto
        # template inputs (semantic aliases: well_type -> well_profile).
        try:
            wiz = self.wizard()
            profile_page = wiz.page(1)
            if hasattr(profile_page, "profile"):
                prof = profile_page.profile()
                profile_aliases = {
                    "well_type": ("well_type", "well_profile"),
                    "environment": ("environment",),
                    "operation": ("operation",),
                    "holes": ("hole_size",),
                }
                # semantic aliases: profile taxonomy -> template taxonomy
                shape_aliases = {
                    "Vertical": ("Vertical",),
                    "Deviated": ("Deviated", "Directional J-Type",
                                 "Directional", "J-Shape"),
                    "Horizontal": ("Horizontal", "L-Shape"),
                    "ERD": ("ERD", "Extended Reach"),
                    "HPHT": ("HPHT",),
                    "Multi-lateral": ("Multi-lateral", "Multilateral"),
                    "Drilling": ("Drilling",),
                    "Workover": ("Workover",),
                    "Onshore": ("Onshore", "Land"),
                    "Offshore": ("Offshore", "Jack-up", "Semi-submersible",
                                 "Drillship"),
                }
                for pkey, val in prof.items():
                    if not val:
                        continue
                    candidates = (val,) + shape_aliases.get(val, ())
                    for tkey in profile_aliases.get(pkey, (pkey,)):
                        w = self.widgets.get(tkey)
                        if w is None:
                            continue
                        if isinstance(w, QLineEdit) and \
                                not w.text().strip():
                            w.setText(val)
                        elif isinstance(w, QComboBox):
                            # widgets are rebuilt on every page entry, so
                            # the combo sits at its default index — prefill
                            # only then, never clobber a user choice
                            for cand in candidates:
                                idx = w.findText(cand)
                                if idx >= 0 and w.currentIndex() == 0:
                                    w.setCurrentIndex(idx)
                                    break
                        elif isinstance(w, QTextEdit) and \
                                not w.toPlainText().strip():
                            w.setPlainText(val)
        except Exception:
            pass

        # Web research tool for field/formation introduction
        web_row = QHBoxLayout()
        self.btn_web = QPushButton("🌐  Web Research — Field / Formation Introduction")
        self.btn_web.setMaximumWidth(360)
        self.btn_web.clicked.connect(self._open_web)
        web_row.addWidget(self.btn_web)
        web_row.addStretch()
        self.form_lay.addLayout(web_row)

        # Field-knowledge enrichment (internal library + ML retrieval)
        kn_group = QGroupBox("📚  Field Knowledge & Intelligence (ML)")
        kn_lay = QVBoxLayout(kn_group)
        self.chk_enrich = QCheckBox(
            "Enrich the output with proven content from the real operations "
            "library (checklists, steps, procedures) — company names removed")
        self.chk_enrich.setChecked(True)
        kn_lay.addWidget(self.chk_enrich)
        krow = QHBoxLayout()
        krow.addWidget(QLabel("Enrichment level:"))
        self.enrich_level = QComboBox()
        self.enrich_level.addItems(["Moderate (recommended)", "Brief", "Detailed"])
        krow.addWidget(self.enrich_level)
        krow.addStretch()
        kn_lay.addLayout(krow)
        knote = QLabel(
            "🧠 Retrieval engine: TF-IDF (built-in ML) ranks the most relevant "
            "field content for this document type; semantic embeddings are "
            "used automatically when 'sentence-transformers' is installed.")
        knote.setWordWrap(True)
        knote.setStyleSheet("color:#8a8a9a;font-size:10px;")
        kn_lay.addWidget(knote)

        self.chk_rope = QCheckBox(
            "Include industry-standard field checklists (Rig Operations "
            "Performance Execution — BHA, bit, casing, cementing, tripping, "
            "well control, stuck pipe, H2S, etc.)")
        self.chk_rope.setChecked(True)
        kn_lay.addWidget(self.chk_rope)

        llm_row = QHBoxLayout()
        self.btn_llm = QPushButton("🤖  LLM Settings (rewrite with AI)")
        self.btn_llm.setMaximumWidth(280)
        self.btn_llm.clicked.connect(self._open_llm)
        llm_row.addWidget(self.btn_llm)
        self.lbl_llm = QLabel("LLM: off")
        self.lbl_llm.setStyleSheet("color:#8a8a9a;font-size:10px;")
        llm_row.addWidget(self.lbl_llm)
        llm_row.addStretch()
        kn_lay.addLayout(llm_row)
        self.form_lay.addWidget(kn_group)

        # Deep Engineering Verification (ROP calibration, HB hydraulics,
        # triaxial casing, compressibility-aware surge/swab)
        deep_group = QGroupBox("🔬  Deep Engineering Verification")
        deep_lay = QVBoxLayout(deep_group)
        self.chk_deep = QCheckBox(
            "Include DEEP ENGINEERING VERIFICATION section — ROP model "
            "(Bourgoyne-Young with offset calibration), Herschel-Bulkley "
            "annular pressure loss, triaxial (von Mises) casing check, "
            "compressibility-aware surge/swab")
        self.chk_deep.setChecked(True)
        deep_lay.addWidget(self.chk_deep)
        rop_row = QHBoxLayout()
        self.btn_rop = QPushButton("🧮  ROP Calibration from Offset Data…")
        self.btn_rop.setMaximumWidth(320)
        self.btn_rop.clicked.connect(self._open_rop_calib)
        rop_row.addWidget(self.btn_rop)
        self.lbl_rop = QLabel("No calibration yet — ROP section will list "
                              "the model with a data request.")
        self.lbl_rop.setStyleSheet("color:#8a8a9a;font-size:10px;")
        self.lbl_rop.setWordWrap(True)
        rop_row.addWidget(self.lbl_rop, 1)
        deep_lay.addLayout(rop_row)
        self._rop_calib: Optional[dict] = None
        self.form_lay.addWidget(deep_group)
        self.form_lay.addStretch()

    def _open_rop_calib(self):
        dlg = ROPCalibrationDialog(self, current=self._rop_calib)
        if dlg.exec() == QDialog.Accepted and dlg.result:
            self._rop_calib = dlg.result
            n = self._rop_calib["n_points"]
            mape = self._rop_calib.get("mape_pct", 0.0)
            self.lbl_rop.setText(
                f"✔ Calibrated from {n} offset point(s), K = "
                f"{self._rop_calib['k']:.4g}, mean abs. error ≈ {mape}% — "
                f"the ROP prediction table will be included in the document.")
        else:
            self.lbl_rop.setText(
                "Calibration cancelled — ROP section will list the model "
                "with a data request.")

    def _open_llm(self):
        try:
            from wizard_llm import LLMSettingsDialog, set_backend
            dlg = LLMSettingsDialog(self)
            if dlg.exec() == QDialog.Accepted:
                backend, key = dlg.get_values()
                set_backend(backend, key)
                self.lbl_llm.setText(
                    f"LLM: {backend if backend != 'none' else 'off'}")
        except Exception as e:
            QMessageBox.warning(self, "LLM Settings", str(e))

    def _open_web(self):
        """Open the web research dialog and store inserted markdown."""
        try:
            from wizard_web import run_web_research
            tdef = self._selected_template()
            key = tdef.key if tdef else ""
            text = run_web_research(key, self)
            if text:
                self.web_notes = text
                QMessageBox.information(
                    self, "Inserted",
                    "The research text will be added to the document as a "
                    "'Field & Formation Introduction (Web Research)' section "
                    "with source links.")
        except Exception as e:
            QMessageBox.warning(self, "Web Research", f"Not available: {e}")

    def _selected_template(self) -> Optional[TemplateDef]:
        wiz = self.wizard()
        page0 = wiz.page(0)
        key = page0.selected_key() if hasattr(page0, "selected_key") else ""
        if hasattr(page0, "_by_key"):
            return page0._by_key(key)
        return None

    def validatePage(self) -> bool:
        missing = []
        for key, spec in self._specs.items():
            if spec.required and not _get_field_value(spec, self.widgets[key]):
                missing.append(spec.label)
        if missing:
            QMessageBox.warning(
                self, "Required Inputs",
                "Please fill the required fields:\n\n• " + "\n• ".join(missing))
            return False
        return True

    def values(self) -> Dict[str, str]:
        out = {}
        for key, spec in self._specs.items():
            val = _get_field_value(spec, self.widgets[key])
            if spec.type == "table" and val:
                val = table_to_md_rows(val, spec.columns)
            out[key] = val
        if self.web_notes:
            out["web_notes"] = self.web_notes
        return out


class _RiskReviewPage(QWizardPage):
    """Page 3: automatic risk review of the document before output."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("5. Risk Review (automatic)")
        self.setSubTitle(
            "The software automatically checks the operations in your "
            "document against the drilling risk knowledge base, shows the "
            "key risks and asks you a few confirmation questions. The "
            "result is added to the document before final output.")

        lay = QVBoxLayout(self)

        # Analyze button + summary
        row = QHBoxLayout()
        self.btn_analyze = QPushButton("🔍  Analyze the Document for Risks")
        self.btn_analyze.setMinimumHeight(38)
        self.btn_analyze.clicked.connect(self._analyze)
        row.addWidget(self.btn_analyze)

        self.summary = QLabel("Not analyzed yet.")
        self.summary.setWordWrap(True)
        self.summary.setStyleSheet("color:#4fc3f7;font-weight:bold;")
        row.addWidget(self.summary, 1)
        lay.addLayout(row)

        # Risks table
        lay.addWidget(QLabel("Key Risks Identified (Critical / High):"))
        self.risk_table = QTableWidget()
        self.risk_table.setColumnCount(5)
        self.risk_table.setHorizontalHeaderLabels(
            ["Risk", "Severity", "Probability", "NPT (hrs)", "Primary Mitigation"])
        self.risk_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.risk_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.Stretch)
        self.risk_table.setAlternatingRowColors(True)
        self.risk_table.setMinimumHeight(150)
        lay.addWidget(self.risk_table)

        # Questions
        self.q_label = QLabel("Confirmation questions — tick the ones already "
                              "addressed in the program:")
        self.q_label.setWordWrap(True)
        lay.addWidget(self.q_label)

        self.q_list = QListWidget()
        self.q_list.setMinimumHeight(120)
        lay.addWidget(self.q_list)

        # Forgotten items
        self.f_label = QLabel("")
        self.f_label.setWordWrap(True)
        lay.addWidget(self.f_label)

        # Include toggle
        self.chk_include = QCheckBox(
            "✅ Include the Risk Assessment & Contingency Plan section in the "
            "generated document")
        self.chk_include.setChecked(True)
        lay.addWidget(self.chk_include)

        self._results: Optional[Dict] = None
        self._questions: List[str] = []

    # -- analysis ----------------------------------------------------------
    def _selected_template(self):
        wiz = self.wizard()
        page0 = wiz.page(0)
        key = page0.selected_key() if hasattr(page0, "selected_key") else ""
        if hasattr(page0, "_by_key"):
            return page0._by_key(key)
        return None

    def _analyze(self):
        try:
            from wizard_risk import (build_sequence_from_document,
                                     run_risk_analysis,
                                     generate_risk_questions)
        except Exception as e:
            QMessageBox.warning(self, "Risk Module", f"Not available: {e}")
            return

        wiz = self.wizard()
        tdef = self._selected_template()
        if tdef is None:
            return
        sec_page = wiz.page(2)
        in_page = wiz.page(3)

        values = dict(in_page.values()) if hasattr(in_page, "values") else {}
        selected = sec_page.selected_heads() if hasattr(sec_page, "selected_heads") else []

        # Build the (filled, section-selected) markdown
        from wizard_engine import fill_template, render_selected
        md = fill_template(tdef, values)
        md = render_selected(md, selected)

        sequence = build_sequence_from_document(md, values)
        results = run_risk_analysis(sequence)

        self._results = results
        self._sequence = sequence
        self._values = values

        sev = results.get("severity_counts", {})
        crit = sev.get(RiskLevel.CRITICAL, 0)
        high = sev.get(RiskLevel.HIGH, 0)
        med = sev.get(RiskLevel.MEDIUM, 0)
        npt = results.get("total_expected_npt", 0)
        self.summary.setText(
            f"✓ {len(results.get('risks', []))} risks identified — "
            f"{crit} critical, {high} high, {med} medium | "
            f"expected NPT {npt:,.0f} hrs")

        # Risk table
        risks = results.get("risks", [])
        top = [r for r in risks
               if r.severity in (RiskLevel.CRITICAL, RiskLevel.HIGH)][:15]
        self.risk_table.setRowCount(len(top))
        for row, r in enumerate(top):
            plan = r.contingency_plans[0].action if r.contingency_plans else "—"
            self.risk_table.setItem(row, 0, QTableWidgetItem(r.problem))
            self.risk_table.setItem(row, 1, QTableWidgetItem(r.severity.value))
            self.risk_table.setItem(row, 2, QTableWidgetItem(f"{r.probability:.0%}"))
            self.risk_table.setItem(row, 3, QTableWidgetItem(f"{r.npt_hours:.0f}"))
            self.risk_table.setItem(row, 4, QTableWidgetItem(plan))
        self.risk_table.resizeRowsToContents()

        # Questions
        self._questions = generate_risk_questions(results, values)
        self.q_list.clear()
        for q in self._questions:
            item = QListWidgetItem(q)
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(Qt.Checked)
            item.setSizeHint(QSize(0, 40))
            self.q_list.addItem(item)

        # Forgotten items
        forgotten = results.get("forgotten_items", [])
        if forgotten:
            self.f_label.setText(
                f"☑️ {len(forgotten)} best-practice reminders will be added "
                f"to the risk section (first {min(len(forgotten), 12)}).")
        else:
            self.f_label.setText("")

    def answers(self) -> Dict[str, bool]:
        out = {}
        for i in range(self.q_list.count()):
            item = self.q_list.item(i)
            out[item.text()] = (item.checkState() == Qt.Checked)
        return out

    def initializePage(self):
        # auto-analyze once when the page is entered
        if self._results is None:
            QApplication.setOverrideCursor(Qt.WaitCursor)
            try:
                self._analyze()
            finally:
                QApplication.restoreOverrideCursor()

    def validatePage(self) -> bool:
        return True


class _OptionsPage(QWizardPage):
    """Page 4: document metadata, formatting & output location."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("6. Document details, formatting & output")
        self.setSubTitle(
            "Set the document control info, choose the font and page "
            "layout, then pick where to save the Word file.")

        # Scrollable content area (page holds a lot of settings)
        scroll = QScrollArea(self)
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        container = QWidget()
        lay = QVBoxLayout(container)
        lay.setContentsMargins(4, 4, 4, 4)
        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.addWidget(scroll)
        scroll.setWidget(container)

        # -- Document control -------------------------------------------------
        g_doc = QGroupBox("Document Control")
        fl = QFormLayout(g_doc)
        fl.setSpacing(6)
        self.prepared = QLineEdit()
        self.reviewed = QLineEdit()
        self.approved = QLineEdit()
        self.revision = QLineEdit("01")
        self.date = QLineEdit(datetime.now().strftime("%d-%B-%Y"))
        self.doc_number = QLineEdit()
        self.doc_number.setPlaceholderText("e.g. DRL-PRG-2026-001 (optional)")

        fl.addRow("Prepared By:", self.prepared)
        fl.addRow("Reviewed By:", self.reviewed)
        fl.addRow("Approved By:", self.approved)
        fl.addRow("Revision:", self.revision)
        fl.addRow("Date:", self.date)
        fl.addRow("Document Number:", self.doc_number)
        lay.addWidget(g_doc)

        # -- Formatting -------------------------------------------------------
        g_fmt = QGroupBox("Formatting & Page Layout")
        ff = QFormLayout(g_fmt)
        ff.setSpacing(6)

        self.font = QComboBox()
        self.font.addItems(["Calibri", "Arial", "Times New Roman", "Cambria",
                            "Georgia", "Segoe UI", "Consolas"])
        self.font_size = QSpinBox()
        self.font_size.setRange(9, 14)
        self.font_size.setValue(11)
        ff.addRow("Font:", self.font)
        ff.addRow("Font Size:", self.font_size)

        self.page_size = QComboBox()
        self.page_size.addItems(["A4", "Letter"])
        self.orientation = QComboBox()
        self.orientation.addItems(["Portrait", "Landscape"])
        ff.addRow("Page Size:", self.page_size)
        ff.addRow("Orientation:", self.orientation)

        mrow = QHBoxLayout()
        self.m_left = QDoubleSpinBox(); self.m_left.setRange(0.5, 6.0); self.m_left.setValue(2.5)
        self.m_right = QDoubleSpinBox(); self.m_right.setRange(0.5, 6.0); self.m_right.setValue(2.0)
        self.m_top = QDoubleSpinBox(); self.m_top.setRange(0.5, 6.0); self.m_top.setValue(2.0)
        self.m_bottom = QDoubleSpinBox(); self.m_bottom.setRange(0.5, 6.0); self.m_bottom.setValue(2.0)
        for w in (self.m_left, self.m_right, self.m_top, self.m_bottom):
            w.setSuffix(" cm"); w.setDecimals(1)
        mrow.addWidget(QLabel("L:")); mrow.addWidget(self.m_left)
        mrow.addWidget(QLabel("R:")); mrow.addWidget(self.m_right)
        mrow.addWidget(QLabel("T:")); mrow.addWidget(self.m_top)
        mrow.addWidget(QLabel("B:")); mrow.addWidget(self.m_bottom)
        mrow.addStretch()
        ff.addRow("Margins:", mrow)

        self.cover = QCheckBox(); self.cover.setChecked(True)
        self.toc = QCheckBox(); self.toc.setChecked(True)
        ff.addRow("Cover page:", self.cover)
        ff.addRow("Table of contents:", self.toc)

        self.header_text = QLineEdit()
        self.header_text.setPlaceholderText("e.g. DRL-PRG-2026-001 | Rev 01")
        self.footer_text = QLineEdit()
        self.footer_text.setPlaceholderText("e.g. Company Confidential")
        ff.addRow("Header text:", self.header_text)
        ff.addRow("Footer text:", self.footer_text)
        lay.addWidget(g_fmt)

        # -- Cost & Pricing (CBS) ---------------------------------------------
        g_cbs = QGroupBox("💰 Cost & Pricing (CBS)")
        cv = QVBoxLayout(g_cbs)
        row_cbs = QHBoxLayout()
        self.chk_cbs = QCheckBox(
            "Include COST BREAKDOWN STRUCTURE / AFE section")
        self.chk_cbs.setChecked(False)
        self.chk_cbs.toggled.connect(
            lambda on: self.btn_cbs_sel.setEnabled(on))
        row_cbs.addWidget(self.chk_cbs)
        row_cbs.addStretch(1)
        self.btn_cbs_sel = QPushButton("📋 Select Goods & Services…")
        self.btn_cbs_sel.setEnabled(False)
        self.btn_cbs_sel.clicked.connect(self._open_cbs_selection)
        row_cbs.addWidget(self.btn_cbs_sel)
        cv.addLayout(row_cbs)
        self.lbl_cbs_sel = QLabel(
            "No items selected — the full catalog will be used "
            "(with quantities from the '💰 Cost & Pricing' tab).")
        self.lbl_cbs_sel.setWordWrap(True)
        self.lbl_cbs_sel.setStyleSheet("color: #8a8a9a; font-size: 10px;")
        cv.addWidget(self.lbl_cbs_sel)
        self._cbs_selection: Optional[list] = None
        lay.addWidget(g_cbs)

        # -- Drilling Problems (prevention & response) -----------------------
        g_prob = QGroupBox("🛟 Drilling Problems — Prevention & Response")
        pv = QVBoxLayout(g_prob)
        prow = QHBoxLayout()
        self.chk_problems = QCheckBox(
            "Include DRILLING PROBLEM PREVENTION & RESPONSE section")
        self.chk_problems.setChecked(False)
        self.chk_problems.toggled.connect(
            lambda on: self.btn_problems_sel.setEnabled(on))
        prow.addWidget(self.chk_problems)
        prow.addStretch(1)
        self.btn_problems_sel = QPushButton("🛟 Select Problems…")
        self.btn_problems_sel.setEnabled(False)
        self.btn_problems_sel.clicked.connect(self._open_problems_selection)
        prow.addWidget(self.btn_problems_sel)
        pv.addLayout(prow)
        self.lbl_problems_sel = QLabel(
            "No problems selected — the section will be skipped.")
        self.lbl_problems_sel.setWordWrap(True)
        self.lbl_problems_sel.setStyleSheet("color: #8a8a9a; font-size: 10px;")
        pv.addWidget(self.lbl_problems_sel)
        self._problems_selection: Optional[list] = None
        lay.addWidget(g_prob)

        # -- Output -----------------------------------------------------------
        out_group = QGroupBox("Output File")
        ol = QVBoxLayout(out_group)
        row = QHBoxLayout()
        self.path = QLineEdit()
        self.path.setPlaceholderText("Select output folder...")
        btn = QPushButton("📁 Browse...")
        btn.clicked.connect(self._browse)
        row.addWidget(self.path)
        row.addWidget(btn)
        ol.addLayout(row)
        self.fname = QLineEdit()
        self.fname.setPlaceholderText("File name (auto-filled)")
        ol.addWidget(self.fname)
        lay.addWidget(out_group)

        self.registerField("preparedBy", self.prepared)
        self.registerField("reviewedBy", self.reviewed)
        self.registerField("approvedBy", self.approved)
        self.registerField("revision", self.revision)
        self.registerField("docDate", self.date)

    def _open_problems_selection(self):
        """باز کردن دیالوگ انتخاب مشکلات حفاری"""
        try:
            from drilling_problems_ui import ProblemsDialog
            dlg = ProblemsDialog(self)
            if dlg.exec() and dlg.selection:
                self._problems_selection = dlg.selection
                n = len(dlg.selection)
                crit = sum(1 for p in dlg.selection
                           if p.severity in ("Critical", "High"))
                self.lbl_problems_sel.setText(
                    f"✔ {n} problem(s) selected ({crit} Critical/High). "
                    f"Open again to change.")
            else:
                self._problems_selection = None
                self.lbl_problems_sel.setText(
                    "Selection cancelled — the section will be skipped.")
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Problems Selection", str(e))

    def _open_cbs_selection(self):
        """باز کردن دیالوگ انتخاب کالا/سرویس برای سکشن CBS"""
        try:
            from cbs_ui import CbsSelectionDialog
            dlg = CbsSelectionDialog(self)
            if dlg.exec() and dlg.selection:
                self._cbs_selection = dlg.selection
                n = len(dlg.selection)
                total = sum(s["price"] * s["qty"] for s in dlg.selection)
                self.lbl_cbs_sel.setText(
                    f"✔ {n} item(s) selected — estimated "
                    f"{total:,.2f} (excl. day-rates/contingency). "
                    f"Open again to change.")
            else:
                self._cbs_selection = None
                self.lbl_cbs_sel.setText(
                    "Selection cancelled — the full catalog will be used.")
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "CBS Selection", str(e))

    def output_options(self) -> dict:
        return {
            "font": self.font.currentText(),
            "font_size": float(self.font_size.value()),
            "page": self.page_size.currentText(),
            "orientation": self.orientation.currentText(),
            "margin_left": self.m_left.value(),
            "margin_right": self.m_right.value(),
            "margin_top": self.m_top.value(),
            "margin_bottom": self.m_bottom.value(),
            "cover": self.cover.isChecked(),
            "toc": self.toc.isChecked(),
            "header_text": self.header_text.text().strip(),
            "footer_text": self.footer_text.text().strip(),
            "document_number": self.doc_number.text().strip(),
            "include_cbs": self.chk_cbs.isChecked(),
            "cbs_selection": self._cbs_selection,
            "include_problems": self.chk_problems.isChecked(),
            "problems_selection": self._problems_selection,
        }

    def initializePage(self):
        tdef = self._selected_template()
        if tdef:
            base = tdef.name.replace(" ", "_").replace("/", "-")
            self.fname.setText(f"{base}.docx")
        default_dir = str(Path.cwd() / "projects" / "exports")
        if not self.path.text():
            self.path.setText(default_dir)

    def _selected_template(self):
        wiz = self.wizard()
        page0 = wiz.page(0)
        key = page0.selected_key() if hasattr(page0, "selected_key") else ""
        if hasattr(page0, "_by_key"):
            return page0._by_key(key)
        return None

    def _browse(self):
        d = QFileDialog.getExistingDirectory(self, "Select Output Folder",
                                              self.path.text() or str(Path.cwd()))
        if d:
            self.path.setText(d)

    def output_path(self) -> str:
        folder = self.path.text().strip() or str(Path.cwd())
        fname = self.fname.text().strip() or "Generated_Document.docx"
        if not fname.lower().endswith(".docx"):
            fname += ".docx"
        return str(Path(folder) / fname)


class _GeneratePage(QWizardPage):
    """Page 5: generate the document."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("7. Generate")
        self.setSubTitle("Generate the Word document now.")

        lay = QVBoxLayout(self)
        self.status = QLabel("Ready to generate.")
        self.status.setWordWrap(True)
        lay.addWidget(self.status)

        self.progress = QProgressBar()
        self.progress.setValue(0)
        lay.addWidget(self.progress)

        btns = QHBoxLayout()
        self.btn_gen = QPushButton("⚙️  GENERATE WORD DOCUMENT")
        self.btn_gen.setMinimumHeight(44)
        self.btn_gen.clicked.connect(self._generate)
        btns.addWidget(self.btn_gen)

        self.btn_open = QPushButton("📂 Open Document")
        self.btn_open.setEnabled(False)
        self.btn_open.clicked.connect(self._open_doc)
        btns.addWidget(self.btn_open)

        self.btn_folder = QPushButton("🗂️ Open Folder")
        self.btn_folder.setEnabled(False)
        self.btn_folder.clicked.connect(self._open_folder)
        btns.addWidget(self.btn_folder)
        lay.addLayout(btns)
        lay.addStretch()

        self._last_path = ""

    def initializePage(self):
        self._last_path = ""
        self.btn_open.setEnabled(False)
        self.btn_folder.setEnabled(False)
        self.progress.setValue(0)
        tdef = self._selected_template()
        if tdef:
            wiz = self.wizard()
            sec_page = wiz.page(2)
            n_sel = len(sec_page.selected_heads()) if hasattr(sec_page, "selected_heads") else 0
            self.status.setText(
                f"Document: <b>{tdef.icon} {tdef.name}</b><br>"
                f"Sections selected: <b>{n_sel}</b><br>"
                f"Press GENERATE to create the Word document.")

    def _selected_template(self):
        wiz = self.wizard()
        page0 = wiz.page(0)
        key = page0.selected_key() if hasattr(page0, "selected_key") else ""
        if hasattr(page0, "_by_key"):
            return page0._by_key(key)
        return None

    def _generate(self):
        tdef = self._selected_template()
        if tdef is None:
            QMessageBox.warning(self, "No Document Type",
                                "Select a document type first.")
            return

        wiz = self.wizard()
        sec_page = wiz.page(2)      # sections
        in_page = wiz.page(3)       # inputs
        opt_page = wiz.page(5)      # options

        values = dict(in_page.values()) if hasattr(in_page, "values") else {}

        # Document metadata
        values.setdefault("prepared_by", self.field("preparedBy") or "")
        values.setdefault("reviewed_by", self.field("reviewedBy") or "")
        values.setdefault("approved_by", self.field("approvedBy") or "")
        values.setdefault("revision", self.field("revision") or "01")
        values.setdefault("doc_date", self.field("docDate") or "")

        meta = {
            "date": values.get("doc_date", ""),
            "document_number": opt_page.output_options().get("document_number", ""),
            "revision": values.get("revision", ""),
            "prepared_by": values.get("prepared_by", ""),
            "reviewed_by": values.get("reviewed_by", ""),
            "approved_by": values.get("approved_by", ""),
            "title": neutralize_text(tdef.name,
                                     str(values.get("operator", "") or ""),
                                     str(values.get("contractor", "") or "")),
            "operator": str(values.get("operator", "") or
                            values.get("operator_name", "") or ""),
            "contractor": str(values.get("contractor", "") or
                              values.get("contractor_name", "") or ""),
        }

        self.progress.setValue(15)
        QApplication.processEvents()

        try:
            # 1) Fill placeholders with user values
            md = fill_template(tdef, values)
            self.progress.setValue(35)
            QApplication.processEvents()

            # 2) Keep only the user-selected sections
            selected = sec_page.selected_heads() if hasattr(sec_page, "selected_heads") else []
            md = render_selected(md, selected)

            # 2b) Well profile section (basis of design) from page 1
            try:
                profile_page = wiz.page(1)
                if hasattr(profile_page, "profile_markdown"):
                    pmd = profile_page.profile_markdown()
                    if pmd:
                        md = pmd + "\n\n---\n\n" + md
            except Exception:
                import traceback
                traceback.print_exc()

            # 3) Web research section (field/formation introduction)
            web_notes = values.get("web_notes", "")
            if web_notes:
                md += ("\n\n## FIELD & FORMATION INTRODUCTION (WEB RESEARCH)\n\n"
                       + web_notes + "\n")
            self.progress.setValue(55)
            QApplication.processEvents()

            # 4) Risk review section (automatic, before final output)
            risk_page = wiz.page(4)
            if (hasattr(risk_page, "_results") and risk_page._results
                    and getattr(risk_page, "chk_include", None)
                    and risk_page.chk_include.isChecked()):
                try:
                    from wizard_risk import build_risk_section_md
                    answers = risk_page.answers()
                    seq = getattr(risk_page, "_sequence", "")
                    risk_md = build_risk_section_md(
                        risk_page._results, answers, seq)
                    if risk_md:
                        md = md.rstrip() + "\n\n---\n\n" + risk_md
                except Exception as e:
                    import traceback
                    traceback.print_exc()
            # 4b) Field-knowledge enrichment (ML retrieval + optional LLM rewrite)
            try:
                if hasattr(in_page, "chk_enrich") and in_page.chk_enrich.isChecked():
                    from wizard_knowledge import (enrich_template,
                                                 get_chunks_for)
                    from wizard_llm import rewrite_chunks
                    level = {
                        "Brief": "brief",
                        "Moderate (recommended)": "moderate",
                        "Detailed": "detailed",
                    }.get(in_page.enrich_level.currentText(), "moderate")
                    op_name = meta.get("operator", "")
                    con_name = meta.get("contractor", "")
                    profile = {}
                    try:
                        profile_page = wiz.page(1)
                        if hasattr(profile_page, "profile"):
                            profile = profile_page.profile()
                    except Exception:
                        pass
                    chunks = get_chunks_for(tdef.key, level,
                                            profile=profile or None)
                    if chunks:
                        kn = rewrite_chunks(
                            chunks, tdef.name,
                            max_input_chars=4000 if level == "brief" else 5500,
                            operator_name=op_name, contractor_name=con_name)
                    else:
                        kn = ""
                    if kn:
                        if kn.startswith("*AI-assisted rewrite"):
                            kn = ("## FIELD KNOWLEDGE ENRICHMENT "
                                  "(AI-ASSISTED, FROM REAL OPERATIONS "
                                  "LIBRARY)\n\n" + kn)
                        else:
                            # LLM unavailable -> verbatim excerpts;
                            # the label must be honest about it
                            kn = ("## FIELD KNOWLEDGE ENRICHMENT "
                                  "(FROM REAL OPERATIONS LIBRARY)\n\n"
                                  "Excerpts below were retrieved by the "
                                  "ML engine from the internal "
                                  "field-document library (verbatim; "
                                  "company and well names removed). "
                                  "Enable the AI assistant in the "
                                  "wizard for a professional "
                                  "paraphrase.\n\n" + kn)
                    if kn:
                        md = md.rstrip() + "\n\n---\n\n" + kn
            except Exception:
                import traceback
                traceback.print_exc()
            self.progress.setValue(60)
            QApplication.processEvents()

            # 4c) ROPE field checklists
            try:
                if hasattr(in_page, "chk_rope") and in_page.chk_rope.isChecked():
                    from wizard_rope import get_rope_checklists
                    rope_md = get_rope_checklists(
                        tdef.key, level,
                        meta.get("operator", ""),
                        meta.get("contractor", ""))
                    if rope_md:
                        md = md.rstrip() + "\n\n---\n\n" + rope_md
            except Exception:
                import traceback
                traceback.print_exc()
            # 4d0) Drilling problems prevention & response section
            try:
                opts4d0 = opt_page.output_options()
                if opts4d0.get("include_problems") and \
                        opts4d0.get("problems_selection"):
                    from drilling_problems_db import build_problems_markdown
                    prob_md = build_problems_markdown(
                        opts4d0["problems_selection"],
                        meta.get("operator", ""))
                    if prob_md:
                        md = md.rstrip() + "\n\n---\n\n" + prob_md
            except Exception:
                import traceback
                traceback.print_exc()
            # 4d) CBS / AFE cost section (user-selected goods & services)
            try:
                opts4d = opt_page.output_options()
                if opts4d.get("include_cbs"):
                    from cbs_db import (CBSDatabase, CbsItem,
                                        build_cbs_markdown,
                                        get_time_breakdown_summary)
                    cdb = CBSDatabase()
                    sel = opts4d.get("cbs_selection")
                    if sel:
                        cbs_items = [CbsItem(
                            name=s["name"], unit=s["unit"],
                            unit_price=s["price"], qty=s["qty"])
                            for s in sel]
                    else:
                        cbs_items = cdb.get_items()
                    tb = get_time_breakdown_summary()
                    total_days = 0.0
                    try:
                        total_days = float(values.get("total_days") or 0)
                    except (TypeError, ValueError):
                        total_days = 0.0
                    if total_days <= 0:
                        # Batch X — cross-project guard: only fall back to
                        # the time-breakdown project when its well matches
                        # this document's well (or no well is named).
                        _doc_well = str(values.get("well_name") or
                                        "").strip().lower()
                        _tb_well = str(tb.get("well_name") or
                                       "").strip().lower()
                        if not _doc_well or not _tb_well or \
                                _doc_well == _tb_well:
                            total_days = tb.get("total_days", 0.0)
                    depth_key = ("target_depth" if values.get("target_depth")
                                 else "depth" if values.get("depth")
                                 else "depth_m")
                    try:
                        depth_m = float(values.get(depth_key) or 0)
                    except (TypeError, ValueError):
                        depth_m = 0.0
                    cbs_md = build_cbs_markdown(
                        cbs_items,
                        total_days=total_days,
                        well_depth_m=depth_m,
                        well_name=values.get("well_name", ""),
                        operator=meta.get("operator", ""),
                        currency=cdb.get_currency())
                    if cbs_md:
                        if not opts4d.get("cbs_selection"):
                            # Batch Y — full-catalog fallback must be
                            # declared in the document, never silent
                            cbs_md = ("> ⚠️ **No goods/services were "
                                      "selected by the user — the FULL "
                                      "catalog is shown below as a "
                                      "drafting aid. Prices are editable "
                                      "defaults and must be confirmed "
                                      "before issue as an AFE.**\n\n"
                                      + cbs_md)
                        md = md.rstrip() + "\n\n---\n\n" + cbs_md
            except Exception:
                import traceback
                traceback.print_exc()
            self.progress.setValue(65)
            QApplication.processEvents()

            # 5) Neutralize any hard-coded company/brand names
            md = neutralize_text(md, meta.get("operator", ""),
                                  meta.get("contractor", ""))

            # 5b) Engineering validation & compliance (P0 audit item)
            try:
                from validation_engine import (validate_well_data,
                                               findings_markdown,
                                               blocking_findings)
                from engineering_dependency import dependency_markdown
                findings = validate_well_data(values)
                crit = blocking_findings(findings)
                if crit:
                    # CRITICAL findings block export unless formally accepted
                    detail = "\n".join(
                        f"• [{f.level}] {f.code}: {f.message}" for f in crit)
                    ret = QMessageBox.warning(
                        self, "⚠️ CRITICAL Engineering Findings",
                        f"The design has {len(crit)} CRITICAL finding(s) "
                        f"that must be resolved before release:\n\n{detail}\n\n"
                        "Export anyway (with justification note in the "
                        "document)?",
                        QMessageBox.Yes | QMessageBox.No,
                        QMessageBox.No)
                    if ret != QMessageBox.Yes:
                        self.status.setText(
                            "❌ Generation blocked by CRITICAL validation "
                            "findings. Resolve them and regenerate.")
                        return
                    # record the override in the audit log
                    try:
                        from audit_log import log_action
                        log_action("validation_override", meta.get("operator"),
                                   values.get("well_name", ""),
                                   f"{len(crit)} CRITICAL overridden")
                    except Exception:
                        pass
                    md = (md.rstrip() + "\n\n---\n\n" +
                          findings_markdown(findings, meta.get("operator", "")) +
                          "\n*CRITICAL findings formally accepted on "
                          f"{values.get('doc_date') or 'generation date'}.*\n")
                else:
                    md = md.rstrip() + "\n\n---\n\n" + \
                        findings_markdown(findings, meta.get("operator", ""))

                # dependency impact note (only when meaningful inputs exist)
                dep_keys = [k for k in ("mud_weight", "hole_size",
                                        "casing_size", "casing_depth",
                                        "formation_pressure", "td_depth",
                                        "bop_wp", "h2s", "depth_m")
                            if values.get(k)]
                if dep_keys:
                    dmd = dependency_markdown(dep_keys)
                    if dmd:
                        md = md.rstrip() + "\n\n---\n\n" + dmd

                # program readiness score (audit P1 — completeness before
                # approval), always appended
                try:
                    from operations_engine import readiness_markdown as rm_
                    rmd_ = rm_(values, meta.get("operator", ""))
                    if rmd_:
                        md = md.rstrip() + "\n\n---\n\n" + rmd_
                except Exception:
                    pass
            except Exception:
                import traceback
                traceback.print_exc()

            # 5c) Planning intelligence: offsets, equipment compatibility,
            #     Monte-Carlo schedule/cost, risk decision matrix
            try:
                from planning_intelligence import (OffsetIntelligence,
                                                   equipment_compatibility,
                                                   compatibility_markdown,
                                                   monte_carlo_time,
                                                   monte_carlo_cost,
                                                   monte_carlo_markdown)
                from risk_decision import find_decisions, decision_markdown
                from document_compliance import (compliance_check,
                                                 compliance_markdown)

                # offset wells (field/type known?)
                if values.get("field_name"):
                    oi = OffsetIntelligence()
                    omd = oi.offset_markdown(
                        field=values.get("field_name", ""),
                        well_type=values.get("well_type", ""),
                        operation=values.get("operation", ""))
                    if omd:
                        md = md.rstrip() + "\n\n---\n\n" + omd

                # equipment compatibility (hole/casing sizes known?)
                if values.get("hole_size") or values.get("casing_size"):
                    comp = equipment_compatibility(
                        hole_size=values.get("hole_size", ""),
                        casing_size=values.get("casing_size", ""),
                        bha_od=values.get("bha_od", ""),
                        bit_size=values.get("bit_size", ""),
                        liner_size=values.get("liner_size", ""),
                        tubing_size=values.get("tubing_size", ""),
                        bop_wp_psi=float(values.get("bop_wp") or 0),
                        max_surface_pressure_psi=float(
                            values.get("masp") or 0),
                        mud_weight_ppg=float(values.get("mud_weight") or 0))
                    cmd = compatibility_markdown(comp)
                    if cmd and any(f["level"] != "INFO" for f in comp):
                        md = md.rstrip() + "\n\n---\n\n" + cmd

                # Monte Carlo schedule/cost (duration known?)
                days = float(values.get("total_days") or 0)
                if days > 0:
                    tr = monte_carlo_time(days, days * 0.85, days * 1.25)
                    cr = monte_carlo_cost(float(values.get("total_cost") or 0)
                                          or days * 100000)
                    mcm = monte_carlo_markdown(tr, cr)
                    if mcm:
                        md = md.rstrip() + "\n\n---\n\n" + mcm

                # risk decision matrix (risk page results available?)
                try:
                    rp = wiz.page(4)
                    if hasattr(rp, "_results") and rp._results:
                        risk_txt = " ".join(
                            r.problem for r in rp._results.get("risks", [])[:10])
                        dcs = find_decisions(risk_txt)
                        if dcs:
                            dmd2 = decision_markdown(dcs)
                            if dmd2:
                                md = md.rstrip() + "\n\n---\n\n" + dmd2
                except Exception:
                    pass

                # diagnostic decision trees: stuck pipe + fishing (audit P1)
                try:
                    from engineering_decisions import (stuck_pipe_markdown,
                                                       fishing_markdown)
                    dmd3 = stuck_pipe_markdown(values,
                                               meta.get("operator", ""))
                    if dmd3:
                        dmd3 = neutralize_text(dmd3,
                                               meta.get("operator", ""),
                                               meta.get("contractor", ""))
                        md = md.rstrip() + "\n\n---\n\n" + dmd3
                    dmd4 = fishing_markdown(values,
                                            meta.get("operator", ""))
                    if dmd4:
                        dmd4 = neutralize_text(dmd4,
                                               meta.get("operator", ""),
                                               meta.get("contractor", ""))
                        md = md.rstrip() + "\n\n---\n\n" + dmd4
                except Exception:
                    pass

                # standards compliance matrix (always)
                try:
                    from standards_engine import compliance_markdown as scm
                    smd = scm(values, meta.get("operator", ""))
                    if smd:
                        md = md.rstrip() + "\n\n---\n\n" + smd
                except Exception:
                    pass

                # document compliance report card (always)
                comp_rep = compliance_check(tdef.key, md)
                cmr = compliance_markdown(comp_rep, meta.get("operator", ""))
                if cmr:
                    md = md.rstrip() + "\n\n---\n\n" + cmr
            except Exception:
                import traceback
                traceback.print_exc()

            # 5d) Deep Engineering Verification (ROP calibration, HB
            #     hydraulics, triaxial, surge/swab) — audit P1
            try:
                if hasattr(in_page, "chk_deep") and in_page.chk_deep.isChecked():
                    from engineering_deep import deep_verify_markdown
                    rop_calib = getattr(in_page, "_rop_calib", None)
                    dmd = deep_verify_markdown(values, rop_calib,
                                               meta.get("operator", ""))
                    if dmd:
                        dmd = neutralize_text(dmd, meta.get("operator", ""),
                                              meta.get("contractor", ""))
                        md = md.rstrip() + "\n\n---\n\n" + dmd
            except Exception:
                import traceback
                traceback.print_exc()

            # 5e) Engineering Calculation Register — every number traceable
            #     to its formula + standard (buyer Q1, audit P1)
            try:
                from engineering_register import (compute_register,
                                                  register_markdown)
                rows = compute_register(values)
                rmd = register_markdown(rows, meta.get("operator", ""))
                if rmd:
                    rmd = neutralize_text(rmd, meta.get("operator", ""),
                                          meta.get("contractor", ""))
                    md = md.rstrip() + "\n\n---\n\n" + rmd
            except Exception:
                import traceback
                traceback.print_exc()

            # 5f) FINAL PLACEHOLDER AUDIT — no raw placeholder may reach
            #     the Word file (Batch X).  Any {{key}}/{key} left after
            #     every enrichment stage means the document is incomplete.
            unresolved = scan_unresolved_placeholders(md)
            if unresolved:
                QMessageBox.warning(
                    self, "⚠️ Unresolved Parameters",
                    "The document still contains unresolved parameter "
                    "placeholder(s):\n\n• " +
                    "\n• ".join(sorted(unresolved)) +
                    "\n\nThese are parameters the software could not fill "
                    "from your inputs. Fix the inputs and regenerate — "
                    "the export is blocked so no incomplete Word file is "
                    "produced.")
                self.status.setText(
                    "❌ Export blocked: unresolved parameter placeholders.")
                return

            # 6) Render to Word with the chosen formatting
            options = opt_page.output_options() if hasattr(opt_page, "output_options") else {}
            out = opt_page.output_path()
            Path(out).parent.mkdir(parents=True, exist_ok=True)
            ok = md_to_docx(md, out, meta, options)
            self.progress.setValue(100)
            self._last_path = out
            if ok:
                self.status.setText(
                    f"✅ <b>Document generated successfully!</b><br>"
                    f"📄 {out}")
                self.btn_open.setEnabled(True)
                self.btn_folder.setEnabled(True)
            else:
                self.status.setText("❌ Generation failed.")
        except Exception as e:
            import traceback
            traceback.print_exc()
            self.status.setText(f"❌ Error: {e}")
            QMessageBox.critical(self, "Generation Error", str(e))

    def _open_doc(self):
        if self._last_path and Path(self._last_path).exists():
            QDesktopServices.openUrl(QUrl.fromLocalFile(self._last_path))

    def _open_folder(self):
        if self._last_path:
            QDesktopServices.openUrl(
                QUrl.fromLocalFile(str(Path(self._last_path).parent)))



# ============================================================================
# WIZARD DARK THEME (consistent with the main app palette)
# ============================================================================

WIZARD_STYLE = """
QWizard {
    background-color: #16213e;
}
QWidget {
    background-color: #16213e;
    color: #e0e0e0;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 12px;
}
QWizard QLabel { color: #c0ccd8; }
QLabel#wizBig { color: #e94560; font-size: 15px; font-weight: bold; }
QLabel#wizSub { color: #8899aa; font-size: 11px; }

QGroupBox {
    border: 2px solid #0f3460;
    border-radius: 8px;
    margin-top: 12px;
    padding-top: 18px;
    font-weight: bold;
    color: #e94560;
    background: #1a1a2e;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 12px;
    padding: 0 8px;
    color: #e94560;
    font-size: 12px;
}

QLineEdit, QDoubleSpinBox, QSpinBox, QComboBox, QDateEdit, QTextEdit {
    background-color: #1a1a2e;
    border: 1px solid #0f3460;
    border-radius: 4px;
    padding: 6px 10px;
    color: #e0e0e0;
    min-height: 26px;
    selection-background-color: #e94560;
}
QLineEdit:focus, QDoubleSpinBox:focus, QSpinBox:focus, QComboBox:focus,
QTextEdit:focus {
    border: 2px solid #e94560;
}
QComboBox QAbstractItemView {
    background: #1a1a2e; color: #e0e0e0;
    selection-background-color: #e94560; selection-color: #ffffff;
}

QListWidget, QTreeWidget, QTableWidget {
    background: #1a1a2e;
    alternate-background-color: #16213e;
    border: 1px solid #0f3460;
    border-radius: 6px;
    color: #e0e0e0;
    outline: none;
}
QListWidget::item { padding: 6px; border-radius: 4px; }
QListWidget::item:selected { background: #0f3460; color: #ffffff; }
QListWidget::item:hover { background: #1a2744; }
QListWidget::item:checked { background: #0f3460; }
QHeaderView::section {
    background-color: #0f3460; color: #ffffff;
    font-weight: bold; padding: 4px; border: 1px solid #1a1a2e;
}

QCheckBox { spacing: 8px; color: #c0ccd8; }
QCheckBox::indicator { width: 16px; height: 16px;
    border: 2px solid #0f3460; border-radius: 4px; background: #0d1525; }
QCheckBox::indicator:checked { background: #e94560; border-color: #e94560; }

QPushButton {
    background-color: #0f3460; color: #ffffff;
    border: none; border-radius: 6px;
    padding: 8px 18px; font-weight: bold; font-size: 12px;
}
QPushButton:hover { background-color: #e94560; }
QPushButton:disabled { background-color: #2c3e50; color: #7a7a8a; }
QPushButton#primary { background-color: #e94560; }
QPushButton#primary:hover { background-color: #ff6b81; }

QScrollArea { background: transparent; border: none; }
QScrollBar:vertical { background: #16213e; width: 10px; margin: 0; }
QScrollBar::handle:vertical { background: #0f3460; border-radius: 5px; min-height: 30px; }
QScrollBar::handle:vertical:hover { background: #e94560; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
QScrollBar:horizontal { background: #16213e; height: 10px; margin: 0; }
QScrollBar::handle:horizontal { background: #0f3460; border-radius: 5px; min-width: 30px; }
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0; }

QWizard QLineEdit, QWizard QComboBox { min-height: 30px; }
"""


class GeneratorWizard(QWizard):
    """The universal program & procedure generator wizard."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("🧙  Program & Procedure Generator Wizard")
        self.setMinimumSize(1100, 760)
        self.setWizardStyle(QWizard.ModernStyle)
        self.setOption(QWizard.NoBackButtonOnStartPage, True)
        self.setOption(QWizard.NoCancelButtonOnLastPage, False)
        self.setOption(QWizard.HaveHelpButton, False)
        self.setStyleSheet(WIZARD_STYLE)

        from wizard_library import ALL_TEMPLATES
        from wizard_procedures import PROCEDURE_TEMPLATES
        from wizard_offshore import OFFSHORE_TEMPLATES
        templates: List[TemplateDef] = list(ALL_TEMPLATES) + \
            list(PROCEDURE_TEMPLATES) + list(OFFSHORE_TEMPLATES)
        # Master procedures (one consolidated procedure per operation)
        try:
            from wizard_master import build_master_templates
            templates += build_master_templates()
        except Exception:
            pass

        self.addPage(_TemplatePage(templates))   # 0 type
        self.addPage(_WellProfilePage())         # 1 well profile (new)
        self.addPage(_SectionsPage())            # 2 sections
        self.addPage(_InputsPage())              # 3 inputs
        self.addPage(_RiskReviewPage())          # 4 risk review
        self.addPage(_OptionsPage())             # 5 options
        self.addPage(_GeneratePage())            # 6 generate

        self.setStartId(0)


def run_wizard(parent=None) -> Optional[str]:
    """Launch the wizard; returns the generated document path if any."""
    wiz = GeneratorWizard(parent)
    if wiz.exec() == QDialog.Accepted:
        gen_page = wiz.page(6)
        if hasattr(gen_page, "_last_path") and gen_page._last_path:
            return gen_page._last_path
    return None
