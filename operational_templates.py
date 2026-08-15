# ============================================================================
# DRILLING PROGRAM & PROCEDURE GENERATOR - PROFESSIONAL EDITION
# Version 3.0
# File: operational_templates.py
# Template Library + UI Dialog - All in one file
# ============================================================================

from __future__ import annotations

import copy
import json
import logging
import threading
import uuid
from dataclasses import dataclass, field, asdict
from datetime import datetime
from enum import Enum
from typing import Any, Callable, Dict, List, Optional

logger = logging.getLogger(__name__)


# ============================================================================
# CONSTANTS
# ============================================================================

class Defaults:
    TEMPLATE_VERSION: str = "1.0"
    MAX_STEP_DURATION_HOURS: float = 720.0
    DEFAULT_RISK: str = "Medium"


class TimeC:
    """Time constants for all operations"""
    # Pre-spud
    MEETING = 2.0
    RIG_UP = 8.0
    BOP_INSTALL = 6.0
    MUD_PREP = 4.0
    SPUD = 4.0
    # Drilling
    SURFACE_DRILL = 24.0
    CIRCULATE = 2.0
    WIPER_TRIP = 4.0
    # Casing
    CASING_RUN = 8.0
    WOC = 12.0
    PRESSURE_TEST = 1.0
    # BOP
    BOP_RIG_UP = 8.0
    BOP_CONNECT = 2.0
    BOP_FUNC_TEST = 2.0
    BOP_LP_TEST = 2.0
    BOP_HP_TEST = 3.0
    BOP_CHOKE_TEST = 2.0
    BOP_DRILL = 1.0
    # Cementing
    PRE_CEMENT = 1.0
    CEMENT_LINE_TEST = 1.5
    PUMP_SPACER = 0.5
    PUMP_LEAD = 1.5
    PUMP_TAIL = 1.0
    DISPLACE = 1.5
    # Well testing
    PRE_TEST = 2.0
    RIG_UP_TEST = 12.0
    RIH_TEST = 8.0
    IFP = 6.0
    ISIP = 12.0
    MAIN_FLOW = 24.0
    FINAL_BU = 24.0
    KILL_POOH = 12.0
    RIG_DOWN_TEST = 8.0
    # Tripping
    TRIP_OUT = 6.0
    TRIP_IN = 6.0
    # Logging
    RIG_UP_LOG = 2.0
    LOG_RUN = 6.0
    RIG_DOWN_LOG = 2.0
    # Coring
    CORE_PICKUP = 3.0
    CORE_RIH = 4.0
    CORE_RUN = 12.0
    CORE_POOH = 6.0
    CORE_HANDLE = 3.0
    # Fishing
    FISH_ASSESS = 1.0
    FISH_TOOL = 2.0
    FISH_RIH = 5.0
    FISH_WORK = 4.0
    FISH_POOH = 5.0
    # Mud conversion
    MUD_MEETING = 1.0
    MUD_DISPLACE = 3.0
    MUD_MIX = 8.0
    MUD_CIRCULATE = 4.0
    MUD_CONFIRM = 1.0
    # Abandonment
    PA_REVIEW = 2.0
    PA_PLUG1 = 4.0
    PA_WOC = 8.0
    PA_PLUG2 = 4.0
    PA_SURF_PLUG = 4.0
    PA_CUT = 6.0
    PA_RESTORE = 4.0
    # Horizontal
    HORIZ_PICKUP = 4.0
    HORIZ_BUILD = 24.0
    HORIZ_LAND = 4.0
    HORIZ_DRILL = 72.0
    HORIZ_SWEEP = 8.0
    HORIZ_CLEANUP = 12.0


# ============================================================================
# ENUMS
# ============================================================================

class TemplateCategory(Enum):
    SURFACE_HOLE     = "Surface Hole Section"
    INTERMEDIATE_HOLE= "Intermediate Hole Section"
    PRODUCTION_HOLE  = "Production Hole Section"
    COMPLETION       = "Completion Operations"
    WORKOVER         = "Workover Operations"
    FISHING          = "Fishing Operations"
    SIDETRACK        = "Sidetrack Operations"
    HORIZONTAL       = "Horizontal / ERD Section"
    CEMENT           = "Cementing Operations"
    TESTING          = "Well Testing Operations"
    LOGGING          = "Logging Operations"
    CASING_RUNNING   = "Casing Running Operations"
    TRIPPING         = "Tripping Operations"
    BOP_TEST         = "BOP Test Operations"
    MUD_ENGINEERING  = "Mud Engineering Operations"
    DIRECTIONAL      = "Directional Drilling Operations"
    CORING           = "Coring Operations"
    ABANDONMENT      = "Well Abandonment / P&A"


class RiskLevel(Enum):
    LOW      = "Low"
    MEDIUM   = "Medium"
    HIGH     = "High"
    CRITICAL = "Critical"


class UnitSystem(Enum):
    FIELD  = "Field (ft, psi, bbl)"
    METRIC = "Metric (m, kPa, m3)"
    MIXED  = "Mixed"


class InputType(Enum):
    TEXT      = "text"
    NUMBER    = "number"
    COMBO     = "combo"
    CHECKBOX  = "checkbox"
    DEPTH     = "depth"
    PRESSURE  = "pressure"
    MUDWEIGHT = "mudweight"


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class TemplateInput:
    key: str
    label: str
    input_type: InputType = InputType.TEXT
    default_value: Any = ""
    options: List[str] = field(default_factory=list)
    unit: str = ""
    min_value: Optional[float] = None
    max_value: Optional[float] = None
    required: bool = False
    tooltip: str = ""

    def to_dict(self) -> Dict:
        return {
            "key": self.key,
            "label": self.label,
            "type": self.input_type.value,
            "default": self.default_value,
            "options": self.options,
            "unit": self.unit,
            "min": self.min_value,
            "max": self.max_value,
            "required": self.required,
            "tooltip": self.tooltip,
        }


@dataclass
class OperationalStep:
    step_number: int
    description: str
    duration_hours: float
    responsible: str = "Drilling Supervisor"
    safety_note: str = ""
    equipment: List[str] = field(default_factory=list)
    parameters: Dict[str, Any] = field(default_factory=dict)
    risk_level: RiskLevel = RiskLevel.LOW
    checkpoint: bool = False

    def __post_init__(self):
        if self.duration_hours < 0:
            raise ValueError(
                f"Negative duration: {self.duration_hours}")

    def resolve_description(self) -> str:
        desc = self.description
        for k, v in self.parameters.items():
            desc = desc.replace("{" + k + "}", str(v))
        return desc

    def to_dict(self) -> Dict:
        d = asdict(self)
        d['risk_level'] = self.risk_level.value
        return d


@dataclass
class PhaseTemplate:
    phase_id: str
    phase_name: str
    category: TemplateCategory
    description: str
    steps: List[OperationalStep] = field(default_factory=list)
    prerequisites: List[str] = field(default_factory=list)
    expected_duration_hours: float = 0.0
    risk_level: RiskLevel = RiskLevel.MEDIUM
    mud_properties: Dict[str, Any] = field(default_factory=dict)
    casing_specs: Dict[str, Any] = field(default_factory=dict)
    bha_components: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)
    inputs: List[TemplateInput] = field(default_factory=list)

    def total_duration(self) -> float:
        return sum(s.duration_hours for s in self.steps)

    def effective_duration(self) -> float:
        if self.expected_duration_hours > 0:
            return self.expected_duration_hours
        return self.total_duration()

    def get_checkpoints(self) -> List[OperationalStep]:
        return [s for s in self.steps if s.checkpoint]

    def get_high_risk_steps(self) -> List[OperationalStep]:
        return [s for s in self.steps
                if s.risk_level in
                (RiskLevel.HIGH, RiskLevel.CRITICAL)]

    def to_dict(self) -> Dict:
        return {
            'phase_id': self.phase_id,
            'phase_name': self.phase_name,
            'category': self.category.value,
            'description': self.description,
            'steps': [s.to_dict() for s in self.steps],
            'prerequisites': self.prerequisites,
            'expected_duration_hours': self.effective_duration(),
            'risk_level': self.risk_level.value,
            'mud_properties': self.mud_properties,
            'casing_specs': self.casing_specs,
            'bha_components': self.bha_components,
            'notes': self.notes,
            'inputs': [i.to_dict() for i in self.inputs],
        }


@dataclass
class WellTemplate:
    template_id: str
    template_name: str
    well_type: str
    company: str = ""
    field_name: str = ""
    rig_name: str = ""
    phases: List[PhaseTemplate] = field(default_factory=list)
    unit_system: UnitSystem = UnitSystem.FIELD
    created_date: str = field(
        default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    version: str = Defaults.TEMPLATE_VERSION
    approved_by: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    global_inputs: List[TemplateInput] = field(default_factory=list)

    def total_well_duration(self) -> float:
        return sum(p.effective_duration() for p in self.phases)

    def get_all_inputs(self) -> List[TemplateInput]:
        all_inputs = list(self.global_inputs)
        seen = {i.key for i in all_inputs}
        for phase in self.phases:
            for inp in phase.inputs:
                if inp.key not in seen:
                    all_inputs.append(inp)
                    seen.add(inp.key)
        return all_inputs

    def to_dict(self) -> Dict:
        return {
            'template_id': self.template_id,
            'template_name': self.template_name,
            'well_type': self.well_type,
            'company': self.company,
            'field_name': self.field_name,
            'rig_name': self.rig_name,
            'phases': [p.to_dict() for p in self.phases],
            'unit_system': self.unit_system.value,
            'total_duration_hours': self.total_well_duration(),
            'total_duration_days': round(
                self.total_well_duration() / 24, 1),
            'created_date': self.created_date,
            'version': self.version,
            'approved_by': self.approved_by,
            'metadata': self.metadata,
            'global_inputs': [i.to_dict()
                              for i in self.global_inputs],
        }


# ============================================================================
# PHASE BUILDER
# ============================================================================

class PhaseBuilder:
    """Builder pattern for PhaseTemplate - eliminates duplication"""

    def __init__(self, phase_id: str, phase_name: str,
                 category: TemplateCategory,
                 description: str):
        self._phase = PhaseTemplate(
            phase_id=phase_id,
            phase_name=phase_name,
            category=category,
            description=description)
        self._step_counter = 0

    def risk(self, level: RiskLevel) -> "PhaseBuilder":
        self._phase.risk_level = level
        return self

    def prereqs(self, *items: str) -> "PhaseBuilder":
        self._phase.prerequisites = list(items)
        return self

    def mud(self, **kwargs) -> "PhaseBuilder":
        self._phase.mud_properties = kwargs
        return self

    def casing(self, **kwargs) -> "PhaseBuilder":
        self._phase.casing_specs = kwargs
        return self

    def bha(self, *components: str) -> "PhaseBuilder":
        self._phase.bha_components = list(components)
        return self

    def inp(self, template_input: TemplateInput
            ) -> "PhaseBuilder":
        self._phase.inputs.append(template_input)
        return self

    def step(self, description: str,
             duration_hours: float,
             responsible: str = "Drilling Supervisor",
             safety_note: str = "",
             equipment: List[str] = None,
             parameters: Dict[str, Any] = None,
             risk_level: RiskLevel = RiskLevel.LOW,
             checkpoint: bool = False
             ) -> "PhaseBuilder":
        self._step_counter += 1
        self._phase.steps.append(OperationalStep(
            step_number=self._step_counter,
            description=description,
            duration_hours=duration_hours,
            responsible=responsible,
            safety_note=safety_note,
            equipment=equipment or [],
            parameters=parameters or {},
            risk_level=risk_level,
            checkpoint=checkpoint))
        return self

    def build(self) -> PhaseTemplate:
        return copy.deepcopy(self._phase)


# ============================================================================
# TEMPLATE LIBRARY - Thread-Safe Singleton
# ============================================================================

class TemplateLibrary:
    """Template library - Thread-Safe Singleton"""

    _instance: Optional[TemplateLibrary] = None
    _lock: threading.Lock = threading.Lock()

    def __new__(cls) -> TemplateLibrary:
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    inst = super().__new__(cls)
                    inst._initialized = False
                    cls._instance = inst
        return cls._instance

    def __init__(self):
        if self._initialized:
            return
        self._templates: Dict[str, WellTemplate] = {}
        self._phases: Dict[str, PhaseTemplate] = {}
        self._rlock = threading.RLock()
        self._load_default_phases()
        self._load_default_templates()
        self._initialized = True

    @classmethod
    def reset(cls) -> None:
        with cls._lock:
            cls._instance = None

    # ----------------------------------------------------------------
    # GLOBAL INPUTS (shared across templates)
    # ----------------------------------------------------------------

    @staticmethod
    def _global_inputs() -> List[TemplateInput]:
        return [
            TemplateInput(
                key="well_name", label="Well Name",
                input_type=InputType.TEXT, required=True),
            TemplateInput(
                key="field_name", label="Field Name",
                input_type=InputType.TEXT),
            TemplateInput(
                key="operator", label="Operator",
                input_type=InputType.TEXT),
            TemplateInput(
                key="rig_name", label="Rig Name",
                input_type=InputType.TEXT),
            TemplateInput(
                key="td_md", label="TD (MD)",
                input_type=InputType.DEPTH,
                unit="ft", required=True),
            TemplateInput(
                key="td_tvd", label="TD (TVD)",
                input_type=InputType.DEPTH, unit="ft"),
            TemplateInput(
                key="unit_system", label="Unit System",
                input_type=InputType.COMBO,
                options=["Field (ft, psi, bbl)",
                         "Metric (m, kPa, m3)"],
                default_value="Field (ft, psi, bbl)"),
        ]

    # ----------------------------------------------------------------
    # PHASE BUILDERS
    # ----------------------------------------------------------------

    def _build_rig_up_spud(self) -> PhaseTemplate:
        return (
            PhaseBuilder("RIG_UP_SPUD",
                         "Rig Up & Spud Well",
                         TemplateCategory.SURFACE_HOLE,
                         "Rig up and spud the well")
            .prereqs("Rig mobilization complete",
                     "Location approved",
                     "Permits obtained")
            .risk(RiskLevel.MEDIUM)
            .mud(type="Gel/Water Spud Mud",
                 weight_ppg=8.7,
                 funnel_viscosity="40-50 sec/qt")
            .inp(TemplateInput(
                key="mud_weight", label="Spud Mud Weight",
                input_type=InputType.MUDWEIGHT,
                default_value=8.7, unit="ppg",
                min_value=8.5, max_value=9.5,
                required=True))
            .inp(TemplateInput(
                key="conductor_depth",
                label="Conductor Setting Depth",
                input_type=InputType.DEPTH,
                default_value=100, unit="ft",
                required=True))
            .inp(TemplateInput(
                key="bit_size", label="Spud Bit Size",
                input_type=InputType.COMBO,
                options=["26\"", "36\"", "17-1/2\""],
                default_value="26\""))
            .step("Conduct pre-spud meeting with all "
                  "personnel. Review well program, safety "
                  "procedures and emergency response plan.",
                  TimeC.MEETING,
                  responsible="Company Man / Tool Pusher",
                  safety_note="All personnel sign attendance",
                  checkpoint=True)
            .step("Rig up and test all surface equipment: "
                  "mud pumps, rotary system, hoisting, "
                  "and circulating system.",
                  TimeC.RIG_UP,
                  responsible="Tool Pusher",
                  equipment=["Mud Pumps",
                             "Rotary/Top Drive",
                             "Drawworks"],
                  safety_note="LOTO procedures required")
            .step("Install and test BOP stack per API RP 53. "
                  "Function test all rams and annular.",
                  TimeC.BOP_INSTALL,
                  responsible="Drilling Supervisor",
                  equipment=["BOP Stack",
                             "Accumulator",
                             "Choke Manifold"],
                  safety_note="Company rep must witness tests",
                  risk_level=RiskLevel.HIGH,
                  checkpoint=True)
            .step("Mix and prepare spud mud. "
                  "Target MW: {mud_weight} ppg.",
                  TimeC.MUD_PREP,
                  responsible="Mud Engineer",
                  parameters={"mud_weight": 8.7,
                               "funnel_viscosity": 45,
                               "ph": 9.5})
            .step("Spud well with {bit_size} bit. "
                  "Drill to {conductor_depth} ft.",
                  TimeC.SPUD,
                  responsible="Driller",
                  equipment=["Drill Bit",
                             "Drill Collars",
                             "Top Drive"],
                  parameters={"bit_size": "26\"",
                               "conductor_depth": 100,
                               "rpm": 60})
            .build()
        )

    def _build_surface_hole(self) -> PhaseTemplate:
        return (
            PhaseBuilder("SURFACE_HOLE",
                         "Drill Surface Hole Section",
                         TemplateCategory.SURFACE_HOLE,
                         "Drill surface hole from conductor "
                         "shoe to surface casing setting depth")
            .risk(RiskLevel.MEDIUM)
            .mud(type="Native/Bentonite Mud",
                 weight_ppg="8.8-9.2",
                 ph="9.0-10.0")
            .bha("17-1/2\" Bit", "17\" Stabilizer",
                 "8\" Drill Collars x12",
                 "5\" HWDP x10", "5\" Drill Pipe")
            .inp(TemplateInput(
                key="bit_size", label="Hole Size",
                input_type=InputType.COMBO,
                options=["26\"", "17-1/2\"", "12-1/4\""],
                default_value="17-1/2\"", required=True))
            .inp(TemplateInput(
                key="start_depth", label="Start Depth",
                input_type=InputType.DEPTH,
                default_value=100, unit="ft", required=True))
            .inp(TemplateInput(
                key="end_depth", label="End Depth (TD)",
                input_type=InputType.DEPTH,
                default_value=2000, unit="ft", required=True))
            .inp(TemplateInput(
                key="wob", label="WOB Range",
                input_type=InputType.TEXT,
                default_value="15,000-30,000 lbs"))
            .inp(TemplateInput(
                key="rpm", label="RPM Range",
                input_type=InputType.TEXT,
                default_value="80-120"))
            .inp(TemplateInput(
                key="flow_rate", label="Flow Rate",
                input_type=InputType.NUMBER,
                default_value=900, unit="GPM"))
            .step("Pick up {bit_size} BHA. Make up and "
                  "verify all connections.",
                  3.0, responsible="Driller",
                  parameters={"bit_size": "17-1/2\""})
            .step("RIH to conductor shoe. Establish "
                  "circulation. Verify returns.",
                  1.0, responsible="Driller",
                  safety_note="Watch for tight spots")
            .step("Drill surface hole from {start_depth} ft "
                  "to {end_depth} ft. WOB: {wob}, "
                  "RPM: {rpm}, Flow: {flow_rate} GPM.",
                  TimeC.SURFACE_DRILL,
                  responsible="Driller",
                  parameters={"start_depth": 100,
                               "end_depth": 2000,
                               "wob": "15,000-30,000 lbs",
                               "rpm": "80-120",
                               "flow_rate": "800-1000"},
                  safety_note="Monitor for shallow gas. "
                              "Keep diverter ready.")
            .step("At TD, circulate 2x hole volumes. "
                  "Monitor returns for cuttings clearance.",
                  TimeC.CIRCULATE, responsible="Driller")
            .step("Perform wiper trip to {wiper_depth} ft "
                  "and back to bottom.",
                  TimeC.WIPER_TRIP, responsible="Driller",
                  parameters={"wiper_depth": 500},
                  safety_note="Monitor for flow while tripping")
            .step("POOH with BHA. Inspect bit per IADC.",
                  3.0, responsible="Driller",
                  checkpoint=True)
            .build()
        )

    def _build_surface_casing(self) -> PhaseTemplate:
        return (
            PhaseBuilder("SURFACE_CASING",
                         "Run & Cement Surface Casing",
                         TemplateCategory.CASING_RUNNING,
                         "Run surface casing and cement")
            .prereqs("Surface hole drilled to TD",
                     "Hole conditioned",
                     "Casing on location")
            .risk(RiskLevel.HIGH)
            .inp(TemplateInput(
                key="casing_size", label="Casing Size",
                input_type=InputType.COMBO,
                options=["30\"", "20\"", "13-3/8\"",
                         "9-5/8\"", "7\""],
                default_value="20\"", required=True))
            .inp(TemplateInput(
                key="setting_depth",
                label="Setting Depth",
                input_type=InputType.DEPTH,
                default_value=2000, unit="ft", required=True))
            .inp(TemplateInput(
                key="lead_density",
                label="Lead Slurry Density",
                input_type=InputType.MUDWEIGHT,
                default_value=13.5, unit="ppg"))
            .inp(TemplateInput(
                key="lead_volume",
                label="Lead Cement Volume",
                input_type=InputType.NUMBER,
                default_value=250, unit="bbl"))
            .inp(TemplateInput(
                key="tail_density",
                label="Tail Slurry Density",
                input_type=InputType.MUDWEIGHT,
                default_value=15.8, unit="ppg"))
            .inp(TemplateInput(
                key="tail_volume",
                label="Tail Cement Volume",
                input_type=InputType.NUMBER,
                default_value=100, unit="bbl"))
            .inp(TemplateInput(
                key="woc_hours", label="WOC Time",
                input_type=InputType.NUMBER,
                default_value=12, unit="hrs"))
            .step("Rig up casing crew. Verify tally. "
                  "Inspect all casing joints.",
                  2.0, responsible="Casing Supervisor",
                  checkpoint=True)
            .step("Run {casing_size} casing to {setting_depth} ft. "
                  "Install centralizers per program.",
                  TimeC.CASING_RUN,
                  responsible="Casing Crew",
                  parameters={"casing_size": "20\"",
                               "setting_depth": 2000},
                  safety_note="Max running speed 3 ft/sec")
            .step("Circulate 2x hole volumes. "
                  "Condition mud for cementing.",
                  TimeC.CIRCULATE,
                  responsible="Driller / Mud Engineer")
            .step("Pump lead cement {lead_volume} bbl @ "
                  "{lead_density} ppg. "
                  "Tail {tail_volume} bbl @ {tail_density} ppg.",
                  3.0, responsible="Cementing Engineer",
                  parameters={"lead_volume": 250,
                               "lead_density": 13.5,
                               "tail_volume": 100,
                               "tail_density": 15.8},
                  risk_level=RiskLevel.HIGH,
                  checkpoint=True)
            .step("Displace cement. Bump plug. "
                  "Hold pressure for verification.",
                  1.5, responsible="Cementing Engineer",
                  safety_note="Do not exceed max pump pressure")
            .step("WOC {woc_hours} hours. "
                  "Monitor casing pressure.",
                  TimeC.WOC, responsible="Drilling Supervisor",
                  parameters={"woc_hours": 12})
            .step("Pressure test casing. Record results.",
                  TimeC.PRESSURE_TEST,
                  responsible="Drilling Supervisor",
                  risk_level=RiskLevel.HIGH,
                  checkpoint=True)
            .build()
        )

    def _build_intermediate_hole(self) -> PhaseTemplate:
        return (
            PhaseBuilder("INTERMEDIATE_HOLE",
                         "Drill Intermediate Hole Section",
                         TemplateCategory.INTERMEDIATE_HOLE,
                         "Drill intermediate section with "
                         "directional control")
            .risk(RiskLevel.HIGH)
            .mud(type="KCl Polymer / PHPA",
                 weight_ppg="10.5-12.5",
                 pv_cp="15-25")
            .inp(TemplateInput(
                key="bit_size", label="Bit Size",
                input_type=InputType.COMBO,
                options=["17-1/2\"", "12-1/4\"", "8-1/2\""],
                default_value="12-1/4\"", required=True))
            .inp(TemplateInput(
                key="start_depth", label="Start Depth",
                input_type=InputType.DEPTH,
                default_value=2020, unit="ft", required=True))
            .inp(TemplateInput(
                key="end_depth", label="End Depth (TD)",
                input_type=InputType.DEPTH,
                default_value=8500, unit="ft", required=True))
            .inp(TemplateInput(
                key="kop", label="KOP Depth",
                input_type=InputType.DEPTH,
                default_value=3000, unit="ft"))
            .inp(TemplateInput(
                key="build_rate", label="Build Rate",
                input_type=InputType.NUMBER,
                default_value=2.0, unit="deg/100ft",
                min_value=0.5, max_value=10.0))
            .inp(TemplateInput(
                key="target_inc",
                label="Target Inclination",
                input_type=InputType.NUMBER,
                default_value=30, unit="deg",
                min_value=0, max_value=90))
            .inp(TemplateInput(
                key="fit_target", label="FIT Target",
                input_type=InputType.MUDWEIGHT,
                default_value=14.0, unit="ppg EMW"))
            .step("Drill out shoe and {cement_length} ft "
                  "new formation.",
                  2.0, responsible="Driller",
                  parameters={"cement_length": 20},
                  checkpoint=True)
            .step("Conduct FIT/LOT at shoe. "
                  "Target: {fit_target} ppg EMW.",
                  1.0, responsible="Drilling Supervisor",
                  parameters={"fit_target": 14.0},
                  risk_level=RiskLevel.HIGH,
                  checkpoint=True)
            .step("Pick up {bit_size} directional BHA "
                  "with MWD/LWD.",
                  4.0, responsible="Directional Driller",
                  parameters={"bit_size": "12-1/4\""})
            .step("Drill from {start_depth} ft to {end_depth} ft. "
                  "KOP at {kop} ft, build {build_rate} deg/100ft, "
                  "target {target_inc} deg.",
                  72.0,
                  responsible="Driller / DD",
                  parameters={"start_depth": 2020,
                               "end_depth": 8500,
                               "kop": 3000,
                               "build_rate": 2.0,
                               "target_inc": 30},
                  safety_note="Monitor for kick continuously")
            .step("Wiper trips every {trip_interval} ft.",
                  8.0, responsible="Driller",
                  parameters={"trip_interval": 1500})
            .step("Pump {pill_volume} bbl hi-vis sweep. "
                  "Circulate bottoms up at TD.",
                  4.0, responsible="Mud Engineer",
                  parameters={"pill_volume": 60})
            .step("POOH with BHA. Run wireline logs.",
                  22.0, responsible="Driller")
            .build()
        )

    def _build_bop_test(self) -> PhaseTemplate:
        return (
            PhaseBuilder("BOP_TEST",
                         "BOP Installation & Testing",
                         TemplateCategory.BOP_TEST,
                         "Install BOP and test per API RP 53")
            .risk(RiskLevel.CRITICAL)
            .inp(TemplateInput(
                key="working_pressure",
                label="BOP Working Pressure",
                input_type=InputType.PRESSURE,
                default_value=10000, unit="psi",
                required=True))
            .inp(TemplateInput(
                key="low_pressure",
                label="Low Test Pressure",
                input_type=InputType.PRESSURE,
                default_value=250, unit="psi"))
            .inp(TemplateInput(
                key="high_pressure",
                label="High Test Pressure",
                input_type=InputType.PRESSURE,
                default_value=5000, unit="psi"))
            .step("Nipple up BOP stack. "
                  "Verify all connections and flanges.",
                  TimeC.BOP_RIG_UP,
                  responsible="Tool Pusher",
                  risk_level=RiskLevel.CRITICAL)
            .step("Connect accumulator. Check precharge.",
                  TimeC.BOP_CONNECT,
                  responsible="BOP Technician")
            .step("Function test all BOP elements: "
                  "annular, pipe rams, blind/shear rams, "
                  "HCR valve, kill/choke valves.",
                  TimeC.BOP_FUNC_TEST,
                  responsible="Drilling Supervisor",
                  risk_level=RiskLevel.CRITICAL,
                  checkpoint=True)
            .step("Low pressure test all rams at "
                  "{low_pressure} psi / 5 min.",
                  TimeC.BOP_LP_TEST,
                  responsible="Drilling Supervisor",
                  parameters={"low_pressure": 250,
                               "acceptance": "Zero loss"},
                  risk_level=RiskLevel.HIGH)
            .step("High pressure test all rams at "
                  "{high_pressure} psi / 10 min.",
                  TimeC.BOP_HP_TEST,
                  responsible="Drilling Supervisor",
                  parameters={"high_pressure": 5000,
                               "acceptance": "<10 psi in 10 min"},
                  risk_level=RiskLevel.CRITICAL,
                  checkpoint=True)
            .step("Test choke/kill line and manifold.",
                  TimeC.BOP_CHOKE_TEST,
                  responsible="Drilling Supervisor")
            .step("Conduct BOP drill with all crews. "
                  "Target: <30 sec closing time.",
                  TimeC.BOP_DRILL,
                  responsible="Tool Pusher",
                  safety_note="All must know their positions")
            .build()
        )

    def _build_production_hole(self) -> PhaseTemplate:
        return (
            PhaseBuilder("PRODUCTION_HOLE",
                         "Drill Production Hole Section",
                         TemplateCategory.PRODUCTION_HOLE,
                         "Drill reservoir section")
            .risk(RiskLevel.HIGH)
            .mud(type="OBM / SBM",
                 weight_ppg="11.5-13.0",
                 oil_water_ratio="80:20")
            .inp(TemplateInput(
                key="bit_size", label="Bit Size",
                input_type=InputType.COMBO,
                options=["12-1/4\"", "8-1/2\"", "6-1/8\""],
                default_value="8-1/2\"", required=True))
            .inp(TemplateInput(
                key="start_depth", label="Start Depth",
                input_type=InputType.DEPTH,
                default_value=8520, unit="ft", required=True))
            .inp(TemplateInput(
                key="end_depth", label="End Depth (TD)",
                input_type=InputType.DEPTH,
                default_value=12000, unit="ft", required=True))
            .inp(TemplateInput(
                key="lot_target", label="LOT Target",
                input_type=InputType.MUDWEIGHT,
                default_value=16.0, unit="ppg EMW"))
            .inp(TemplateInput(
                key="reservoir_top",
                label="Expected Reservoir Top",
                input_type=InputType.DEPTH,
                default_value=11200, unit="ft"))
            .step("Drill out intermediate shoe. Rat hole.",
                  2.0, responsible="Driller")
            .step("LOT/FIT at shoe. Target: {lot_target} ppg.",
                  1.5, responsible="Drilling Supervisor",
                  parameters={"lot_target": 16.0},
                  risk_level=RiskLevel.HIGH, checkpoint=True)
            .step("Pick up {bit_size} BHA with RSS/LWD.",
                  4.0, responsible="Directional Driller",
                  parameters={"bit_size": "8-1/2\""})
            .step("Drill from {start_depth} ft to {end_depth} ft.",
                  48.0, responsible="Driller / DD",
                  parameters={"start_depth": 8520,
                               "end_depth": 12000},
                  safety_note="Near reservoir - monitor gas closely")
            .step("Monitor LWD for reservoir entry at "
                  "{reservoir_top} ft.",
                  4.0, responsible="Geologist / DS",
                  parameters={"reservoir_top": 11200},
                  risk_level=RiskLevel.CRITICAL,
                  checkpoint=True)
            .step("Circulate and condition mud for casing.",
                  4.0, responsible="Mud Engineer")
            .step("POOH. Run wireline logs.",
                  24.0, responsible="Logging Engineer")
            .build()
        )

    def _build_well_testing(self) -> PhaseTemplate:
        return (
            PhaseBuilder("WELL_TESTING",
                         "Drill Stem Test / Well Testing",
                         TemplateCategory.TESTING,
                         "Well test for reservoir evaluation")
            .prereqs("Production section drilled",
                     "Logs evaluated",
                     "Test program approved")
            .risk(RiskLevel.CRITICAL)
            .inp(TemplateInput(
                key="packer_depth",
                label="Packer Setting Depth",
                input_type=InputType.DEPTH,
                default_value=11000, unit="ft", required=True))
            .inp(TemplateInput(
                key="initial_flow",
                label="Initial Flow Period",
                input_type=InputType.NUMBER,
                default_value=6, unit="hrs"))
            .inp(TemplateInput(
                key="main_flow", label="Main Flow Period",
                input_type=InputType.NUMBER,
                default_value=24, unit="hrs"))
            .inp(TemplateInput(
                key="final_shut",
                label="Final Shut-In Period",
                input_type=InputType.NUMBER,
                default_value=24, unit="hrs"))
            .step("Pre-test safety meeting. "
                  "Review H2S contingency.",
                  TimeC.PRE_TEST,
                  responsible="Company Man",
                  risk_level=RiskLevel.CRITICAL,
                  checkpoint=True)
            .step("Rig up test equipment: DST string, "
                  "separator, flare boom.",
                  TimeC.RIG_UP_TEST,
                  responsible="Well Test Engineer")
            .step("RIH. Set packer at {packer_depth} ft.",
                  TimeC.RIH_TEST,
                  responsible="Driller / WTE",
                  parameters={"packer_depth": 11000})
            .step("Initial flow {initial_flow} hrs. "
                  "Initial shut-in 1-2 hrs.",
                  TimeC.IFP + 2.0,
                  responsible="Well Test Engineer",
                  parameters={"initial_flow": 6},
                  risk_level=RiskLevel.CRITICAL)
            .step("Main flow {main_flow} hrs. "
                  "Record rates, pressures, GOR.",
                  TimeC.MAIN_FLOW,
                  responsible="Well Test Engineer",
                  parameters={"main_flow": 24})
            .step("Final shut-in {final_shut} hrs.",
                  TimeC.FINAL_BU,
                  responsible="Well Test Engineer",
                  parameters={"final_shut": 24})
            .step("Kill well. Reverse circulate. POOH.",
                  TimeC.KILL_POOH,
                  responsible="Drilling Supervisor",
                  risk_level=RiskLevel.HIGH)
            .step("Rig down test equipment. "
                  "Prepare preliminary report.",
                  TimeC.RIG_DOWN_TEST,
                  responsible="Well Test Engineer")
            .build()
        )

    def _build_primary_cement(self) -> PhaseTemplate:
        return (
            PhaseBuilder("CEMENT_PRIMARY",
                         "Primary Cementing",
                         TemplateCategory.CEMENT,
                         "Primary cement job for casing")
            .risk(RiskLevel.HIGH)
            .inp(TemplateInput(
                key="lead_volume",
                label="Lead Cement Volume",
                input_type=InputType.NUMBER,
                default_value=200, unit="bbl"))
            .inp(TemplateInput(
                key="lead_density",
                label="Lead Slurry Density",
                input_type=InputType.MUDWEIGHT,
                default_value=13.5, unit="ppg"))
            .inp(TemplateInput(
                key="tail_volume",
                label="Tail Cement Volume",
                input_type=InputType.NUMBER,
                default_value=80, unit="bbl"))
            .inp(TemplateInput(
                key="tail_density",
                label="Tail Slurry Density",
                input_type=InputType.MUDWEIGHT,
                default_value=15.8, unit="ppg"))
            .inp(TemplateInput(
                key="displacement_volume",
                label="Displacement Volume",
                input_type=InputType.NUMBER,
                default_value=300, unit="bbl"))
            .inp(TemplateInput(
                key="bump_pressure",
                label="Plug Bump Pressure",
                input_type=InputType.PRESSURE,
                default_value=1200, unit="psi"))
            .inp(TemplateInput(
                key="woc_hours", label="WOC Time",
                input_type=InputType.NUMBER,
                default_value=12, unit="hrs"))
            .step("Pre-job meeting. Review cement program.",
                  TimeC.PRE_CEMENT,
                  responsible="Cementing Engineer",
                  checkpoint=True)
            .step("Rig up cement unit. "
                  "Pressure test lines.",
                  TimeC.CEMENT_LINE_TEST,
                  responsible="Cementing Engineer")
            .step("Pump spacer/wash.",
                  TimeC.PUMP_SPACER,
                  responsible="Cementing Engineer")
            .step("Drop bottom plug. "
                  "Pump lead cement: {lead_volume} bbl "
                  "@ {lead_density} ppg.",
                  TimeC.PUMP_LEAD,
                  responsible="Cementing Engineer",
                  parameters={"lead_volume": 200,
                               "lead_density": 13.5})
            .step("Drop top plug. "
                  "Pump tail: {tail_volume} bbl "
                  "@ {tail_density} ppg.",
                  TimeC.PUMP_TAIL,
                  responsible="Cementing Engineer",
                  parameters={"tail_volume": 80,
                               "tail_density": 15.8})
            .step("Displace {displacement_volume} bbl. "
                  "Bump plug @ {bump_pressure} psi.",
                  TimeC.DISPLACE,
                  responsible="Cementing Engineer",
                  parameters={"displacement_volume": 300,
                               "bump_pressure": 1200},
                  checkpoint=True)
            .step("WOC {woc_hours} hrs. "
                  "Monitor casing pressure.",
                  TimeC.WOC,
                  responsible="Drilling Supervisor",
                  parameters={"woc_hours": 12})
            .build()
        )

    def _build_trip_out(self) -> PhaseTemplate:
        return (
            PhaseBuilder("TRIP_OUT",
                         "Trip Out of Hole",
                         TemplateCategory.TRIPPING,
                         "Standard POOH procedure")
            .risk(RiskLevel.MEDIUM)
            .inp(TemplateInput(
                key="fill_interval",
                label="Fill Hole Interval",
                input_type=InputType.NUMBER,
                default_value=5,
                tooltip="Fill every N stands"))
            .step("Circulate bottoms up. "
                  "Monitor for gas.",
                  1.0, responsible="Driller")
            .step("POOH. Fill hole every {fill_interval} stands. "
                  "Monitor trip tank.",
                  TimeC.TRIP_OUT,
                  responsible="Driller",
                  parameters={"fill_interval": 5},
                  safety_note="Stop if well does not take fill")
            .step("Clear rotary. Set slips. "
                  "Flow check 15 min.",
                  0.25, responsible="Driller")
            .build()
        )

    def _build_trip_in(self) -> PhaseTemplate:
        return (
            PhaseBuilder("TRIP_IN",
                         "Trip In Hole",
                         TemplateCategory.TRIPPING,
                         "Standard RIH procedure")
            .risk(RiskLevel.MEDIUM)
            .step("Verify BHA make-up. Record torques.",
                  0.5, responsible="Driller",
                  checkpoint=True)
            .step("RIH. Monitor weight for tight spots. "
                  "Do not force past ledges.",
                  TimeC.TRIP_IN, responsible="Driller",
                  safety_note="Pull back if overpull >30 klbs")
            .step("At bottom: circulate, check fill, "
                  "take survey.",
                  1.5, responsible="Driller")
            .build()
        )

    def _build_coring(self) -> PhaseTemplate:
        return (
            PhaseBuilder("CORING",
                         "Conventional Coring Operations",
                         TemplateCategory.CORING,
                         "Core acquisition in reservoir")
            .risk(RiskLevel.HIGH)
            .inp(TemplateInput(
                key="core_bit_size",
                label="Core Bit Size",
                input_type=InputType.COMBO,
                options=["8-1/2\"", "6\"", "4-3/4\""],
                default_value="8-1/2\"", required=True))
            .inp(TemplateInput(
                key="core_depth",
                label="Core Top Depth",
                input_type=InputType.DEPTH,
                default_value=11500, unit="ft",
                required=True))
            .inp(TemplateInput(
                key="core_length",
                label="Core Barrel Length",
                input_type=InputType.NUMBER,
                default_value=60, unit="ft"))
            .step("Pick up {core_bit_size} coring assembly.",
                  TimeC.CORE_PICKUP,
                  responsible="Core Technician",
                  parameters={"core_bit_size": "8-1/2\""})
            .step("RIH to {core_depth} ft.",
                  TimeC.CORE_RIH,
                  responsible="Driller",
                  parameters={"core_depth": 11500})
            .step("Core {core_length} ft. "
                  "WOB: 5-15 klbs, RPM: 40-80.",
                  TimeC.CORE_RUN,
                  responsible="Driller / Core Tech",
                  parameters={"core_length": 60},
                  safety_note="Do not exceed WOB - jamming risk")
            .step("POOH. Extract and mark core. "
                  "Preserve per program.",
                  TimeC.CORE_POOH + TimeC.CORE_HANDLE,
                  responsible="Core Tech / Geologist")
            .build()
        )

    def _build_fishing(self) -> PhaseTemplate:
        return (
            PhaseBuilder("FISHING",
                         "Fishing Operations",
                         TemplateCategory.FISHING,
                         "Recover fish from wellbore")
            .risk(RiskLevel.HIGH)
            .inp(TemplateInput(
                key="fish_type", label="Fish Type",
                input_type=InputType.COMBO,
                options=["Drill Collars", "Drill Pipe",
                         "MWD Tool", "Bit", "Junk"],
                default_value="Drill Collars", required=True))
            .inp(TemplateInput(
                key="fish_depth",
                label="Top of Fish Depth",
                input_type=InputType.DEPTH,
                default_value=8500, unit="ft", required=True))
            .step("Assess: {fish_type} at {fish_depth} ft.",
                  TimeC.FISH_ASSESS,
                  responsible="Fishing Engineer",
                  parameters={"fish_type": "Drill Collars",
                               "fish_depth": 8500},
                  checkpoint=True)
            .step("Select fishing tool. Make up string.",
                  TimeC.FISH_TOOL,
                  responsible="Fishing Engineer")
            .step("RIH. Tag fish. Engage tool.",
                  TimeC.FISH_RIH,
                  responsible="Driller / FE")
            .step("Apply pull and jar. Work fish free.",
                  TimeC.FISH_WORK,
                  responsible="Driller",
                  risk_level=RiskLevel.HIGH)
            .step("POOH with fish. Inspect. "
                  "Assess hole condition.",
                  TimeC.FISH_POOH,
                  responsible="Driller / FE")
            .build()
        )

    def _build_mud_conversion(self) -> PhaseTemplate:
        return (
            PhaseBuilder("MUD_CONVERSION",
                         "Mud Conversion (WBM to OBM)",
                         TemplateCategory.MUD_ENGINEERING,
                         "Convert mud system from WBM to OBM")
            .risk(RiskLevel.MEDIUM)
            .inp(TemplateInput(
                key="mw", label="Target Mud Weight",
                input_type=InputType.MUDWEIGHT,
                default_value=12.0, unit="ppg", required=True))
            .inp(TemplateInput(
                key="owr", label="Oil-Water Ratio",
                input_type=InputType.COMBO,
                options=["70:30", "75:25", "80:20",
                         "85:15"],
                default_value="80:20"))
            .step("Pre-conversion meeting. "
                  "Review OBM handling and PPE.",
                  TimeC.MUD_MEETING,
                  responsible="Mud Engineer",
                  checkpoint=True)
            .step("Displace WBM from hole with spacer. "
                  "Transfer WBM to storage.",
                  TimeC.MUD_DISPLACE,
                  responsible="Mud Engineer")
            .step("Mix OBM: MW {mw} ppg, OWR {owr}.",
                  TimeC.MUD_MIX,
                  responsible="Mud Engineer",
                  parameters={"mw": 12.0, "owr": "80:20"})
            .step("Circulate OBM. Verify properties.",
                  TimeC.MUD_CIRCULATE,
                  responsible="Mud Engineer")
            .step("Confirm stable properties. "
                  "Set up monitoring program.",
                  TimeC.MUD_CONFIRM,
                  responsible="Mud Engineer",
                  checkpoint=True)
            .build()
        )

    def _build_abandonment(self) -> PhaseTemplate:
        return (
            PhaseBuilder("ABANDONMENT",
                         "Plug & Abandon Well",
                         TemplateCategory.ABANDONMENT,
                         "P&A per regulatory requirements")
            .risk(RiskLevel.HIGH)
            .inp(TemplateInput(
                key="plug1_base",
                label="Plug 1 Base Depth",
                input_type=InputType.DEPTH,
                default_value=11000, unit="ft", required=True))
            .inp(TemplateInput(
                key="plug1_top",
                label="Plug 1 Top Depth",
                input_type=InputType.DEPTH,
                default_value=10500, unit="ft"))
            .inp(TemplateInput(
                key="woc_hours", label="WOC Time",
                input_type=InputType.NUMBER,
                default_value=8, unit="hrs"))
            .step("Review P&A program. "
                  "Obtain regulatory approval.",
                  TimeC.PA_REVIEW,
                  responsible="Company Man",
                  checkpoint=True)
            .step("Set cement plug #1 (reservoir isolation) "
                  "{plug1_base} ft to {plug1_top} ft.",
                  TimeC.PA_PLUG1,
                  responsible="Cementing Engineer",
                  parameters={"plug1_base": 11000,
                               "plug1_top": 10500})
            .step("WOC {woc_hours} hrs. "
                  "Tag and pressure test cement.",
                  TimeC.PA_WOC,
                  responsible="Drilling Supervisor",
                  parameters={"woc_hours": 8},
                  checkpoint=True)
            .step("Set cement plug #2 (intermediate).",
                  TimeC.PA_PLUG2,
                  responsible="Cementing Engineer")
            .step("Set surface cement plug.",
                  TimeC.PA_SURF_PLUG,
                  responsible="Cementing Engineer")
            .step("Cut and pull casing. Weld cap.",
                  TimeC.PA_CUT,
                  responsible="Casing Crew")
            .step("Restore location. "
                  "File abandonment report.",
                  TimeC.PA_RESTORE,
                  responsible="Company Man",
                  checkpoint=True)
            .build()
        )

    def _build_horizontal_section(self) -> PhaseTemplate:
        return (
            PhaseBuilder("HORIZONTAL_SECTION",
                         "Horizontal Section Drilling",
                         TemplateCategory.HORIZONTAL,
                         "Drill horizontal section with RSS")
            .risk(RiskLevel.HIGH)
            .mud(type="High-Performance OBM",
                 weight_ppg="10.5-11.5",
                 oil_water_ratio="85:15")
            .inp(TemplateInput(
                key="kop_depth", label="KOP Depth",
                input_type=InputType.DEPTH,
                default_value=9000, unit="ft", required=True))
            .inp(TemplateInput(
                key="build_rate", label="Build Rate",
                input_type=InputType.NUMBER,
                default_value=6.0, unit="deg/100ft",
                min_value=1.0, max_value=15.0))
            .inp(TemplateInput(
                key="lateral_end",
                label="Lateral TD (MD)",
                input_type=InputType.DEPTH,
                default_value=15000, unit="ft",
                required=True))
            .inp(TemplateInput(
                key="target_tvd",
                label="Landing TVD",
                input_type=InputType.DEPTH,
                default_value=10500, unit="ft",
                required=True))
            .step("Pick up horizontal BHA with RSS and LWD.",
                  TimeC.HORIZ_PICKUP,
                  responsible="Directional Driller")
            .step("RIH to {kop_depth} ft. "
                  "Begin building at {build_rate} deg/100ft.",
                  TimeC.HORIZ_BUILD,
                  responsible="Directional Driller",
                  parameters={"kop_depth": 9000,
                               "build_rate": 6.0})
            .step("Land at {target_tvd} ft TVD. "
                  "Adjust for target zone.",
                  TimeC.HORIZ_LAND,
                  responsible="DD / Geologist",
                  parameters={"target_tvd": 10500},
                  risk_level=RiskLevel.HIGH,
                  checkpoint=True)
            .step("Drill horizontal to {lateral_end} ft MD.",
                  TimeC.HORIZ_DRILL,
                  responsible="Driller / DD",
                  parameters={"lateral_end": 15000},
                  safety_note="Monitor ECD - narrow MW window")
            .step("Hi-vis sweeps every 3 stands.",
                  TimeC.HORIZ_SWEEP,
                  responsible="Mud Engineer")
            .step("Circulate clean. Wiper trip to shoe.",
                  TimeC.HORIZ_CLEANUP,
                  responsible="Driller",
                  safety_note="High drag - monitor stuck pipe")
            .build()
        )

    def _build_wireline_logging(self) -> PhaseTemplate:
        return (
            PhaseBuilder("WIRELINE_LOGGING",
                         "Wireline Logging Operations",
                         TemplateCategory.LOGGING,
                         "Run wireline logging suite")
            .risk(RiskLevel.MEDIUM)
            .inp(TemplateInput(
                key="log_bottom",
                label="Log Bottom Depth",
                input_type=InputType.DEPTH,
                default_value=12000, unit="ft", required=True))
            .inp(TemplateInput(
                key="log_top", label="Log Top Depth",
                input_type=InputType.DEPTH,
                default_value=2000, unit="ft"))
            .inp(TemplateInput(
                key="run1_tools", label="Run 1 Tools",
                input_type=InputType.TEXT,
                default_value="GR + Resistivity + SP"))
            .inp(TemplateInput(
                key="run2_tools", label="Run 2 Tools",
                input_type=InputType.TEXT,
                default_value="Density + Neutron + Sonic"))
            .step("Rig up logging unit. Calibrate tools.",
                  TimeC.RIG_UP_LOG,
                  responsible="Logging Engineer")
            .step("Run 1: {run1_tools}. "
                  "Log {log_bottom} ft to {log_top} ft.",
                  TimeC.LOG_RUN,
                  responsible="Logging Engineer",
                  parameters={"run1_tools": "GR+RES+SP",
                               "log_bottom": 12000,
                               "log_top": 2000})
            .step("Run 2: {run2_tools}.",
                  TimeC.LOG_RUN,
                  responsible="Logging Engineer",
                  parameters={"run2_tools": "DEN+NEU+SON"})
            .step("Rig down. QC log data. Transmit.",
                  TimeC.RIG_DOWN_LOG,
                  responsible="Logging Engineer")
            .build()
        )

    # ----------------------------------------------------------------
    # LOAD ALL PHASES
    # ----------------------------------------------------------------

    def _load_default_phases(self) -> None:
        builders: Dict[str, Callable] = {
            "RIG_UP_SPUD":        self._build_rig_up_spud,
            "SURFACE_HOLE":       self._build_surface_hole,
            "SURFACE_CASING":     self._build_surface_casing,
            "INTERMEDIATE_HOLE":  self._build_intermediate_hole,
            "BOP_TEST":           self._build_bop_test,
            "PRODUCTION_HOLE":    self._build_production_hole,
            "WELL_TESTING":       self._build_well_testing,
            "CEMENT_PRIMARY":     self._build_primary_cement,
            "TRIP_OUT":           self._build_trip_out,
            "TRIP_IN":            self._build_trip_in,
            "CORING":             self._build_coring,
            "FISHING":            self._build_fishing,
            "MUD_CONVERSION":     self._build_mud_conversion,
            "ABANDONMENT":        self._build_abandonment,
            "HORIZONTAL_SECTION": self._build_horizontal_section,
            "WIRELINE_LOGGING":   self._build_wireline_logging,
        }
        with self._rlock:
            for pid, builder in builders.items():
                try:
                    self._phases[pid] = builder()
                except Exception as e:
                    logger.error(
                        f"Phase build failed {pid}: {e}")

    def _load_default_templates(self) -> None:
        def make(tid: str, name: str,
                 well_type: str,
                 phase_ids: List[str]) -> WellTemplate:
            phases = [
                copy.deepcopy(self._phases[pid])
                for pid in phase_ids
                if pid in self._phases
            ]
            return WellTemplate(
                template_id=tid,
                template_name=name,
                well_type=well_type,
                phases=phases,
                global_inputs=self._global_inputs())

        defs = [
            ("TMPL_VERT_EXPLORATION",
             "Standard Vertical Exploration Well",
             "Exploration",
             ["RIG_UP_SPUD", "SURFACE_HOLE",
              "SURFACE_CASING", "BOP_TEST",
              "INTERMEDIATE_HOLE", "CEMENT_PRIMARY",
              "PRODUCTION_HOLE", "WIRELINE_LOGGING",
              "WELL_TESTING"]),
            ("TMPL_DEV_DEVELOPMENT",
             "Deviated Development Well",
             "Development",
             ["RIG_UP_SPUD", "SURFACE_HOLE",
              "SURFACE_CASING", "BOP_TEST",
              "INTERMEDIATE_HOLE", "CEMENT_PRIMARY",
              "PRODUCTION_HOLE", "WIRELINE_LOGGING"]),
            ("TMPL_HORIZONTAL",
             "Horizontal Well with Extended Lateral",
             "Development",
             ["RIG_UP_SPUD", "SURFACE_HOLE",
              "SURFACE_CASING", "BOP_TEST",
              "INTERMEDIATE_HOLE", "CEMENT_PRIMARY",
              "MUD_CONVERSION", "HORIZONTAL_SECTION",
              "WIRELINE_LOGGING"]),
        ]

        with self._rlock:
            for tid, name, wtype, pids in defs:
                try:
                    self._templates[tid] = make(
                        tid, name, wtype, pids)
                except Exception as e:
                    logger.error(
                        f"Template build failed {tid}: {e}")

    # ----------------------------------------------------------------
    # PUBLIC API
    # ----------------------------------------------------------------

    def get_all_phase_ids(self) -> List[str]:
        with self._rlock:
            return list(self._phases.keys())

    def get_phase(self, phase_id: str
                  ) -> Optional[PhaseTemplate]:
        with self._rlock:
            p = self._phases.get(phase_id)
            return copy.deepcopy(p) if p else None

    def get_all_template_ids(self) -> List[str]:
        with self._rlock:
            return list(self._templates.keys())

    def get_template(self, template_id: str
                     ) -> Optional[WellTemplate]:
        with self._rlock:
            t = self._templates.get(template_id)
            return copy.deepcopy(t) if t else None

    def get_phases_by_category(
            self, category: TemplateCategory
    ) -> List[PhaseTemplate]:
        with self._rlock:
            return [copy.deepcopy(p)
                    for p in self._phases.values()
                    if p.category == category]

    def create_custom_template(
            self, template_name: str, well_type: str,
            phase_ids: List[str], company: str = "",
            field_name: str = "",
            rig_name: str = "") -> WellTemplate:
        tid = f"TMPL_CUSTOM_{uuid.uuid4().hex[:8].upper()}"
        phases = [self.get_phase(pid)
                  for pid in phase_ids
                  if self.get_phase(pid) is not None]
        tmpl = WellTemplate(
            template_id=tid,
            template_name=template_name,
            well_type=well_type,
            company=company,
            field_name=field_name,
            rig_name=rig_name,
            phases=[p for p in phases if p],
            global_inputs=self._global_inputs())
        with self._rlock:
            self._templates[tid] = tmpl
        return copy.deepcopy(tmpl)

    def export_template_json(
            self, template_id: str) -> Optional[str]:
        tmpl = self.get_template(template_id)
        if tmpl:
            return json.dumps(
                tmpl.to_dict(), indent=2,
                ensure_ascii=False)
        return None

    def import_template_json(
            self, json_str: str) -> Optional[WellTemplate]:
        try:
            data = json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON: {e}")
            return None

        required = {"template_id", "template_name",
                    "well_type", "phases"}
        if not required.issubset(data.keys()):
            logger.error(
                f"Missing keys: "
                f"{required - data.keys()}")
            return None

        try:
            phases = []
            for p_data in data.get('phases', []):
                steps = []
                for s_data in p_data.get('steps', []):
                    steps.append(OperationalStep(
                        step_number=s_data['step_number'],
                        description=s_data['description'],
                        duration_hours=s_data['duration_hours'],
                        responsible=s_data.get('responsible', ''),
                        safety_note=s_data.get('safety_note', ''),
                        equipment=s_data.get('equipment', []),
                        parameters=s_data.get('parameters', {}),
                        risk_level=RiskLevel(
                            s_data.get('risk_level', 'Low')),
                        checkpoint=s_data.get('checkpoint', False)))

                inputs = []
                for i_data in p_data.get('inputs', []):
                    inputs.append(TemplateInput(
                        key=i_data['key'],
                        label=i_data['label'],
                        input_type=InputType(
                            i_data.get('type', 'text')),
                        default_value=i_data.get('default', ''),
                        options=i_data.get('options', []),
                        unit=i_data.get('unit', ''),
                        required=i_data.get('required', False)))

                phases.append(PhaseTemplate(
                    phase_id=p_data['phase_id'],
                    phase_name=p_data['phase_name'],
                    category=TemplateCategory(
                        p_data.get('category',
                                   'Surface Hole Section')),
                    description=p_data.get('description', ''),
                    steps=steps,
                    prerequisites=p_data.get('prerequisites', []),
                    expected_duration_hours=p_data.get(
                        'expected_duration_hours', 0),
                    risk_level=RiskLevel(
                        p_data.get('risk_level', 'Medium')),
                    inputs=inputs))

            tmpl = WellTemplate(
                template_id=data['template_id'],
                template_name=data['template_name'],
                well_type=data['well_type'],
                company=data.get('company', ''),
                phases=phases,
                version=data.get('version', '1.0'))

            with self._rlock:
                self._templates[tmpl.template_id] = tmpl
            return copy.deepcopy(tmpl)

        except (KeyError, ValueError) as e:
            logger.error(f"Import error: {e}")
            return None

    def get_template_summary(
            self, template_id: str) -> Optional[Dict]:
        tmpl = self.get_template(template_id)
        if not tmpl:
            return None
        total_steps = sum(
            len(p.steps) for p in tmpl.phases)
        checkpoints = sum(
            len(p.get_checkpoints())
            for p in tmpl.phases)
        critical = [p.phase_name for p in tmpl.phases
                    if p.risk_level in
                    (RiskLevel.HIGH, RiskLevel.CRITICAL)]
        all_inputs = tmpl.get_all_inputs()
        return {
            "template_name": tmpl.template_name,
            "well_type": tmpl.well_type,
            "total_phases": len(tmpl.phases),
            "total_steps": total_steps,
            "total_hours": round(
                tmpl.total_well_duration(), 1),
            "total_days": round(
                tmpl.total_well_duration() / 24, 1),
            "critical_phases": critical,
            "checkpoints": checkpoints,
            "phase_list": [p.phase_name
                           for p in tmpl.phases],
            "total_inputs": len(all_inputs),
            "required_inputs": sum(
                1 for i in all_inputs if i.required),
        }


# ============================================================================
# PARAMETER RESOLVER
# ============================================================================

class ParameterResolver:

    @staticmethod
    def find_missing(step: OperationalStep
                     ) -> List[str]:
        import re
        desc = step.description
        placeholders = re.findall(r'\{(\w+)\}', desc)
        return [p for p in placeholders
                if p not in step.parameters]

    @staticmethod
    def resolve_step(
            step: OperationalStep,
            extra: Dict[str, Any] = None
    ) -> OperationalStep:
        resolved = copy.deepcopy(step)
        if extra:
            resolved.parameters.update(extra)
        resolved.description = (
            resolved.resolve_description())
        return resolved

    @staticmethod
    def resolve_phase(
            phase: PhaseTemplate,
            extra: Dict[str, Any] = None
    ) -> PhaseTemplate:
        resolved = copy.deepcopy(phase)
        resolved.steps = [
            ParameterResolver.resolve_step(s, extra)
            for s in resolved.steps]
        return resolved

    @staticmethod
    def resolve_template(
            template: WellTemplate,
            user_inputs: Dict[str, Any] = None
    ) -> WellTemplate:
        resolved = copy.deepcopy(template)
        resolved.phases = [
            ParameterResolver.resolve_phase(
                p, user_inputs)
            for p in resolved.phases]
        return resolved


# ============================================================================
# WORD EXPORTER
# ============================================================================

class TemplateWordExporter:
    """Export template to Word - 4 formats"""

    FORMATS = {
        "standard":     "Standard A4 Report",
        "compact":      "Compact (smaller fonts)",
        "checklist":    "Checklist Table Format",
        "presentation": "A4 Landscape Presentation",
    }

    def __init__(self, template: WellTemplate,
                 user_inputs: Dict[str, Any] = None):
        self.template = template
        self.user_inputs = user_inputs or {}
        self._resolved: Optional[WellTemplate] = None

    def _get_resolved(self) -> WellTemplate:
        if self._resolved is None:
            self._resolved = ParameterResolver.resolve_template(
                self.template, self.user_inputs)
        return self._resolved

    def export(self, file_path: str,
               format_key: str = "standard") -> bool:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            return False

        fmt_map = {
            "standard":     self._export_standard,
            "compact":      self._export_compact,
            "checklist":    self._export_checklist,
            "presentation": self._export_presentation,
        }
        fn = fmt_map.get(
            format_key, self._export_standard)
        try:
            fn(file_path)
            return True
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return False

    def _setup_doc(self, margin: float = 2.0):
        from docx import Document
        from docx.shared import Cm
        doc = Document()
        s = doc.sections[0]
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.top_margin = Cm(margin)
        s.bottom_margin = Cm(margin)
        s.left_margin = Cm(margin + 0.5)
        s.right_margin = Cm(margin)
        return doc

    def _add_cover(self, doc, tmpl: WellTemplate):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} '
            f'w:fill="0C2D48" w:val="clear"/>'))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"  {tmpl.template_name.upper()}  ")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")

        rows = [
            ("Well", self.user_inputs.get('well_name', '____')),
            ("Field", self.user_inputs.get(
                'field_name', tmpl.field_name or '____')),
            ("Operator", self.user_inputs.get(
                'operator', tmpl.company or '____')),
            ("Well Type", tmpl.well_type),
            ("Total Duration",
             f"{tmpl.total_well_duration() / 24:.1f} days"),
            ("Generated",
             datetime.now().strftime("%d-%b-%Y")),
        ]

        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = 'Table Grid'

        for i, (k, v) in enumerate(rows):
            row = tbl.rows[i]
            lbl = row.cells[0]
            lbl.text = ""
            p2 = lbl.paragraphs[0]
            r2 = p2.add_run(k + ":")
            r2.bold = True
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x0C, 0x2D, 0x48)
            from docx.oxml import parse_xml
            shd = parse_xml(
                f'<w:shd {nsdecls("w")} '
                f'w:fill="EBF5FB"/>')
            lbl._tc.get_or_add_tcPr().append(shd)

            val_c = row.cells[1]
            val_c.text = ""
            p3 = val_c.paragraphs[0]
            r3 = p3.add_run(str(v))
            r3.font.size = Pt(10)

        doc.add_page_break()

    def _add_phase(self, doc, phase: PhaseTemplate,
                   idx: int, font_size: int = 10):
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        colors = {
            RiskLevel.LOW:      "27AE60",
            RiskLevel.MEDIUM:   "1B4F72",
            RiskLevel.HIGH:     "E67E22",
            RiskLevel.CRITICAL: "C0392B",
        }
        color = colors.get(phase.risk_level, "1B4F72")

        ph = doc.add_paragraph()
        pPr = ph._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} '
            f'w:fill="{color}" w:val="clear"/>'))
        r = ph.add_run(
            f"  Phase {idx}: {phase.phase_name}  "
            f"[{phase.risk_level.value.upper()} RISK]")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        info = doc.add_paragraph()
        ir = info.add_run(
            f"Category: {phase.category.value}  |  "
            f"Duration: {phase.effective_duration():.1f} hrs  |  "
            f"Steps: {len(phase.steps)}  |  "
            f"Checkpoints: {len(phase.get_checkpoints())}")
        ir.font.size = Pt(9)
        ir.font.italic = True
        ir.font.color.rgb = RGBColor(0x56, 0x6C, 0x73)

        if phase.prerequisites:
            pp = doc.add_paragraph()
            rp = pp.add_run(
                "Prerequisites: " +
                ", ".join(phase.prerequisites))
            rp.font.size = Pt(9)
            rp.font.italic = True
            rp.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)

        doc.add_paragraph("")

        for step in phase.steps:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_after = Pt(3)

            sr = p.add_run(f"{step.step_number}. ")
            sr.bold = True
            sr.font.size = Pt(font_size)
            sr.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

            dr = p.add_run(step.description)
            dr.font.size = Pt(font_size)

            dur_r = p.add_run(
                f"  [{step.duration_hours:.1f}h]")
            dur_r.font.size = Pt(font_size - 2)
            dur_r.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)
            dur_r.font.italic = True

            resp_r = p.add_run(f"  -- {step.responsible}")
            resp_r.font.size = Pt(font_size - 1)
            resp_r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

            if step.safety_note:
                sn = doc.add_paragraph()
                sn.paragraph_format.left_indent = Cm(1.0)
                sn.paragraph_format.space_after = Pt(1)
                snr = sn.add_run(f"WARNING: {step.safety_note}")
                snr.font.size = Pt(font_size - 1)
                snr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                snr.font.italic = True

            if step.checkpoint:
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Cm(1.0)
                cpr = cp.add_run(
                    "CHECKPOINT - Supervisor approval required")
                cpr.font.size = Pt(font_size - 1)
                cpr.font.bold = True
                cpr.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

        if phase.bha_components:
            bp = doc.add_paragraph()
            br1 = bp.add_run("BHA: ")
            br1.bold = True
            br1.font.size = Pt(font_size - 1)
            br2 = bp.add_run(
                " -> ".join(phase.bha_components))
            br2.font.size = Pt(font_size - 1)
            br2.font.color.rgb = RGBColor(0x56, 0x6C, 0x73)

        doc.add_paragraph("")
        doc.add_page_break()

    def _export_standard(self, file_path: str):
        doc = self._setup_doc(2.0)
        tmpl = self._get_resolved()
        self._add_cover(doc, tmpl)

        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        # TOC
        th = doc.add_paragraph()
        pPr = th._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} '
            f'w:fill="0C2D48" w:val="clear"/>'))
        th.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = th.add_run("  TABLE OF CONTENTS  ")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph("")

        for i, phase in enumerate(tmpl.phases, 1):
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Cm(0.3)
            r1 = tp.add_run(f"{i}. ")
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = tp.add_run(phase.phase_name)
            r2.font.size = Pt(10)
            r3 = tp.add_run(
                f"  ({phase.effective_duration():.1f} hrs)")
            r3.font.size = Pt(9)
            r3.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

        doc.add_page_break()

        for i, phase in enumerate(tmpl.phases, 1):
            self._add_phase(doc, phase, i, font_size=10)

        doc.save(file_path)

    def _export_compact(self, file_path: str):
        doc = self._setup_doc(1.5)
        tmpl = self._get_resolved()
        self._add_cover(doc, tmpl)
        for i, phase in enumerate(tmpl.phases, 1):
            self._add_phase(doc, phase, i, font_size=8)
        doc.save(file_path)

    def _export_checklist(self, file_path: str):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        doc = self._setup_doc(1.5)
        tmpl = self._get_resolved()
        self._add_cover(doc, tmpl)

        for phase in tmpl.phases:
            ph = doc.add_paragraph()
            pPr = ph._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:shd {nsdecls("w")} '
                f'w:fill="0C2D48" w:val="clear"/>'))
            r = ph.add_run(f"  {phase.phase_name}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            if not phase.steps:
                continue

            headers = ["#", "Operation Description",
                       "Duration", "Responsible",
                       "Done", "Init."]
            tbl = doc.add_table(
                rows=len(phase.steps) + 1, cols=6)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            hrow = tbl.rows[0]
            for ci, h in enumerate(headers):
                cell = hrow.cells[ci]
                cell.text = ""
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(h)
                cr.bold = True
                cr.font.size = Pt(8)
                cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} '
                    f'w:fill="0C2D48"/>')
                cell._tc.get_or_add_tcPr().append(shd)

            widths = [0.3, 4.2, 0.7, 1.5, 0.5, 0.6]
            for row in tbl.rows:
                for ci, w in enumerate(widths):
                    row.cells[ci].width = Inches(w)

            for si, step in enumerate(phase.steps):
                row = tbl.rows[si + 1]
                vals = [
                    str(step.step_number),
                    step.description[:120] + (
                        "..." if len(step.description) > 120
                        else ""),
                    f"{step.duration_hours:.1f}h",
                    step.responsible[:25],
                    "[ ]", "",
                ]
                for ci, val in enumerate(vals):
                    cell = row.cells[ci]
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    cp.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                        if ci == 1
                        else WD_ALIGN_PARAGRAPH.CENTER)
                    cr = cp.add_run(val)
                    cr.font.size = Pt(8)

                    if step.checkpoint:
                        shd = parse_xml(
                            f'<w:shd {nsdecls("w")} '
                            f'w:fill="FEF9E7"/>')
                        cell._tc.get_or_add_tcPr(
                        ).append(shd)
                    elif si % 2 == 0:
                        shd = parse_xml(
                            f'<w:shd {nsdecls("w")} '
                            f'w:fill="EBF5FB"/>')
                        cell._tc.get_or_add_tcPr(
                        ).append(shd)

            sig = doc.add_paragraph()
            sig.add_run(
                "Sign-Off: Driller: _________  "
                "Toolpusher: _________  "
                "Co. Man: _________  "
                "Date: _________"
            ).font.size = Pt(9)

            doc.add_paragraph("")
            doc.add_page_break()

        doc.save(file_path)

    def _export_presentation(self, file_path: str):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        doc = Document()
        s = doc.sections[0]
        s.page_height = Cm(21.0)
        s.page_width = Cm(29.7)
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)

        tmpl = self._get_resolved()
        self._add_cover(doc, tmpl)

        for phase in tmpl.phases:
            ph = doc.add_paragraph()
            pPr = ph._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:shd {nsdecls("w")} '
                f'w:fill="0C2D48" w:val="clear"/>'))
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ph.add_run(
                f"  {phase.phase_name.upper()}  ")
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            doc.add_paragraph("")

            for step in phase.steps:
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Cm(0.5)
                sp.paragraph_format.space_after = Pt(6)
                sr = sp.add_run(
                    f"Step {step.step_number}: "
                    f"{step.description}")
                sr.font.size = Pt(13)
                if step.checkpoint:
                    sr.bold = True
                    sr.font.color.rgb = RGBColor(
                        0xE9, 0x45, 0x60)

                dur = doc.add_paragraph()
                dur.paragraph_format.left_indent = Cm(1.5)
                dr = dur.add_run(
                    f"Time: {step.duration_hours:.1f}h  |  "
                    f"{step.responsible}")
                dr.font.size = Pt(11)
                dr.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)
                dr.font.italic = True

            doc.add_page_break()

        doc.save(file_path)


# ============================================================================
# CONVENIENCE FUNCTIONS
# ============================================================================

def get_template_library() -> TemplateLibrary:
    return TemplateLibrary()


def list_available_templates() -> List[Dict]:
    lib = get_template_library()
    result = []
    for tid in lib.get_all_template_ids():
        summary = lib.get_template_summary(tid)
        if summary:
            summary['template_id'] = tid
            result.append(summary)
    return result


def list_available_phases() -> List[Dict]:
    lib = get_template_library()
    result = []
    for pid in lib.get_all_phase_ids():
        phase = lib.get_phase(pid)
        if phase:
            result.append({
                'phase_id': pid,
                'phase_name': phase.phase_name,
                'category': phase.category.value,
                'steps_count': len(phase.steps),
                'duration_hours': phase.effective_duration(),
                'risk_level': phase.risk_level.value,
                'checkpoints': len(phase.get_checkpoints()),
                'inputs_count': len(phase.inputs),
            })
    return result


def export_template_to_word(
        template_id: str, file_path: str,
        format_key: str = "standard",
        user_inputs: Dict[str, Any] = None) -> bool:
    lib = get_template_library()
    tmpl = lib.get_template(template_id)
    if not tmpl:
        return False
    exporter = TemplateWordExporter(tmpl, user_inputs)
    return exporter.export(file_path, format_key)


# ============================================================================
# UI DIALOG (PySide6) - Integrated
# ============================================================================

try:
    from PySide6.QtWidgets import (
        QDialog, QVBoxLayout, QHBoxLayout,
        QLabel, QPushButton, QListWidget,
        QListWidgetItem, QSplitter, QWidget,
        QScrollArea, QMessageBox, QTabWidget,
        QGroupBox, QFormLayout, QLineEdit,
        QComboBox, QDoubleSpinBox, QSpinBox,
        QCheckBox, QFileDialog, QApplication,
        QTextEdit, QProgressBar, QInputDialog,
        QAbstractItemView, QTreeWidget,
        QTreeWidgetItem, QHeaderView, QFrame
    )
    from PySide6.QtCore import (
        Qt, Signal, QSize, QThread, QObject)
    from PySide6.QtGui import QColor, QFont
    HAS_QT = True
except ImportError:
    HAS_QT = False


if HAS_QT:
    TEMPLATE_STYLE = """
    QDialog, QWidget {
        background-color: #0a0f1a;
        color: #e0e0e0;
        font-family: 'Segoe UI';
        font-size: 11px;
    }
    QLabel { color: #c0ccd8; }
    QLabel#title {
        color: #e94560;
        font-size: 20px;
        font-weight: bold;
    }
    QLabel#subtitle { color: #8899aa; font-size: 10px; }
    QLabel#section {
        color: #e94560;
        font-size: 12px;
        font-weight: bold;
    }
    QGroupBox {
        border: 2px solid #1a2744;
        border-radius: 8px;
        margin-top: 12px;
        padding-top: 18px;
        font-weight: bold;
        color: #e94560;
        background: #0d1525;
    }
    QGroupBox::title {
        subcontrol-origin: margin;
        left: 12px;
        padding: 0 8px;
    }
    QLineEdit, QDoubleSpinBox, QSpinBox, QComboBox {
        background: #0d1525;
        border: 1px solid #1a2744;
        border-radius: 4px;
        padding: 5px 8px;
        color: #e0e0e0;
        min-height: 26px;
    }
    QLineEdit:focus, QComboBox:focus,
    QDoubleSpinBox:focus {
        border: 2px solid #e94560;
    }
    QComboBox QAbstractItemView {
        background: #0d1525;
        color: #e0e0e0;
        selection-background-color: #e94560;
    }
    QCheckBox { spacing: 8px; }
    QCheckBox::indicator {
        width: 16px; height: 16px;
        border: 2px solid #0f3460;
        border-radius: 3px;
        background: #0a0f1a;
    }
    QCheckBox::indicator:checked {
        background: #e94560;
        border-color: #e94560;
    }
    QPushButton {
        background: #0f3460;
        color: #fff;
        border: none;
        border-radius: 6px;
        padding: 8px 18px;
        font-weight: bold;
        min-height: 32px;
    }
    QPushButton:hover { background: #1a5276; }
    QPushButton#gen_btn {
        background: qlineargradient(
            x1:0, y1:0, x2:1, y2:0,
            stop:0 #e94560, stop:1 #c0392b);
        font-size: 13px;
        padding: 12px 30px;
        min-height: 42px;
    }
    QPushButton#gen_btn:hover {
        background: qlineargradient(
            x1:0, y1:0, x2:1, y2:0,
            stop:0 #ff6b81, stop:1 #e74c3c);
    }
    QPushButton#add_btn { background: #27ae60; }
    QPushButton#del_btn { background: #c0392b; }
    QPushButton#cancel_btn {
        background: #1a2744;
        color: #a0b0c0;
        border: 1px solid #2c3e50;
    }
    QTabWidget::pane {
        border: 1px solid #1a2744;
        background: #0d1525;
    }
    QTabBar::tab {
        background: #0a0f1a;
        color: #8899aa;
        padding: 7px 16px;
        font-size: 10px;
    }
    QTabBar::tab:selected {
        background: #1a2744;
        color: #e94560;
        border-bottom: 2px solid #e94560;
    }
    QListWidget, QTreeWidget {
        background: #0d1525;
        border: 2px solid #1a2744;
        border-radius: 6px;
        color: #d0d8e0;
        outline: none;
    }
    QListWidget::item, QTreeWidget::item {
        padding: 4px 6px;
    }
    QListWidget::item:selected,
    QTreeWidget::item:selected {
        background: #1a2744;
        border: 1px solid #e94560;
        color: #fff;
    }
    QTextEdit {
        background: #0d1525;
        border: 1px solid #1a2744;
        border-radius: 4px;
        color: #c0ccd8;
        padding: 8px;
        font-size: 10px;
    }
    QProgressBar {
        border: 1px solid #1a2744;
        border-radius: 5px;
        text-align: center;
        color: #fff;
        background: #0a0f1a;
        min-height: 22px;
    }
    QProgressBar::chunk {
        background: qlineargradient(
            x1:0, y1:0, x2:1, y2:0,
            stop:0 #e94560, stop:1 #0f3460);
        border-radius: 4px;
    }
    QScrollArea { border: none; background: transparent; }
    QSplitter::handle {
        background: #1a2744;
        width: 2px;
    }
    """

    class InputWidget:
        """Factory for input widgets"""

        @staticmethod
        def create(inp: TemplateInput) -> QWidget:
            t = inp.input_type
            if t == InputType.TEXT:
                w = QLineEdit()
                w.setPlaceholderText(
                    inp.tooltip or inp.label)
                if inp.default_value:
                    w.setText(str(inp.default_value))
                return w
            elif t in (InputType.NUMBER, InputType.DEPTH,
                       InputType.PRESSURE,
                       InputType.MUDWEIGHT):
                w = QDoubleSpinBox()
                ranges = {
                    InputType.NUMBER: (0.0, 99999.0, 1),
                    InputType.DEPTH: (0.0, 50000.0, 0),
                    InputType.PRESSURE: (0.0, 30000.0, 0),
                    InputType.MUDWEIGHT: (5.0, 22.0, 2),
                }
                mn, mx, dec = ranges.get(
                    t, (0.0, 99999.0, 1))
                w.setMinimum(inp.min_value or mn)
                w.setMaximum(inp.max_value or mx)
                w.setDecimals(dec)
                if inp.unit:
                    w.setSuffix(f" {inp.unit}")
                if inp.default_value:
                    try:
                        w.setValue(float(inp.default_value))
                    except (ValueError, TypeError):
                        pass
                if inp.tooltip:
                    w.setToolTip(inp.tooltip)
                return w
            elif t == InputType.COMBO:
                w = QComboBox()
                w.addItems(inp.options)
                if (inp.default_value
                        and inp.default_value in inp.options):
                    w.setCurrentText(str(inp.default_value))
                return w
            elif t == InputType.CHECKBOX:
                w = QCheckBox()
                w.setChecked(bool(inp.default_value))
                if inp.tooltip:
                    w.setToolTip(inp.tooltip)
                return w
            else:
                w = QLineEdit()
                if inp.default_value:
                    w.setText(str(inp.default_value))
                return w

        @staticmethod
        def get_value(widget: QWidget) -> Any:
            if isinstance(widget, QLineEdit):
                return widget.text()
            if isinstance(widget, QDoubleSpinBox):
                return widget.value()
            if isinstance(widget, QSpinBox):
                return widget.value()
            if isinstance(widget, QComboBox):
                return widget.currentText()
            if isinstance(widget, QCheckBox):
                return widget.isChecked()
            return ""

        @staticmethod
        def set_value(widget: QWidget, value: Any):
            try:
                if isinstance(widget, QLineEdit):
                    widget.setText(str(value))
                elif isinstance(widget,
                                (QDoubleSpinBox, QSpinBox)):
                    widget.setValue(float(value))
                elif isinstance(widget, QComboBox):
                    idx = widget.findText(str(value))
                    if idx >= 0:
                        widget.setCurrentIndex(idx)
                elif isinstance(widget, QCheckBox):
                    widget.setChecked(bool(value))
            except (ValueError, TypeError):
                pass

    class InputForm(QScrollArea):
        """Auto-generated input form from TemplateInput list"""

        valuesChanged = Signal()

        def __init__(self, inputs: List[TemplateInput],
                     parent=None):
            super().__init__(parent)
            self.setWidgetResizable(True)
            self._inputs = inputs
            self._widgets: Dict[str, QWidget] = {}

            container = QWidget()
            layout = QFormLayout(container)
            layout.setSpacing(8)
            layout.setContentsMargins(12, 12, 12, 12)
            layout.setLabelAlignment(Qt.AlignRight)

            for inp in inputs:
                widget = InputWidget.create(inp)
                lbl_text = inp.label + (" *" if inp.required else "")
                lbl = QLabel(lbl_text)
                if inp.required:
                    lbl.setStyleSheet(
                        "color: #e94560; font-weight: bold;")
                layout.addRow(lbl, widget)
                self._widgets[inp.key] = widget

            if any(i.required for i in inputs):
                note = QLabel("* Required fields")
                note.setStyleSheet(
                    "color: #e94560; font-size: 9px;")
                layout.addRow("", note)

            self.setWidget(container)

        def get_values(self) -> Dict[str, Any]:
            return {k: InputWidget.get_value(w)
                    for k, w in self._widgets.items()}

        def validate(self) -> List[str]:
            errors = []
            for inp in self._inputs:
                if inp.required:
                    w = self._widgets.get(inp.key)
                    if w:
                        val = InputWidget.get_value(w)
                        if not val and val != 0:
                            errors.append(
                                f"'{inp.label}' is required")
            return errors

        def reset(self):
            for inp in self._inputs:
                if inp.key in self._widgets:
                    InputWidget.set_value(
                        self._widgets[inp.key],
                        inp.default_value)

    class ExportWorker(QObject):
        """Background Word export worker"""
        finished = Signal(bool, str)
        progress = Signal(int)

        def __init__(self, template: WellTemplate,
                     file_path: str, format_key: str,
                     user_inputs: Dict[str, Any]):
            super().__init__()
            self.template = template
            self.file_path = file_path
            self.format_key = format_key
            self.user_inputs = user_inputs

        def run(self):
            try:
                self.progress.emit(10)
                exporter = TemplateWordExporter(
                    self.template, self.user_inputs)
                self.progress.emit(30)
                ok = exporter.export(
                    self.file_path, self.format_key)
                self.progress.emit(100)
                self.finished.emit(ok, self.file_path)
            except Exception as e:
                self.finished.emit(False, str(e))

    class OperationalTemplateDialog(QDialog):
        """Main template selector dialog"""

        templateApplied = Signal(object)

        def __init__(self, parent=None):
            super().__init__(parent)
            self.setWindowTitle(
                "Operational Template Library")
            self.setMinimumSize(1250, 800)
            self.setStyleSheet(TEMPLATE_STYLE)

            self.lib = TemplateLibrary()
            self._current: Optional[WellTemplate] = None
            self._form: Optional[InputForm] = None
            self._thread: Optional[QThread] = None

            self._build_ui()
            self._load_templates()

        def _build_ui(self):
            layout = QVBoxLayout(self)
            layout.setSpacing(8)
            layout.setContentsMargins(12, 12, 12, 12)

            # Title
            hdr = QHBoxLayout()
            t = QLabel("Operational Template Library")
            t.setObjectName("title")
            hdr.addWidget(t)
            hdr.addStretch()
            sub = QLabel("Select template, configure parameters, generate Word")
            sub.setObjectName("subtitle")
            hdr.addWidget(sub)
            layout.addLayout(hdr)

            # Main splitter
            splitter = QSplitter(Qt.Horizontal)

            # Left panel
            left = QWidget()
            ll = QVBoxLayout(left)
            ll.setContentsMargins(0, 0, 0, 0)

            lbl = QLabel("Available Templates:")
            lbl.setObjectName("section")
            ll.addWidget(lbl)

            self.template_list = QListWidget()
            self.template_list.setMinimumWidth(300)
            self.template_list.currentItemChanged.connect(
                self._on_template_selected)
            ll.addWidget(self.template_list)

            # Custom builder
            grp = QGroupBox("Custom Template Builder")
            gl = QVBoxLayout()
            gl.addWidget(QLabel("Select phases:"))
            self.phase_list = QListWidget()
            self.phase_list.setSelectionMode(
                QAbstractItemView.ExtendedSelection)
            self.phase_list.setMaximumHeight(160)
            gl.addWidget(self.phase_list)
            self._load_phase_list()
            btn_create = QPushButton("Create Custom Template")
            btn_create.setObjectName("add_btn")
            btn_create.clicked.connect(self._create_custom)
            gl.addWidget(btn_create)
            grp.setLayout(gl)
            ll.addWidget(grp)

            # JSON Import/Export
            ie = QHBoxLayout()
            btn_imp = QPushButton("Import JSON")
            btn_imp.clicked.connect(self._import)
            ie.addWidget(btn_imp)
            btn_exp = QPushButton("Export JSON")
            btn_exp.clicked.connect(self._export_json)
            ie.addWidget(btn_exp)
            ll.addLayout(ie)

            splitter.addWidget(left)

            # Right panel (tabs)
            self.tabs = QTabWidget()

            # Tab 1: Overview
            self.overview_text = QTextEdit()
            self.overview_text.setReadOnly(True)
            self.tabs.addTab(self.overview_text, "Overview")

            # Tab 2: Parameters
            self.params_container = QWidget()
            self.params_layout = QVBoxLayout(
                self.params_container)
            self.params_layout.addWidget(
                QLabel("Select a template to configure"))
            self.tabs.addTab(
                self.params_container, "Parameters")

            # Tab 3: Phase Tree
            self.phase_tree = QTreeWidget()
            self.phase_tree.setHeaderLabels([
                "Phase / Step", "Duration",
                "Risk", "Responsible", "CP"])
            hdr = self.phase_tree.header()
            hdr.setSectionResizeMode(
                0, QHeaderView.Stretch)
            self.tabs.addTab(
                self.phase_tree, "Phase Tree")

            splitter.addWidget(self.tabs)
            splitter.setSizes([380, 870])
            layout.addWidget(splitter)

            # Bottom bar
            bot = QHBoxLayout()
            bot.addStretch()

            self.progress = QProgressBar()
            self.progress.setMinimumWidth(200)
            self.progress.setVisible(False)
            bot.addWidget(self.progress)

            bot.addWidget(QLabel("Format:"))
            self.fmt_combo = QComboBox()
            self.fmt_combo.addItems([
                "standard -- Full A4 Report",
                "compact -- Compact Format",
                "checklist -- Checklist Table",
                "presentation -- Landscape Presentation",
            ])
            self.fmt_combo.setMinimumWidth(250)
            bot.addWidget(self.fmt_combo)

            btn_cancel = QPushButton("Close")
            btn_cancel.setObjectName("cancel_btn")
            btn_cancel.clicked.connect(self.reject)
            bot.addWidget(btn_cancel)

            self.btn_gen = QPushButton(
                "GENERATE WORD DOCUMENT")
            self.btn_gen.setObjectName("gen_btn")
            self.btn_gen.setEnabled(False)
            self.btn_gen.clicked.connect(self._generate)
            bot.addWidget(self.btn_gen)
            bot.addStretch()
            layout.addLayout(bot)

        def _load_templates(self):
            self.template_list.clear()
            for t in list_available_templates():
                item = QListWidgetItem()
                item.setData(Qt.UserRole, t['template_id'])
                item.setText(
                    f"{t['template_name']}\n"
                    f"     {t['total_days']:.1f} days  |  "
                    f"{t['total_phases']} phases  |  "
                    f"{t['checkpoints']} checkpoints")
                item.setSizeHint(QSize(0, 55))
                self.template_list.addItem(item)

        def _load_phase_list(self):
            self.phase_list.clear()
            cat_icons = {
                "Surface Hole Section": "S",
                "Intermediate Hole Section": "I",
                "Production Hole Section": "P",
                "Cementing Operations": "C",
                "BOP Test Operations": "B",
                "Tripping Operations": "T",
                "Logging Operations": "L",
                "Coring Operations": "X",
                "Fishing Operations": "F",
                "Mud Engineering Operations": "M",
                "Well Testing Operations": "W",
                "Well Abandonment / P&A": "A",
                "Horizontal / ERD Section": "H",
            }
            for p in list_available_phases():
                icon = cat_icons.get(p['category'], "-")
                item = QListWidgetItem(
                    f"[{icon}] {p['phase_name']}\n"
                    f"     {p['duration_hours']:.1f}h  |  "
                    f"{p['steps_count']} steps  |  "
                    f"{p['risk_level']}")
                item.setSizeHint(QSize(0, 42))
                item.setData(Qt.UserRole, p['phase_id'])
                self.phase_list.addItem(item)

        def _on_template_selected(self, current, previous):
            if not current:
                return
            tid = current.data(Qt.UserRole)
            tmpl = self.lib.get_template(tid)
            if not tmpl:
                return
            self._current = tmpl
            self.btn_gen.setEnabled(True)
            self._update_overview(tmpl)
            self._update_params(tmpl)
            self._update_tree(tmpl)

        def _update_overview(self, tmpl: WellTemplate):
            html = (
                f"<h3 style='color:#e94560'>"
                f"{tmpl.template_name}</h3>"
                f"<p style='color:#8899aa'>"
                f"Well Type: {tmpl.well_type}  |  "
                f"Duration: {tmpl.total_well_duration()/24:.1f} days"
                f"</p>")

            for i, phase in enumerate(tmpl.phases, 1):
                colors = {
                    'Low': '#27ae60', 'Medium': '#f1c40f',
                    'High': '#e67e22', 'Critical': '#e74c3c'}
                c = colors.get(
                    phase.risk_level.value, '#8899aa')
                html += (
                    f"<p>"
                    f"<span style='color:{c};"
                    f"font-weight:bold'>{i}. "
                    f"{phase.phase_name}</span>"
                    f"<span style='color:#8899aa;"
                    f"font-size:9px'>"
                    f"  {phase.effective_duration():.1f}h  "
                    f"|  {len(phase.steps)} steps</span>"
                    f"</p>")

            self.overview_text.setHtml(html)

        def _update_params(self, tmpl: WellTemplate):
            # Clear
            while self.params_layout.count():
                item = self.params_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()

            inputs = tmpl.get_all_inputs()
            if not inputs:
                self.params_layout.addWidget(
                    QLabel("No configurable parameters"))
                self._form = None
                return

            lbl = QLabel(
                f"{len(inputs)} parameters "
                f"({sum(1 for i in inputs if i.required)}"
                f" required)")
            lbl.setStyleSheet("color: #8899aa;")
            self.params_layout.addWidget(lbl)

            self._form = InputForm(inputs)
            self.params_layout.addWidget(self._form)

            btn_r = QPushButton("Reset to Defaults")
            btn_r.clicked.connect(
                lambda: self._form.reset()
                if self._form else None)
            self.params_layout.addWidget(btn_r)

        def _update_tree(self, tmpl: WellTemplate):
            self.phase_tree.clear()
            risk_icons = {
                RiskLevel.LOW: "[L]",
                RiskLevel.MEDIUM: "[M]",
                RiskLevel.HIGH: "[H]",
                RiskLevel.CRITICAL: "[!]",
            }
            risk_colors = {
                RiskLevel.LOW: QColor("#1a3320"),
                RiskLevel.MEDIUM: QColor("#332a00"),
                RiskLevel.HIGH: QColor("#331a00"),
                RiskLevel.CRITICAL: QColor("#330000"),
            }

            for i, phase in enumerate(tmpl.phases, 1):
                pi = QTreeWidgetItem()
                icon = risk_icons.get(
                    phase.risk_level, "[?]")
                pi.setText(0,
                           f"{icon} Phase {i}: "
                           f"{phase.phase_name}")
                pi.setText(1,
                           f"{phase.effective_duration():.1f}h")
                pi.setText(2, phase.risk_level.value)
                pi.setText(3, "")
                pi.setText(4, "")

                bg = risk_colors.get(
                    phase.risk_level, QColor("#0d1525"))
                for c in range(5):
                    pi.setBackground(c, bg)

                for step in phase.steps:
                    si = QTreeWidgetItem(pi)
                    si.setText(0,
                               f"  {step.step_number}. "
                               f"{step.description[:55]}"
                               + ("..." if len(step.description) > 55
                                  else ""))
                    si.setText(1,
                               f"{step.duration_hours:.1f}h")
                    si.setText(2,
                               step.risk_level.value
                               if step.risk_level != RiskLevel.LOW
                               else "")
                    si.setText(3, step.responsible[:20])
                    si.setText(4, "CP" if step.checkpoint else "")
                    if step.checkpoint:
                        for c in range(5):
                            si.setForeground(
                                c, QColor("#e94560"))

                self.phase_tree.addTopLevelItem(pi)
            self.phase_tree.expandAll()

        def _create_custom(self):
            selected = self.phase_list.selectedItems()
            if not selected:
                QMessageBox.warning(self, "No Phases",
                                    "Select at least one phase.")
                return
            name, ok = QInputDialog.getText(
                self, "Custom Template", "Template Name:",
                text="My Custom Template")
            if not ok or not name:
                return
            well_type, ok2 = QInputDialog.getItem(
                self, "Well Type", "Select:",
                ["Development", "Exploration",
                 "Appraisal", "Workover"],
                0, False)
            if not ok2:
                return
            phase_ids = [item.data(Qt.UserRole)
                         for item in selected]
            tmpl = self.lib.create_custom_template(
                name, well_type, phase_ids)
            item = QListWidgetItem()
            item.setData(Qt.UserRole, tmpl.template_id)
            item.setText(
                f"[Custom] {tmpl.template_name}\n"
                f"     {tmpl.total_well_duration()/24:.1f} days  "
                f"|  {len(tmpl.phases)} phases")
            item.setSizeHint(QSize(0, 55))
            self.template_list.addItem(item)
            self.template_list.setCurrentItem(item)
            QMessageBox.information(
                self, "Created",
                f"Custom template '{name}' created!")

        def _import(self):
            path, _ = QFileDialog.getOpenFileName(
                self, "Import Template", "",
                "JSON Files (*.json)")
            if not path:
                return
            try:
                with open(path, 'r',
                          encoding='utf-8') as f:
                    json_str = f.read()
                tmpl = self.lib.import_template_json(
                    json_str)
                if not tmpl:
                    QMessageBox.critical(
                        self, "Failed",
                        "Invalid template format.")
                    return
                self._load_templates()
                QMessageBox.information(
                    self, "Imported",
                    f"'{tmpl.template_name}' imported!")
            except Exception as e:
                QMessageBox.critical(
                    self, "Error", str(e))

        def _export_json(self):
            if not self._current:
                QMessageBox.warning(
                    self, "No Template",
                    "Select a template first.")
                return
            path, _ = QFileDialog.getSaveFileName(
                self, "Export Template",
                f"{self._current.template_name}.json",
                "JSON Files (*.json)")
            if not path:
                return
            json_str = self.lib.export_template_json(
                self._current.template_id)
            if not json_str:
                return
            try:
                with open(path, 'w',
                          encoding='utf-8') as f:
                    f.write(json_str)
                QMessageBox.information(
                    self, "Exported",
                    f"Exported to:\n{path}")
            except Exception as e:
                QMessageBox.critical(
                    self, "Error", str(e))

        def _generate(self):
            if not self._current:
                return

            if self._form:
                errors = self._form.validate()
                if errors:
                    QMessageBox.warning(
                        self, "Missing Fields",
                        "Fill required fields:\n\n" +
                        "\n".join(f"- {e}" for e in errors))
                    self.tabs.setCurrentIndex(1)
                    return

            fmt_text = self.fmt_combo.currentText()
            fmt_key = fmt_text.split(" -- ")[0].strip()

            user_inputs = {}
            if self._form:
                user_inputs = self._form.get_values()

            well_name = user_inputs.get('well_name', '')
            default_name = (
                f"Template_{self._current.template_name.replace(' ','_')}"
                + (f"_{well_name.replace(' ','_')}"
                   if well_name else "")
                + f"_{fmt_key}.docx")

            path, _ = QFileDialog.getSaveFileName(
                self, "Save Document", default_name,
                "Word Documents (*.docx)")
            if not path:
                return

            self.btn_gen.setEnabled(False)
            self.progress.setVisible(True)
            self.progress.setValue(0)

            self._thread = QThread()
            worker = ExportWorker(
                self._current, path,
                fmt_key, user_inputs)
            worker.moveToThread(self._thread)

            self._thread.started.connect(worker.run)
            worker.progress.connect(
                self.progress.setValue)
            worker.finished.connect(
                lambda ok, fp:
                self._on_done(ok, fp))
            worker.finished.connect(
                self._thread.quit)
            self._thread.finished.connect(
                self._thread.deleteLater)
            self._thread.start()

        def _on_done(self, success: bool,
                     file_path: str):
            self.progress.setVisible(False)
            self.btn_gen.setEnabled(True)
            if success:
                QMessageBox.information(
                    self, "Done",
                    f"Document generated!\n\n"
                    f"File: {file_path}")
                if os.name == 'nt':
                    try:
                        import os as _os
                        _os.startfile(file_path)
                    except Exception:
                        pass
            else:
                QMessageBox.critical(
                    self, "Failed",
                    f"Export failed:\n{file_path}")

        def closeEvent(self, event):
            if (self._thread
                    and self._thread.isRunning()):
                self._thread.quit()
                self._thread.wait(3000)
            super().closeEvent(event)

    def show_template_dialog(parent=None):
        """Quick access function"""
        dialog = OperationalTemplateDialog(parent)
        dialog.exec()