# ============================================================================
# DOCUMENT INTEGRITY & QA SUITE — Batch X
# File: tests/test_integrity.py
#
# Locks the fixes from the external technical reviews:
#   1. CRITICAL-override dialog actually offers Yes/No (dead-code bug)
#   2. Numeric inputs left empty are NOT treated as 0
#   3. Unified placeholder syntax {{x}} and {x}
#   4. Final unresolved-placeholder audit (no raw placeholders in Word)
#   5. Validation/Standards see UI aliases (fracture_gradient vs _ppg)
#   6. Depth units handled canonically in feet (no double conversion)
#   7. Procedure DB: user values substituted into steps; unresolved
#      parameters reported; defaults marked "(default)"
#   8. Time Breakdown cross-project contamination guard
#
# Run:  LD_LIBRARY_PATH=/tmp/glstubs PYTHONPATH=. QT_QPA_PLATFORM=offscreen \
#       python3 tests/test_integrity.py
# ============================================================================

import os
import re
import sqlite3
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_PASS = 0
_FAIL = 0


def ok(cond, label, extra=""):
    global _PASS, _FAIL
    if cond:
        _PASS += 1
    else:
        _FAIL += 1
        print(f"  ✘ {label} {extra}")


SINGLE_BRACE = re.compile(r"(?<!\{)\{([a-zA-Z0-9_]{2,})\}(?!\})")


def test_override_dialog_buttons():
    print("\n[1] CRITICAL OVERRIDE DIALOG — Yes/No offered")
    from PySide6.QtWidgets import QMessageBox
    import wizard_engine as we
    # patch warning to capture the buttons argument
    captured = {}
    orig = we.QMessageBox.warning

    def fake_warning(parent, title, text, *args, **kwargs):
        captured["buttons"] = args[0] if args else kwargs.get("buttons")
        captured["default"] = args[1] if len(args) > 1 else \
            kwargs.get("defaultButton")
        return QMessageBox.Yes
    we.QMessageBox.warning = staticmethod(fake_warning)
    try:
        # run the wizard generation with critical data (MW > FG)
        from PySide6.QtWidgets import QApplication
        app = QApplication.instance() or QApplication([])
        from wizard_engine import GeneratorWizard
        from tests.defaults import build_default_values
        from generation_pipeline import all_templates
        from PySide6.QtWidgets import (QLineEdit, QComboBox,
                                       QDoubleSpinBox, QSpinBox,
                                       QCheckBox, QTextEdit)
        wiz = GeneratorWizard()
        wiz.show()
        app.processEvents()
        for pid in range(1, 7):
            wiz.setCurrentId(pid)
            app.processEvents()
        p3 = wiz.page(3)
        vals = build_default_values(all_templates())
        for key, w in p3.widgets.items():
            if isinstance(w, QLineEdit):
                w.setText(str(vals.get(key, "Generic")))
            elif isinstance(w, QDoubleSpinBox):
                w.setValue(float(vals.get(key, 100)))
            elif isinstance(w, QSpinBox):
                w.setValue(int(vals.get(key, 100)))
            elif isinstance(w, QComboBox):
                w.setCurrentIndex(0)
            elif isinstance(w, QCheckBox):
                w.setChecked(True)
            elif isinstance(w, QTextEdit):
                w.setPlainText(str(vals.get(key, "Generic text.")))
        # force a CRITICAL: MW 17 > FG 16
        w = p3.widgets.get("mud_weight")
        if isinstance(w, QDoubleSpinBox):
            w.setValue(17)
        wf = p3.widgets.get("fracture_gradient")
        if isinstance(wf, QDoubleSpinBox):
            wf.setValue(16)
        tmp = tempfile.mkdtemp(prefix="drl_ovr_")
        p5 = wiz.page(5)
        p5.path.setText(tmp)
        p5.fname.setText("ovr.docx")
        p6 = wiz.page(6)
        p6.initializePage()
        p6._generate()
        app.processEvents()
    finally:
        we.QMessageBox.warning = orig
    ok(captured.get("buttons") == QMessageBox.Yes | QMessageBox.No,
       "warning called with Yes|No buttons",
       f"got {captured.get('buttons')}")
    # with Yes returned, generation proceeds (file exists) — proves the
    # override path is not dead code
    ok(os.path.exists(os.path.join(tmp, "ovr.docx")) or
       p6._last_path and os.path.exists(p6._last_path),
       "override with Yes proceeds to export")


def test_numeric_empty():
    print("\n[2] NUMERIC EMPTY ≠ 0")
    import wizard_engine as we
    spec = we.InputSpec("test_pressure", "Test Pressure", "number",
                        unit="psi", required=True)
    w = we._build_field(spec)
    # untouched -> empty string
    val = we._get_field_value(spec, w)
    ok(val == "", f"untouched numeric reads '' (got {val!r})")
    # after entering a value -> the value
    w.setValue(5000)
    ok(we._get_field_value(spec, w) == "5000.0",
       "entered numeric reads value")
    # explicit default 0 still reads as 0
    spec0 = we.InputSpec("water_depth", "Water Depth", "number",
                         unit="ft", default="0")
    w0 = we._build_field(spec0)
    ok(we._get_field_value(spec0, w0) == "0.0",
       "explicit default 0 preserved")


def test_unified_placeholders():
    print("\n[3] UNIFIED PLACEHOLDER SYNTAX {{x}} AND {x}")
    from wizard_engine import fill_template, scan_unresolved_placeholders
    from wizard_engine import TemplateDef
    td = TemplateDef(
        key="t", name="T", kind="Procedure", icon="", description="",
        inputs=[], markdown="RIH speed: {rih_speed} m/min, "
                            "max {{set_pressure}} psi.")
    out = fill_template(td, {"rih_speed": "8", "set_pressure": "3500"})
    ok("{rih_speed}" not in out and "{{set_pressure}}" not in out,
       "both syntaxes replaced", out)
    ok("RIH speed: 8 m/min, max 3500 psi." in out, "values in place", out)
    # unresolved audit
    ok(scan_unresolved_placeholders(out) == [], "no unresolved after fill")
    ok(scan_unresolved_placeholders("a {missing} b {{gone}}") ==
       ["gone", "missing"], "audit lists both syntaxes")


def test_placeholder_audit_templates():
    print("\n[4] FINAL AUDIT — no unresolved placeholders in 51 templates")
    from generation_pipeline import (all_templates, build_document_markdown,
                                     generate_document)
    from wizard_engine import scan_unresolved_placeholders
    from tests.defaults import build_default_values
    tpl = all_templates()
    vals = build_default_values(tpl)
    bad = []
    for td in tpl:
        md = build_document_markdown(td, vals)
        un = scan_unresolved_placeholders(md)
        if un:
            bad.append((td.key, un))
    ok(not bad, f"all templates clean ({len(tpl)})",
       f"bad: {bad[:3]}" if bad else "")
    # generate_document reports the audit too
    rep = generate_document(tpl[0], vals, {}, {}, "/tmp/_audit_t.docx")
    ok(rep["unresolved_placeholders"] == [], "report carries audit")


def test_validation_aliases():
    print("\n[5] VALIDATION / STANDARDS — UI aliases resolved")
    from validation_engine import validate_well_data
    from standards_engine import compliance_matrix
    # fracture_gradient (UI name) must trigger the MW>FG CRITICAL
    fs = validate_well_data({"mud_weight": "17", "fracture_gradient": "16",
                             "td_depth": "10000"})
    ok("ENG-MW-FG" in [f.code for f in fs], "fracture_gradient alias")
    # td_depth is feet — casing 12000 > TD 10000 must flag (no double
    # conversion of feet to meters)
    fs2 = validate_well_data({"td_depth": "10000", "casing_depth": "12000"})
    ok("LOGIC-CASING-TD" in [f.code for f in fs2],
       "casing>TD flagged with feet units")
    # standards sees the alias
    rows = compliance_matrix({"ecd": "17.2", "fracture_gradient": "16",
                              "td_depth": "10000"})
    md = [r for r in rows if r["rule_id"] == "STD-MD-002"][0]
    ok(md["status"] == "FAIL", "STD-MD-002 uses fracture_gradient alias")


def test_procedure_substitution():
    print("\n[6] PROCEDURE DB — value substitution + unresolved report")
    from procedures_db import (ProcedureDatabase, generate_procedures_docx,
                               _subst)
    db = ProcedureDatabase()
    tmp = tempfile.mkdtemp(prefix="drl_int_")
    # find a procedure whose steps contain {placeholders}
    hit = None
    for p in db.get_all_procedures(active_only=True):
        rec = db.get_procedure(p.id)
        keys = []
        for s in rec.steps:
            for m in SINGLE_BRACE.finditer(s.text or ""):
                keys.append(m.group(1))
        if keys:
            hit = (rec, sorted(set(keys)))
            break
    ok(hit is not None, "procedure with {placeholder} steps exists")
    if not hit:
        db.close()
        return
    rec, keys = hit
    # without values -> unresolved reported
    out1 = os.path.join(tmp, "a.docx")
    r1 = generate_procedures_docx(db, [rec.id], out1)
    ok(rec.name in (r1.get("unresolved") or {}),
       "unresolved reported without values",
       str(r1.get("unresolved")))
    # with all inputs filled -> substituted, zero unresolved
    ins = db.get_inputs(rec.id)
    vals = {rec.id: {i["input_key"]: (i.get("input_default") or "42")
                     for i in ins}}
    out2 = os.path.join(tmp, "b.docx")
    r2 = generate_procedures_docx(db, [rec.id], out2, input_values=vals)
    ok(not (r2.get("unresolved") or {}).get(rec.name),
       "zero unresolved when inputs filled",
       str(r2.get("unresolved")))
    from docx import Document
    d = Document(out2)
    text = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                text += "\n" + c.text
    ok(not SINGLE_BRACE.search(text), "no raw {key} in Word")
    ok("(default)" in text or "42" in text, "defaults marked / substituted")
    # _subst unit check
    out, un = _subst("Run at {rih_speed} m/min", {"rih_speed": "8"})
    ok(out == "Run at 8 m/min" and un == [], "subst basic")
    out2s, un2 = _subst("x {a} y {{b}} z", {"a": "1"})
    ok(out2s == "x 1 y {{b}} z" and un2 == ["b"], "subst partial + report")
    db.close()


def test_time_breakdown_guard():
    print("\n[7] TIME BREAKDOWN — cross-project guard")
    from generation_pipeline import (template_by_key, generate_document,
                                     all_templates, build_document_markdown)
    from tests.defaults import build_default_values
    vals = build_default_values(all_templates())
    td = template_by_key("drilling_program")
    # document for a different well than the TB project
    vals2 = dict(vals, well_name="Completely Different Well")
    md = build_document_markdown(td, vals2)
    ok("TIME BREAKDOWN SUMMARY" not in md,
       "TB section skipped for mismatched well")
    # same well (or no well name) -> included
    vals3 = dict(vals, well_name="")
    md3 = build_document_markdown(td, vals3)
    ok("TIME BREAKDOWN SUMMARY" in md3,
       "TB section included when no well named")


def test_template_audit():
    print("\n[8] TEMPLATE & PROCEDURE AUDIT — zero FAIL")
    import tests.template_audit as ta
    report = ta.run_audit()
    ok(report["stats"]["total"] >= 230, "audited >= 230 items",
       f"got {report['stats']['total']}")
    ok(report["stats"]["fail"] == 0, "zero FAIL items",
       f"{report['stats']['fail']} FAIL")
    ok(report["stats"]["pass"] >= 150, ">= 150 PASS",
       f"{report['stats']['pass']} PASS")


def test_combo_not_editable():
    print("\n[9] COMBO WITH OPTIONS IS NOT FREE-TEXT")
    import wizard_engine as we
    spec = we.InputSpec("kill_method", "Kill Method", "combo",
                        options=["Wait and Weight", "Driller's Method"])
    w = we._build_field(spec)
    ok(hasattr(w, "isEditable"), "is a combo")
    ok(not w.isEditable(), "combo with options locked")
    spec2 = we.InputSpec("free", "Free Text", "combo", options=[])
    w2 = we._build_field(spec2)
    ok(w2.isEditable(), "combo without options editable")


def test_standards_applicability_column():
    print("\n[10] STANDARDS — applicability shown in document")
    from standards_engine import compliance_markdown
    md = compliance_markdown({"bop_wp": "10000", "masp": "2000"})
    ok("Applicability" in md, "applicability column present")
    ok("API RP 53" in md, "standard shown")


def test_consistency_check():
    print("\n[11] CROSS-DOCUMENT CONSISTENCY")
    from engineering_consistency import (consistency_check,
                                         consistency_markdown)
    bad = """
| Total Depth | 10000 ft |
| Casing Depth | 12000 ft |
| Mud Weight | 17 ppg |
| Fracture Gradient | 16 ppg |
| ECD | 16.5 ppg |
| BOP Working Pressure | 5000 psi |
| MASP | 7000 psi |
"""
    f = consistency_check(bad)
    codes = {x["code"] for x in f}
    ok("CONS-CASING-TD" in codes, "casing > TD flagged")
    ok("CONS-MW-FG" in codes, "MW > FG flagged")
    ok("CONS-ECD-FG" in codes, "ECD > FG flagged")
    ok("CONS-BOP-MASP" in codes, "BOP < MASP flagged")
    clean = """
| Total Depth | 10000 ft |
| Casing Depth | 8000 ft |
| Mud Weight | 12 ppg |
| Fracture Gradient | 16 ppg |
| ECD | 12.5 ppg |
| BOP Working Pressure | 10000 psi |
| MASP | 3000 psi |
"""
    ok(consistency_check(clean) == [], "clean doc no findings")
    ok(consistency_markdown(f) != "", "section rendered for findings")
    ok(consistency_markdown([]) == "", "no section when clean")


def test_engineering_ranges():
    print("\n[12] ENGINEERING RANGE VALIDATION")
    from validation_engine import validate_well_data
    fs = validate_well_data({"mud_weight": "50000", "td_depth": "10000"})
    codes = [x.code for x in fs]
    ok("SCHEMA-RANGE-MUD_WEIGHT" in codes, "absurd MW flagged")
    ok("SCHEMA-RANGE-MUD_WEIGHT" not in [
        x.code for x in validate_well_data(
            {"mud_weight": "12", "td_depth": "10000"})],
       "normal MW not flagged")
    fs2 = validate_well_data({"bop_wp": "500", "td_depth": "10000"})
    ok("SCHEMA-RANGE-BOP_WP" in [x.code for x in fs2],
       "BOP below API range flagged")


def test_witsml_wizard_prefill():
    print("\n[14] WITSML IMPORT → WIZARD PREFILL")
    from PySide6.QtWidgets import QApplication
    app = QApplication.instance() or QApplication([])
    from wizard_engine import GeneratorWizard
    from PySide6.QtWidgets import QLineEdit, QTextEdit, QComboBox
    xml = """<?xml version="1.0"?><wells xmlns="http://www.witsml.org/schemas/1series">
<well><name>Imported Well</name><field>Imported Field</field>
<wellbore><trajectory><trajectoryStation><md uom='ft'>0</md>
<inclination uom='deg'>0</inclination><azimuth uom='deg'>0</azimuth>
</trajectoryStation></trajectory></wellbore></well></wells>"""
    from witsml_import import parse_witsml
    parsed = parse_witsml(xml)
    init = {"well_name": parsed["well_name"], "field": parsed["field"],
            "trajectory_table": parsed["trajectory_table"]}
    wiz = GeneratorWizard(initial_values=init)
    wiz.show()
    app.processEvents()
    wiz.setCurrentId(3)
    app.processEvents()
    p3 = wiz.page(3)
    w = p3.widgets.get("well_name")
    if isinstance(w, QLineEdit):
        ok(w.text() == "Imported Well", "well_name prefilled from WITSML",
           w.text())
    else:
        ok(True, "well_name not an input in this template (skipped)")
    w2 = p3.widgets.get("trajectory_table")
    if w2 is not None and hasattr(w2, "toPlainText"):
        ok("Imported Well" in wiz._initial_values, "wizard holds values")
    wiz.close()


def test_well_model_integration():
    print("\n[15] WELL MODEL — persist on generation, upsert by name")
    from generation_pipeline import template_by_key, generate_document
    from well_model import load_well_values, WellDatabase
    import tempfile
    td = template_by_key("drilling_program")
    tmp = tempfile.mkdtemp(prefix="drl_wm_")
    import uuid
    _wn = f"Well INT-{uuid.uuid4().hex[:6]}"
    vals = {"well_name": _wn, "field_name": "Field X",
            "operator": "the Operator", "mud_weight": "12.5",
            "td_depth": "10000", "hole_size": "12.25",
            "casing_size": "9.625", "casing_depth": "8000",
            "mud_type": "OBM"}
    out = os.path.join(tmp, "w.docx")
    ok(generate_document(td, vals, {}, {}, out)["ok"], "generated")
    v = load_well_values(well_name=_wn)
    ok(v.get("well_name") == _wn, "well loaded by name")
    ok(v.get("mud_weight") == "12.5", "mud weight round-trip")
    # second generation with different MW -> same well, latest value
    vals2 = dict(vals, mud_weight="13.5")
    generate_document(td, vals2, {}, {}, out)
    v2 = load_well_values(well_name=_wn)
    ok(v2.get("mud_weight") == "13.5", "latest revision loaded")
    db = WellDatabase()
    mine = [w for w in db.list_wells() if w.get("well_name") == _wn]
    ok(len(mine) == 1, "no duplicate wells (upsert by name)",
       f"{len(mine)} for {_wn}")
    db.close()


def test_section_presence_heading():
    print("\n[13] SECTION PRESENCE — heading-based, not phrase-based")
    from document_compliance import compliance_check
    # a stray mention of 'casing' must NOT satisfy the Casing section
    md = ("## 1. SCOPE\nThis program covers casing, mud and cement "
          "operations for the well.\n## VALIDATION\nok")
    rep = compliance_check("drilling_program", md)
    ok("casing" in rep["missing_sections"],
       "phrase mention does not count as section",
       str(rep["missing_sections"]))
    # a real heading satisfies it
    md2 = ("## CASING PROGRAM\n| Size | Depth |\n|---|---|\n| 9.625 | 8000 |\n"
           "## VALIDATION\nok")
    rep2 = compliance_check("drilling_program", md2)
    ok("casing" not in rep2["missing_sections"],
       "real heading counts", str(rep2["missing_sections"]))


if __name__ == "__main__":
    test_override_dialog_buttons()
    test_numeric_empty()
    test_unified_placeholders()
    test_placeholder_audit_templates()
    test_validation_aliases()
    test_procedure_substitution()
    test_time_breakdown_guard()
    test_template_audit()
    test_combo_not_editable()
    test_standards_applicability_column()
    test_consistency_check()
    test_engineering_ranges()
    test_section_presence_heading()
    test_witsml_wizard_prefill()
    test_well_model_integration()
    print("\n" + "=" * 60)
    print(f"RESULT: {_PASS} passed, {_FAIL} failed")
    sys.exit(1 if _FAIL else 0)
