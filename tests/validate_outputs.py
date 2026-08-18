# ============================================================================
# COMPREHENSIVE OUTPUT VALIDATION SUITE
# File: tests/validate_outputs.py
#
# Goal: for EVERY capability of the application, feed default data,
# generate the real output, and validate that the output is
#   - well-formed (valid file format, opens with the official readers)
#   - structurally correct (sections present, tables are real tables,
#     cover/TOC present, no markdown/HTML artifacts)
#   - content-correct (no unfilled placeholders, no leaked entities,
#     no garbage, non-trivial text)
#
# Validated capabilities:
#   1. All 51 wizard templates -> Word documents
#   2. Procedures database -> Word document (headless export)
#   3. Well Engineering Report -> Word document
#   4. CBS / AFE section -> markdown + Word
#   5. Drilling problems section -> markdown + Word
#   6. Risk decision matrix -> markdown + Word
#   7. Excel statistical report -> workbook
#   8. WITSML export -> XML (parsed)
#   9. ROPE checklists -> markdown
#  10. Knowledge enrichment (LLM-off fallback) -> markdown
#  11. Time breakdown summary section -> Word
#  12. REST API generate endpoint -> base64 docx
#
# Run:  LD_LIBRARY_PATH=/tmp/glstubs PYTHONPATH=. QT_QPA_PLATFORM=offscreen \
#       python3 tests/validate_outputs.py
# ============================================================================

import base64
import os
import re
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

_PASS = 0
_FAIL = 0
_FAILURES = []


def ok(cond, label, extra=""):
    global _PASS, _FAIL
    if cond:
        _PASS += 1
    else:
        _FAIL += 1
        _FAILURES.append(f"{label}: {extra}")
        print(f"  ✘ {label} {extra}")


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

ARTIFACTS = [
    (re.compile(r"\{\{"), "unfilled placeholder {{"),
    (re.compile(r"\}\}"), "unfilled placeholder }}"),
    (re.compile(r"\*\*"), "markdown bold **"),
    (re.compile(r"^#{1,6}\s"), "markdown heading #"),
    (re.compile(r"^```"), "code fence"),
    (re.compile(r"&amp;"), "html entity &amp;"),
    (re.compile(r"&lt;"), "html entity &lt;"),
    (re.compile(r"&gt;"), "html entity &gt;"),
    (re.compile(r"&copy;"), "html entity &copy;"),
    (re.compile(r"\[To Be Filled\]"), "unfilled [To Be Filled]"),
    (re.compile(r"^\[(?:[^\]\n]{1,40})\]\(http"), "markdown link"),
    (re.compile(r"^---+$"), "markdown hr ---"),
    (re.compile(r"^\|.*\|$"), "raw markdown table pipe"),
    (re.compile(r":\s*None\s*$"), "literal None value"),
    (re.compile(r"nan\b", re.IGNORECASE), "literal nan"),
    (re.compile(r"^=\s*$"), "stray equals"),
]


def docx_text(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def scan_artifacts(text: str) -> list:
    hits = []
    for pat, label in ARTIFACTS:
        m = pat.search(text)
        if m:
            hits.append(f"{label} ({m.group(0)[:30]!r})")
    return hits


def leak_scan(text: str) -> list:
    from wizard_engine import (OPERATOR_NAMES, SERVICE_NAMES, WELL_PATTERNS)
    pats = [p for p, _ in WELL_PATTERNS] + OPERATOR_NAMES + SERVICE_NAMES
    hits = []
    for p in pats:
        if re.search(p, str(text), re.IGNORECASE):
            hits.append(p)
    return hits


def gen_options():
    return {"font": "Calibri", "font_size": 11.0, "page": "A4",
            "orientation": "Portrait", "margin_left": 2.5,
            "margin_right": 2.0, "margin_top": 2.0, "margin_bottom": 2.0,
            "cover": True, "toc": True, "header_text": "",
            "footer_text": ""}


def gen_meta(tdef_name, operator="the Operator",
             contractor="the Service Company"):
    return {"title": tdef_name, "operator": operator,
            "contractor": contractor, "date": "17-August-2026",
            "revision": "01", "prepared_by": "Engineer",
            "reviewed_by": "Lead Engineer", "approved_by": "Drilling Manager",
            "document_number": ""}


# ---------------------------------------------------------------------------
# 1. all 51 wizard templates
# ---------------------------------------------------------------------------

def validate_templates(tmp):
    print("\n[1] ALL 51 TEMPLATES → WORD (format + content + structure)")
    from generation_pipeline import all_templates, generate_document
    from tests.defaults import build_default_values
    tpl = all_templates()
    if len(tpl) < 51:
        print(f"  ⚠ only {len(tpl)} templates found (master procedures DB "
              f"missing?) — run: python3 bootstrap.py")
    vals = build_default_values(tpl)
    for i, td in enumerate(tpl, 1):
        out = os.path.join(tmp, f"tpl_{td.key}.docx")
        rep = generate_document(td, vals, gen_meta(td.name), gen_options(),
                                out)
        label = f"[{i:02d}] {td.key}"
        if not rep["ok"] or not os.path.exists(out):
            ok(False, label, "generation failed")
            continue
        try:
            text = docx_text(out)
        except Exception as e:
            ok(False, label, f"docx unreadable: {e}")
            continue
        # format
        ok(len(text) > 1500, label, f"text too short ({len(text)})")
        hits = scan_artifacts(text)
        ok(not hits, label, f"artifacts: {hits[:3]}" if hits else "")
        # structure
        for sec in ("VALIDATION & COMPLIANCE", "PROGRAM READINESS SCORE",
                    "STANDARDS COMPLIANCE MATRIX",
                    "DOCUMENT COMPLIANCE REPORT",
                    "ENGINEERING CALCULATION REGISTER",
                    "DEEP ENGINEERING VERIFICATION"):
            ok(sec in text, label, f"missing section {sec}")
        from docx import Document
        d = Document(out)
        ok(len(d.tables) >= 1, label, "no tables in docx")
        ok("TABLE OF CONTENTS" in text, label, "no TOC")
        # content integrity
        leaks = leak_scan(text)
        ok(not leaks, label, f"LEAKS {leaks[:3]}" if leaks else "")
        ok(rep["register_rows"] > 0, label,
           f"register empty ({rep['register_rows']})")
        # no doubled paragraph separators / empty doc
        non_empty = [p for p in d.paragraphs if p.text.strip()]
        ok(len(non_empty) > 20, label,
           f"only {len(non_empty)} non-empty paragraphs")


# ---------------------------------------------------------------------------
# 2. procedures DB -> Word
# ---------------------------------------------------------------------------

def validate_procedures_word(tmp):
    print("\n[2] PROCEDURES DATABASE → WORD (headless export)")
    from procedures_db import ProcedureDatabase, generate_procedures_docx
    db = ProcedureDatabase()
    procs = db.get_all_procedures(active_only=True)
    ids = [p.id for p in procs[:5]]
    if len(procs) < 60:
        ok(False, "procedures present",
           f"got {len(procs)} — run: python3 bootstrap.py "
           f"(seeds 186+ procedures)")
        db.close()
        return
    ok(len(procs) >= 60, "procedures present", f"got {len(procs)}")
    if len(procs) < 180:
        print(f"  ⚠ only {len(procs)} procedures — run: python3 "
              f"bootstrap.py for the full library")
    out = os.path.join(tmp, "procedures.docx")
    res = generate_procedures_docx(db, ids, out)
    ok(res["ok"], "export ok", str(res))
    if res["ok"]:
        text = docx_text(out)
        ok("DRILLING OPERATIONS" in text, "cover text")
        ok("CHECKLIST" in text, "checklist section")
        hits = scan_artifacts(text)
        ok(not hits, "no artifacts", f"{hits[:3]}" if hits else "")
        leaks = leak_scan(text)
        ok(not leaks, "no leaks", f"{leaks[:3]}" if leaks else "")
        from docx import Document
        d = Document(out)
        ok(len(d.tables) >= 5, "tables present", f"{len(d.tables)}")
        ok(len([p for p in d.paragraphs if p.text.strip()]) > 40,
           "content paragraphs")
    db.close()


# ---------------------------------------------------------------------------
# 3. well report
# ---------------------------------------------------------------------------

def validate_well_report(tmp):
    print("\n[3] WELL ENGINEERING REPORT → WORD")
    from well_report import generate_well_report_docx, _demo_values
    vals = _demo_values()
    vals.update({"operator": "the Operator", "contractor":
                 "the Service Company"})
    out = os.path.join(tmp, "well_report.docx")
    ok(generate_well_report_docx(vals, out), "report generated")
    text = docx_text(out)
    for sec in ("WELL PROFILE", "ENGINEERING VALIDATION",
                "PROGRAM READINESS", "STANDARDS COMPLIANCE",
                "DOCUMENT COMPLIANCE"):
        ok(sec in text, f"section {sec}")
    hits = scan_artifacts(text)
    ok(not hits, "no artifacts", f"{hits[:3]}" if hits else "")
    leaks = leak_scan(text)
    ok(not leaks, "no leaks", f"{leaks[:3]}" if leaks else "")


# ---------------------------------------------------------------------------
# 4-6. CBS / problems / risk sections
# ---------------------------------------------------------------------------

def validate_sections(tmp):
    print("\n[4-6] CBS / PROBLEMS / RISK SECTIONS → markdown + Word")
    from cbs_db import CBSDatabase, build_cbs_markdown
    cdb = CBSDatabase()
    items = cdb.get_items()[:12]
    cbs_md = build_cbs_markdown(items, total_days=45, well_depth_m=3050,
                                well_name="Well A", operator="the Operator",
                                currency="USD")
    ok("COST BREAKDOWN" in cbs_md.upper() or "AFE" in cbs_md.upper(),
       "CBS heading")
    ok("|" in cbs_md, "CBS has tables")
    ok("the Operator" in cbs_md, "operator label")

    from drilling_problems_db import ProblemDatabase, build_problems_markdown
    pdb = ProblemDatabase()
    probs = pdb.all()[:5]
    prob_md = build_problems_markdown(probs, "the Operator")
    ok("PREVENTION" in prob_md.upper() or "RESPONSE" in prob_md.upper(),
       "problems heading")
    ok("- " in prob_md, "problems has bullet actions")

    from risk_decision import find_decisions, decision_markdown
    dcs = find_decisions("kick lost circulation stuck pipe H2S")
    ok(len(dcs) >= 3, f"risk decisions found ({len(dcs)})")
    d_md = decision_markdown(dcs)
    ok("|" in d_md, "risk matrix has tables")

    for name, md in (("cbs", cbs_md), ("problems", prob_md),
                     ("risk", d_md)):
        from wizard_engine import md_to_docx
        out = os.path.join(tmp, f"sec_{name}.docx")
        ok(md_to_docx(md, out, gen_meta(name), gen_options()), f"{name} docx")
        text = docx_text(out)
        hits = scan_artifacts(text)
        ok(not hits, f"{name} no artifacts", f"{hits[:3]}" if hits else "")
        leaks = leak_scan(text)
        ok(not leaks, f"{name} no leaks", f"{leaks[:3]}" if leaks else "")


# ---------------------------------------------------------------------------
# 7. excel report
# ---------------------------------------------------------------------------

def validate_excel(tmp):
    print("\n[7] STATISTICAL REPORT → EXCEL")
    import reporting
    out = os.path.join(tmp, "report.xlsx")
    n = reporting.export_report_excel(out)
    if n < 7:
        ok(False, f"sheets >= 7 (got {n}) — run: python3 bootstrap.py")
        return
    ok(n >= 7, f"sheets >= 7 (got {n})")
    from openpyxl import load_workbook
    wb = load_workbook(out)
    ok("Library" in wb.sheetnames, "Library sheet")
    ok(wb["Library"].max_row >= 700, "Library rows",
       f"{wb['Library'].max_row}")
    ok("Procedures-Category" in wb.sheetnames, "procedures sheet")
    ws = wb["Procedures-Category"]
    ok(ws.max_row >= 15, "procedure categories rows",
       f"{ws.max_row}")


# ---------------------------------------------------------------------------
# 8. WITSML
# ---------------------------------------------------------------------------

def validate_witsml(tmp):
    print("\n[8] WITSML EXPORT → XML")
    from witsml_export import build_witsml, build_json
    from xml.dom import minidom
    vals = {"well_name": "Well A", "field_name": "Field X",
            "operator": "the Operator", "mud_weight": "12",
            "depth": "10000",
            "trajectory_table": (
                "| MD (ft) | Inc (°) | Az (°) |\n|---|---|---|\n"
                "| 0 | 0 | 90 |\n| 5000 | 35 | 90 |\n"
                "| 10000 | 35 | 90 |")}
    xml = build_witsml(vals)
    dom = minidom.parseString(xml)   # raises if malformed
    ok(len(dom.getElementsByTagName("trajectoryStation")) == 3,
       "trajectory stations in XML")
    ok(dom.getElementsByTagName("name")[0].firstChild.data == "Well A",
       "well name")
    j = build_json(vals)
    ok(j["well"]["name"] == "Well A", "json handoff")
    ok(j["basis"]["mud_weight"] == "12", "json basis values")


# ---------------------------------------------------------------------------
# 9-10. ROPE + enrichment
# ---------------------------------------------------------------------------

def validate_rope_enrichment():
    print("\n[9-10] ROPE CHECKLISTS + KNOWLEDGE ENRICHMENT (LLM off)")
    from wizard_rope import get_rope_checklists
    from wizard_knowledge import get_chunks_for
    from wizard_llm import rewrite_chunks
    for tk in ("drilling_program", "cementing_procedure",
               "stuck_pipe_procedure"):
        rope = get_rope_checklists(tk, "moderate", "the Operator",
                                   "the Service Company")
        ok(len(rope) > 100, f"ROPE {tk} non-empty ({len(rope)})")
        if rope:
            leaks = leak_scan(rope)
            ok(not leaks, f"ROPE {tk} no leaks",
               f"{leaks[:3]}" if leaks else "")
    chunks = get_chunks_for("drilling_program", "moderate",
                            max_docs=2, max_chunks=8)
    enr = rewrite_chunks(chunks, "Drilling Program")
    ok(len(chunks) > 0, f"chunks retrieved ({len(chunks)})")
    ok(len(enr) > 200, f"enrichment text non-trivial ({len(enr)})")
    if enr:
        leaks = leak_scan(enr)
        ok(not leaks, "enrichment no leaks", f"{leaks[:3]}" if leaks else "")
        ok("......" not in enr, "no TOC dots")


# ---------------------------------------------------------------------------
# 11. time breakdown section
# ---------------------------------------------------------------------------

def validate_time_breakdown(tmp):
    print("\n[11] TIME BREAKDOWN SUMMARY SECTION IN DOCUMENT")
    from generation_pipeline import (template_by_key, generate_document)
    from tests.defaults import build_default_values
    from generation_pipeline import all_templates
    tb_db = Path.home() / ".drilling_program" / "time_breakdown.db"
    if not tb_db.exists():
        print("  ⚠ time_breakdown.db missing — section skipped. Run: "
              "python3 bootstrap.py (seeds 167-row project, 131.82 days)")
        return
    vals = build_default_values(all_templates())
    vals["well_name"] = "Example Well"  # matches the seeded TB project
    td = template_by_key("drilling_program")
    out = os.path.join(tmp, "tb.docx")
    rep = generate_document(td, vals, gen_meta(td.name), gen_options(), out)
    ok(rep["ok"], "doc generated")
    text = docx_text(out)
    ok("TIME BREAKDOWN SUMMARY" in text, "TB summary section present")
    ok("131.82" in text or "Total planned" in text, "TB total shown")


# ---------------------------------------------------------------------------
# 12. API generate
# ---------------------------------------------------------------------------

def validate_api(tmp):
    print("\n[12] REST API /generate → base64 docx")
    from fastapi.testclient import TestClient
    from api_server import create_app, get_api_key
    from tests.defaults import build_default_values
    from generation_pipeline import all_templates
    app = create_app(auth_enabled=True)
    client = TestClient(app)
    H = {"X-API-Key": get_api_key()}
    vals = build_default_values(all_templates())
    r = client.post("/api/generate",
                    json={"template_key": "drilling_program",
                          "values": vals},
                    headers=H)
    ok(r.status_code == 200, "generate 200", str(r.status_code))
    if r.status_code == 200:
        docx_bytes = base64.b64decode(r.json()["docx_base64"])
        ok(docx_bytes[:2] == b"PK", "docx zip header")
        p = os.path.join(tmp, "api.docx")
        with open(p, "wb") as f:
            f.write(docx_bytes)
        text = docx_text(p)
        ok("VALIDATION & COMPLIANCE" in text, "api doc sections")
        hits = scan_artifacts(text)
        ok(not hits, "api doc no artifacts", f"{hits[:3]}" if hits else "")


# ---------------------------------------------------------------------------
def main():
    tmp = tempfile.mkdtemp(prefix="drl_validate_")
    validate_templates(tmp)
    validate_procedures_word(tmp)
    validate_well_report(tmp)
    validate_sections(tmp)
    validate_excel(tmp)
    validate_witsml(tmp)
    validate_rope_enrichment()
    validate_time_breakdown(tmp)
    validate_api(tmp)
    print("\n" + "=" * 64)
    print(f"RESULT: {_PASS} passed, {_FAIL} failed")
    if _FAILURES:
        print("Failures (first 20):")
        for f in _FAILURES[:20]:
            print("  -", f[:160])
    print("=" * 64)
    sys.exit(1 if _FAIL else 0)


if __name__ == "__main__":
    main()
