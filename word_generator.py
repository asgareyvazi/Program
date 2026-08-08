# ============================================================================
# DRILLING PROGRAM & PROCEDURE GENERATOR - PROFESSIONAL EDITION
# Version 3.0
# File: word_generator.py
# Word Document Generator - All formats + Enhanced Tables + Helpers
# ============================================================================

import os
import math
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL)
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. "
          "Run: pip install python-docx")


# ============================================================================
# COLORS - Single source of truth
# ============================================================================

class DocColors:
    """All color definitions in one place"""

    # RGBColor objects
    DARK_NAVY  = RGBColor(0x0C, 0x2D, 0x48)
    NAVY       = RGBColor(0x0F, 0x34, 0x60)
    STEEL_BLUE = RGBColor(0x1B, 0x4F, 0x72)
    MED_BLUE   = RGBColor(0x21, 0x61, 0x8C)
    LIGHT_BLUE = RGBColor(0x2E, 0x86, 0xC1)
    RED        = RGBColor(0xE9, 0x45, 0x60)
    DARK_RED   = RGBColor(0xC0, 0x39, 0x2B)
    ORANGE     = RGBColor(0xE6, 0x7E, 0x22)
    GREEN      = RGBColor(0x27, 0xAE, 0x60)
    GOLD       = RGBColor(0xF3, 0x9C, 0x12)
    BLACK      = RGBColor(0x17, 0x20, 0x2A)
    DARK_GRAY  = RGBColor(0x2C, 0x3E, 0x50)
    GRAY       = RGBColor(0x85, 0x92, 0x9E)
    WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

    # Hex strings for XML shading
    HEADER_BG      = "0C2D48"
    SUBHEADER_BG   = "1B4F72"
    ROW_ALT_BG     = "EBF5FB"
    ROW_NORMAL_BG  = "FFFFFF"
    CAUTION_BG     = "FDEDEC"
    WARNING_BG     = "FEF9E7"
    SUCCESS_BG     = "EAFAF1"
    INFO_BG        = "EBF5FB"

    # Section backgrounds
    CONDUCTOR_BG   = "E8F8F5"
    SURFACE_BG     = "EBF5FB"
    INTER_BG       = "FEF9E7"
    PROD_BG        = "FDEDEC"
    LINER_BG       = "F5EEF8"

    @classmethod
    def section_bg(cls, section_name: str) -> str:
        """Get background color for casing section"""
        name = section_name.lower()
        if 'conductor' in name:
            return cls.CONDUCTOR_BG
        if 'surface' in name:
            return cls.SURFACE_BG
        if 'intermediate' in name:
            return cls.INTER_BG
        if 'production' in name:
            return cls.PROD_BG
        if 'liner' in name:
            return cls.LINER_BG
        return cls.ROW_ALT_BG

    @classmethod
    def risk_color_hex(cls, level: str) -> str:
        """Get hex color for risk level"""
        mapping = {
            'Low':      "27AE60",
            'Medium':   "F1C40F",
            'High':     "E67E22",
            'Critical': "C0392B",
        }
        return mapping.get(level, "1B4F72")


# ============================================================================
# DATA FORMATTER - Consistent formatting
# ============================================================================

class DataFormatter:
    """Consistent data formatting - eliminates scattered f-strings"""

    @staticmethod
    def depth(value: float, unit: str = "ft") -> str:
        if not value or value <= 0:
            return "-"
        return f"{value:,.0f} {unit}"

    @staticmethod
    def depth_range(start: float, end: float,
                    unit: str = "ft") -> str:
        if start <= 0 and end <= 0:
            return "-"
        return f"{start:,.0f}-{end:,.0f} {unit}"

    @staticmethod
    def pressure(value: float, unit: str = "psi") -> str:
        if not value or value <= 0:
            return "-"
        return f"{value:,.0f} {unit}"

    @staticmethod
    def mud_weight(value: float, unit: str = "ppg") -> str:
        if not value or value <= 0:
            return "-"
        return f"{value:.1f} {unit}"

    @staticmethod
    def mud_weight_range(start: float, end: float,
                         unit: str = "ppg") -> str:
        return f"{start:.1f}-{end:.1f} {unit}"

    @staticmethod
    def temperature(value: float, unit: str = "F") -> str:
        if not value:
            return "-"
        return f"{value:.0f} {unit}"

    @staticmethod
    def percent(value: float) -> str:
        if not value:
            return "-"
        return f"{value:.1f}%"

    @staticmethod
    def dimension(value: float, unit: str = '"') -> str:
        if not value or value <= 0:
            return "-"
        return f"{value}{unit}"

    @staticmethod
    def safe_float(text: str, default: float = 0.0) -> float:
        if not text or str(text).strip() in ('-', '', 'N/A'):
            return default
        try:
            return float(str(text).replace(',', '').strip())
        except (ValueError, TypeError):
            return default


# ============================================================================
# SECTION NUMBERING
# ============================================================================

class SectionNumbering:
    """Manages section numbering across document"""

    def __init__(self, start: int = 0):
        self._main = start
        self._sub = 0
        self._subsub = 0

    def next_main(self, title: str) -> str:
        self._main += 1
        self._sub = 0
        self._subsub = 0
        return f"{self._main}. {title}"

    def next_sub(self, title: str) -> str:
        self._sub += 1
        self._subsub = 0
        return f"{self._main}.{self._sub} {title}"

    def next_subsub(self, title: str) -> str:
        self._subsub += 1
        return f"{self._main}.{self._sub}.{self._subsub} {title}"

    @property
    def current_main(self) -> int:
        return self._main


# ============================================================================
# REPORT FORMAT DEFINITIONS
# ============================================================================

class ReportFormat:
    """Output format definitions"""
    STANDARD     = "standard"
    COMPACT      = "compact"
    EXECUTIVE    = "executive"
    LANDSCAPE    = "landscape"
    FIELD        = "field"
    PRESENTATION = "presentation"

    DESCRIPTIONS = {
        STANDARD:     "Full Drilling Program (150-200 pages)",
        COMPACT:      "Compact Version (80-100 pages)",
        EXECUTIVE:    "Executive Summary (15-20 pages)",
        LANDSCAPE:    "Data Tables Landscape Format",
        FIELD:        "Field Reference Card (simplified)",
        PRESENTATION: "Presentation Format (large fonts)",
    }

    PAGE_CONFIGS = {
        STANDARD: dict(
            height=29.7, width=21.0,
            top=2.0, bottom=2.0,
            left=2.5, right=2.0,
            landscape=False),
        COMPACT: dict(
            height=29.7, width=21.0,
            top=1.5, bottom=1.5,
            left=2.0, right=1.5,
            landscape=False),
        EXECUTIVE: dict(
            height=29.7, width=21.0,
            top=2.5, bottom=2.5,
            left=3.0, right=2.5,
            landscape=False),
        LANDSCAPE: dict(
            height=21.0, width=29.7,
            top=1.5, bottom=1.5,
            left=1.5, right=1.5,
            landscape=True),
        FIELD: dict(
            height=29.7, width=21.0,
            top=1.5, bottom=1.5,
            left=1.5, right=1.5,
            landscape=False),
        PRESENTATION: dict(
            height=21.0, width=29.7,
            top=2.0, bottom=2.0,
            left=2.5, right=2.5,
            landscape=True),
    }

    @classmethod
    def apply_page_setup(cls, doc, fmt: str = STANDARD):
        """Apply page settings to document"""
        cfg = cls.PAGE_CONFIGS.get(
            fmt, cls.PAGE_CONFIGS[cls.STANDARD])
        section = doc.sections[0]
        section.page_height = Cm(cfg["height"])
        section.page_width = Cm(cfg["width"])
        section.top_margin = Cm(cfg["top"])
        section.bottom_margin = Cm(cfg["bottom"])
        section.left_margin = Cm(cfg["left"])
        section.right_margin = Cm(cfg["right"])


# ============================================================================
# DOCUMENT STYLE MANAGER
# ============================================================================

class DocumentStyleManager:
    """Manages document styles"""

    def __init__(self, doc):
        self.doc = doc
        self.setup_styles()

    def setup_styles(self):
        styles = self.doc.styles

        self._modify_style(
            'Heading 1', font_size=18,
            color=DocColors.WHITE, bold=True,
            space_before=Pt(18), space_after=Pt(8))

        self._modify_style(
            'Heading 2', font_size=14,
            color=DocColors.DARK_NAVY, bold=True,
            space_before=Pt(14), space_after=Pt(6))

        self._modify_style(
            'Heading 3', font_size=12,
            color=DocColors.STEEL_BLUE, bold=True,
            space_before=Pt(10), space_after=Pt(4))

        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.font.color.rgb = DocColors.BLACK
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.15

        for style_name in ('ProcedureStep', 'TableCaption',
                           'NoteStyle'):
            try:
                s = styles.add_style(
                    style_name, WD_STYLE_TYPE.PARAGRAPH)
                s.font.name = 'Calibri'
                s.font.size = Pt(9)
            except ValueError:
                pass

    def _modify_style(self, name: str, font_size: int,
                      color: RGBColor, bold: bool = False,
                      space_before: Any = None,
                      space_after: Any = None):
        try:
            style = self.doc.styles[name]
            style.font.name = 'Calibri'
            style.font.size = Pt(font_size)
            style.font.color.rgb = color
            style.font.bold = bold
            if space_before:
                style.paragraph_format.space_before = space_before
            if space_after:
                style.paragraph_format.space_after = space_after
        except Exception:
            pass


# ============================================================================
# TABLE HELPER
# ============================================================================

class TableHelper:
    """Professional table builder"""

    @staticmethod
    def set_cell_shading(cell, color_hex: str):
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} '
            f'w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def format_header_cell(cell, text: str,
                           font_size: int = 9):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.WHITE
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        TableHelper.set_cell_shading(
            cell, DocColors.HEADER_BG)

    @staticmethod
    def format_subheader_cell(cell, text: str,
                              font_size: int = 9):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.WHITE
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        TableHelper.set_cell_shading(
            cell, DocColors.SUBHEADER_BG)

    @staticmethod
    def format_data_cell(cell, text: str,
                         font_size: int = 9,
                         bold: bool = False,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color: RGBColor = None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = alignment
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'
        run.font.color.rgb = (
            color if color else DocColors.BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @staticmethod
    def format_label_cell(cell, text: str,
                          font_size: int = 9):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.DARK_NAVY
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        TableHelper.set_cell_shading(
            cell, DocColors.ROW_ALT_BG)

    @staticmethod
    def format_value_cell(cell, text: str,
                          font_size: int = 9):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.font.size = Pt(font_size)
        run.font.color.rgb = DocColors.BLACK
        run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @staticmethod
    def set_row_height(row, height_cm: float = 0.7):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(
            f'<w:trHeight {nsdecls("w")} '
            f'w:val="{int(height_cm * 567)}" '
            f'w:hRule="atLeast"/>')
        trPr.append(trHeight)

    @staticmethod
    def create_professional_table(
            doc, headers: List[str],
            data: List[List[str]],
            col_widths: List[float] = None,
            caption: str = None,
            alt_row_shading: bool = True):
        """Create styled professional table"""
        num_cols = len(headers)
        num_rows = len(data) + 1

        table = doc.add_table(
            rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        if col_widths:
            for i, width in enumerate(col_widths):
                if i < num_cols:
                    for row in table.rows:
                        row.cells[i].width = Inches(width)

        header_row = table.rows[0]
        TableHelper.set_row_height(header_row, 0.8)
        for i, h in enumerate(headers):
            TableHelper.format_header_cell(
                header_row.cells[i], h)

        for row_idx, row_data in enumerate(data):
            row = table.rows[row_idx + 1]
            TableHelper.set_row_height(row, 0.6)
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    TableHelper.format_data_cell(
                        cell, str(cell_text))
                    if (alt_row_shading
                            and row_idx % 2 == 0):
                        TableHelper.set_cell_shading(
                            cell, DocColors.ROW_ALT_BG)

        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = DocColors.GRAY

        return table

    @staticmethod
    def create_key_value_table(
            doc, data: List[Tuple[str, str]],
            title: str = None,
            num_cols: int = 2):
        """Key-value pair table"""
        if not data:
            return None

        if num_cols == 4:
            mid = math.ceil(len(data) / 2)
            table = doc.add_table(rows=mid, cols=4)
        else:
            table = doc.add_table(
                rows=len(data), cols=2)

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        if num_cols == 4:
            for i in range(mid):
                row = table.rows[i]
                TableHelper.set_row_height(row, 0.6)
                left_i = i
                if left_i < len(data):
                    k, v = data[left_i]
                    TableHelper.format_label_cell(
                        row.cells[0], str(k))
                    TableHelper.format_value_cell(
                        row.cells[1], str(v))
                right_i = i + mid
                if right_i < len(data):
                    k, v = data[right_i]
                    TableHelper.format_label_cell(
                        row.cells[2], str(k))
                    TableHelper.format_value_cell(
                        row.cells[3], str(v))
                else:
                    TableHelper.format_label_cell(
                        row.cells[2], "")
                    TableHelper.format_value_cell(
                        row.cells[3], "")
        else:
            for i, (k, v) in enumerate(data):
                row = table.rows[i]
                TableHelper.set_row_height(row, 0.6)
                TableHelper.format_label_cell(
                    row.cells[0], str(k))
                TableHelper.format_value_cell(
                    row.cells[1], str(v))

        return table


# ============================================================================
# ENHANCED TABLE BUILDER
# ============================================================================

class EnhancedTableBuilder:
    """Advanced tables with color-coded data"""

    @staticmethod
    def _shd(color_hex: str):
        return parse_xml(
            f'<w:shd {nsdecls("w")} '
            f'w:fill="{color_hex}" w:val="clear"/>')

    @classmethod
    def _cell(cls, cell, text: str,
              font_size: int = 8,
              bold: bool = False,
              color: RGBColor = None,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              bg: str = None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = color
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg:
            cell._tc.get_or_add_tcPr().append(
                cls._shd(bg))

    @classmethod
    def formation_table(cls, doc, formations: list):
        """Formation tops with PP-based color coding"""
        headers = [
            "Formation", "Lithology",
            "MD Top\n(ft)", "MD Bot\n(ft)",
            "PP\n(ppg EMW)", "FG\n(ppg EMW)",
            "Temp\n(F)", "Drillability",
            "Remarks"
        ]
        rows = len(formations) + 1
        tbl = doc.add_table(rows=rows, cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cls._cell(hrow.cells[i], h,
                      font_size=8, bold=True,
                      color=DocColors.WHITE,
                      bg=DocColors.HEADER_BG)

        for idx, ft in enumerate(formations):
            row = tbl.rows[idx + 1]
            pp = ft.pore_pressure_bottom
            if pp > 15.0:
                bg = "FDEDEC"
            elif pp > 11.0:
                bg = "FEF9E7"
            else:
                bg = "EBF5FB" if idx % 2 == 0 else "FFFFFF"

            vals = [
                (ft.name, WD_ALIGN_PARAGRAPH.LEFT),
                (ft.formation_type, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{ft.md_top:,.0f}", WD_ALIGN_PARAGRAPH.CENTER),
                (f"{ft.md_bottom:,.0f}", WD_ALIGN_PARAGRAPH.CENTER),
                (f"{ft.pore_pressure_top:.1f}-{ft.pore_pressure_bottom:.1f}",
                 WD_ALIGN_PARAGRAPH.CENTER),
                (f"{ft.fracture_gradient_top:.1f}-{ft.fracture_gradient_bottom:.1f}",
                 WD_ALIGN_PARAGRAPH.CENTER),
                (f"{ft.temperature_bottom:.0f}", WD_ALIGN_PARAGRAPH.CENTER),
                (ft.drillability, WD_ALIGN_PARAGRAPH.CENTER),
                ((ft.remarks or "")[:60], WD_ALIGN_PARAGRAPH.LEFT),
            ]
            for i, (val, align) in enumerate(vals):
                cls._cell(row.cells[i], val,
                          font_size=8, align=align, bg=bg)
        return tbl

    @classmethod
    def casing_summary_table(cls, doc, casings: list):
        """Casing design with DF color coding"""
        headers = [
            "Section", "Hole\n(in)", "Csg OD\n(in)",
            "ID\n(in)", "Wt\n(ppf)", "Grade",
            "Connection", "MD\n(ft)", "TVD\n(ft)",
            "TOC\n(ft)", "Burst\n(psi)", "Collapse\n(psi)",
            "Tension\n(klbs)", "DF\nBurst",
            "DF\nCollapse", "DF\nTension"
        ]
        rows = len(casings) + 1
        tbl = doc.add_table(rows=rows, cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cls._cell(hrow.cells[i], h,
                      font_size=7, bold=True,
                      color=DocColors.WHITE,
                      bg=DocColors.HEADER_BG)

        def df_color(val, min_df):
            if val <= 0:
                return None
            if val < min_df:
                return DocColors.DARK_RED
            if val < min_df * 1.1:
                return DocColors.ORANGE
            return DocColors.GREEN

        for idx, cd in enumerate(casings):
            row = tbl.rows[idx + 1]
            bg = DocColors.section_bg(cd.section_name)

            vals = [
                cd.section_name,
                f"{cd.hole_size}\"",
                f"{cd.casing_od}\"",
                f"{cd.casing_id}\"",
                f"{cd.casing_weight}",
                cd.casing_grade,
                cd.casing_connection,
                f"{cd.setting_depth_md:,.0f}",
                f"{cd.setting_depth_tvd:,.0f}",
                f"{cd.top_of_cement_md:,.0f}",
                f"{cd.burst_rating:,.0f}",
                f"{cd.collapse_rating:,.0f}",
                f"{cd.tensile_rating/1000:.0f}",
                f"{cd.min_design_factor_burst:.2f}",
                f"{cd.min_design_factor_collapse:.2f}",
                f"{cd.min_design_factor_tension:.2f}",
            ]
            colors = [None] * 13 + [
                df_color(cd.min_design_factor_burst, 1.10),
                df_color(cd.min_design_factor_collapse, 1.10),
                df_color(cd.min_design_factor_tension, 1.60),
            ]
            for i, (val, col) in enumerate(zip(vals, colors)):
                align = (WD_ALIGN_PARAGRAPH.LEFT
                         if i in (0, 6)
                         else WD_ALIGN_PARAGRAPH.CENTER)
                cls._cell(row.cells[i], val,
                          font_size=7,
                          bold=(col is not None),
                          color=col, align=align, bg=bg)
        return tbl

    @classmethod
    def cement_summary_table(cls, doc,
                              cements: list):
        """Cement job summary table"""
        headers = [
            "Section", "Csg OD\n(in)",
            "Hole\n(in)", "Shoe\n(ft)",
            "TOC\n(ft)",
            "Lead\nDens\n(ppg)",
            "Lead\nVol\n(bbl)",
            "Tail\nDens\n(ppg)",
            "Tail\nVol\n(bbl)",
            "Disp\n(bbl)", "WOC\n(hrs)",
            "Bump\n(psi)", "Excess\n(%)", "CBL?"
        ]
        rows = len(cements) + 1
        tbl = doc.add_table(rows=rows, cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cls._cell(hrow.cells[i], h,
                      font_size=7, bold=True,
                      color=DocColors.WHITE,
                      bg=DocColors.HEADER_BG)

        bgs = ["E8F8F5", "EBF5FB", "FEF9E7", "FDEDEC"]
        for idx, cd in enumerate(cements):
            row = tbl.rows[idx + 1]
            bg = bgs[idx % len(bgs)]
            vals = [
                cd.section_name,
                f"{cd.casing_od}\"",
                f"{cd.hole_size}\"",
                f"{cd.shoe_depth_md:,.0f}",
                f"{cd.toc_md:,.0f}",
                f"{cd.lead_slurry_density:.1f}",
                f"{cd.lead_slurry_volume:.0f}",
                f"{cd.tail_slurry_density:.1f}",
                f"{cd.tail_slurry_volume:.0f}",
                f"{cd.displacement_volume:.0f}",
                f"{cd.woc_time:.0f}",
                f"{cd.plug_bump_pressure:,.0f}",
                f"{cd.excess_percentage:.0f}",
                "Yes" if cd.cbl_cbil_required else "No",
            ]
            for i, val in enumerate(vals):
                align = (WD_ALIGN_PARAGRAPH.LEFT
                         if i == 0
                         else WD_ALIGN_PARAGRAPH.CENTER)
                color = None
                if i == 13 and val == "Yes":
                    color = DocColors.GREEN
                cls._cell(row.cells[i], val,
                          font_size=7, color=color,
                          align=align, bg=bg)
        return tbl

    @classmethod
    def mud_comparison_table(cls, doc, muds: list):
        """Mud properties comparison table"""
        properties = [
            ("Mud Type",       "mud_type",          "str"),
            ("MW In (ppg)",    "mud_weight_in",      "float"),
            ("MW Out (ppg)",   "mud_weight_out",     "float"),
            ("FV (s/qt)",      "funnel_viscosity",   "float"),
            ("PV (cP)",        "plastic_viscosity",  "float"),
            ("YP (lb/100ft)",  "yield_point",        "float"),
            ("Gel 10s",        "gel_strength_10s",   "float"),
            ("API FL (ml)",    "fluid_loss",         "float"),
            ("HTHP FL (ml)",   "hthp_fluid_loss",    "float"),
            ("pH",             "ph",                 "float"),
            ("Cl (ppm)",       "chlorides",          "float"),
            ("OWR",            "oil_water_ratio",    "str"),
            ("ES (V)",         "electrical_stability","float"),
            ("ECD Shoe (ppg)", "ecd_at_shoe",        "float"),
        ]
        cols = len(muds) + 1
        rows = len(properties) + 1
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hrow = tbl.rows[0]
        cls._cell(hrow.cells[0], "Property",
                  font_size=8, bold=True,
                  color=DocColors.WHITE,
                  bg=DocColors.HEADER_BG)
        bgs = ["E8F8F5", "EBF5FB", "FEF9E7", "FDEDEC"]
        for i, mp in enumerate(muds):
            cls._cell(hrow.cells[i + 1],
                      f"{mp.section_name}\n({mp.hole_size}\")",
                      font_size=8, bold=True,
                      color=DocColors.WHITE,
                      bg=DocColors.SUBHEADER_BG)

        for row_i, (label, attr, typ) in enumerate(properties):
            row = tbl.rows[row_i + 1]
            cls._cell(row.cells[0], label,
                      font_size=8, bold=True,
                      align=WD_ALIGN_PARAGRAPH.LEFT,
                      bg="EBF5FB")
            for col_i, mp in enumerate(muds):
                bg = bgs[col_i % len(bgs)]
                val = getattr(mp, attr, "")
                if typ == "float":
                    if isinstance(val, (int, float)) and val > 0:
                        display = f"{val:.1f}"
                    else:
                        display = "-"
                else:
                    display = str(val)[:25] if val else "-"
                cls._cell(row.cells[col_i + 1], display,
                          font_size=8, bg=bg)
        return tbl

    @classmethod
    def time_gantt_table(cls, doc,
                          time_estimates: list):
        """Gantt-style time estimate table"""
        max_days = max(
            (te.cumulative_days
             for te in time_estimates), default=1)
        num_bars = 10

        headers = (["Section", "Operation",
                    "Days", "Cum."] +
                   [f"W{i+1}" for i in range(num_bars)])
        rows = len(time_estimates) + 1
        tbl = doc.add_table(
            rows=rows, cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cls._cell(hrow.cells[i], h,
                      font_size=7, bold=True,
                      color=DocColors.WHITE,
                      bg=DocColors.HEADER_BG)

        bgs = ["E8F8F5", "EBF5FB", "FEF9E7",
               "FDEDEC", "F5EEF8", "E8F4FD"]

        for idx, te in enumerate(time_estimates):
            row = tbl.rows[idx + 1]
            bg = bgs[idx % len(bgs)]
            cls._cell(row.cells[0], te.section_name,
                      font_size=7,
                      align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg)
            cls._cell(row.cells[1],
                      te.operation[:30],
                      font_size=7,
                      align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg)
            cls._cell(row.cells[2],
                      f"{te.total_section_days:.1f}",
                      font_size=7, bg=bg)
            cls._cell(row.cells[3],
                      f"{te.cumulative_days:.1f}",
                      font_size=7, bold=True, bg=bg)

            start_pct = ((te.cumulative_days
                          - te.total_section_days)
                         / max_days)
            end_pct = te.cumulative_days / max_days

            for col_i in range(num_bars):
                col_start = col_i / num_bars
                col_end = (col_i + 1) / num_bars
                cell = row.cells[col_i + 4]
                if (start_pct <= col_start
                        and col_end <= end_pct):
                    cls._cell(cell, "",
                               font_size=7, bg="1B4F72")
                elif (start_pct < col_end
                        and col_start < end_pct):
                    cls._cell(cell, "",
                               font_size=7, bg="D4E6F1")
                else:
                    cls._cell(cell, "", font_size=7, bg=bg)
        return tbl


# ============================================================================
# MAIN WORD DOCUMENT GENERATOR
# ============================================================================

class DrillingProgramWordGenerator:
    """Main Word document generator"""

    def __init__(self, project):
        self.project = project
        self.doc = Document()
        self.style_mgr = DocumentStyleManager(self.doc)
        self.table_helper = TableHelper()
        self.section_counter = 0
        self.table_counter = 0
        self.figure_counter = 0
        self._numbering = SectionNumbering()

        from engineering_calculations import (
            ProcedureGenerator, CalculationEngine)
        self.proc_gen = ProcedureGenerator(project)
        self.calc_engine = CalculationEngine(project)

    def generate(self, file_path: str,
                 progress=None):
        """Generate complete document"""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed")

        self._setup_document()
        self._update_progress(progress, 5)

        self._create_cover_page()
        self._update_progress(progress, 10)

        self._create_revision_history()
        self._update_progress(progress, 12)

        self._create_table_of_contents()
        self._update_progress(progress, 14)

        self._create_abbreviations()
        self._update_progress(progress, 16)

        self._create_executive_summary()
        self._update_progress(progress, 18)

        self._create_well_information_section()
        self._update_progress(progress, 25)

        self._create_rig_specification_section()
        self._update_progress(progress, 28)

        self._create_formation_prognosis_section()
        self._update_progress(progress, 33)

        self._create_hazard_analysis_section()
        self._update_progress(progress, 37)

        self._create_casing_design_section()
        self._update_progress(progress, 43)

        self._create_mud_program_section()
        self._update_progress(progress, 49)

        self._create_bha_section()
        self._update_progress(progress, 53)

        self._create_hydraulics_section()
        self._update_progress(progress, 56)

        self._create_cement_section()
        self._update_progress(progress, 61)

        self._create_directional_section()
        self._update_progress(progress, 65)

        self._create_bop_well_control_section()
        self._update_progress(progress, 69)

        self._create_evaluation_section()
        self._update_progress(progress, 72)

        self._create_time_estimate_section()
        self._update_progress(progress, 76)

        self._create_all_procedures_section()
        self._update_progress(progress, 90)

        self._create_emergency_section()
        self._update_progress(progress, 94)

        self._create_appendices()
        self._update_progress(progress, 97)

        self._add_headers_footers()

        self.doc.save(file_path)
        self._update_progress(progress, 100)

    def _update_progress(self, progress, value: int):
        if progress:
            try:
                progress.setValue(value)
                from PySide6.QtWidgets import QApplication
                QApplication.processEvents()
            except Exception:
                pass

    def _setup_document(self):
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    def _add_section_heading(self, text: str,
                              level: int = 1):
        if level == 1:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:shd {nsdecls("w")} '
                f'w:fill="{DocColors.HEADER_BG}" '
                f'w:val="clear"/>'))
            run = p.add_run(f"  {text}")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = DocColors.WHITE
            run.font.name = 'Calibri'

        elif level == 2:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" '
                f'w:space="1" w:color="1B4F72"/>'
                f'</w:pBdr>')
            pPr.append(pBdr)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = DocColors.STEEL_BLUE
            run.font.name = 'Calibri'

        elif level == 3:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DocColors.MED_BLUE
            run.font.name = 'Calibri'

    def _add_colored_line(self, color: RGBColor,
                           width: int = 2):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        color_hex = (f"{color[0]:02x}"
                     f"{color[1]:02x}"
                     f"{color[2]:02x}")
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" '
            f'w:sz="{width * 8}" '
            f'w:space="1" w:color="{color_hex}"/>'
            f'</w:pBdr>')
        pPr.append(pBdr)

    def _add_note(self, text: str,
                   note_type: str = "NOTE"):
        configs = {
            "WARNING":   ("FDEDEC", DocColors.RED, "WARNING: "),
            "CAUTION":   ("FEF9E7", DocColors.ORANGE, "CAUTION: "),
            "IMPORTANT": ("FDEDEC", DocColors.DARK_RED, "IMPORTANT: "),
            "NOTE":      ("EBF5FB", DocColors.STEEL_BLUE, "NOTE: "),
        }
        bg, color, prefix = configs.get(
            note_type, configs["NOTE"])

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} '
            f'w:fill="{bg}" w:val="clear"/>'))

        run_label = p.add_run(prefix)
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = color
        run_label.font.name = 'Calibri'

        run_text = p.add_run(text)
        run_text.font.size = Pt(9)
        run_text.font.name = 'Calibri'

    def _next_table(self) -> int:
        self.table_counter += 1
        return self.table_counter

    def _next_figure(self) -> int:
        self.figure_counter += 1
        return self.figure_counter

    def _add_procedure_steps(self, steps: List[str]):
        for step in steps:
            if not step.strip():
                self.doc.add_paragraph("")
                continue

            stripped = step.lstrip()
            indent_level = len(step) - len(stripped)

            if stripped.upper() == stripped and len(stripped) > 5:
                self._add_section_heading(stripped, level=3)
                continue

            p = self.doc.add_paragraph()
            indent = 0.3 + indent_level * 0.15
            p.paragraph_format.left_indent = Cm(indent)
            p.paragraph_format.space_after = Pt(2)

            parts = stripped.split(' ', 1)
            if (len(parts) == 2
                    and '.' in parts[0]
                    and any(c.isdigit()
                            for c in parts[0])):
                run_num = p.add_run(parts[0] + " ")
                run_num.bold = True
                run_num.font.size = Pt(10)
                run_num.font.color.rgb = DocColors.STEEL_BLUE
                run_num.font.name = 'Calibri'
                run_text = p.add_run(parts[1])
                run_text.font.size = Pt(10)
                run_text.font.name = 'Calibri'
            else:
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

    # ----------------------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------------------

    def _create_cover_page(self):
        ci = self.project.company_info
        wi = self.project.well_info

        self._add_colored_line(DocColors.DARK_NAVY, 3)
        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        cls_para = self.doc.add_paragraph()
        cls_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cls_para.add_run(
            f"  {ci.classification.upper()}  ")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = DocColors.RED

        self.doc.add_paragraph("")

        if ci.operator_name:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(ci.operator_name.upper())
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = DocColors.DARK_NAVY
            run.font.name = 'Calibri'

        self.doc.add_paragraph("")

        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("DRILLING PROGRAM")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = DocColors.NAVY
        run.font.name = 'Calibri'

        sub_para = self.doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run("& OPERATIONAL PROCEDURES")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = DocColors.STEEL_BLUE
        run.font.name = 'Calibri'

        self.doc.add_paragraph("")
        self._add_colored_line(DocColors.RED, 2)
        self.doc.add_paragraph("")

        cover_data = []
        if ci.well_name:
            cover_data.append(("Well Name", ci.well_name))
        if ci.field_name:
            cover_data.append(("Field", ci.field_name))
        if ci.country:
            cover_data.append(("Country / Region",
                               f"{ci.country} / {ci.region}"))
        if ci.rig_name:
            cover_data.append(("Rig",
                               f"{ci.rig_name} ({ci.rig_type})"))
        if wi.well_type:
            cover_data.append(("Well Type", wi.well_type))
        if wi.total_depth_md > 0:
            cover_data.append(("Planned TD (MD/TVD)",
                               f"{wi.total_depth_md:,.0f} ft / "
                               f"{wi.total_depth_tvd:,.0f} ft"))
        if wi.target_formation:
            cover_data.append(("Target Formation",
                               wi.target_formation))

        if cover_data:
            table = self.doc.add_table(
                rows=len(cover_data), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            for i, (key, val) in enumerate(cover_data):
                row = table.rows[i]
                TableHelper.format_label_cell(
                    row.cells[0], key, font_size=11)
                cell_val = row.cells[1]
                cell_val.text = ""
                p = cell_val.paragraphs[0]
                r = p.add_run(str(val))
                r.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = DocColors.DARK_NAVY
                r.font.name = 'Calibri'

        self.doc.add_paragraph("")
        self.doc.add_paragraph("")

        ctrl_data = [
            ("Document Number",
             ci.document_number or "DRL-PRG-001"),
            ("Revision", ci.revision or "0"),
            ("Date", datetime.now().strftime("%d-%B-%Y")),
            ("Prepared By", ci.prepared_by),
            ("Reviewed By", ci.reviewed_by),
            ("Approved By", ci.approved_by),
        ]

        ctrl_table = self.doc.add_table(
            rows=len(ctrl_data), cols=2)
        ctrl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ctrl_table.style = 'Table Grid'
        for i, (key, val) in enumerate(ctrl_data):
            row = ctrl_table.rows[i]
            TableHelper.format_label_cell(
                row.cells[0], key, font_size=10)
            TableHelper.format_value_cell(
                row.cells[1], val or "", font_size=10)

        self.doc.add_paragraph("")
        if ci.contractor_name:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                f"Drilling Contractor: {ci.contractor_name}")
            run.font.size = Pt(10)
            run.font.color.rgb = DocColors.GRAY

        self._add_colored_line(DocColors.DARK_NAVY, 3)

        disc = self.doc.add_paragraph()
        disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disc.add_run(
            f"This document is {ci.classification}. "
            f"Unauthorized distribution is prohibited.")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = DocColors.RED

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # REVISION HISTORY
    # ----------------------------------------------------------------

    def _create_revision_history(self):
        self._add_section_heading("REVISION HISTORY", 1)
        headers = ["Rev.", "Date", "Description",
                   "Prepared By", "Reviewed By", "Approved By"]
        data = [[
            self.project.company_info.revision or "0",
            datetime.now().strftime("%d-%b-%Y"),
            "Initial Issue",
            self.project.company_info.prepared_by or "",
            self.project.company_info.reviewed_by or "",
            self.project.company_info.approved_by or "",
        ]]
        TableHelper.create_professional_table(
            self.doc, headers, data,
            col_widths=[0.5, 1.0, 2.5, 1.2, 1.2, 1.2])
        self.doc.add_paragraph("")

        sig_table = self.doc.add_table(rows=4, cols=4)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.style = 'Table Grid'
        for i, h in enumerate(
                ["Role", "Name", "Signature", "Date"]):
            TableHelper.format_header_cell(
                sig_table.rows[0].cells[i], h)
        roles = [
            ("Prepared By", self.project.company_info.prepared_by),
            ("Reviewed By", self.project.company_info.reviewed_by),
            ("Approved By", self.project.company_info.approved_by),
        ]
        for i, (role, name) in enumerate(roles):
            row = sig_table.rows[i + 1]
            TableHelper.set_row_height(row, 1.5)
            TableHelper.format_data_cell(
                row.cells[0], role,
                alignment=WD_ALIGN_PARAGRAPH.LEFT)
            TableHelper.format_data_cell(
                row.cells[1], name or "")
            TableHelper.format_data_cell(row.cells[2], "")
            TableHelper.format_data_cell(row.cells[3], "")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # TABLE OF CONTENTS
    # ----------------------------------------------------------------

    def _create_table_of_contents(self):
        self._add_section_heading("TABLE OF CONTENTS", 1)

        toc_items = [
            ("1.", "EXECUTIVE SUMMARY"),
            ("2.", "WELL INFORMATION"),
            ("3.", "RIG SPECIFICATIONS"),
            ("4.", "FORMATION PROGNOSIS"),
            ("5.", "HAZARD ANALYSIS"),
            ("6.", "CASING DESIGN"),
            ("7.", "DRILLING FLUID PROGRAM"),
            ("8.", "BHA & DRILL STRING DESIGN"),
            ("9.", "HYDRAULICS ANALYSIS"),
            ("10.", "CEMENTING PROGRAM"),
            ("11.", "DIRECTIONAL DRILLING PLAN"),
            ("12.", "BOP & WELL CONTROL"),
            ("13.", "EVALUATION PROGRAM"),
            ("14.", "TIME vs DEPTH ESTIMATE"),
            ("15.", "DETAILED OPERATING PROCEDURES"),
            ("16.", "EMERGENCY PROCEDURES"),
            ("17.", "APPENDICES"),
        ]

        table = self.doc.add_table(
            rows=len(toc_items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (num, title) in enumerate(toc_items):
            row = table.rows[i]
            cell = row.cells[0]
            cell.text = ""
            p = cell.paragraphs[0]
            is_main = not num.startswith("   ")
            indent = 0 if is_main else 0.5
            p.paragraph_format.left_indent = Cm(indent)

            run_num = p.add_run(num.strip() + "  ")
            run_num.bold = is_main
            run_num.font.size = Pt(10 if is_main else 9)
            run_num.font.color.rgb = DocColors.DARK_NAVY
            run_num.font.name = 'Calibri'

            run_title = p.add_run(title)
            run_title.bold = is_main
            run_title.font.size = Pt(10 if is_main else 9)
            run_title.font.name = 'Calibri'

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # ABBREVIATIONS
    # ----------------------------------------------------------------

    def _create_abbreviations(self):
        self._add_section_heading("ABBREVIATIONS & DEFINITIONS", 1)

        abbreviations = [
            ("AFE", "Authorization for Expenditure"),
            ("API", "American Petroleum Institute"),
            ("BHA", "Bottom Hole Assembly"),
            ("BHP", "Bottom Hole Pressure"),
            ("BHT", "Bottom Hole Temperature"),
            ("BOP", "Blowout Preventer"),
            ("CBL", "Cement Bond Log"),
            ("DC",  "Drill Collar"),
            ("DLS", "Dog Leg Severity"),
            ("DP",  "Drill Pipe"),
            ("DST", "Drill Stem Test"),
            ("ECD", "Equivalent Circulating Density"),
            ("EMW", "Equivalent Mud Weight"),
            ("FIT", "Formation Integrity Test"),
            ("FV",  "Funnel Viscosity"),
            ("GPM", "Gallons Per Minute"),
            ("HSE", "Health, Safety & Environment"),
            ("HWDP","Heavy Weight Drill Pipe"),
            ("IADC","International Association of Drilling Contractors"),
            ("ICP", "Initial Circulating Pressure"),
            ("KOP", "Kick-Off Point"),
            ("LCM", "Lost Circulation Material"),
            ("LOT", "Leak-Off Test"),
            ("LWD", "Logging While Drilling"),
            ("MAASP","Maximum Allowable Annular Surface Pressure"),
            ("MD",  "Measured Depth"),
            ("MW",  "Mud Weight"),
            ("MWD", "Measurement While Drilling"),
            ("NACE","National Association of Corrosion Engineers"),
            ("NPT", "Non-Productive Time"),
            ("OBM", "Oil-Based Mud"),
            ("PDC", "Polycrystalline Diamond Compact"),
            ("POOH","Pull Out Of Hole"),
            ("PV",  "Plastic Viscosity"),
            ("PWD", "Pressure While Drilling"),
            ("RIH", "Run In Hole"),
            ("ROP", "Rate of Penetration"),
            ("RPM", "Revolutions Per Minute"),
            ("RSS", "Rotary Steerable System"),
            ("SBM", "Synthetic-Based Mud"),
            ("SICP","Shut-In Casing Pressure"),
            ("SIDPP","Shut-In Drill Pipe Pressure"),
            ("SPM", "Strokes Per Minute"),
            ("SPP", "Stand Pipe Pressure"),
            ("TD",  "Total Depth"),
            ("TFA", "Total Flow Area"),
            ("TOC", "Top of Cement"),
            ("TVD", "True Vertical Depth"),
            ("VDL", "Variable Density Log"),
            ("WBM", "Water-Based Mud"),
            ("WOB", "Weight on Bit"),
            ("WOC", "Wait on Cement"),
            ("YP",  "Yield Point"),
        ]

        mid = len(abbreviations) // 2 + len(abbreviations) % 2
        left = abbreviations[:mid]
        right = abbreviations[mid:]

        table = self.doc.add_table(rows=mid + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for i, h in enumerate(
                ["Abbrev.", "Definition",
                 "Abbrev.", "Definition"]):
            TableHelper.format_header_cell(
                table.rows[0].cells[i], h, font_size=8)

        for i in range(mid):
            row = table.rows[i + 1]
            TableHelper.set_row_height(row, 0.5)
            if i < len(left):
                TableHelper.format_data_cell(
                    row.cells[0], left[i][0],
                    font_size=8, bold=True)
                TableHelper.format_data_cell(
                    row.cells[1], left[i][1],
                    font_size=8,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
            if i < len(right):
                TableHelper.format_data_cell(
                    row.cells[2], right[i][0],
                    font_size=8, bold=True)
                TableHelper.format_data_cell(
                    row.cells[3], right[i][1],
                    font_size=8,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
            if i % 2 == 0:
                for c in range(4):
                    TableHelper.set_cell_shading(
                        row.cells[c], DocColors.ROW_ALT_BG)

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # ----------------------------------------------------------------

    def _create_executive_summary(self):
        ci = self.project.company_info
        wi = self.project.well_info
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. EXECUTIVE SUMMARY", 1)

        summary = (
            f"This Drilling Program has been prepared for the "
            f"drilling of well {ci.well_name or '[Well Name]'} "
            f"located in {ci.field_name or '[Field]'} field, "
            f"{ci.country or '[Country]'}. "
            f"The well is a {wi.well_type} {wi.well_profile} "
            f"well targeting the "
            f"{wi.target_formation or '[Formation]'} "
            f"formation at a planned total depth of "
            f"{wi.total_depth_md:,.0f} ft MD "
            f"({wi.total_depth_tvd:,.0f} ft TVD).")

        p = self.doc.add_paragraph()
        run = p.add_run(summary)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

        self._add_section_heading("Key Well Parameters", 2)

        key_params = [
            ("Operator", ci.operator_name or ""),
            ("Well Name", ci.well_name or ""),
            ("Field", ci.field_name or ""),
            ("Well Type", wi.well_type),
            ("Profile", wi.well_profile),
            ("TD (MD)", DataFormatter.depth(wi.total_depth_md)),
            ("TD (TVD)", DataFormatter.depth(wi.total_depth_tvd)),
            ("Target", wi.target_formation or ""),
            ("Rig", f"{ci.rig_name} ({ci.rig_type})"),
            ("KB Elevation",
             f"{wi.kb_elevation:,.1f} ft"),
        ]

        if wi.water_depth > 0:
            key_params.append(
                ("Water Depth",
                 DataFormatter.depth(wi.water_depth)))
        if wi.expected_h2s_concentration > 0:
            key_params.append(
                ("H2S Expected",
                 DataFormatter.percent(
                     wi.expected_h2s_concentration)))
        if wi.nace_required:
            key_params.append(
                ("NACE Required", "Yes"))

        TableHelper.create_key_value_table(
            self.doc, key_params, num_cols=4)

        # H2S warning
        if wi.expected_h2s_concentration > 0:
            self._add_note(
                f"SOUR SERVICE WELL: H2S expected "
                f"{wi.expected_h2s_concentration:.1f}%. "
                f"All materials shall comply with "
                f"NACE MR-0175 / ISO 15156.",
                "WARNING")

        if self.project.casing_design:
            self._add_section_heading(
                "Casing Program Overview", 2)
            headers = [
                "Section", "Hole Size",
                "Casing OD", "Grade",
                "Setting Depth MD", "TOC"
            ]
            data = [
                [cd.section_name,
                 f'{cd.hole_size}"',
                 f'{cd.casing_od}"',
                 cd.casing_grade,
                 DataFormatter.depth(cd.setting_depth_md),
                 DataFormatter.depth(cd.top_of_cement_md)]
                for cd in self.project.casing_design
            ]
            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=(f"Table {self._next_table()}: "
                         f"Casing Program Summary"))

        if self.project.time_estimates:
            total_days = max(
                te.cumulative_days
                for te in self.project.time_estimates)
            self._add_section_heading(
                "Estimated Duration", 2)
            p = self.doc.add_paragraph()
            run = p.add_run(
                f"Estimated Total Duration: "
                f"{total_days:.1f} Days")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = DocColors.DARK_NAVY

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # WELL INFORMATION
    # ----------------------------------------------------------------

    def _create_well_information_section(self):
        self.section_counter += 1
        ci = self.project.company_info
        wi = self.project.well_info

        self._add_section_heading(
            f"{self.section_counter}. WELL INFORMATION", 1)

        self._add_section_heading(
            f"{self.section_counter}.1 General Well Data", 2)

        general_data = [
            ("Operator", ci.operator_name),
            ("Drilling Contractor", ci.contractor_name),
            ("Well Name", ci.well_name),
            ("Well Number", ci.well_number),
            ("Field Name", ci.field_name),
            ("Block / License", ci.block_license),
            ("Country", ci.country),
            ("Region / Area", ci.region),
            ("Rig Name", ci.rig_name),
            ("Rig Type", ci.rig_type),
            ("Well Type", wi.well_type),
            ("Well Profile", wi.well_profile),
            ("Planned Spud Date", ci.spud_date),
            ("TD (MD)", DataFormatter.depth(wi.total_depth_md)),
            ("TD (TVD)", DataFormatter.depth(wi.total_depth_tvd)),
            ("KB Elevation", f"{wi.kb_elevation:,.1f} ft"),
            ("Ground Elevation",
             f"{wi.ground_elevation:,.1f} ft"),
        ]
        if wi.water_depth > 0:
            general_data.append(
                ("Water Depth",
                 DataFormatter.depth(wi.water_depth)))

        TableHelper.create_key_value_table(
            self.doc, general_data, num_cols=4)
        self.doc.add_paragraph("")

        self._add_section_heading(
            f"{self.section_counter}.2 Location & Coordinates", 2)

        location_data = [
            ("Surface Latitude", wi.surface_latitude),
            ("Surface Longitude", wi.surface_longitude),
            ("Target Latitude", wi.target_latitude),
            ("Target Longitude", wi.target_longitude),
            ("Coordinate System", wi.coordinate_system),
            ("Magnetic Declination",
             f"{wi.magnetic_declination:+.2f}"),
            ("Grid Convergence",
             f"{wi.grid_convergence:+.2f}"),
        ]
        TableHelper.create_key_value_table(
            self.doc, location_data, num_cols=4)
        self.doc.add_paragraph("")

        self._add_section_heading(
            f"{self.section_counter}.3 "
            f"Target / Reservoir Information", 2)

        reservoir_data = [
            ("Target Formation", wi.target_formation),
            ("Target Zone", wi.target_zone),
            ("Reservoir Pressure",
             DataFormatter.pressure(
                 wi.expected_reservoir_pressure)),
            ("Reservoir Temperature",
             DataFormatter.temperature(
                 wi.expected_reservoir_temperature)),
            ("H2S Expected",
             DataFormatter.percent(
                 wi.expected_h2s_concentration)),
            ("CO2 Expected",
             DataFormatter.percent(
                 wi.expected_co2_concentration)),
            ("NACE Required",
             "Yes" if wi.nace_required else "No"),
            ("Wellhead Type", wi.wellhead_type),
            ("Christmas Tree", wi.xmas_tree_type),
        ]
        TableHelper.create_key_value_table(
            self.doc, reservoir_data, num_cols=4)

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # RIG SPECIFICATIONS
    # ----------------------------------------------------------------

    def _create_rig_specification_section(self):
        self.section_counter += 1
        rs = self.project.rig_spec

        self._add_section_heading(
            f"{self.section_counter}. RIG SPECIFICATIONS", 1)

        rig_data = [
            ("Rig Name", rs.rig_name),
            ("Rig Type", rs.rig_type),
            ("Rig Contractor", rs.rig_contractor),
            ("Max Hook Load",
             DataFormatter.force(rs.max_hook_load)),
            ("Drawworks Power",
             DataFormatter.pressure(
                 rs.drawworks_power).replace("psi", "HP")),
            ("Top Drive",
             f"{'Yes' if rs.top_drive else 'No'} - "
             f"{rs.top_drive_model}"),
            ("Top Drive Torque",
             f"{rs.top_drive_torque:,.0f} ft-lbs"),
            ("Max RPM", f"{rs.max_rotary_speed:,.0f}"),
            ("Derrick Height",
             DataFormatter.depth(rs.derrick_height)),
            ("Rotary Table",
             f'{rs.rotary_table_size}" opening'),
        ]
        TableHelper.create_key_value_table(
            self.doc, rig_data, num_cols=4)

        self._add_section_heading("Mud Pumps", 2)
        pump_headers = [
            "Parameter", "Pump #1", "Pump #2", "Pump #3"]
        pump_data = [
            ["Type", rs.mud_pump_1_type,
             rs.mud_pump_2_type, rs.mud_pump_3_type],
            ["HP", f"{rs.mud_pump_1_hp:,.0f}",
             f"{rs.mud_pump_2_hp:,.0f}",
             f"{rs.mud_pump_3_hp:,.0f}"],
            ["Liner",
             f'{rs.mud_pump_1_liner}"',
             f'{rs.mud_pump_2_liner}"', ""],
            ["Max Pressure",
             DataFormatter.pressure(
                 rs.mud_pump_1_max_pressure),
             DataFormatter.pressure(
                 rs.mud_pump_2_max_pressure), ""],
            ["Max Flow",
             f"{rs.mud_pump_1_max_flow:,.0f} GPM",
             f"{rs.mud_pump_2_max_flow:,.0f} GPM", ""],
        ]
        TableHelper.create_professional_table(
            self.doc, pump_headers, pump_data,
            caption=(f"Table {self._next_table()}: "
                     f"Mud Pump Specifications"))

        self._add_section_heading("Utilities", 2)
        pit_data = [
            ("Total Pit Volume",
             f"{rs.pit_volume_total:,.0f} bbl"),
            ("Active Pit Volume",
             f"{rs.pit_volume_active:,.0f} bbl"),
            ("Shale Shakers",
             f"{rs.shale_shaker_count} units"),
            ("Degasser", rs.degasser_type),
            ("Centrifuge", rs.centrifuge),
            ("Generators", rs.generators),
            ("Total Power", f"{rs.total_power:,.0f} kW"),
            ("Crane", f"{rs.crane_capacity:,.0f} tons"),
            ("Accommodation",
             f"{rs.accommodation} persons"),
        ]
        TableHelper.create_key_value_table(
            self.doc, pit_data, num_cols=4)

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # FORMATION PROGNOSIS
    # ----------------------------------------------------------------

    def _create_formation_prognosis_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. FORMATION PROGNOSIS", 1)

        if not self.project.formation_tops:
            self.doc.add_paragraph(
                "No formation data entered.")
            self.doc.add_page_break()
            return

        self._add_section_heading("Formation Tops", 2)

        # Use enhanced table
        EnhancedTableBuilder.formation_table(
            self.doc, self.project.formation_tops)

        p_cap = self.doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(
            f"Table {self._next_table()}: "
            f"Formation Tops Prognosis - "
            f"Red=Overpressure, Yellow=Elevated PP")
        r_cap.font.size = Pt(8)
        r_cap.font.italic = True
        r_cap.font.color.rgb = DocColors.GRAY

        self.doc.add_paragraph("")

        self._add_section_heading(
            "Formation Descriptions", 2)
        for ft in self.project.formation_tops:
            self._add_section_heading(
                f"{ft.name} ({ft.formation_type})", 3)

            desc = (
                f"Expected from "
                f"{ft.md_top:,.0f} ft to "
                f"{ft.md_bottom:,.0f} ft MD. "
                f"Pore pressure: "
                f"{ft.pore_pressure_top:.1f}-"
                f"{ft.pore_pressure_bottom:.1f} ppg EMW. "
                f"Fracture gradient: "
                f"{ft.fracture_gradient_top:.1f}-"
                f"{ft.fracture_gradient_bottom:.1f} ppg EMW. "
                f"Temperature: "
                f"{ft.temperature_top:.0f}-"
                f"{ft.temperature_bottom:.0f} F. "
                f"Drillability: {ft.drillability}.")

            p = self.doc.add_paragraph()
            run = p.add_run(desc)
            run.font.size = Pt(10)

            if ft.remarks:
                self._add_note(ft.remarks, "INFO")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # HAZARD ANALYSIS
    # ----------------------------------------------------------------

    def _create_hazard_analysis_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"HAZARD ANALYSIS & RISK ASSESSMENT", 1)

        if self.project.hazards:
            headers = [
                "Hazard", "Depth\n(ft MD)",
                "Severity", "Probability",
                "Description", "Mitigation", "Contingency"
            ]
            data = [
                [h.hazard_type,
                 f"{h.md_top:,.0f}-{h.md_bottom:,.0f}",
                 h.severity, h.probability,
                 h.description, h.mitigation,
                 h.contingency]
                for h in self.project.hazards
            ]
            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=(f"Table {self._next_table()}: "
                         f"Hazard Analysis Summary"))
        else:
            self._add_note(
                "No hazard data entered. "
                "Refer to geological prognosis.", "INFO")

        self._add_section_heading(
            "Risk Assessment Matrix", 2)
        self._create_risk_matrix()

        self.doc.add_page_break()

    def _create_risk_matrix(self):
        table = self.doc.add_table(rows=6, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        headers = ["Severity/Prob",
                   "Negligible", "Minor",
                   "Moderate", "Major", "Catastrophic"]
        for i, h in enumerate(headers):
            TableHelper.format_header_cell(
                table.rows[0].cells[i], h, 8)

        probs = ["Almost Certain", "Likely",
                 "Possible", "Unlikely", "Rare"]
        risk_grid = [
            ['M', 'H', 'H', 'E', 'E'],
            ['M', 'M', 'H', 'H', 'E'],
            ['L', 'M', 'M', 'H', 'E'],
            ['L', 'L', 'M', 'M', 'H'],
            ['L', 'L', 'L', 'M', 'M'],
        ]
        risk_colors = {
            'L': "27AE60", 'M': "F1C40F",
            'H': "E67E22", 'E': "E74C3C"
        }
        risk_labels = {
            'L': 'LOW', 'M': 'MED',
            'H': 'HIGH', 'E': 'EXTREME'
        }

        for row_idx, prob in enumerate(probs):
            row = table.rows[row_idx + 1]
            TableHelper.format_subheader_cell(
                row.cells[0], prob, 8)
            for col_idx in range(5):
                cell = row.cells[col_idx + 1]
                risk = risk_grid[row_idx][col_idx]
                color = risk_colors[risk]
                TableHelper.set_cell_shading(cell, color)
                text_color = (
                    DocColors.WHITE
                    if risk in ('H', 'E')
                    else DocColors.BLACK)
                TableHelper.format_data_cell(
                    cell, risk_labels[risk],
                    font_size=8, bold=True,
                    color=text_color)

    # ----------------------------------------------------------------
    # CASING DESIGN
    # ----------------------------------------------------------------

    def _create_casing_design_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. CASING DESIGN", 1)

        if not self.project.casing_design:
            self.doc.add_paragraph(
                "No casing design data entered.")
            self.doc.add_page_break()
            return

        self._add_section_heading(
            "Casing Program Summary", 2)

        EnhancedTableBuilder.casing_summary_table(
            self.doc, self.project.casing_design)

        p_cap = self.doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(
            f"Table {self._next_table()}: "
            f"Casing Design - DF color: "
            f"Green=OK, Orange=Marginal, Red=Fail")
        r_cap.font.size = Pt(8)
        r_cap.font.italic = True
        r_cap.font.color.rgb = DocColors.GRAY

        self.doc.add_paragraph("")

        self._add_section_heading(
            "Casing String Details", 2)

        for cd in self.project.casing_design:
            self._add_section_heading(
                f"{cd.section_name}: "
                f"{cd.casing_od}\" {cd.casing_grade} "
                f"@ {cd.setting_depth_md:,.0f} ft MD", 3)

            props = [
                ("Casing OD", f'{cd.casing_od}"'),
                ("Casing ID", f'{cd.casing_id}"'),
                ("Weight", f"{cd.casing_weight} ppf"),
                ("Grade", cd.casing_grade),
                ("Connection", cd.casing_connection),
                ("Setting Depth MD",
                 DataFormatter.depth(cd.setting_depth_md)),
                ("Setting Depth TVD",
                 DataFormatter.depth(cd.setting_depth_tvd)),
                ("TOC MD",
                 DataFormatter.depth(cd.top_of_cement_md)),
                ("Drift ID", f'{cd.drift_id}"'),
                ("Burst Rating",
                 DataFormatter.pressure(cd.burst_rating)),
                ("Collapse Rating",
                 DataFormatter.pressure(cd.collapse_rating)),
                ("Tensile Rating",
                 f"{cd.tensile_rating:,.0f} lbs"),
                ("Float Shoe", cd.float_shoe_type),
                ("Float Collar Depth",
                 DataFormatter.depth(cd.float_collar_depth)),
                ("Centralizer Type", cd.centralizer_type),
                ("Centralizer Spacing",
                 f"{cd.centralizer_spacing:.0f} ft"),
            ]
            TableHelper.create_key_value_table(
                self.doc, props, num_cols=4)

            if cd.remarks:
                self._add_note(cd.remarks, "INFO")
            self.doc.add_paragraph("")

        if self.project.well_info.nace_required:
            self._add_note(
                "SOUR SERVICE: All casing and connections "
                "shall comply with NACE MR-0175 / ISO 15156.",
                "WARNING")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # MUD PROGRAM
    # ----------------------------------------------------------------

    def _create_mud_program_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"DRILLING FLUID PROGRAM", 1)

        if not self.project.mud_programs:
            self.doc.add_paragraph(
                "No mud program data entered.")
            self.doc.add_page_break()
            return

        self._add_section_heading(
            "Mud Properties Comparison", 2)

        EnhancedTableBuilder.mud_comparison_table(
            self.doc, self.project.mud_programs)

        p_cap = self.doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(
            f"Table {self._next_table()}: "
            f"Drilling Fluid Properties by Section")
        r_cap.font.size = Pt(8)
        r_cap.font.italic = True
        r_cap.font.color.rgb = DocColors.GRAY

        self.doc.add_paragraph("")

        self._add_section_heading(
            "Mud Systems Detail", 2)

        for mp in self.project.mud_programs:
            self._add_section_heading(
                f"{mp.section_name} - {mp.mud_type} "
                f"({DataFormatter.mud_weight_range(mp.mud_weight_in, mp.mud_weight_out)})",
                3)

            props = [
                ("Section", mp.section_name),
                ("Hole Size", f'{mp.hole_size}"'),
                ("Depth Range",
                 DataFormatter.depth_range(
                     mp.depth_from, mp.depth_to)),
                ("Mud Type", mp.mud_type),
                ("MW In", DataFormatter.mud_weight(
                    mp.mud_weight_in)),
                ("MW Out", DataFormatter.mud_weight(
                    mp.mud_weight_out)),
                ("FV", f"{mp.funnel_viscosity:.0f} s/qt"),
                ("PV", f"{mp.plastic_viscosity:.0f} cP"),
                ("YP", f"{mp.yield_point:.0f} lb/100ft"),
                ("Gel 10s/10m/30m",
                 f"{mp.gel_strength_10s:.0f} / "
                 f"{mp.gel_strength_10m:.0f} / "
                 f"{mp.gel_strength_30m:.0f}"),
                ("API FL", f"{mp.fluid_loss:.1f} ml"),
                ("HTHP FL",
                 f"{mp.hthp_fluid_loss:.1f} ml"),
                ("Total Volume",
                 f"{mp.total_volume_required:,.0f} bbl"),
                ("ECD Shoe",
                 DataFormatter.mud_weight(mp.ecd_at_shoe)),
                ("ECD TD",
                 DataFormatter.mud_weight(mp.ecd_at_td)),
            ]

            if mp.oil_water_ratio:
                props.append(("OWR", mp.oil_water_ratio))
            if mp.electrical_stability > 0:
                props.append(("ES",
                              f"{mp.electrical_stability:.0f} V"))
            if mp.key_additives:
                props.append(("Key Additives",
                              mp.key_additives))

            TableHelper.create_key_value_table(
                self.doc, props, num_cols=4)

            if mp.remarks:
                self._add_note(mp.remarks, "INFO")
            self.doc.add_paragraph("")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # BHA
    # ----------------------------------------------------------------

    def _create_bha_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"BHA & DRILL STRING DESIGN", 1)

        if not self.project.bha_designs:
            self.doc.add_paragraph(
                "No BHA design data entered.")
            self.doc.add_page_break()
            return

        self._add_section_heading("BHA Summary", 2)

        for bha in self.project.bha_designs:
            self._add_section_heading(
                f"BHA #{bha.bha_number} - "
                f"{bha.section_name} Section "
                f"({bha.hole_size}\" Hole)", 3)

            bha_info = [
                ("BHA Type", bha.bha_type),
                ("Bit Type", bha.bit_type),
                ("Bit Size", f'{bha.bit_size}"'),
                ("Bit Manufacturer", bha.bit_manufacturer),
                ("Bit Model", bha.bit_model),
                ("Nozzles (TFA)", bha.bit_nozzles),
            ]
            if bha.motor_type:
                bha_info += [
                    ("Motor Type", bha.motor_type),
                    ("Motor OD", f'{bha.motor_od}"'),
                    ("Motor Bend", f"{bha.motor_bend}"),
                ]
            if bha.rss_type:
                bha_info.append(("RSS", bha.rss_type))
            if bha.mwd_type:
                bha_info.append(("MWD", bha.mwd_type))
            if bha.lwd_sensors:
                bha_info.append(("LWD", bha.lwd_sensors))

            bha_info += [
                ("WOB", bha.recommended_wob),
                ("RPM", bha.recommended_rpm),
                ("Flow Rate", bha.recommended_flow_rate),
                ("Max Torque", bha.recommended_torque),
            ]
            if bha.remarks:
                bha_info.append(("Remarks", bha.remarks))

            TableHelper.create_key_value_table(
                self.doc, bha_info, num_cols=4)
            self.doc.add_paragraph("")

        if self.project.drilling_parameters:
            self._add_section_heading(
                "Drilling Parameters", 2)
            headers = [
                "Section", "Hole\n(in)",
                "WOB\n(klbs)", "RPM",
                "Flow\n(GPM)", "Max Torque\n(ft-lbs)",
                "ROP Avg\n(ft/hr)", "Max ECD\n(ppg)"
            ]
            data = [
                [dp.section_name,
                 f"{dp.hole_size}",
                 f"{dp.wob_min}-{dp.wob_max}",
                 f"{dp.rpm_min:.0f}-{dp.rpm_max:.0f}",
                 f"{dp.flow_rate_min:.0f}-"
                 f"{dp.flow_rate_max:.0f}",
                 f"{dp.torque_max:,.0f}",
                 f"{dp.rop_average:.0f}",
                 f"{dp.max_ecd:.1f}"
                 if dp.max_ecd > 0 else "Monitor"]
                for dp in self.project.drilling_parameters
            ]
            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=(f"Table {self._next_table()}: "
                         f"Drilling Parameters Summary"))

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # HYDRAULICS
    # ----------------------------------------------------------------

    def _create_hydraulics_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"HYDRAULICS ANALYSIS", 1)

        p = self.doc.add_paragraph()
        run = p.add_run(
            "Hydraulics calculations per API RP 13D. "
            "ECD must not exceed fracture gradient. "
            "Minimum annular velocity: 120 ft/min "
            "vertical, 150-180 ft/min deviated.")
        run.font.size = Pt(10)

        if self.project.mud_programs:
            headers = [
                "Section", "Hole\n(in)",
                "MW\n(ppg)", "ECD Shoe\n(ppg)",
                "ECD TD\n(ppg)"
            ]
            data = [
                [mp.section_name,
                 f"{mp.hole_size}",
                 DataFormatter.mud_weight(mp.mud_weight_out),
                 DataFormatter.mud_weight(mp.ecd_at_shoe),
                 DataFormatter.mud_weight(mp.ecd_at_td)]
                for mp in self.project.mud_programs
            ]
            TableHelper.create_professional_table(
                self.doc, headers, data,
                caption=(f"Table {self._next_table()}: "
                         f"Hydraulics Summary"))

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # CEMENT
    # ----------------------------------------------------------------

    def _create_cement_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"CEMENTING PROGRAM", 1)

        if not self.project.cement_design:
            self.doc.add_paragraph(
                "No cement design data entered.")
            self.doc.add_page_break()
            return

        self._add_section_heading(
            "Cement Program Summary", 2)

        EnhancedTableBuilder.cement_summary_table(
            self.doc, self.project.cement_design)

        p_cap = self.doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(
            f"Table {self._next_table()}: "
            f"Cement Program Summary All Sections")
        r_cap.font.size = Pt(8)
        r_cap.font.italic = True
        r_cap.font.color.rgb = DocColors.GRAY

        self.doc.add_paragraph("")

        for cd in self.project.cement_design:
            self._add_section_heading(
                f"{cd.section_name} Cement Job - "
                f"{cd.casing_od}\" Casing", 3)

            params = [
                ("Lead Type", cd.lead_slurry_type),
                ("Lead Density",
                 DataFormatter.mud_weight(
                     cd.lead_slurry_density)),
                ("Lead Volume",
                 f"{cd.lead_slurry_volume:.0f} bbl"),
                ("Lead TT",
                 f"{cd.lead_slurry_thickening_time:.0f} hrs"),
                ("Lead CS",
                 DataFormatter.pressure(
                     cd.lead_slurry_compressive_strength)),
                ("Tail Type", cd.tail_slurry_type),
                ("Tail Density",
                 DataFormatter.mud_weight(
                     cd.tail_slurry_density)),
                ("Tail Volume",
                 f"{cd.tail_slurry_volume:.0f} bbl"),
                ("Spacer Type", cd.spacer_type),
                ("Spacer Volume",
                 f"{cd.spacer_volume:.0f} bbl"),
                ("Displacement Vol",
                 f"{cd.displacement_volume:.0f} bbl"),
                ("Displacement Rate",
                 f"{cd.displacement_rate:.0f} bpm"),
                ("WOC Time",
                 f"{cd.woc_time:.0f} hrs"),
                ("Bump Pressure",
                 DataFormatter.pressure(
                     cd.plug_bump_pressure)),
                ("Excess",
                 DataFormatter.percent(
                     cd.excess_percentage)),
                ("CBL/CBIL",
                 "Required"
                 if cd.cbl_cbil_required
                 else "Not Required"),
            ]
            TableHelper.create_key_value_table(
                self.doc, params, num_cols=4)

            if cd.cement_additives:
                p = self.doc.add_paragraph()
                r1 = p.add_run("Additives: ")
                r1.bold = True
                r1.font.size = Pt(9)
                r2 = p.add_run(cd.cement_additives)
                r2.font.size = Pt(9)

            if cd.remarks:
                self._add_note(
                    cd.remarks,
                    "WARNING"
                    if "critical" in cd.remarks.lower()
                    else "INFO")

            self.doc.add_paragraph("")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # DIRECTIONAL
    # ----------------------------------------------------------------

    def _create_directional_section(self):
        self.section_counter += 1
        dp = self.project.directional_plan

        self._add_section_heading(
            f"{self.section_counter}. "
            f"DIRECTIONAL DRILLING PLAN", 1)

        self._add_section_heading(
            "Directional Summary", 2)

        dir_data = [
            ("Survey Tool", dp.survey_tool),
            ("Survey Interval",
             f"{dp.survey_frequency:.0f} ft"),
            ("KOP (MD)",
             DataFormatter.depth(dp.kickoff_point_md)),
            ("KOP (TVD)",
             DataFormatter.depth(dp.kickoff_point_tvd)),
            ("Build Rate",
             f"{dp.build_rate:.2f} deg/100ft"),
            ("Hold Inclination",
             f"{dp.hold_inclination:.1f}"),
            ("Hold Azimuth", f"{dp.hold_azimuth:.1f}"),
            ("Target Inclination",
             f"{dp.target_inclination:.1f}"),
            ("Target Azimuth",
             f"{dp.target_azimuth:.1f}"),
            ("Max DLS",
             f"{dp.max_dls:.2f} deg/100ft"),
            ("Horizontal Displacement",
             DataFormatter.depth(dp.horizontal_displacement)),
            ("Vertical Section",
             DataFormatter.depth(dp.vertical_section)),
        ]
        TableHelper.create_key_value_table(
            self.doc, dir_data, num_cols=4)

        if dp.wellpath_data:
            self._add_section_heading(
                "Planned Wellpath", 2)
            wp_headers = [
                "MD\n(ft)", "TVD\n(ft)",
                "Inc\n(deg)", "Azi\n(deg)",
                "DLS\n(deg/100ft)",
                "N/S\n(ft)", "E/W\n(ft)", "VS\n(ft)"
            ]
            wp_data = [
                [str(wp.get('md', 0)),
                 str(wp.get('tvd', 0)),
                 str(wp.get('inclination', 0)),
                 str(wp.get('azimuth', 0)),
                 str(wp.get('dls', 0)),
                 str(wp.get('ns', '')),
                 str(wp.get('ew', '')),
                 str(wp.get('vs', ''))]
                for wp in dp.wellpath_data
            ]
            TableHelper.create_professional_table(
                self.doc, wp_headers, wp_data,
                caption=(f"Table {self._next_table()}: "
                         f"Planned Wellpath Data"))

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # BOP & WELL CONTROL
    # ----------------------------------------------------------------

    def _create_bop_well_control_section(self):
        self.section_counter += 1
        bop = self.project.bop_stack
        wc = self.project.well_control

        self._add_section_heading(
            f"{self.section_counter}. "
            f"BOP & WELL CONTROL", 1)

        self._add_section_heading(
            "BOP Stack Configuration", 2)

        bop_data = [
            ("BOP Type", bop.bop_type),
            ("Working Pressure",
             DataFormatter.pressure(bop.working_pressure)),
            ("Bore Size", f'{bop.bore_size}"'),
            ("Manufacturer", bop.manufacturer),
            ("Model", bop.model),
            ("Annular Preventer",
             f'{bop.annular_preventer_size}" @ '
             f'{bop.annular_preventer_wp:,.0f} psi'),
            ("Pipe Ram Size", bop.pipe_ram_size),
            ("Blind Rams",
             "Yes" if bop.blind_ram else "No"),
            ("Shear Rams",
             "Yes" if bop.shear_ram else "No"),
            ("Variable Bore Rams",
             "Yes" if bop.variable_bore_ram else "No"),
            ("Kill Line", f'{bop.kill_line_size}" ID'),
            ("Choke Line", f'{bop.choke_line_size}" ID'),
            ("Choke Manifold WP",
             DataFormatter.pressure(bop.choke_manifold_wp)),
            ("Accumulator",
             f"{bop.accumulator_capacity:,.0f} gal @ "
             f"{bop.accumulator_precharge:,.0f} psi"),
            ("Diverter Size", f'{bop.diverter_size}"'),
            ("Function Test Freq",
             bop.function_test_frequency),
            ("Pressure Test Freq",
             bop.pressure_test_frequency),
            ("Low Test Pressure",
             DataFormatter.pressure(
                 bop.bop_test_pressure_low)),
            ("High Test Pressure",
             DataFormatter.pressure(
                 bop.bop_test_pressure_high)),
        ]
        TableHelper.create_key_value_table(
            self.doc, bop_data, num_cols=4)

        self._add_section_heading(
            "Well Control Data", 2)

        wc_data = [
            ("Kill Method", wc.kill_method),
            ("MAASP",
             DataFormatter.pressure(wc.maasp_surface)),
            ("Kick Tolerance",
             f"{wc.kick_tolerance:.0f} bbl"),
            ("Pit Gain Alarm",
             f"{wc.pit_gain_action_level:.0f} bbl"),
            ("Slow Pump Rate #1",
             f"{wc.slow_pump_rate_1:.0f} SPM @ "
             f"{wc.slow_pump_pressure_1:,.0f} psi"),
            ("Slow Pump Rate #2",
             f"{wc.slow_pump_rate_2:.0f} SPM @ "
             f"{wc.slow_pump_pressure_2:,.0f} psi"),
        ]
        TableHelper.create_key_value_table(
            self.doc, wc_data, num_cols=4)

        if wc.emergency_contacts:
            self._add_section_heading(
                "Emergency Contacts", 2)
            p = self.doc.add_paragraph(
                wc.emergency_contacts)
            for run in p.runs:
                run.font.size = Pt(10)

        if wc.h2s_action_levels:
            self._add_note(
                wc.h2s_action_levels, "WARNING")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # EVALUATION
    # ----------------------------------------------------------------

    def _create_evaluation_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"EVALUATION PROGRAM", 1)

        self._add_note(
            "Evaluation program includes wireline logging, "
            "LWD/MWD data, coring, formation testing, "
            "and cement evaluation. "
            "Coordinate with geology and reservoir teams.",
            "NOTE")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # TIME ESTIMATE
    # ----------------------------------------------------------------

    def _create_time_estimate_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"TIME vs DEPTH ESTIMATE", 1)

        if not self.project.time_estimates:
            self.doc.add_paragraph(
                "No time estimate data entered.")
            self.doc.add_page_break()
            return

        self._add_section_heading(
            "Gantt Schedule View", 2)

        EnhancedTableBuilder.time_gantt_table(
            self.doc, self.project.time_estimates)

        p_cap = self.doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(
            f"Table {self._next_table()}: "
            f"Time vs Depth - Dark blue = planned operation")
        r_cap.font.size = Pt(8)
        r_cap.font.italic = True
        r_cap.font.color.rgb = DocColors.GRAY

        self.doc.add_paragraph("")

        self._add_section_heading(
            "Operations Summary", 2)

        headers = [
            "Section", "Operation",
            "Depth From\n(ft)", "Depth To\n(ft)",
            "Section\nDays", "Cumulative\nDays"
        ]
        data = [
            [te.section_name, te.operation,
             DataFormatter.depth(te.depth_from)
             if te.depth_from > 0 else "-",
             DataFormatter.depth(te.depth_to)
             if te.depth_to > 0 else "-",
             f"{te.total_section_days:.1f}",
             f"{te.cumulative_days:.1f}"]
            for te in self.project.time_estimates
        ]
        TableHelper.create_professional_table(
            self.doc, headers, data,
            caption=(f"Table {self._next_table()}: "
                     f"Time vs Depth Estimate (AFE)"))

        total_days = max(
            te.cumulative_days
            for te in self.project.time_estimates)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(
            f"Estimated Total Duration: "
            f"{total_days:.1f} Days")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = DocColors.DARK_NAVY

        self._add_note(
            "Time estimate includes contingency. "
            "Actual duration may vary.",
            "NOTE")

        self.doc.add_page_break()

    # ----------------------------------------------------------------
    # PROCEDURES
    # ----------------------------------------------------------------

    def _create_all_procedures_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"DETAILED OPERATING PROCEDURES", 1)

        p = self.doc.add_paragraph()
        run = p.add_run(
            "This section contains detailed step-by-step "
            "procedures for all major drilling operations. "
            "All procedures are based on industry best practices "
            "and API standards.")
        run.font.size = Pt(10)

        self.doc.add_paragraph("")
        procedures = self.proc_gen.generate_all_procedures()
        sub = 1

        proc_map = [
            ('pre_spud', "Pre-Spud Checklist"),
            ('conductor', "Conductor Setting Procedure"),
        ]
        for key, title in proc_map:
            if key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{sub} {title}", 2)
                self._add_procedure_steps(procedures[key])
                self.doc.add_page_break()
                sub += 1

        for casing in self.project.casing_design:
            sk = casing.section_name.lower().replace(' ', '_')
            for key, label in [
                (f'drill_{sk}',
                 f"Drilling - {casing.section_name}"),
                (f'trip_{sk}',
                 f"Tripping - {casing.section_name}"),
                (f'casing_run_{sk}',
                 f"Casing Running - {casing.casing_od}\""),
                (f'cement_{sk}',
                 f"Cementing - {casing.section_name}"),
            ]:
                if key in procedures:
                    self._add_section_heading(
                        f"{self.section_counter}.{sub} {label}", 2)
                    self._add_procedure_steps(procedures[key])
                    self.doc.add_page_break()
                    sub += 1

        fixed_procs = [
            ('bop_nipple_up', "BOP Nipple Up"),
            ('bop_test', "BOP Test"),
            ('lot_fit', "LOT / FIT"),
            ('well_control', "Well Control"),
            ('directional', "Directional Drilling"),
            ('logging', "Wireline Logging"),
            ('stuck_pipe', "Stuck Pipe"),
            ('lost_circulation', "Lost Circulation"),
            ('wellbore_cleanup', "Wellbore Cleanup"),
        ]
        for key, label in fixed_procs:
            if key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{sub} {label}", 2)
                self._add_procedure_steps(procedures[key])
                self.doc.add_page_break()
                sub += 1

    # ----------------------------------------------------------------
    # EMERGENCY PROCEDURES
    # ----------------------------------------------------------------

    def _create_emergency_section(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. "
            f"EMERGENCY PROCEDURES", 1)

        procedures = self.proc_gen.generate_all_procedures()

        emerg_procs = [
            ('h2s_emergency', "1", "H2S Emergency"),
            ('kick_drill', "2", "Well Control Drill"),
            ('fishing', "3", "Fishing Operations"),
            ('hse', "4", "HSE Requirements"),
            ('abandonment', "5", "Well Abandonment"),
        ]
        for key, num, label in emerg_procs:
            if key in procedures:
                self._add_section_heading(
                    f"{self.section_counter}.{num} {label}", 2)
                self._add_procedure_steps(procedures[key])
                self.doc.add_page_break()

    # ----------------------------------------------------------------
    # APPENDICES
    # ----------------------------------------------------------------

    def _create_appendices(self):
        self.section_counter += 1

        self._add_section_heading(
            f"{self.section_counter}. APPENDICES", 1)

        appendices = [
            "Appendix A: Wellbore Schematic",
            "Appendix B: BHA Diagrams",
            "Appendix C: Casing Design Summary",
            "Appendix D: Directional Survey Program",
            "Appendix E: Cement Program Details",
            "Appendix F: Mud Program Details",
            "Appendix G: Hydraulics Calculations",
            "Appendix H: Kill Sheet",
            "Appendix I: Emergency Contact List",
            "Appendix J: Rig Floor Poster",
        ]

        for app in appendices:
            p = self.doc.add_paragraph()
            run = p.add_run(app)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DocColors.DARK_NAVY
            run.font.name = 'Calibri'

            p2 = self.doc.add_paragraph()
            run2 = p2.add_run(
                "  [Attached separately]")
            run2.font.size = Pt(9)
            run2.font.italic = True
            run2.font.color.rgb = DocColors.GRAY
            self.doc.add_paragraph("")

        self.doc.add_page_break()
        ci = self.project.company_info
        end_para = self.doc.add_paragraph()
        end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_para.paragraph_format.space_before = Pt(80)

        self._add_colored_line(DocColors.DARK_NAVY, 3)

        end_run = end_para.add_run(
            "\n\n--- END OF DRILLING PROGRAM ---\n\n")
        end_run.bold = True
        end_run.font.size = Pt(16)
        end_run.font.color.rgb = DocColors.DARK_NAVY

        end_run2 = end_para.add_run(
            f"{ci.operator_name}\n"
            f"Well: {ci.well_name}\n"
            f"Document: {ci.document_number}\n"
            f"Revision: {ci.revision}\n"
            f"Date: {datetime.now().strftime('%d-%B-%Y')}\n")
        end_run2.font.size = Pt(10)
        end_run2.font.color.rgb = DocColors.GRAY

        self._add_colored_line(DocColors.DARK_NAVY, 3)

    # ----------------------------------------------------------------
    # HEADERS & FOOTERS
    # ----------------------------------------------------------------

    def _add_headers_footers(self):
        ci = self.project.company_info
        for section in self.doc.sections:
            try:
                header = section.header
                header.is_linked_to_previous = False

                if header.paragraphs:
                    hp = header.paragraphs[0]
                else:
                    hp = header.add_paragraph()

                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hp.clear()

                r1 = hp.add_run(
                    ci.operator_name or "Operator")
                r1.font.size = Pt(8)
                r1.font.color.rgb = DocColors.NAVY
                r1.font.bold = True

                r2 = hp.add_run("  |  ")
                r2.font.size = Pt(8)

                r3 = hp.add_run(
                    f"Drilling Program - "
                    f"{ci.well_name or 'Well'}")
                r3.font.size = Pt(8)
                r3.font.color.rgb = DocColors.DARK_NAVY
                r3.font.bold = True

                r4 = hp.add_run("  |  ")
                r4.font.size = Pt(8)

                r5 = hp.add_run(
                    f"{ci.document_number or 'DRL-PRG-001'} "
                    f"Rev.{ci.revision or '0'}")
                r5.font.size = Pt(8)
                r5.font.color.rgb = DocColors.GRAY

                try:
                    pPr = hp._p.get_or_add_pPr()
                    pBdr = parse_xml(
                        f'<w:pBdr {nsdecls("w")}>'
                        f'<w:bottom w:val="single" '
                        f'w:sz="4" w:space="1" '
                        f'w:color="0F3460"/>'
                        f'</w:pBdr>')
                    pPr.append(pBdr)
                except Exception:
                    pass

                footer = section.footer
                footer.is_linked_to_previous = False

                if footer.paragraphs:
                    fp = footer.paragraphs[0]
                else:
                    fp = footer.add_paragraph()

                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fp.clear()

                rf1 = fp.add_run(
                    ci.classification.upper())
                rf1.font.size = Pt(7)
                rf1.font.color.rgb = DocColors.RED
                rf1.font.bold = True

                rf2 = fp.add_run(
                    f"  |  Printed: "
                    f"{datetime.now().strftime('%d-%b-%Y')}"
                    f"  |  Page ")
                rf2.font.size = Pt(7)
                rf2.font.color.rgb = DocColors.GRAY

                try:
                    from docx.oxml.ns import nsdecls as ns
                    rp1 = fp.add_run()
                    rp1._r.append(parse_xml(
                        f'<w:fldChar {ns("w")} '
                        f'w:fldCharType="begin"/>'))
                    rp1.font.size = Pt(7)

                    rp2 = fp.add_run()
                    rp2._r.append(parse_xml(
                        f'<w:instrText {ns("w")} '
                        f'xml:space="preserve">'
                        f' PAGE </w:instrText>'))
                    rp2.font.size = Pt(7)

                    rp3 = fp.add_run()
                    rp3._r.append(parse_xml(
                        f'<w:fldChar {ns("w")} '
                        f'w:fldCharType="end"/>'))
                    rp3.font.size = Pt(7)
                except Exception:
                    pass

            except Exception as e:
                logger.warning(
                    f"Header/footer error: {e}")