# ============================================================================
# TEMPLATE REGRESSION SUITE — all wizard templates end-to-end
# File: tests/regression_templates.py
# Audit item (Testing): integration + document regression.
#
# Drives the real generation pipeline (fill_template -> render_selected ->
# neutralize -> validation/readiness/standards/compliance -> deep engineering
# -> calculation register -> Word render) for EVERY template in the wizard
# (51 templates) and asserts that each output document contains the four
# governance sections plus the engineering calculation register.
#
# Run:  LD_LIBRARY_PATH=/tmp/glstubs PYTHONPATH=. QT_QPA_PLATFORM=offscreen \
#       python3 tests/regression_templates.py
# Exit code 0 = all templates pass.
# ============================================================================

import os
import re
import sys
import tempfile
import traceback

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from wizard_engine import extract_sections, neutralize_text
from wizard_library import ALL_TEMPLATES
from wizard_procedures import PROCEDURE_TEMPLATES
from wizard_offshore import OFFSHORE_TEMPLATES
from wizard_master import build_master_templates
from generation_pipeline import generate_document

PASS = 0
FAIL = 0
FAILURES = []


# Generic but realistic values — cover the union of wizard input keys so
# that every placeholder is filled and validation produces few findings.
def generic_values():
    v = {
        # identity — neutral, no company/field names
        "operator": "the Operator", "contractor": "the Service Company",
        "well_name": "Well A", "field": "Field X", "field_name": "Field X",
        "reservoir": "Main Reservoir", "target_formation": "Main Reservoir",
        "rig": "Rig 1", "rig_name": "Rig 1", "rig_type": "Land Rig",
        "country": "", "province": "",
        # geometry
        "hole_size": "12.25", "bit_size": "12.25", "casing_size": "9.625",
        "casing_od": "9.625", "casing_wall": "0.472",
        "casing_yield": "110000", "casing_depth": "8000",
        "pipe_od": "5", "bha_od": "8", "tubing_size": "3.5",
        "depth": "10000", "depth_m": "3050", "td_depth": "10000",
        "target_depth": "10000", "td_md": "10000", "td_tvd": "9800",
        "total_depth": "10000", "total_depth_md": "10000",
        "vertical_td": "9500", "landing_depth": "9800", "landing_tvd": "9600",
        "kop": "3000", "kop_depth": "3000", "water_depth": "0",
        "kb_elevation": "30", "rt_elevation": "30", "ground_elevation": "0",
        # fluids
        "mud_weight": "12", "mud_type": "Oil Based Mud",
        "yield_point": "20", "plastic_viscosity": "25",
        "flow_rate": "500", "pump_rate": "500", "circ_rate": "500",
        "ann_velocity": "100", "n_index": "0.6", "k_index": "120",
        "yield_stress": "8", "trip_speed": "60",
        "kill_mw": "12.5", "kill_fluid_weight": "12.5",
        "ecd": "12.4", "ecd_max": "12.6", "maasp": "2000",
        "fg": "16", "fracture_gradient": "16", "frac_gradient": "16",
        "formation_pressure": "11", "pressure_gradient": "0.57",
        "sidpp": "400", "sidpip": "200", "pit_gain": "20",
        "bop_wp": "10000", "bop_rating": "10000",
        # casing / cement
        "casing_program": "20 in @ 500 ft; 13.375 in @ 4000 ft; 9.625 in "
                          "@ 10000 ft",
        "casing_table": "|Size|Depth|Grade|Weight|\n"
                        "|---|---|---|---|\n|9.625|8000|L80|47|",
        "cement_type": "Class G", "slurry_density": "15.8",
        "slurry_volume": "500", "lead_volume": "300", "tail_volume": "200",
        "lead_density": "14.5", "tail_density": "16.2",
        "excess": "30", "shoe_track": "60", "toc": "4000",
        "centralizers": "Centralized every 60 ft", "float_equipment": "Yes",
        "woc": "8", "woc_time": "8",
        # BHA / bit
        "bha": "12.25 in PDC bit + 9.5 in motor + 8 in stabilizer",
        "bha_desc": "12.25 in PDC bit; 9.5 in mud motor; 8 in stabilizer; "
                    "MWD; 6.5 in DC; HWDP; 5 in DP",
        "bit_type": "PDC", "nozzles": "3 x 14", "tfa": "1.2",
        "bit_hsi": "3.5", "hsi": "3.5", "bit_pressure_drop": "800",
        "jet_impact": "2.5", "nozzle_velocity": "350",
        "wob": "25", "rpm": "100", "rotate_rpm": "100", "rop": "30",
        "rop_target": "30",
        # directional
        "well_profile": "J-Shape", "well_type": "Development",
        "build_rate": "3", "max_inc": "35", "hold_inclination": "35",
        "hold_azimuth": "90", "max_dls": "3.5", "horizontal_displacement": "1500",
        "survey_interval": "100", "survey_tool": "MWD",
        "anti_collision": "Separation factor > 1.5 maintained",
        "trajectory_table": "|MD|Inc|Azi|\n|---|---|---|\n|1000|5|90|\n|5000|35|90|",
        # well control
        "bop_stack": "Ram BOP stack with annular; 2 pipe rams; blind shear",
        "pipe_rams": "2 x 5 in pipe rams; blind shear ram",
        "choke": "2 x 4 in remote chokes", "choke_manifold": "10K",
        "kill_plan": "Wait and Weight method", "kill_method": "Wait and Weight",
        "primary_method": "Wait and Weight", "secondary_method": "Bullheading",
        "kill_volume": "400", "kill_time": "2", "kill_depth": "8000",
        "kill_hydrostatic": "5200", "kill_surface_pressure": "2000",
        "kill_fluid": "12.5 ppg OBM", "kill_barriers": "BOP rams; kill fluid",
        "icp": "1500", "fcp": "800", "slow_pump_rates": "0.5 bbl/min @ 400 psi",
        "shut_in_procedure": "1. Stop rotation; 2. pick up off bottom; "
                             "3. shut BOP; 4. record SIDPP/SICP",
        # time & cost
        "total_days": "45", "time_days": "45", "duration": "45",
        "total_cost": "12000000", "cost_per_m": "3800",
        "npt_contingency": "8", "npt_percent": "10",
        # misc
        "h2s": "No", "classification": "Normal", "environment": "Onshore",
        "operation": "Drilling", "operation_type": "Drilling",
        "job_type": "Drilling", "workover_type": "Workover",
        "interval": "Main reservoir", "packer": "Retrievable packer",
        "rate_loss": "10 bbl/hr", "short_history": "Offset wells drilled "
                    "without major problems", "fish_description": "Fish",
        "wellhead_type": "Compact wellhead", "water_depth": "0",
        "plug_depth": "5000", "plugs_table": "|Plug|Depth|Test|\n"
                    "|---|---|---|\n|P1|5000|500 psi|",
        "problem": "Intermittent losses observed in offset wells",
    }
    return v


def all_templates():
    return list(ALL_TEMPLATES) + list(PROCEDURE_TEMPLATES) + \
        list(OFFSHORE_TEMPLATES) + build_master_templates()


def run_template(tdef, out_dir, values):
    """Generate via the shared headless pipeline (same as the REST API)."""
    out = os.path.join(out_dir, f"{tdef.key}.docx")
    report = generate_document(
        tdef, values,
        meta={"prepared_by": "Engineer", "reviewed_by": "Lead Engineer",
              "approved_by": "Drilling Manager", "document_number": "",
              "date": "17-August-2026", "revision": "01"},
        options={"font": "Calibri", "font_size": 11.0, "page": "A4",
                 "orientation": "Portrait", "margin_left": 2.5,
                 "margin_right": 2.0, "margin_top": 2.0,
                 "margin_bottom": 2.0, "cover": True, "toc": True,
                 "header_text": "", "footer_text": ""},
        out_path=out)
    if not report["ok"] or not os.path.exists(out):
        raise RuntimeError("docx generation failed")
    return out


def docx_text(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


REQUIRED_SECTIONS = [
    ("VALIDATION & COMPLIANCE", "validation"),
    ("PROGRAM READINESS SCORE", "readiness"),
    ("STANDARDS COMPLIANCE MATRIX", "standards"),
    ("DOCUMENT COMPLIANCE REPORT", "compliance"),
    ("ENGINEERING CALCULATION REGISTER", "register"),
]


def main():
    global PASS, FAIL
    templates = all_templates()
    print(f"Regression: {len(templates)} templates\n")
    out_dir = tempfile.mkdtemp(prefix="tmpl_reg_")
    values = generic_values()
    for i, tdef in enumerate(templates, 1):
        label = f"[{i:02d}/{len(templates)}] {tdef.icon} {tdef.name}"
        try:
            out = run_template(tdef, out_dir, values)
            text = docx_text(out).upper()
            missing = [name for name, _ in REQUIRED_SECTIONS
                       if name not in text]
            if missing:
                raise RuntimeError("missing sections: " + ", ".join(missing))
            # entity scrub: no blacklisted tokens may survive (word boundary)
            leak = leak_scan(text)
            if leak:
                raise RuntimeError("leak: " + ", ".join(leak))
            PASS += 1
            print(f"  ✔ {label}")
        except Exception as e:
            FAIL += 1
            FAILURES.append(f"{tdef.key}: {e}")
            print(f"  ✘ {label}\n    {e}")
            traceback.print_exc()
    print("\n" + "=" * 60)
    print(f"RESULT: {PASS}/{len(templates)} templates PASS")
    if FAIL:
        print("FAILURES:")
        for f in FAILURES:
            print("  -", f)
        return 1
    return 0


def leak_scan(text: str) -> list:
    """Scan for known company/field tokens with word boundaries."""
    from wizard_engine import (OPERATOR_NAMES, SERVICE_NAMES, WELL_PATTERNS)
    leaks = []
    for pat in OPERATOR_NAMES + SERVICE_NAMES:
        p = re.compile(pat, re.IGNORECASE)
        if p.search(text):
            leaks.append(pat)
    for pat, _repl in WELL_PATTERNS:
        p = re.compile(pat, re.IGNORECASE)
        if p.search(text):
            leaks.append(pat)
    return leaks


if __name__ == "__main__":
    sys.exit(main())
