# ============================================================================
# CBS DATABASE MODULE — Cost Breakdown Structure
# File: cbs_db.py
# Editable price database for drilling goods & services (AFE / CBS).
# Catalog seeded from a generic industry catalog (prices are editable
# defaults; the user's own price tables can be imported via
# seed_azns_prices.py / the CBS tab). No well names or company names
# are stored — the catalog is fully general.
# ============================================================================

import sqlite3
import json
import os
import re
from datetime import datetime
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Dict, Optional

APP_DIR = Path.home() / ".drilling_program"
DEFAULT_DB = str(APP_DIR / "cbs.db")


# ============================================================================
# GENERALIZATION — keep the catalog free of well names / company names /
# reservoir names. Runs on every load so the database stays general.
# ============================================================================

# (regex, replacement) — applied to item names, descriptions, sources.
GENERALIZE_RULES = [
    # well codes & project names
    (r"\bAZNS\s*[-–]?\s*[A-Z0-9]*\b", ""),
    (r"\bF-?20\b", ""),
    (r"\bPAD[- ]?\d+\b", ""),
    (r"\bSIAH MAKAN\b", ""),
    (r"\bSI-?\d+\b", ""),
    (r"\bAZR[- ]?\d+\b", ""),
    (r"\bG-?\d+\b", ""),
    # company / brand names (defensive; the wizard neutralize also handles it)
    (r"\bMSA\b", ""),
    (r"\bOEOC\b", ""),
    (r"\bNIOC\b", ""),
    (r"\bNISOC\b", ""),
    (r"\bPEDEC\b", ""),
    (r"\bPEDCO\b", ""),
    (r"\bSLB\b", ""),
    (r"\bSchlumberger\b", ""),
    (r"\bHalliburton\b", ""),
    (r"\bBaker Hughes\b", ""),
    (r"\bBaker\b", "thread-lock"),
    (r"\bHALCO\b", ""),
    (r"\bAnadrill\b", ""),
    (r"\bWeatherford\b", ""),
    (r"\bKEPCO\b", ""),
    (r"\bNDCO\b", ""),
    (r"\bIOOC\b", ""),
    (r"\bNICO\b", ""),
    (r"\bMI SWACO\b", ""),
    (r"\bBaroid\b", ""),
    (r"\bTOTCO\b", ""),
    (r"\bSPD\s*\d*\b", ""),
    (r"\bKangan\b", ""),
    (r"\bYaran\b", ""),
    (r"\bShah Deniz\b", "the field"),
    (r"\bSDX-?\d*\b", ""),
    (r"\bSD A-03\b", "the well"),
    # major-operator reference manuals (knowledge sources — never in output)
    (r"\bADCO\b", ""),
    (r"\bAbu Dhabi Company for Onshore Oil Operations\b", ""),
    (r"\bSaudi Aramco\b", ""),
    (r"\bAramco\b", ""),
    (r"\bNimir\b", ""),
    (r"\bPETROM\b", ""),
    (r"\bExxonMobil\b", ""),
    (r"\bExxon\b", ""),
    (r"\bMobil\b", ""),
    (r"\bIADC\b", ""),
    (r"\bSteve Devereux\b", ""),
    (r"\bDevereux\b", ""),
    (r"\bWell Control School\b", ""),
    (r"\bAberdeen\b", ""),
    (r"\bHouston, Texas\b", ""),
    (r"\bMI\b(?!-)", ""),
    (r"\bMi-cide\b", "biocide"),
    (r"\bDril-Quip\b", ""),
    (r"\bBaker \(Brown\)\b", ""),
    (r"\bBrown HMC\b", ""),
    (r"\bHSR\b", ""),
    (r"\bHydril\b", ""),
    (r"\bTIW\b", ""),
    (r"\bBJ Services\b", ""),
    (r"\bHalliburton\b", ""),
    (r"\bRTTS\b", ""),
    (r"\bIngram-Cactus\b", ""),
    (r"\bCentric\b", ""),
    (r"\bBobcat\b", ""),
    (r"\bJohnson\b", ""),
    (r"\bArrow\b", ""),
    (r"\bHurricane\b", ""),
    (r"\bEZ Drill\b", ""),
    (r"\bEZ\b", ""),
    (r"\bBrown\b", ""),
    (r"\bBP\b", ""),
    # reservoir/formation-specific well types -> generic descriptors
    (r"\bfor\s+Fahliyan\s+wells?\b", "for HPHT wells"),
    (r"\bfor\s+Sarvak\s+wells?\b", "for standard wells"),
    (r"\bfor\s+Water Disposal\s+wells?\b", "for water-disposal wells"),
    (r"\bfor\s+Kazhdumi\s+wells?\b", "for wells"),
    (r"\bfor\s+Gadvan\s+wells?\b", "for wells"),
    (r"\bFahliyan\b", "HPHT"),
    (r"\bSarvak\b", ""),
    (r"\bKazhdumi\b", ""),
    (r"\bGadvan\b", ""),
    (r"\bin\s+[A-Za-z]+,\s*[A-Za-z]+ and HPHT wells?\b", "wells"),
    (r"\bvertical wells?\b", "vertical wells"),
    (r"\bHorizontal wells?\b", "horizontal wells"),
    (r"\bdirectional wells?\b", "directional wells"),
]

_WS = re.compile(r"\s{2,}")


def generalize_text(text: str) -> str:
    """حذف نام چاه/شرکت/مخزن از متن — خروجی همیشه جنرال"""
    if not text:
        return text
    out = text
    for pat, repl in GENERALIZE_RULES:
        out = re.sub(pat, repl, out, flags=re.IGNORECASE)
    out = _WS.sub(" ", out)
    out = out.replace(" ,", ",").replace("  ", " ")
    return out.strip(" ,.-–—")


# ============================================================================
# DATA MODEL
# ============================================================================

@dataclass
class CbsItem:
    """یک قلم کالا/خدمت در ساختار شکست هزینه (CBS)"""
    id: int = 0
    code: str = ""                # e.g. "1.1.1"
    category: str = ""            # e.g. "Drilling Services"
    name: str = ""                # e.g. "PDC Bit 8-1/2\""
    description: str = ""
    unit: str = "each"            # each / set / day / m / m3 / bbl / kg / trip
    unit_price: float = 0.0       # USD — default, editable
    qty: float = 0.0
    is_service: int = 0           # 1 = service (day/rate based), 0 = commodity
    source: str = ""              # reference document (generic, no well/company names)
    price_source: str = "TBD"     # TBD / Price Table / User
    active: int = 1


# ============================================================================
# SEED CATALOG — names & services from Azadegan field documents
# Prices are 0.00 (TBD) until a price table is loaded or the user edits.
# ============================================================================

def _i(code, category, name, desc, unit, price=0.0, qty=0.0, svc=0, src="Default catalog"):
    return CbsItem(code=code, category=category, name=name, description=desc,
                   unit=unit, unit_price=price, qty=qty, is_service=svc, source=src)

SEED_CATALOG: List[CbsItem] = [
    # ------------------------------------------------------------------ 1. RIG
    _i("1.1.1", "1. Rig & Dayrates", "Drilling Rig Day Rate",
       "Heavy land rig (e.g. 2000 HP) — day rate", "day", svc=1),
    _i("1.1.2", "1. Rig & Dayrates", "Spread Cost",
       "Crew, camp, catering, power, water, fuel — per day", "day", svc=1),
    _i("1.1.3", "1. Rig & Dayrates", "Rig Mobilization",
       "Mob of rig & ancillary equipment to location", "lump sum", svc=1),
    _i("1.1.4", "1. Rig & Dayrates", "Rig Demobilization",
       "Demob after well completion", "lump sum", svc=1),
    _i("1.1.5", "1. Rig & Dayrates", "Rig Move (Skidding)",
       "Skidding / rig move between locations", "trip", svc=1),
    _i("1.1.6", "1. Rig & Dayrates", "Diesel Fuel",
       "Fuel for rig power generation", "liter"),
    _i("1.1.7", "1. Rig & Dayrates", "Drill Water",
       "Drill water supply", "m3"),
    _i("1.1.8", "1. Rig & Dayrates", "Potable Water",
       "Domestic water supply", "m3"),
    _i("1.1.9", "1. Rig & Dayrates", "Camp & Catering",
       "Accommodation & catering — per head per day", "day", svc=1),
    _i("1.1.10", "1. Rig & Dayrates", "Waste Management Service",
       "Cuttings & waste treatment, discharge water specs per contract", "day", svc=1),
    _i("1.1.11", "1. Rig & Dayrates", "Helicopter / Logistics Support",
       "Crew change & priority logistics", "hour", svc=1),
    _i("1.1.12", "1. Rig & Dayrates", "Transportation (Trucks)",
       "Trucking of equipment & supplies", "trip", svc=1),

    # --------------------------------------------------- 2. DRILLING SERVICES
    _i("2.1.1", "2. Drilling Services", "Directional Drilling Package",
       "MOTOR/RSS + MWD/LWD incl. engineers — day rate", "day", svc=1),
    _i("2.1.2", "2. Drilling Services", "MWD/LWD Tools",
       "Measurement / logging while drilling tool rental", "day", svc=1),
    _i("2.1.3", "2. Drilling Services", "Rotary Steerable System (RSS)",
       "RSS BHA package incl. field engineers", "day", svc=1),
    _i("2.1.4", "2. Drilling Services", "Drill Bit — PDC",
       "PDC bit per size (16\", 12-1/4\", 8-1/2\", 6\")", "each"),
    _i("2.1.5", "2. Drilling Services", "Drill Bit — Tricone/Roller",
       "Roller cone bit per size", "each"),
    _i("2.1.6", "2. Drilling Services", "Mud Logging Unit",
       "Gas chromatograph, sensors, geologist — day rate", "day", svc=1),
    _i("2.1.7", "2. Drilling Services", "Surveying (Totco / Single Shot)",
       "Directional surveys incl. gyro where required", "each", svc=1),
    _i("2.1.8", "2. Drilling Services", "Coring Service",
       "Conventional core barrel + core handling", "run", svc=1),
    _i("2.1.9", "2. Drilling Services", "Fishing Tools & Services",
       "Overshot, jars, mills, fishing engineer", "day", svc=1),
    _i("2.1.10", "2. Drilling Services", "Solids Control Equipment",
       "Shale shakers, centrifuges, degasser rental", "day", svc=1),
    _i("2.1.11", "2. Drilling Services", "Downhole Motors / Turbo",
       "Positive displacement motor rental", "day", svc=1),
    _i("2.1.12", "2. Drilling Services", "Junk Sub / String Tools Rental",
       "Drill string accessories rental", "day", svc=1),

    # ------------------------------------------------------ 3. DRILLING FLUIDS
    _i("3.1.1", "3. Drilling Fluids & Chemicals", "Water Base Mud System",
       "WBM incl. additives — per m3 (per hole section)", "m3"),
    _i("3.1.2", "3. Drilling Fluids & Chemicals", "Oil Base Mud System",
       "OBM / SBM incl. base oil & additives", "m3"),
    _i("3.1.3", "3. Drilling Fluids & Chemicals", "Hi-Vis / Sweep Pills",
       "High viscosity sweeps & pills", "m3"),
    _i("3.1.4", "3. Drilling Fluids & Chemicals", "Barite",
       "Weighting material", "kg"),
    _i("3.1.5", "3. Drilling Fluids & Chemicals", "Bentonite",
       "Viscosifier", "kg"),
    _i("3.1.6", "3. Drilling Fluids & Chemicals", "Caustic Soda (NaOH)",
       "Alkalinity control", "kg"),
    _i("3.1.7", "3. Drilling Fluids & Chemicals", "Soda Ash (Na2CO3)",
       "Calcium removal", "kg"),
    _i("3.1.8", "3. Drilling Fluids & Chemicals", "LCM (Lost Circulation Materials)",
       "Mica, nut plug, fibers, etc.", "kg"),
    _i("3.1.9", "3. Drilling Fluids & Chemicals", "KCl (Potassium Chloride)",
       "Shale inhibition additive", "kg"),
    _i("3.1.10", "3. Drilling Fluids & Chemicals", "PAC / CMC Polymers",
       "Fluid loss control polymers", "kg"),
    _i("3.1.11", "3. Drilling Fluids & Chemicals", "Xanthan Gum (XC)",
       "Viscosifier", "kg"),
    _i("3.1.12", "3. Drilling Fluids & Chemicals", "Mud Testing / Lab",
       "Daily mud checks & lab support", "day", svc=1),

    # ----------------------------------------------------------- 4. CEMENTING
    _i("4.1.1", "4. Cementing", "Cementing Service (per job)",
       "Cement unit + crew + engineer — per casing/plug job", "job", svc=1),
    _i("4.1.2", "4. Cementing", "Class G Cement",
       "API Class G (HSR) cement", "ton"),
    _i("4.1.3", "4. Cementing", "Cement Additives",
       "Retarders, accelerators, dispersants, fluid loss", "kg"),
    _i("4.1.4", "4. Cementing", "Cement Plug Job (Balanced)",
       "Balanced cement plug — service & material", "job", svc=1),
    _i("4.1.5", "4. Cementing", "Cementing Head & Accessories",
       "Cement head, plugs, wiper darts", "job", svc=1),
    _i("4.1.6", "4. Cementing", "CIT / FIT Test",
       "Casing integrity / formation integrity test", "job", svc=1),

    # ----------------------------------------------------------- 5. TUBULARS
    _i("5.1.1", "5. Tubulars & Connectors", "Casing 20\" (e.g. K-55, BTC)",
       "Conductor / surface casing per size & grade", "m"),
    _i("5.1.2", "5. Tubulars & Connectors", "Casing 13-3/8\" (e.g. N-80)",
       "Surface casing per size & grade", "m"),
    _i("5.1.3", "5. Tubulars & Connectors", "Casing 9-5/8\" (e.g. L-80 / N-80)",
       "Intermediate casing per size & grade", "m"),
    _i("5.1.4", "5. Tubulars & Connectors", "Liner 7\" (e.g. L-80 / 13Cr)",
       "Production liner per size & grade", "m"),
    _i("5.1.5", "5. Tubulars & Connectors", "Tubing 4-1/2\" (e.g. L-80, premium)",
       "Completion tubing incl. premium connections", "m"),
    _i("5.1.6", "5. Tubulars & Connectors", "Premium Connections",
       "Threading / premium connection premium", "joint"),
    _i("5.1.7", "5. Tubulars & Connectors", "Centralizers & Stop Collars",
       "Casing centralization accessories", "each"),
    _i("5.1.8", "5. Tubulars & Connectors", "Crossover Subs & Pup Joints",
       "Drill string / casing accessories", "each"),
    _i("5.1.9", "5. Tubulars & Connectors", "Wear Sleeves / Pipe Protection",
       "Casing & drill pipe protection", "each"),

    # ------------------------------------------------------- 6. COMPLETION
    _i("6.1.1", "6. Completion Equipment", "ESP String (Pump + Motor + Seal)",
       "Electrical submersible pump package per spec", "set"),
    _i("6.1.2", "6. Completion Equipment", "ESP Cable & Accessories",
       "Power cable, pothead, bands", "set"),
    _i("6.1.3", "6. Completion Equipment", "Production Packer",
       "Retrievable / permanent packer", "each"),
    _i("6.1.4", "6. Completion Equipment", "TRSV / TRSSSV",
       "Surface controlled subsurface safety valve", "each"),
    _i("6.1.5", "6. Completion Equipment", "SSD / Sliding Sleeve",
       "Selective sliding sleeve", "each"),
    _i("6.1.6", "6. Completion Equipment", "Nipple Profile & NO-GO",
       "Landing nipples & plugs", "each"),
    _i("6.1.7", "6. Completion Equipment", "Xmas Tree & Wellhead",
       "Production tree incl. installation", "set"),
    _i("6.1.8", "6. Completion Equipment", "Completion / Workover Fluids",
       "Completion brine, kill fluid", "m3"),
    _i("6.1.9", "6. Completion Equipment", "Slickline Service",
       "Slickline unit + crew (day rate)", "day", svc=1),
    _i("6.1.10", "6. Completion Equipment", "Perforating (TCP / Wireline)",
       "Perforating service incl. guns", "job", svc=1),
    _i("6.1.11", "6. Completion Equipment", "Wellhead Installation Service",
       "Wellhead & casing hanger installation", "job", svc=1),

    # -------------------------------------------------- 7. WELL CONTROL / TEST
    _i("7.1.1", "7. Well Control & Testing", "BOP Stack Rental",
       "Annular + ram BOPs per size/pressure rating", "day", svc=1),
    _i("7.1.2", "7. Well Control & Testing", "BOP Test Service",
       "BOP & choke manifold pressure testing", "test", svc=1),
    _i("7.1.3", "7. Well Control & Testing", "Choke & Kill Manifold",
       "Manifold rental & hook-up", "day", svc=1),
    _i("7.1.4", "7. Well Control & Testing", "Accumulator Unit",
       "BOP closing unit rental", "day", svc=1),
    _i("7.1.5", "7. Well Control & Testing", "DST Service",
       "Drill stem testing incl. tools & crew", "job", svc=1),
    _i("7.1.6", "7. Well Control & Testing", "Wireline Logging (Open Hole)",
       "Full suite open hole logs", "job", svc=1),
    _i("7.1.7", "7. Well Control & Testing", "Wireline Logging (Cased Hole)",
       "CBL/VDL, gyro, cement evaluation", "job", svc=1),
    _i("7.1.8", "7. Well Control & Testing", "Well Testing (Surface)",
       "Surface well test package — day rate", "day", svc=1),
    _i("7.1.9", "7. Well Control & Testing", "Coiled Tubing Service",
       "CT unit + crew — day rate", "day", svc=1),
    _i("7.1.10", "7. Well Control & Testing", "Stimulation / Acidizing",
       "Acidizing & stimulation service", "job", svc=1),

    # -------------------------------------- 8. WASTE MGMT (contract equipment)
    _i("8.1.1", "8. Waste Management Equipment", "Sediment Remover",
       "Waste water pre-treatment (contract table 1)", "set"),
    _i("8.1.2", "8. Waste Management Equipment", "Waste Water Conditioning Tank",
       "Conditioning tank — contract table 1", "set"),
    _i("8.1.3", "8. Waste Management Equipment", "Filter for Oil Removal",
       "Oil removal filter — contract table 1", "set"),
    _i("8.1.4", "8. Waste Management Equipment", "Reaction Tank (Organic Decomposition)",
       "Reaction tank — contract table 1", "set"),
    _i("8.1.5", "8. Waste Management Equipment", "Fiber Ball Filter",
       "Fiber ball filter — contract table 1", "piece"),
    _i("8.1.6", "8. Waste Management Equipment", "Nanometer Filtrating Equipment",
       "Nano filtration — contract table 1", "set"),
    _i("8.1.7", "8. Waste Management Equipment", "Waste Water Pump",
       "Transfer pump — contract table 1", "set"),
    _i("8.1.8", "8. Waste Management Equipment", "Chemical Feeder",
       "Chemical dosing — contract table 2", "set"),
    _i("8.1.9", "8. Waste Management Equipment", "Flocculation & Gel Breaking Tank",
       "Reaction tank — contract table 2", "set"),
    _i("8.1.10", "8. Waste Management Equipment", "Vacuum Dewatering Apparatus",
       "Dewatering — contract table 2", "set"),
    _i("8.1.11", "8. Waste Management Equipment", "Precise Filter",
       "Final filter — contract table 2", "piece"),
    _i("8.1.12", "8. Waste Management Equipment", "COD Remover",
       "COD removal unit — contract table 2", "set"),
    _i("8.1.13", "8. Waste Management Equipment", "Ultra-Filter",
       "Ultra filtration — contract table 2", "set"),
    _i("8.1.14", "8. Waste Management Equipment", "Reverse Osmosis Apparatus",
       "RO unit — contract table 2", "set"),
    _i("8.1.15", "8. Waste Management Equipment", "Mud Transmitting Machine",
       "Sludge transfer — contract table 2", "set"),
    _i("8.1.16", "8. Waste Management Equipment", "Water Quality Examination Apparatus",
       "Lab analysis — contract table 2", "set"),

    # ------------------------------------------------- 9. OVERHEADS & OTHER
    _i("9.1.1", "9. Overheads & Contingency", "Insurance (Well Control / 3rd Party)",
       "Well control & third party insurance", "lump sum", svc=1),
    _i("9.1.2", "9. Overheads & Contingency", "Supervision (Company Staff)",
       "Company supervision — per month", "month", svc=1),
    _i("9.1.3", "9. Overheads & Contingency", "Office / Engineering Support",
       "Base support & engineering", "lump sum", svc=1),
    _i("9.1.4", "9. Overheads & Contingency", "Contingency",
       "Percentage of total (5-15% typical)", "%", qty=10.0),
]


# ============================================================================
# DATABASE
# ============================================================================

class CBSDatabase:
    """دیتابیس قیمت‌ها (Cost Breakdown Structure)"""

    def __init__(self, db_path: str = None):
        self.db_path = db_path or DEFAULT_DB
        APP_DIR.mkdir(exist_ok=True)
        self.conn = sqlite3.connect(self.db_path)
        self.conn.row_factory = sqlite3.Row
        self._create_tables()
        self._seed_if_empty()
        self._scrub_catalog()

    def _create_tables(self):
        self.conn.executescript("""
            CREATE TABLE IF NOT EXISTS cbs_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                code TEXT DEFAULT '',
                category TEXT DEFAULT '',
                name TEXT NOT NULL,
                description TEXT DEFAULT '',
                unit TEXT DEFAULT 'each',
                unit_price REAL DEFAULT 0,
                qty REAL DEFAULT 0,
                is_service INTEGER DEFAULT 0,
                source TEXT DEFAULT '',
                price_source TEXT DEFAULT 'TBD',
                active INTEGER DEFAULT 1
            );
            CREATE TABLE IF NOT EXISTS cbs_settings (
                key TEXT PRIMARY KEY,
                value TEXT
            );
            CREATE INDEX IF NOT EXISTS idx_cbs_cat ON cbs_items(category);
        """)
        self.conn.commit()

    def _seed_if_empty(self):
        cur = self.conn.execute("SELECT COUNT(*) AS c FROM cbs_items")
        if cur.fetchone()["c"] == 0:
            self.conn.executemany(
                """INSERT INTO cbs_items
                   (code, category, name, description, unit, unit_price,
                    qty, is_service, source, price_source)
                   VALUES (?,?,?,?,?,?,?,?,?,?)""",
                [(i.code, i.category, i.name, i.description, i.unit,
                  i.unit_price, i.qty, i.is_service, i.source, i.price_source)
                 for i in SEED_CATALOG])
            self.conn.commit()

    def _scrub_catalog(self):
        """جلوگیری از ورود نام چاه/شرکت به دیتابیس (هر بار اجرا می‌شود)"""
        rows = self.conn.execute(
            "SELECT id, name, description, source, price_source "
            "FROM cbs_items").fetchall()
        changed = 0
        for r in rows:
            n_name = generalize_text(r["name"])
            n_desc = generalize_text(r["description"])
            n_src = generalize_text(r["source"])
            n_psrc = generalize_text(r["price_source"])
            if (n_name, n_desc, n_src, n_psrc) != (
                    r["name"], r["description"], r["source"], r["price_source"]):
                self.conn.execute(
                    "UPDATE cbs_items SET name=?, description=?, source=?, "
                    "price_source=? WHERE id=?",
                    (n_name, n_desc, n_src, n_psrc, r["id"]))
                changed += 1
        if changed:
            self.conn.commit()

    def close(self):
        self.conn.close()

    # ---- CRUD ----

    def get_items(self, category: str = "", active_only: bool = True) -> List[CbsItem]:
        sql = "SELECT * FROM cbs_items WHERE 1=1"
        args = []
        if category:
            sql += " AND category = ?"
            args.append(category)
        if active_only:
            sql += " AND active = 1"
        sql += " ORDER BY code"
        rows = self.conn.execute(sql, args).fetchall()
        return [self._row_to_item(r) for r in rows]

    def get_categories(self) -> List[str]:
        rows = self.conn.execute(
            "SELECT DISTINCT category FROM cbs_items WHERE active=1 ORDER BY code"
        ).fetchall()
        # order by first item code of category
        cats = [r["category"] for r in rows]
        return sorted(cats, key=lambda c: self._cat_code(c))

    def _cat_code(self, category: str) -> str:
        r = self.conn.execute(
            "SELECT MIN(code) AS c FROM cbs_items WHERE category=?",
            (category,)).fetchone()
        return r["c"] or "zz"

    def save_item(self, item: CbsItem):
        if item.id:
            self.conn.execute(
                """UPDATE cbs_items SET code=?, category=?, name=?, description=?,
                   unit=?, unit_price=?, qty=?, is_service=?, source=?,
                   price_source=?, active=? WHERE id=?""",
                (item.code, item.category, item.name, item.description,
                 item.unit, item.unit_price, item.qty, item.is_service,
                 item.source, item.price_source, item.active, item.id))
        else:
            cur = self.conn.execute(
                """INSERT INTO cbs_items
                   (code, category, name, description, unit, unit_price,
                    qty, is_service, source, price_source, active)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                (item.code, item.category, item.name, item.description,
                 item.unit, item.unit_price, item.qty, item.is_service,
                 item.source, item.price_source, item.active))
            item.id = cur.lastrowid
        self.conn.commit()

    def save_items(self, items: List[CbsItem]):
        for it in items:
            self.save_item(it)

    def delete_item(self, item_id: int):
        self.conn.execute("DELETE FROM cbs_items WHERE id=?", (item_id,))
        self.conn.commit()

    def reset_to_defaults(self):
        """بازنشانی کاتالوگ به حالت پیش‌فرض (قیمت‌ها TBD)"""
        self.conn.execute("DELETE FROM cbs_items")
        self.conn.commit()
        self._seed_if_empty()

    # ---- Settings ----

    def get_setting(self, key: str, default: str = "") -> str:
        r = self.conn.execute(
            "SELECT value FROM cbs_settings WHERE key=?", (key,)).fetchone()
        return r["value"] if r else default

    def set_setting(self, key: str, value: str):
        self.conn.execute(
            "INSERT OR REPLACE INTO cbs_settings (key, value) VALUES (?,?)",
            (key, value))
        self.conn.commit()

    def get_currency(self) -> str:
        return self.get_setting("currency", "USD")

    def set_currency(self, cur: str):
        self.set_setting("currency", cur)

    # ---- Calculations ----

    @staticmethod
    def compute_totals(items: List[CbsItem]) -> Dict:
        """محاسبه جمع هر دسته و کل — contingency به‌صورت درصد"""
        cat_totals: Dict[str, float] = {}
        grand = 0.0
        contingency_pct = 0.0
        for it in items:
            amt = it.unit_price * it.qty
            if it.category == "9. Overheads & Contingency" and "contingency" in it.name.lower():
                contingency_pct = it.qty if it.qty else 0.0
                continue
            cat_totals[it.category] = cat_totals.get(it.category, 0.0) + amt
            grand += amt
        contingency_amt = grand * contingency_pct / 100.0
        return {
            "category_totals": cat_totals,
            "subtotal": grand,
            "contingency_pct": contingency_pct,
            "contingency_amt": contingency_amt,
            "total": grand + contingency_amt,
        }

    @staticmethod
    def _row_to_item(r) -> CbsItem:
        return CbsItem(
            id=r["id"], code=r["code"], category=r["category"],
            name=r["name"], description=r["description"], unit=r["unit"],
            unit_price=r["unit_price"], qty=r["qty"], is_service=r["is_service"],
            source=r["source"], price_source=r["price_source"], active=r["active"])


# ============================================================================
# AFE GENERATION (combines CBS prices with Time Breakdown durations)
# ============================================================================

def build_afe(cbs_items: List[CbsItem],
              total_days: float = 0.0,
              rig_day_rate: float = 0.0,
              spread_day_rate: float = 0.0,
              well_depth_m: float = 0.0,
              currency: str = "USD") -> Dict:
    """
    ساخت AFE کامل:
      - هزینه دکل = (دیریت دکل + spread) × مدت (از Time Breakdown)
      - هزینه اقلام CBS = قیمت × تعداد
      - جمع‌بندی دسته‌ها + contingency + هزینه به‌ازای متر
    """
    totals = CBSDatabase.compute_totals(cbs_items)
    rig_cost = (rig_day_rate + spread_day_rate) * total_days
    subtotal = totals["subtotal"] + rig_cost
    contingency_pct = totals["contingency_pct"]
    contingency_amt = subtotal * contingency_pct / 100.0
    grand = subtotal + contingency_amt

    return {
        "total_days": total_days,
        "rig_day_rate": rig_day_rate,
        "spread_day_rate": spread_day_rate,
        "rig_cost": rig_cost,
        "category_totals": totals["category_totals"],
        "items_total": totals["subtotal"],
        "subtotal": subtotal,
        "contingency_pct": contingency_pct,
        "contingency_amt": contingency_amt,
        "total": grand,
        "cost_per_m": grand / well_depth_m if well_depth_m > 0 else 0.0,
        "currency": currency,
    }


def export_afe_docx(path: str, cbs_items: List[CbsItem],
                    total_days: float, rig_day_rate: float,
                    spread_day_rate: float, well_depth_m: float,
                    well_name: str = "", operator: str = "",
                    currency: str = "USD"):
    """خروجی Word برای AFE / Cost Breakdown Structure"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from word_generator import DocColors, TableHelper

    afe = build_afe(cbs_items, total_days, rig_day_rate, spread_day_rate,
                    well_depth_m, currency)

    doc = Document()
    # Title
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COST BREAKDOWN STRUCTURE (CBS) / AFE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph("")

    if well_name:
        doc.add_paragraph(f"Well: {well_name}")
    if operator:
        doc.add_paragraph(f"Operator: {operator}")
    doc.add_paragraph(
        f"Currency: {currency}   |   Total Duration: "
        f"{total_days:.1f} days   |   Depth: "
        f"{well_depth_m:,.0f} m   |   Date: "
        f"{datetime.now().strftime('%Y-%m-%d')}")

    # Dayrate table
    if rig_day_rate > 0 or total_days > 0:
        doc.add_heading("1. Rig Cost (from Time Breakdown)", level=2)
        rig_rows = [
            ["Drilling Rig Day Rate", f"{rig_day_rate:,.2f}", f"{total_days:.1f}",
             f"{rig_day_rate * total_days:,.2f}"],
            ["Spread Cost (crew, camp, fuel, water, waste)",
             f"{spread_day_rate:,.2f}", f"{total_days:.1f}",
             f"{spread_day_rate * total_days:,.2f}"],
        ]
        TableHelper.create_professional_table(
            doc, ["Item", "Rate/Unit", "Days", f"Amount ({currency})"],
            rig_rows, caption="Rig & Dayrate Costs")

    # Items by category
    doc.add_heading("2. Goods & Services (CBS Catalog)", level=2)
    cats: Dict[str, List[CbsItem]] = {}
    for it in cbs_items:
        if it.category == "9. Overheads & Contingency":
            continue
        cats.setdefault(it.category, []).append(it)

    totals = afe["category_totals"]
    for cat in sorted(cats, key=lambda c: min(x.code for x in cats[c])):
        rows = []
        for it in cats[cat]:
            amt = it.unit_price * it.qty
            rows.append([
                it.code, it.name, it.unit,
                f"{it.unit_price:,.2f}" if it.unit_price else "TBD",
                f"{it.qty:,.2f}" if it.qty else "",
                f"{amt:,.2f}" if amt else "—",
            ])
        rows.append(["", f"Subtotal — {cat}", "", "", "",
                     f"{totals.get(cat, 0.0):,.2f}"])
        TableHelper.create_professional_table(
            doc,
            ["Code", "Item", "Unit", "Unit Price", "Qty", f"Amount ({currency})"],
            rows, caption=cat)

    # Summary
    doc.add_heading("3. Cost Summary", level=2)
    summary = [
        ["Rig & Dayrate Cost", f"{afe['rig_cost']:,.2f}"],
        ["Goods & Services Subtotal", f"{afe['items_total']:,.2f}"],
        ["Subtotal", f"{afe['subtotal']:,.2f}"],
        [f"Contingency ({afe['contingency_pct']:.0f}%)",
         f"{afe['contingency_amt']:,.2f}"],
        ["TOTAL ESTIMATED COST", f"{afe['total']:,.2f}"],
        ["Cost per Meter", f"{afe['cost_per_m']:,.2f}"],
    ]
    TableHelper.create_professional_table(
        doc, ["Description", f"Amount ({currency})"], summary)

    note = doc.add_paragraph()
    nrun = note.add_run(
        "Note: Prices are editable defaults. Confirm all unit prices against "
        "the latest price table before issuing the AFE.")
    nrun.font.size = Pt(9)
    nrun.font.italic = True

    doc.save(path)


# ============================================================================
# MARKDOWN SECTION FOR THE WIZARD OUTPUT
# ============================================================================

def build_cbs_markdown(cbs_items: List[CbsItem],
                       total_days: float = 0.0,
                       rig_day_rate: float = 0.0,
                       spread_day_rate: float = 0.0,
                       well_depth_m: float = 0.0,
                       well_name: str = "",
                       operator: str = "",
                       currency: str = "USD") -> str:
    """سکشن Cost Breakdown Structure به‌صورت Markdown برای خروجی ویزارد"""
    afe = build_afe(cbs_items, total_days, rig_day_rate, spread_day_rate,
                    well_depth_m, currency)

    lines = ["## COST BREAKDOWN STRUCTURE (CBS) / AFE", ""]
    if well_name:
        lines.append(f"**Well:** {well_name}")
    if operator:
        lines.append(f"**Operator:** {operator}")
    lines.append(f"**Currency:** {currency}  |  "
                 f"**Total Duration:** {total_days:.1f} days  |  "
                 f"**Depth:** {well_depth_m:,.0f} m")
    lines.append("")

    if rig_day_rate > 0 or total_days > 0:
        lines.append("### 1. Rig Cost (from Time Breakdown)")
        lines.append("")
        lines.append("| Item | Rate/Day | Days | Amount |")
        lines.append("|---|---:|---:|---:|")
        lines.append(f"| Drilling Rig Day Rate | {rig_day_rate:,.2f} | "
                     f"{total_days:.1f} | {rig_day_rate * total_days:,.2f} |")
        lines.append(f"| Spread Cost | {spread_day_rate:,.2f} | "
                     f"{total_days:.1f} | {spread_day_rate * total_days:,.2f} |")
        lines.append("")

    # group items by category (skip contingency row)
    cats: Dict[str, List[CbsItem]] = {}
    for it in cbs_items:
        if "contingency" in it.name.lower():
            continue
        cats.setdefault(it.category, []).append(it)

    if cats:
        lines.append("### 2. Goods & Services")
        lines.append("")
        for cat in sorted(cats, key=lambda c: min(x.code or "zz" for x in cats[c])):
            sub = sum((x.unit_price * x.qty for x in cats[cat]), 0.0)
            lines.append(f"**{cat}**")
            lines.append("")
            lines.append("| Code | Item | Unit | Unit Price | Qty | Amount |")
            lines.append("|---|---|---|---:|---:|---:|")
            for it in cats[cat]:
                amt = it.unit_price * it.qty
                price = f"{it.unit_price:,.2f}" if it.unit_price else "TBD"
                qty = f"{it.qty:,.2f}" if it.qty else ""
                amt_s = f"{amt:,.2f}" if amt else "—"
                lines.append(f"| {it.code or ''} | {it.name} | {it.unit} | "
                             f"{price} | {qty} | {amt_s} |")
            lines.append(f"| | **Subtotal** | | | | **{sub:,.2f}** |")
            lines.append("")

    lines.append("### 3. Cost Summary")
    lines.append("")
    lines.append("| Description | Amount |")
    lines.append("|---:|---:|")
    lines.append(f"| Rig & Dayrate Cost | {afe['rig_cost']:,.2f} |")
    lines.append(f"| Goods & Services Subtotal | {afe['items_total']:,.2f} |")
    lines.append(f"| Subtotal | {afe['subtotal']:,.2f} |")
    lines.append(f"| Contingency ({afe['contingency_pct']:.0f}%) | "
                 f"{afe['contingency_amt']:,.2f} |")
    lines.append(f"| **TOTAL ESTIMATED COST** | **{afe['total']:,.2f}** |")
    lines.append(f"| Cost per Meter | {afe['cost_per_m']:,.2f} |")
    lines.append("")
    lines.append("*Note: prices are editable defaults — confirm against the "
                 "latest price table before issuing the AFE.*")
    return "\n".join(lines)


# ============================================================================
# TIME BREAKDOWN LINK — read latest project from time_breakdown.db
# ============================================================================

def get_time_breakdown_summary(db_path: str = None) -> Dict:
    """
    خواندن آخرین پروژه Time Breakdown برای اتصال خودکار به AFE.
    برمی‌گرداند: name, well, total_days (آخرین cumulative), sections.
    """
    tb_path = db_path or str(Path.home() / ".drilling_program" / "time_breakdown.db")
    result = {"name": "", "well_name": "", "total_days": 0.0,
              "contingency_days": 0.0, "rows": 0, "sections": []}
    if not os.path.exists(tb_path):
        return result
    try:
        conn = sqlite3.connect(tb_path)
        conn.row_factory = sqlite3.Row
        proj = conn.execute(
            "SELECT * FROM tb_projects ORDER BY id DESC LIMIT 1").fetchone()
        if not proj:
            conn.close()
            return result
        result["name"] = proj["name"]
        result["well_name"] = proj["well_name"]
        rows = conn.execute(
            "SELECT * FROM tb_rows WHERE project_id=? AND is_contingency=0 "
            "ORDER BY row_number", (proj["id"],)).fetchall()
        crows = conn.execute(
            "SELECT * FROM tb_rows WHERE project_id=? AND is_contingency=1 "
            "ORDER BY row_number", (proj["id"],)).fetchall()
        conn.close()
        result["rows"] = len(rows)
        if rows:
            result["total_days"] = max((r["cumulative_days"] or 0) for r in rows)
        if crows:
            result["contingency_days"] = max((r["cumulative_days"] or 0) for r in crows)
        result["sections"] = [r["section_name"] for r in rows
                              if r["is_section_header"]]
    except Exception:
        pass
    return result
