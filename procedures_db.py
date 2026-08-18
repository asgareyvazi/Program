# ============================================================================
# PROCEDURES DATABASE MODULE
# File: procedures_db.py
# SQLite database for procedures, checklists, categories
# Full CRUD + Import/Export + Preview + Word Generation
# ============================================================================

import sqlite3
import json
import os
import shutil
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field, asdict
from pathlib import Path


# ============================================================================
# DATA MODELS
# ============================================================================

@dataclass
class ProcedureStep:
    """یک مرحله از پروسیجر — structured execution model (audit P1):
    هر step می‌تواند پیش‌شرط (precondition)، معیار پذیرش (acceptance) و
    نقاط Hold/Witness داشته باشد."""
    id: int = 0
    procedure_id: int = 0
    step_number: int = 0
    indent_level: int = 0  # 0=main, 1=sub, 2=sub-sub
    text: str = ""
    is_header: bool = False
    is_note: bool = False
    is_warning: bool = False
    precondition: str = ""
    acceptance: str = ""
    hold_point: bool = False
    witness_point: bool = False
    role: str = ""          # responsible role (Driller, Supervisor, ...)


@dataclass
class ChecklistItem:
    """یک آیتم چک‌لیست"""
    id: int = 0
    procedure_id: int = 0
    item_number: int = 0
    text: str = ""
    category: str = "General"


@dataclass
class ProcedureRecord:
    """رکورد کامل یک پروسیجر"""
    id: int = 0
    proc_key: str = ""
    name: str = ""
    category_id: int = 0
    category_name: str = ""
    description: str = ""
    has_checklist: bool = True
    is_builtin: bool = True
    is_active: bool = True
    created_date: str = ""
    modified_date: str = ""
    version: int = 1
    # Input definitions (JSON)
    inputs_json: str = "[]"
    # Loaded data
    steps: List[ProcedureStep] = field(default_factory=list)
    checklist: List[ChecklistItem] = field(default_factory=list)
    tags: str = ""
    # Procedure <-> Well / Risk linking (audit P1)
    linked_well_id: str = ""
    linked_section: str = ""
    linked_risk_ids: str = "[]"


@dataclass
class CategoryRecord:
    """دسته‌بندی"""
    id: int = 0
    name: str = ""
    icon: str = "📁"
    sort_order: int = 0
    is_builtin: bool = True


# ============================================================================
# DATABASE MANAGER
# ============================================================================

class ProcedureDatabase:
    """مدیریت دیتابیس پروسیجرها"""

    DB_VERSION = 1

    def __init__(self, db_path: str = None):
        if db_path is None:
            app_dir = Path.home() / ".drilling_program"
            app_dir.mkdir(exist_ok=True)
            db_path = str(app_dir / "procedures.db")

        self.db_path = db_path
        self.conn = None
        self._connect()
        self._create_tables()

        # Self-migrating schema (audit P0): bring the DB to the latest
        # version before use so old databases keep working.
        try:
            from db_migrations import migrate
            migrate(db_path)
        except Exception:
            pass

        # Seed built-in data if empty
        if self._is_empty():
            self._seed_builtin_data()

    def _connect(self):
        self.conn = sqlite3.connect(self.db_path)
        self.conn.row_factory = sqlite3.Row
        self.conn.execute("PRAGMA foreign_keys = ON")
        self.conn.execute("PRAGMA journal_mode = WAL")

    def close(self):
        if self.conn:
            self.conn.close()

    def _create_tables(self):
        self.conn.executescript("""
            CREATE TABLE IF NOT EXISTS db_meta (
                key TEXT PRIMARY KEY,
                value TEXT
            );

            CREATE TABLE IF NOT EXISTS categories (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                icon TEXT DEFAULT '📁',
                sort_order INTEGER DEFAULT 0,
                is_builtin INTEGER DEFAULT 0
            );

            CREATE TABLE IF NOT EXISTS procedures (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                proc_key TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                category_id INTEGER,
                description TEXT DEFAULT '',
                has_checklist INTEGER DEFAULT 1,
                is_builtin INTEGER DEFAULT 0,
                is_active INTEGER DEFAULT 1,
                inputs_json TEXT DEFAULT '[]',
                tags TEXT DEFAULT '',
                created_date TEXT,
                modified_date TEXT,
                version INTEGER DEFAULT 1,
                FOREIGN KEY (category_id) REFERENCES categories(id)
            );

            CREATE TABLE IF NOT EXISTS procedure_steps (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                procedure_id INTEGER NOT NULL,
                step_number INTEGER NOT NULL,
                indent_level INTEGER DEFAULT 0,
                text TEXT NOT NULL,
                is_header INTEGER DEFAULT 0,
                is_note INTEGER DEFAULT 0,
                is_warning INTEGER DEFAULT 0,
                precondition TEXT DEFAULT '',
                acceptance TEXT DEFAULT '',
                hold_point INTEGER DEFAULT 0,
                witness_point INTEGER DEFAULT 0,
                FOREIGN KEY (procedure_id) REFERENCES procedures(id)
                    ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS checklist_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                procedure_id INTEGER NOT NULL,
                item_number INTEGER NOT NULL,
                text TEXT NOT NULL,
                category TEXT DEFAULT 'General',
                FOREIGN KEY (procedure_id) REFERENCES procedures(id)
                    ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_steps_proc
                ON procedure_steps(procedure_id);
            CREATE INDEX IF NOT EXISTS idx_checklist_proc
                ON checklist_items(procedure_id);
            CREATE INDEX IF NOT EXISTS idx_proc_category
                ON procedures(category_id);
            CREATE INDEX IF NOT EXISTS idx_proc_active
                ON procedures(is_active);

            CREATE TABLE IF NOT EXISTS procedure_inputs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                procedure_id INTEGER NOT NULL,
                input_key TEXT NOT NULL,
                input_label TEXT NOT NULL,
                input_type TEXT DEFAULT 'text',
                input_options TEXT DEFAULT '',
                input_default TEXT DEFAULT '',
                input_unit TEXT DEFAULT '',
                input_order INTEGER DEFAULT 0,
                is_required INTEGER DEFAULT 0,
                FOREIGN KEY (procedure_id) REFERENCES procedures(id)
                    ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS procedure_tables (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                procedure_id INTEGER NOT NULL,
                table_name TEXT NOT NULL,
                table_columns TEXT NOT NULL,
                table_data TEXT DEFAULT '[]',
                table_order INTEGER DEFAULT 0,
                FOREIGN KEY (procedure_id) REFERENCES procedures(id)
                    ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_inputs_proc
                ON procedure_inputs(procedure_id);
            CREATE INDEX IF NOT EXISTS idx_tables_proc
                ON procedure_tables(procedure_id);
                
        """)
        self.conn.commit()

    def _is_empty(self) -> bool:
        row = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM categories").fetchone()
        return row['cnt'] == 0

    # ================================================================
    # CATEGORY CRUD
    # ================================================================

    def get_all_categories(self) -> List[CategoryRecord]:
        rows = self.conn.execute(
            "SELECT * FROM categories ORDER BY sort_order, name"
        ).fetchall()
        return [CategoryRecord(
            id=r['id'], name=r['name'], icon=r['icon'],
            sort_order=r['sort_order'], is_builtin=bool(r['is_builtin'])
        ) for r in rows]

    def add_category(self, name: str, icon: str = "📁",
                     sort_order: int = 99) -> int:
        cur = self.conn.execute(
            "INSERT INTO categories (name, icon, sort_order, is_builtin) "
            "VALUES (?, ?, ?, 0)",
            (name, icon, sort_order)
        )
        self.conn.commit()
        return cur.lastrowid

    def update_category(self, cat_id: int, name: str, icon: str):
        self.conn.execute(
            "UPDATE categories SET name=?, icon=? WHERE id=?",
            (name, icon, cat_id)
        )
        self.conn.commit()

    def delete_category(self, cat_id: int):
        # Only delete non-builtin
        self.conn.execute(
            "DELETE FROM categories WHERE id=? AND is_builtin=0",
            (cat_id,))
        self.conn.commit()

    def get_category_id(self, name: str) -> int:
        row = self.conn.execute(
            "SELECT id FROM categories WHERE name=?", (name,)
        ).fetchone()
        return row['id'] if row else 0

    # ================================================================
    # PROCEDURE CRUD
    # ================================================================

    def get_all_procedures(self, active_only: bool = True,
                            category_id: int = None) -> List[ProcedureRecord]:
        query = "SELECT p.*, c.name as cat_name FROM procedures p "
        query += "LEFT JOIN categories c ON p.category_id = c.id "
        conditions = []
        params = []

        if active_only:
            conditions.append("p.is_active = 1")
        if category_id:
            conditions.append("p.category_id = ?")
            params.append(category_id)

        if conditions:
            query += "WHERE " + " AND ".join(conditions)
        query += " ORDER BY c.sort_order, p.name"

        rows = self.conn.execute(query, params).fetchall()
        results = []
        for r in rows:
            rec = ProcedureRecord(
                id=r['id'], proc_key=r['proc_key'], name=r['name'],
                category_id=r['category_id'] or 0,
                category_name=r['cat_name'] or '',
                description=r['description'] or '',
                has_checklist=bool(r['has_checklist']),
                is_builtin=bool(r['is_builtin']),
                is_active=bool(r['is_active']),
                inputs_json=r['inputs_json'] or '[]',
                tags=r['tags'] or '',
                created_date=r['created_date'] or '',
                modified_date=r['modified_date'] or '',
                version=r['version'] or 1,
                linked_well_id=r['linked_well_id']
                if 'linked_well_id' in r.keys() else '',
                linked_section=r['linked_section']
                if 'linked_section' in r.keys() else '',
                linked_risk_ids=r['linked_risk_ids']
                if 'linked_risk_ids' in r.keys() else '[]',
            )
            results.append(rec)
        return results

    def get_procedure(self, proc_id: int) -> Optional[ProcedureRecord]:
        row = self.conn.execute(
            "SELECT p.*, c.name as cat_name FROM procedures p "
            "LEFT JOIN categories c ON p.category_id = c.id "
            "WHERE p.id = ?", (proc_id,)
        ).fetchone()
        if not row:
            return None

        rec = ProcedureRecord(
            id=row['id'], proc_key=row['proc_key'], name=row['name'],
            category_id=row['category_id'] or 0,
            category_name=row['cat_name'] or '',
            description=row['description'] or '',
            has_checklist=bool(row['has_checklist']),
            is_builtin=bool(row['is_builtin']),
            is_active=bool(row['is_active']),
            inputs_json=row['inputs_json'] or '[]',
            tags=row['tags'] or '',
            created_date=row['created_date'] or '',
            modified_date=row['modified_date'] or '',
            version=row['version'] or 1,
            linked_well_id=row['linked_well_id']
            if 'linked_well_id' in row.keys() else '',
            linked_section=row['linked_section']
            if 'linked_section' in row.keys() else '',
            linked_risk_ids=row['linked_risk_ids']
            if 'linked_risk_ids' in row.keys() else '[]',
        )

        # Load steps
        steps = self.conn.execute(
            "SELECT * FROM procedure_steps WHERE procedure_id=? "
            "ORDER BY step_number", (proc_id,)
        ).fetchall()
        rec.steps = [ProcedureStep(
            id=s['id'], procedure_id=s['procedure_id'],
            step_number=s['step_number'],
            indent_level=s['indent_level'],
            text=s['text'],
            is_header=bool(s['is_header']),
            is_note=bool(s['is_note']),
            is_warning=bool(s['is_warning']),
            precondition=s['precondition'] if 'precondition' in s.keys()
            else '',
            acceptance=s['acceptance'] if 'acceptance' in s.keys() else '',
            hold_point=bool(s['hold_point']) if 'hold_point' in s.keys()
            else False,
            witness_point=bool(s['witness_point'])
            if 'witness_point' in s.keys() else False,
            role=s['role'] if 'role' in s.keys() else '',
        ) for s in steps]

        # Load checklist
        items = self.conn.execute(
            "SELECT * FROM checklist_items WHERE procedure_id=? "
            "ORDER BY item_number", (proc_id,)
        ).fetchall()
        rec.checklist = [ChecklistItem(
            id=ci['id'], procedure_id=ci['procedure_id'],
            item_number=ci['item_number'],
            text=ci['text'], category=ci['category']
        ) for ci in items]

        # Load inputs
        inputs = self.conn.execute(
            "SELECT * FROM procedure_inputs WHERE procedure_id=? "
            "ORDER BY input_order", (proc_id,)
        ).fetchall()
        rec.inputs_json = json.dumps([dict(i) for i in inputs])
        
        return rec

    def add_procedure(self, name: str, category_id: int,
                      description: str = "",
                      has_checklist: bool = True,
                      inputs_json: str = "[]",
                      tags: str = "",
                      linked_well_id: str = "",
                      linked_section: str = "",
                      linked_risk_ids: str = "[]") -> int:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        key = name.lower().replace(' ', '_').replace('/', '_')
        key = ''.join(c for c in key if c.isalnum() or c == '_')
        # Make unique
        existing = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM procedures WHERE proc_key=?",
            (key,)).fetchone()
        if existing['cnt'] > 0:
            key = f"{key}_{int(datetime.now().timestamp())}"

        cur = self.conn.execute(
            "INSERT INTO procedures "
            "(proc_key, name, category_id, description, has_checklist, "
            "is_builtin, is_active, inputs_json, tags, "
            "created_date, modified_date, version, "
            "linked_well_id, linked_section, linked_risk_ids) "
            "VALUES (?, ?, ?, ?, ?, 0, 1, ?, ?, ?, ?, 1, ?, ?, ?)",
            (key, name, category_id, description, int(has_checklist),
             inputs_json, tags, now, now,
             linked_well_id or "", linked_section or "",
             linked_risk_ids or "[]")
        )
        self.conn.commit()
        return cur.lastrowid

    # ================================================================
    # PROCEDURE LIFECYCLE (audit P1)
    # Draft -> Technical Review -> HSE Review -> Client Review ->
    # Approved -> Released -> Superseded -> Archived
    # ================================================================

    LIFECYCLE_STATES = ["Draft", "Technical Review", "HSE Review",
                        "Client/Operator Review", "Approved", "Released",
                        "Superseded", "Archived"]

    def set_status(self, proc_id: int, status: str, user: str = "",
                   approved_by: str = "", effective_date: str = ""):
        """Move a procedure through its lifecycle; records the transition
        in the audit log."""
        if status not in self.LIFECYCLE_STATES:
            raise ValueError(f"Invalid status: {status}")
        self.conn.execute(
            "UPDATE procedures SET status=?, modified_date=? WHERE id=?",
            (status, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), proc_id))
        if approved_by:
            self.conn.execute(
                "UPDATE procedures SET approved_by=? WHERE id=?",
                (approved_by, proc_id))
        if effective_date:
            self.conn.execute(
                "UPDATE procedures SET effective_date=? WHERE id=?",
                (effective_date, proc_id))
        self.conn.commit()
        try:
            from audit_log import log_action
            log_action("procedure_status", user, f"proc#{proc_id}",
                       f"status -> {status}", "HIGH" if status in
                       ("Approved", "Released") else "INFO")
        except Exception:
            pass
        return True

    def get_lifecycle(self, proc_id: int) -> Dict:
        r = self.conn.execute(
            "SELECT id, name, status, owner, approved_by, effective_date, "
            "supersedes FROM procedures WHERE id=?",
            (proc_id,)).fetchone()
        return dict(r) if r else {}

    def supersede(self, proc_id: int, new_proc_id: int, user: str = ""):
        """Mark a procedure as superseded by another."""
        self.set_status(proc_id, "Superseded", user)
        self.conn.execute("UPDATE procedures SET supersedes=? WHERE id=?",
                          (str(new_proc_id), proc_id))
        self.conn.commit()
        return True

    def approve(self, proc_id: int, approver: str = "", effective_date: str = ""):
        return self.set_status(proc_id, "Approved", approver,
                               approved_by=approver,
                               effective_date=effective_date or
                               datetime.now().strftime("%Y-%m-%d"))

    def update_procedure(self, proc_id: int, name: str = None,
                         category_id: int = None,
                         description: str = None,
                         has_checklist: bool = None,
                         inputs_json: str = None,
                         tags: str = None,
                         linked_well_id: str = None,
                         linked_section: str = None,
                         linked_risk_ids: str = None):
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        fields = []
        params = []

        if name is not None:
            fields.append("name=?")
            params.append(name)
        if category_id is not None:
            fields.append("category_id=?")
            params.append(category_id)
        if linked_well_id is not None:
            fields.append("linked_well_id=?")
            params.append(linked_well_id)
        if linked_section is not None:
            fields.append("linked_section=?")
            params.append(linked_section)
        if linked_risk_ids is not None:
            fields.append("linked_risk_ids=?")
            params.append(linked_risk_ids)
        if description is not None:
            fields.append("description=?")
            params.append(description)
        if has_checklist is not None:
            fields.append("has_checklist=?")
            params.append(int(has_checklist))
        if inputs_json is not None:
            fields.append("inputs_json=?")
            params.append(inputs_json)
        if tags is not None:
            fields.append("tags=?")
            params.append(tags)

        fields.append("modified_date=?")
        params.append(now)
        fields.append("version = version + 1")

        params.append(proc_id)
        self.conn.execute(
            f"UPDATE procedures SET {', '.join(fields)} WHERE id=?",
            params
        )
        self.conn.commit()

    def delete_procedure(self, proc_id: int):
        # Soft delete for builtin, hard delete for user
        row = self.conn.execute(
            "SELECT is_builtin FROM procedures WHERE id=?",
            (proc_id,)).fetchone()
        if row and row['is_builtin']:
            self.conn.execute(
                "UPDATE procedures SET is_active=0 WHERE id=?",
                (proc_id,))
        else:
            self.conn.execute(
                "DELETE FROM procedures WHERE id=?", (proc_id,))
        self.conn.commit()

    def restore_procedure(self, proc_id: int):
        self.conn.execute(
            "UPDATE procedures SET is_active=1 WHERE id=?",
            (proc_id,))
        self.conn.commit()

    # ================================================================
    # PROCEDURE <-> WELL / RISK LINKING (audit P1)
    # ================================================================

    def link_well(self, proc_id: int, well_id: str = "",
                  section: str = "") -> bool:
        """Link the procedure to a canonical well (wells.db id or name)
        and an optional well section."""
        self.conn.execute(
            "UPDATE procedures SET linked_well_id=?, linked_section=? "
            "WHERE id=?",
            (well_id or "", section or "", proc_id))
        self.conn.commit()
        return True

    def link_risks(self, proc_id: int, risk_ids: List[int]) -> bool:
        """Link the procedure to drilling-problem/risk records."""
        self.conn.execute(
            "UPDATE procedures SET linked_risk_ids=? WHERE id=?",
            (json.dumps([int(i) for i in risk_ids]), proc_id))
        self.conn.commit()
        return True

    def get_links(self, proc_id: int) -> Dict:
        row = self.conn.execute(
            "SELECT linked_well_id, linked_section, linked_risk_ids "
            "FROM procedures WHERE id=?", (proc_id,)).fetchone()
        if not row:
            return {"well_id": "", "section": "", "risk_ids": []}
        try:
            risk_ids = json.loads(row['linked_risk_ids'] or "[]")
        except Exception:
            risk_ids = []
        return {"well_id": row['linked_well_id'] or "",
                "section": row['linked_section'] or "",
                "risk_ids": [int(i) for i in risk_ids]}

    def procedures_for_well(self, well_id: str) -> List[ProcedureRecord]:
        """All procedures linked to a canonical well id/name."""
        rows = self.conn.execute(
            "SELECT p.*, c.name as cat_name FROM procedures p "
            "LEFT JOIN categories c ON p.category_id = c.id "
            "WHERE p.is_active=1 AND p.linked_well_id=? "
            "ORDER BY p.name", (well_id,)).fetchall()
        out = []
        for r in rows:
            out.append(ProcedureRecord(
                id=r['id'], proc_key=r['proc_key'], name=r['name'],
                category_id=r['category_id'] or 0,
                category_name=r['cat_name'] or '',
                description=r['description'] or '',
                has_checklist=bool(r['has_checklist']),
                is_builtin=bool(r['is_builtin']),
                is_active=bool(r['is_active']),
                inputs_json=r['inputs_json'] or '[]',
                tags=r['tags'] or '',
                created_date=r['created_date'] or '',
                modified_date=r['modified_date'] or '',
                version=r['version'] or 1,
                linked_well_id=r['linked_well_id']
                if 'linked_well_id' in r.keys() else '',
                linked_section=r['linked_section']
                if 'linked_section' in r.keys() else '',
                linked_risk_ids=r['linked_risk_ids']
                if 'linked_risk_ids' in r.keys() else '[]'))
        return out

    def duplicate_procedure(self, proc_id: int) -> int:
        rec = self.get_procedure(proc_id)
        if not rec:
            return 0

        new_id = self.add_procedure(
            name=f"{rec.name} (Copy)",
            category_id=rec.category_id,
            description=rec.description,
            has_checklist=rec.has_checklist,
            inputs_json=rec.inputs_json,
            tags=rec.tags
        )

        # Copy steps
        for step in rec.steps:
            self.add_step(new_id, step.text, step.indent_level,
                          step.is_header, step.is_note, step.is_warning)

        # Copy checklist
        for item in rec.checklist:
            self.add_checklist_item(new_id, item.text, item.category)

        return new_id

    # ================================================================
    # STEPS CRUD
    # ================================================================

    def get_steps(self, proc_id: int) -> List[ProcedureStep]:
        rows = self.conn.execute(
            "SELECT * FROM procedure_steps WHERE procedure_id=? "
            "ORDER BY step_number", (proc_id,)
        ).fetchall()
        return [ProcedureStep(
            id=r['id'], procedure_id=r['procedure_id'],
            step_number=r['step_number'],
            indent_level=r['indent_level'],
            text=r['text'],
            is_header=bool(r['is_header']),
            is_note=bool(r['is_note']),
            is_warning=bool(r['is_warning']),
            precondition=r['precondition'] if 'precondition' in r.keys()
            else '',
            acceptance=r['acceptance'] if 'acceptance' in r.keys() else '',
            hold_point=bool(r['hold_point']) if 'hold_point' in r.keys()
            else False,
            witness_point=bool(r['witness_point'])
            if 'witness_point' in r.keys() else False,
            role=r['role'] if 'role' in r.keys() else '',
        ) for r in rows]

    def add_step(self, proc_id: int, text: str,
                 indent_level: int = 0,
                 is_header: bool = False,
                 is_note: bool = False,
                 is_warning: bool = False,
                 precondition: str = "",
                 acceptance: str = "",
                 hold_point: bool = False,
                 witness_point: bool = False,
                 role: str = "") -> int:
        # Get next step number
        row = self.conn.execute(
            "SELECT COALESCE(MAX(step_number), 0) + 1 as next_num "
            "FROM procedure_steps WHERE procedure_id=?",
            (proc_id,)).fetchone()
        next_num = row['next_num']

        cur = self.conn.execute(
            "INSERT INTO procedure_steps "
            "(procedure_id, step_number, indent_level, text, "
            "is_header, is_note, is_warning, precondition, acceptance, "
            "hold_point, witness_point, role) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (proc_id, next_num, indent_level, text,
             int(is_header), int(is_note), int(is_warning),
             precondition or "", acceptance or "",
             int(hold_point), int(witness_point), role or "")
        )
        self.conn.commit()
        return cur.lastrowid

    def update_step(self, step_id: int, text: str = None,
                    indent_level: int = None,
                    is_header: bool = None,
                    is_note: bool = None,
                    is_warning: bool = None,
                    precondition: str = None,
                    acceptance: str = None,
                    hold_point: bool = None,
                    witness_point: bool = None,
                    role: str = None):
        fields = []
        params = []
        if text is not None:
            fields.append("text=?")
            params.append(text)
        if indent_level is not None:
            fields.append("indent_level=?")
            params.append(indent_level)
        if is_header is not None:
            fields.append("is_header=?")
            params.append(int(is_header))
        if is_note is not None:
            fields.append("is_note=?")
            params.append(int(is_note))
        if is_warning is not None:
            fields.append("is_warning=?")
            params.append(int(is_warning))
        if precondition is not None:
            fields.append("precondition=?")
            params.append(precondition)
        if acceptance is not None:
            fields.append("acceptance=?")
            params.append(acceptance)
        if hold_point is not None:
            fields.append("hold_point=?")
            params.append(int(hold_point))
        if witness_point is not None:
            fields.append("witness_point=?")
            params.append(int(witness_point))
        if role is not None:
            fields.append("role=?")
            params.append(role)

        if fields:
            params.append(step_id)
            self.conn.execute(
                f"UPDATE procedure_steps SET {', '.join(fields)} "
                f"WHERE id=?", params)
            self.conn.commit()

    def delete_step(self, step_id: int):
        self.conn.execute(
            "DELETE FROM procedure_steps WHERE id=?", (step_id,))
        self.conn.commit()

    def reorder_steps(self, proc_id: int, step_ids: List[int]):
        for idx, step_id in enumerate(step_ids):
            self.conn.execute(
                "UPDATE procedure_steps SET step_number=? "
                "WHERE id=? AND procedure_id=?",
                (idx + 1, step_id, proc_id))
        self.conn.commit()

    def replace_all_steps(self, proc_id: int,
                           steps: List[Dict]):
        self.conn.execute(
            "DELETE FROM procedure_steps WHERE procedure_id=?",
            (proc_id,))
        for idx, step in enumerate(steps):
            self.conn.execute(
                "INSERT INTO procedure_steps "
                "(procedure_id, step_number, indent_level, text, "
                "is_header, is_note, is_warning, precondition, acceptance, "
                "hold_point, witness_point, role) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (proc_id, idx + 1,
                 step.get('indent', 0),
                 step.get('text', ''),
                 int(step.get('is_header', False)),
                 int(step.get('is_note', False)),
                 int(step.get('is_warning', False)),
                 step.get('precondition', '') or '',
                 step.get('acceptance', '') or '',
                 int(step.get('hold_point', False)),
                 int(step.get('witness_point', False)),
                 step.get('role', '') or '')
            )
        self.conn.commit()

    # ================================================================
    # CHECKLIST CRUD
    # ================================================================

    def get_checklist(self, proc_id: int) -> List[ChecklistItem]:
        rows = self.conn.execute(
            "SELECT * FROM checklist_items WHERE procedure_id=? "
            "ORDER BY item_number", (proc_id,)
        ).fetchall()
        return [ChecklistItem(
            id=r['id'], procedure_id=r['procedure_id'],
            item_number=r['item_number'],
            text=r['text'], category=r['category']
        ) for r in rows]

    def add_checklist_item(self, proc_id: int, text: str,
                            category: str = "General") -> int:
        row = self.conn.execute(
            "SELECT COALESCE(MAX(item_number), 0) + 1 as next_num "
            "FROM checklist_items WHERE procedure_id=?",
            (proc_id,)).fetchone()
        next_num = row['next_num']

        cur = self.conn.execute(
            "INSERT INTO checklist_items "
            "(procedure_id, item_number, text, category) "
            "VALUES (?, ?, ?, ?)",
            (proc_id, next_num, text, category)
        )
        self.conn.commit()
        return cur.lastrowid

    def update_checklist_item(self, item_id: int, text: str = None,
                               category: str = None):
        fields = []
        params = []
        if text is not None:
            fields.append("text=?")
            params.append(text)
        if category is not None:
            fields.append("category=?")
            params.append(category)
        if fields:
            params.append(item_id)
            self.conn.execute(
                f"UPDATE checklist_items SET {', '.join(fields)} "
                f"WHERE id=?", params)
            self.conn.commit()

    def delete_checklist_item(self, item_id: int):
        self.conn.execute(
            "DELETE FROM checklist_items WHERE id=?", (item_id,))
        self.conn.commit()

    def replace_all_checklist(self, proc_id: int,
                               items: List[Dict]):
        self.conn.execute(
            "DELETE FROM checklist_items WHERE procedure_id=?",
            (proc_id,))
        for idx, item in enumerate(items):
            self.conn.execute(
                "INSERT INTO checklist_items "
                "(procedure_id, item_number, text, category) "
                "VALUES (?, ?, ?, ?)",
                (proc_id, idx + 1,
                 item.get('text', ''),
                 item.get('category', 'General'))
            )
        self.conn.commit()

    # ================================================================
    # IMPORT / EXPORT
    # ================================================================

    def export_procedures(self, proc_ids: List[int],
                           file_path: str):
        data = {
            'version': self.DB_VERSION,
            'exported': datetime.now().isoformat(),
            'procedures': []
        }

        for pid in proc_ids:
            rec = self.get_procedure(pid)
            if not rec:
                continue

            proc_data = {
                'name': rec.name,
                'category': rec.category_name,
                'description': rec.description,
                'has_checklist': rec.has_checklist,
                'inputs_json': rec.inputs_json,
                'tags': rec.tags,
                'steps': [
                    {'text': s.text, 'indent': s.indent_level,
                     'is_header': s.is_header,
                     'is_note': s.is_note,
                     'is_warning': s.is_warning}
                    for s in rec.steps
                ],
                'checklist': [
                    {'text': c.text, 'category': c.category}
                    for c in rec.checklist
                ]
            }
            data['procedures'].append(proc_data)

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

    def import_procedures(self, file_path: str) -> int:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        imported = 0
        for proc_data in data.get('procedures', []):
            cat_name = proc_data.get('category', 'Imported')
            cat_id = self.get_category_id(cat_name)
            if not cat_id:
                cat_id = self.add_category(cat_name, "📥")

            proc_id = self.add_procedure(
                name=proc_data.get('name', 'Imported Procedure'),
                category_id=cat_id,
                description=proc_data.get('description', ''),
                has_checklist=proc_data.get('has_checklist', True),
                inputs_json=proc_data.get('inputs_json', '[]'),
                tags=proc_data.get('tags', 'imported')
            )

            # Add steps
            for step in proc_data.get('steps', []):
                self.add_step(
                    proc_id, step.get('text', ''),
                    step.get('indent', 0),
                    step.get('is_header', False),
                    step.get('is_note', False),
                    step.get('is_warning', False)
                )

            # Add checklist
            for item in proc_data.get('checklist', []):
                self.add_checklist_item(
                    proc_id, item.get('text', ''),
                    item.get('category', 'General')
                )

            imported += 1

        return imported

    # ================================================================
    # SEARCH
    # ================================================================

    def search_procedures(self, query: str) -> List[ProcedureRecord]:
        q = f"%{query}%"
        rows = self.conn.execute(
            "SELECT p.*, c.name as cat_name FROM procedures p "
            "LEFT JOIN categories c ON p.category_id = c.id "
            "WHERE p.is_active=1 AND "
            "(p.name LIKE ? OR p.description LIKE ? OR p.tags LIKE ?) "
            "ORDER BY p.name",
            (q, q, q)
        ).fetchall()
        return [ProcedureRecord(
            id=r['id'], proc_key=r['proc_key'], name=r['name'],
            category_id=r['category_id'] or 0,
            category_name=r['cat_name'] or '',
            description=r['description'] or '',
            has_checklist=bool(r['has_checklist']),
            is_builtin=bool(r['is_builtin']),
            is_active=bool(r['is_active']),
            tags=r['tags'] or '',
        ) for r in rows]

    # ================================================================
    # STATISTICS
    # ================================================================

    def get_stats(self) -> Dict:
        total = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM procedures WHERE is_active=1"
        ).fetchone()['cnt']
        builtin = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM procedures "
            "WHERE is_active=1 AND is_builtin=1"
        ).fetchone()['cnt']
        custom = total - builtin
        cats = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM categories"
        ).fetchone()['cnt']
        steps = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM procedure_steps"
        ).fetchone()['cnt']
        items = self.conn.execute(
            "SELECT COUNT(*) as cnt FROM checklist_items"
        ).fetchone()['cnt']

        return {
            'total_procedures': total,
            'builtin': builtin,
            'custom': custom,
            'categories': cats,
            'total_steps': steps,
            'total_checklist_items': items,
        }

    # ================================================================
    # BACKUP & RESTORE
    # ================================================================

    def backup(self, backup_path: str = None):
        if backup_path is None:
            backup_dir = Path(self.db_path).parent / "backups"
            backup_dir.mkdir(exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_path = str(backup_dir / f"procedures_backup_{ts}.db")

        self.conn.commit()
        shutil.copy2(self.db_path, backup_path)
        return backup_path

    def reset_to_builtin(self):
        """ریست به حالت پیش‌فرض"""
        self.conn.executescript("""
            DELETE FROM checklist_items;
            DELETE FROM procedure_steps;
            DELETE FROM procedures;
            DELETE FROM categories;
        """)
        self.conn.commit()
        self._seed_builtin_data()

    # ================================================================
    # SEED BUILT-IN DATA
    # ================================================================

    def _seed_builtin_data(self):
        """درج داده‌های پیش‌فرض"""
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Categories
        categories = [
            ("Pre-Spud & Rig-Up", "🏗️", 1),
            ("Drilling Operations", "⚙️", 2),
            ("Tripping Operations", "🔄", 3),
            ("Casing & Liner", "🔧", 4),
            ("Cementing", "🧱", 5),
            ("Directional Drilling", "🧭", 6),
            ("Drilling Fluids", "🧪", 7),
            ("Formation Evaluation", "📊", 8),
            ("Well Control", "🔴", 9),
            ("BOP Operations", "🛑", 10),
            ("Fishing & Remedial", "🎣", 11),
            ("Completion", "✅", 12),
            ("Abandonment (P&A)", "🚫", 13),
            ("HSE & Emergency", "⚠️", 14),
            ("Special Operations", "⭐", 15),
        ]

        for name, icon, order in categories:
            self.conn.execute(
                "INSERT INTO categories (name, icon, sort_order, is_builtin) "
                "VALUES (?, ?, ?, 1)",
                (name, icon, order))

        self.conn.commit()

        # ---- SEED PROCEDURES ----
        # Pre-Spud
        cat_id = self.get_category_id("Pre-Spud & Rig-Up")
        self._seed_procedure(cat_id, "pre_spud", "Pre-Spud Checklist & Meeting",
            "Complete pre-spud verification and meeting procedure", now,
            steps=[
                (0, True, "PRE-SPUD CHECKLIST AND PROCEDURE"),
                (0, True, "1. DOCUMENTATION VERIFICATION"),
                (1, False, "1.1. Verify all permits and licenses are in place."),
                (1, False, "1.2. Confirm approved Drilling Program on site."),
                (1, False, "1.3. Verify Environmental Impact Assessment (EIA) approval."),
                (1, False, "1.4. Confirm all regulatory notifications submitted."),
                (1, False, "1.5. Verify insurance certificates are current."),
                (1, False, "1.6. Review all contractor certifications."),
                (0, True, "2. LOCATION PREPARATION"),
                (1, False, "2.1. Verify well location coordinates with survey."),
                (1, False, "2.2. Confirm cellar/conductor pit condition."),
                (1, False, "2.3. Verify access road condition and load capacity."),
                (1, False, "2.4. Confirm drainage system and containment berms."),
                (1, False, "2.5. Verify water supply availability."),
                (1, False, "2.6. Confirm waste disposal arrangements."),
                (0, True, "3. RIG ACCEPTANCE"),
                (1, False, "3.1. Complete rig acceptance inspection."),
                (1, False, "3.2. Verify all rig certification documents current."),
                (1, False, "3.3. Confirm third-party inspections (derrick, drawworks, BOP, pumps, cranes)."),
                (1, False, "3.4. Verify maximum hook load capacity."),
                (1, False, "3.5. Conduct ton-mile calculations for drilling line."),
                (0, True, "4. WELL CONTROL EQUIPMENT"),
                (1, False, "4.1. Verify BOP stack configuration matches program."),
                (1, False, "4.2. Confirm BOP rams correct size."),
                (1, False, "4.3. Verify accumulator capacity and precharge pressure."),
                (1, False, "4.4. Function test all BOP rams and annular."),
                (1, False, "4.5. Verify choke and kill manifold integrity."),
                (1, False, "4.6. Verify diverter system installation and function."),
                (1, False, "4.7. Calibrate pit volume totalizer and flow show."),
                (1, False, "4.8. Calibrate gas detection system."),
                (0, True, "5. MUD SYSTEM"),
                (1, False, "5.1. Verify active pit volume and reserve capacity."),
                (1, False, "5.2. Confirm all solids control equipment operational."),
                (1, False, "5.3. Verify mud products inventory per program."),
                (1, False, "5.4. Calibrate pit level indicators."),
                (0, True, "6. SAFETY SYSTEMS"),
                (1, False, "6.1. Verify fire detection and fighting equipment."),
                (1, False, "6.2. Test emergency shutdown (ESD) system."),
                (1, False, "6.3. Verify gas detection system calibrated."),
                (1, False, "6.4. Verify PA system and emergency alarms."),
                (1, False, "6.5. Confirm medical equipment and first aider."),
                (1, False, "6.6. Verify emergency evacuation plan and muster stations."),
                (0, True, "7. PRE-SPUD MEETING"),
                (1, False, "7.1. Conduct Pre-Spud Meeting with ALL crew members."),
                (1, False, "7.2. Review Drilling Program highlights."),
                (1, False, "7.3. Review HSE plan and emergency procedures."),
                (1, False, "7.4. Review well control procedures and responsibilities."),
                (1, False, "7.5. Discuss lessons learned from offset wells."),
                (1, False, "7.6. Ensure all crew signed attendance sheet."),
                (1, False, "7.7. Obtain authorization to spud from Company Representative."),
            ],
            checklist=[
                ("Approved Drilling Program on site", "Documentation"),
                ("All permits and licenses verified", "Documentation"),
                ("Environmental Impact Assessment approved", "Documentation"),
                ("Well Control Manual on site", "Documentation"),
                ("Emergency Response Plan on site", "Documentation"),
                ("Well coordinates verified (survey)", "Location"),
                ("Cellar/conductor pit condition OK", "Location"),
                ("Access road verified", "Location"),
                ("Rig acceptance inspection complete", "Rig"),
                ("Derrick/Mast certification current", "Rig"),
                ("Drawworks & brake inspected", "Rig"),
                ("Top drive inspected", "Rig"),
                ("Mud pumps tested", "Rig"),
                ("BOP stack inspected and correct size rams", "Well Control"),
                ("Accumulator precharge OK", "Well Control"),
                ("Choke/kill manifold tested", "Well Control"),
                ("Diverter system tested", "Well Control"),
                ("PVT and flow show calibrated", "Well Control"),
                ("Gas detection calibrated", "Well Control"),
                ("Shale shakers operational", "Mud System"),
                ("Degasser functional", "Mud System"),
                ("Mud products inventory verified", "Mud System"),
                ("Fire equipment inspected", "Safety"),
                ("ESD system tested", "Safety"),
                ("PA system tested", "Safety"),
                ("Medical equipment verified", "Safety"),
                ("Emergency contacts posted", "Safety"),
                ("All BHA on location", "Materials"),
                ("Casing inventory verified", "Materials"),
                ("Cement and additives verified", "Materials"),
                ("Pre-Spud Meeting conducted", "Meeting"),
                ("All crew signed attendance", "Meeting"),
                ("Authorization to spud obtained", "Meeting"),
            ]
        )

        # Drilling Operations
        cat_id = self.get_category_id("Drilling Operations")
        self._seed_procedure(cat_id, "drilling_ahead", "Drilling Ahead (Per Section)",
            "Detailed drilling procedure for each hole section", now,
            steps=[
                (0, True, "DRILLING PROCEDURE"),
                (0, True, "1. PRE-DRILLING PREPARATION"),
                (1, False, "1.1. Review section program with all crews."),
                (1, False, "1.2. Conduct pre-section toolbox talk / safety meeting."),
                (1, False, "1.3. Prepare mud system per program."),
                (1, False, "1.4. Verify BHA components and bit."),
                (1, False, "1.5. Verify MWD/LWD tools initialized."),
                (1, False, "1.6. Verify solids control equipment operational."),
                (1, False, "1.7. Verify mud logging unit is recording."),
                (0, True, "2. BHA MAKE-UP"),
                (1, False, "2.1. Make up bit - record serial number and nozzles."),
                (1, False, "2.2. Make up BHA per BHA diagram."),
                (1, False, "2.3. Verify all connections made up to recommended torque."),
                (1, False, "2.4. Perform MWD surface check / flow check."),
                (0, True, "3. RUN IN HOLE"),
                (1, False, "3.1. RIH to previous shoe with pumps on."),
                (1, False, "3.2. Record all tight spots."),
                (1, False, "3.3. Tag cement/shoe track and record depth."),
                (1, False, "3.4. Drill out cement and float equipment."),
                (0, True, "4. LOT / FIT"),
                (1, False, "4.1. Perform LOT/FIT at shoe depth."),
                (1, False, "4.2. Record result and calculate EMW."),
                (1, False, "4.3. Report to office and obtain approval."),
                (0, True, "5. DRILLING AHEAD"),
                (1, False, "5.1. Drill hole per program parameters (WOB, RPM, Flow)."),
                (1, False, "5.2. Monitor and record all drilling parameters."),
                (1, False, "5.3. Take surveys per program."),
                (1, False, "5.4. Monitor ECD - do not exceed fracture gradient."),
                (1, False, "5.5. Perform connection gas / background gas monitoring."),
                (1, False, "5.6. Perform wiper trips as required."),
                (0, True, "6. HOLE CONDITION MONITORING"),
                (1, False, "6.1. Monitor for stuck pipe indicators (torque, drag, fill)."),
                (1, False, "6.2. Monitor for lost circulation indicators."),
                (1, False, "6.3. Monitor for kick indicators (pit gain, flow, drilling break)."),
                (1, False, "6.4. Monitor for wellbore stability (cavings, tight hole)."),
                (0, True, "7. REACHING SECTION TD"),
                (1, False, "7.1. Circulate minimum 2x bottoms up."),
                (1, False, "7.2. Take final survey at TD."),
                (1, False, "7.3. Perform short trip to shoe."),
                (1, False, "7.4. Condition mud for casing/logging."),
                (0, True, "8. POOH & BIT GRADING"),
                (1, False, "8.1. POOH pumping out of hole."),
                (1, False, "8.2. Keep hole full at all times."),
                (1, False, "8.3. Grade bit per IADC system."),
                (1, False, "8.4. Lay down and inspect BHA."),
            ],
            checklist=[
                ("Section program reviewed with crew", "Preparation"),
                ("Pre-section safety meeting conducted", "Preparation"),
                ("BHA components verified", "Preparation"),
                ("Bit inspected - S/N recorded", "Preparation"),
                ("MWD/LWD tools initialized", "Preparation"),
                ("Mud system ready per program", "Mud"),
                ("LOT/FIT result meets requirements", "Verification"),
                ("LOT/FIT reported to office", "Verification"),
                ("Surveys taken per program", "Drilling"),
                ("ECD monitored continuously", "Drilling"),
                ("Hole cleaned - 2x bottoms up", "TD Operations"),
                ("Final survey at TD recorded", "TD Operations"),
                ("Short trip completed", "TD Operations"),
                ("Mud conditioned for next operation", "TD Operations"),
                ("Bit graded per IADC", "POOH"),
                ("BHA inspected and documented", "POOH"),
            ]
        )

        # Well Control
        cat_id = self.get_category_id("Well Control")
        self._seed_procedure(cat_id, "shut_in_drilling",
            "Shut-In While Drilling (Hard Shut-In)",
            "Hard shut-in procedure during drilling operations", now,
            steps=[
                (0, True, "HARD SHUT-IN PROCEDURE (WHILE DRILLING)"),
                (0, False, ""),
                (0, True, "UPON POSITIVE KICK INDICATION:"),
                (0, False, ""),
                (0, True, "1. DRILLER - IMMEDIATE ACTIONS"),
                (1, False, "1.1. Pick up off bottom - space out to tool joint above rotary."),
                (1, False, "1.2. STOP PUMPS immediately."),
                (1, False, "1.3. Close HCR valve on choke line."),
                (1, False, "1.4. Open remote choke line valve."),
                (0, False, ""),
                (0, True, "2. CONFIRM KICK"),
                (1, False, "2.1. Observe well - confirm flow."),
                (1, False, "2.2. If well is flowing:"),
                (2, False, "a) CLOSE PIPE RAMS (preferred) or ANNULAR."),
                (2, False, "b) Notify Toolpusher and Company Man."),
                (2, False, "c) Record TIME of shut-in."),
                (0, False, ""),
                (0, True, "3. RECORD DATA"),
                (1, False, "3.1. Allow pressures to stabilize (max 10 minutes)."),
                (1, False, "3.2. Record stabilized SIDPP (Shut-In Drill Pipe Pressure)."),
                (1, False, "3.3. Record stabilized SICP (Shut-In Casing Pressure)."),
                (1, False, "3.4. Record pit gain volume."),
                (1, False, "3.5. Record time and depth."),
                (0, False, ""),
                (0, True, "4. NOTIFICATIONS"),
                (1, False, "4.1. Notify Toolpusher immediately."),
                (1, False, "4.2. Notify Company Man."),
                (1, False, "4.3. Notify office (Drilling Superintendent)."),
                (0, False, ""),
                (0, True, "5. PREPARE FOR KILL"),
                (1, False, "5.1. Calculate kill mud weight:"),
                (2, False, "Kill MW = Current MW + (SIDPP / (0.052 × TVD))"),
                (1, False, "5.2. Calculate ICP and FCP."),
                (1, False, "5.3. Prepare kill sheet."),
                (1, False, "5.4. Prepare pressure schedule."),
                (1, False, "5.5. Begin mixing kill weight mud."),
                (1, False, "5.6. Select kill method (Driller's or W&W)."),
                (0, False, ""),
                (0, True, "⚠️ CRITICAL REMINDERS"),
                (1, False, "• Do NOT move pipe while shut-in (unless stripping)."),
                (1, False, "• Do NOT attempt to circulate without proper kill procedure."),
                (1, False, "• Monitor pressures continuously."),
                (1, False, "• Record all data on kill sheet."),
            ],
            checklist=[
                ("Picked up off bottom to tool joint", "Shut-In"),
                ("Pumps stopped", "Shut-In"),
                ("HCR valve closed", "Shut-In"),
                ("Remote choke valve opened", "Shut-In"),
                ("Well flow confirmed", "Shut-In"),
                ("Pipe rams (or annular) CLOSED", "Shut-In"),
                ("Time of shut-in recorded", "Data"),
                ("SIDPP recorded (stabilized)", "Data"),
                ("SICP recorded (stabilized)", "Data"),
                ("Pit gain volume recorded", "Data"),
                ("Current depth recorded", "Data"),
                ("Current mud weight recorded", "Data"),
                ("Toolpusher notified", "Notifications"),
                ("Company Man notified", "Notifications"),
                ("Office notified", "Notifications"),
                ("Kill mud weight calculated", "Kill Preparation"),
                ("ICP calculated", "Kill Preparation"),
                ("FCP calculated", "Kill Preparation"),
                ("Kill sheet completed", "Kill Preparation"),
                ("Pressure schedule prepared", "Kill Preparation"),
                ("Kill mud mixing started", "Kill Preparation"),
                ("Kill method selected and agreed", "Kill Preparation"),
            ]
        )

        # BOP Operations
        cat_id = self.get_category_id("BOP Operations")
        self._seed_procedure(cat_id, "bop_pressure_test",
            "BOP Pressure Test (API RP 53)",
            "Low and high pressure BOP testing per API RP 53", now,
            steps=[
                (0, True, "BOP PRESSURE TEST PROCEDURE (API RP 53)"),
                (0, False, ""),
                (0, True, "1. TEST REQUIREMENTS"),
                (1, False, "• After initial nipple-up on each casing string"),
                (1, False, "• Before drilling out under new casing shoe"),
                (1, False, "• After any BOP repair or ram change"),
                (1, False, "• After disconnection of pressure-containing component"),
                (1, False, "• Every 14 days (or per regulation)"),
                (0, False, ""),
                (0, True, "2. PREPARATION"),
                (1, False, "2.1. Notify all personnel of BOP test."),
                (1, False, "2.2. Clear non-essential personnel from BOP area."),
                (1, False, "2.3. Verify test pump rigged up and calibrated."),
                (1, False, "2.4. Prepare chart recorder or data acquisition."),
                (1, False, "2.5. Record current mud weight."),
                (0, False, ""),
                (0, True, "3. TEST SEQUENCE"),
                (0, False, ""),
                (1, True, "A. PIPE RAMS TEST"),
                (1, False, "3.1. RIH drill pipe / test mandrel into stack."),
                (1, False, "3.2. Close pipe rams on drill pipe."),
                (1, False, "3.3. Apply LOW test pressure (250 psi) - hold 5 minutes."),
                (1, False, "3.4. Record: start pressure, end pressure, volume."),
                (1, False, "3.5. Increase to HIGH test pressure - hold 10 minutes."),
                (1, False, "3.6. Record results. Bleed off. Open rams."),
                (0, False, ""),
                (1, True, "B. BLIND/SHEAR RAMS TEST"),
                (1, False, "3.7. Remove drill pipe from stack bore."),
                (1, False, "3.8. Close blind/shear rams."),
                (1, False, "3.9. Test LOW (250 psi / 5 min) then HIGH (10 min)."),
                (1, False, "3.10. Record results. Bleed off. Open rams."),
                (0, False, ""),
                (1, True, "C. ANNULAR PREVENTER TEST"),
                (1, False, "3.11. RIH drill pipe into stack."),
                (1, False, "3.12. Close annular on drill pipe."),
                (1, False, "3.13. Test LOW (250 psi / 5 min)."),
                (1, False, "3.14. Test at 70% of annular WP (10 min)."),
                (1, False, "3.15. Record results. Open annular."),
                (0, False, ""),
                (1, True, "D. KILL/CHOKE LINE TEST"),
                (1, False, "3.16. Test kill line to high pressure."),
                (1, False, "3.17. Test choke line and manifold valves."),
                (1, False, "3.18. Record all results."),
                (0, False, ""),
                (0, True, "4. DOCUMENTATION"),
                (1, False, "4.1. Complete BOP test report form."),
                (1, False, "4.2. Record serial numbers of all components."),
                (1, False, "4.3. Attach pressure charts."),
                (1, False, "4.4. Obtain Toolpusher and Company Man sign-off."),
                (0, False, ""),
                (0, True, "5. FAILURE PROCEDURE"),
                (1, False, "5.1. If ANY component fails test:"),
                (2, False, "a) Stop test. Identify failed component."),
                (2, False, "b) Repair or replace."),
                (2, False, "c) Retest failed component."),
                (2, False, "d) Document failure and corrective action."),
                (1, False, "5.2. Do NOT proceed until ALL tests PASS."),
            ],
            checklist=[
                ("All personnel notified of BOP test", "Preparation"),
                ("Non-essential personnel cleared", "Safety"),
                ("Test pump rigged up and calibrated", "Equipment"),
                ("Chart recorder ready", "Equipment"),
                ("Mud weight recorded", "Data"),
                ("Pipe rams - LOW test PASS/FAIL", "Rams"),
                ("Pipe rams - HIGH test PASS/FAIL", "Rams"),
                ("Blind/Shear rams - LOW test PASS/FAIL", "Rams"),
                ("Blind/Shear rams - HIGH test PASS/FAIL", "Rams"),
                ("Annular - LOW test PASS/FAIL", "Annular"),
                ("Annular - 70% WP test PASS/FAIL", "Annular"),
                ("Kill line tested PASS/FAIL", "Lines"),
                ("Choke line tested PASS/FAIL", "Lines"),
                ("Choke manifold valves tested PASS/FAIL", "Manifold"),
                ("BOP test report completed", "Documentation"),
                ("Pressure charts attached", "Documentation"),
                ("Serial numbers recorded", "Documentation"),
                ("Toolpusher sign-off obtained", "Approval"),
                ("Company Man sign-off obtained", "Approval"),
            ]
        )

        # Casing
        cat_id = self.get_category_id("Casing & Liner")
        self._seed_procedure(cat_id, "casing_running", "Running Casing",
            "Complete casing running procedure including accessories", now,
            steps=[
                (0, True, "CASING RUNNING PROCEDURE"),
                (0, True, "1. PRE-RUN PREPARATION"),
                (1, False, "1.1. Verify casing on rack per tally."),
                (1, False, "1.2. Confirm ALL casing drifted (100%)."),
                (1, False, "1.3. Inspect threads on each joint."),
                (1, False, "1.4. Rack casing in running order."),
                (1, False, "1.5. Verify casing accessories on location."),
                (1, False, "1.6. Verify hole condition - wiper trip completed."),
                (1, False, "1.7. Condition mud (YP < 15, Gel 10s < 8, FL < 5)."),
                (1, False, "1.8. Verify fill-up line and rate established."),
                (1, False, "1.9. Prepare thread compound."),
                (1, False, "1.10. Verify cement unit rigged up and tested."),
                (1, False, "1.11. Conduct pre-job safety meeting."),
                (0, True, "2. RUNNING CASING"),
                (1, False, "2.1. Make up float shoe on first joint."),
                (1, False, "2.2. Apply thread compound per specs."),
                (1, False, "2.3. Make up to recommended torque."),
                (1, False, "2.4. RIH and fill with mud."),
                (1, False, "2.5. Install float collar per program."),
                (1, False, "2.6. Continue running - torque each joint."),
                (1, False, "2.7. Install centralizers per program."),
                (1, False, "2.8. Fill casing every 5-10 joints."),
                (1, False, "2.9. Maximum running speed: 3 ft/s."),
                (1, False, "2.10. Record tally and verify depth."),
                (0, True, "3. LANDING CASING"),
                (1, False, "3.1. Land casing at programmed depth."),
                (1, False, "3.2. Verify depth with tally."),
                (1, False, "3.3. Establish circulation."),
                (1, False, "3.4. Circulate minimum 1.5x annular volume."),
                (1, False, "3.5. Record pickup/slackoff/rotating weights."),
            ],
            checklist=[
                ("All casing drifted (100%)", "Preparation"),
                ("Threads inspected", "Preparation"),
                ("Casing racked in running order", "Preparation"),
                ("Float shoe inspected", "Accessories"),
                ("Float collar inspected", "Accessories"),
                ("Centralizers correct size and quantity", "Accessories"),
                ("Stop collars on location", "Accessories"),
                ("Wiper trip completed - no tight spots", "Hole Condition"),
                ("Mud conditioned (YP<15, Gel<8, FL<5)", "Hole Condition"),
                ("Thread compound prepared", "Running"),
                ("Fill-up line tested", "Running"),
                ("Torque gauge calibrated", "Running"),
                ("Pre-job safety meeting conducted", "Safety"),
                ("Cement unit rigged up and tested", "Cementing"),
                ("Tally sheet complete", "QC"),
                ("Casing landed at correct depth", "QC"),
                ("Circulation established - full returns", "QC"),
                ("Weights recorded (PU/SO/RW)", "QC"),
            ]
        )

        # HSE
        cat_id = self.get_category_id("HSE & Emergency")
        self._seed_procedure(cat_id, "h2s_safety",
            "H₂S Safety & Emergency Response",
            "H₂S detection, alarm levels, evacuation, rescue", now,
            steps=[
                (0, True, "H₂S SAFETY AND EMERGENCY RESPONSE"),
                (0, True, "1. EXPOSURE LIMITS"),
                (1, False, "• 10 ppm: ALERT - 8-hr TWA"),
                (1, False, "• 20 ppm: ALARM - Don SCBA"),
                (1, False, "• 50 ppm: DANGER - Evacuate"),
                (1, False, "• 100 ppm: IDLH - Full Emergency"),
                (1, False, "• 500+ ppm: DEATH within minutes"),
                (0, True, "2. PREPARATION"),
                (1, False, "2.1. All personnel H₂S certified."),
                (1, False, "2.2. SCBA at all muster points and rig floor."),
                (1, False, "2.3. Fixed H₂S detectors at all critical locations."),
                (1, False, "2.4. Personal monitors for ALL personnel."),
                (1, False, "2.5. Wind socks at 4 corners."),
                (1, False, "2.6. Escape routes marked (min 2 routes)."),
                (1, False, "2.7. Buddy system enforced."),
                (1, False, "2.8. All materials NACE MR-0175 compliant."),
                (0, True, "3. ALARM RESPONSE"),
                (1, True, "10 ppm - ALERT"),
                (1, False, "3.1. Notify all personnel. SCBA accessible."),
                (1, False, "3.2. Check wind direction. Monitor closely."),
                (1, True, "20 ppm - ALARM"),
                (1, False, "3.3. Sound alarm. Non-essential evacuate."),
                (1, False, "3.4. Essential personnel don SCBA."),
                (1, True, "50 ppm - DANGER"),
                (1, False, "3.5. Full evacuation to upwind muster."),
                (1, False, "3.6. Account for ALL personnel."),
                (1, True, "100+ ppm - EMERGENCY (IDLH)"),
                (1, False, "3.7. Full emergency response."),
                (1, False, "3.8. Contact emergency services."),
                (0, True, "4. RESCUE"),
                (1, False, "⚠️ NEVER attempt rescue WITHOUT SCBA!"),
                (1, False, "4.1. Approach from UPWIND."),
                (1, False, "4.2. Remove victim to fresh air."),
                (1, False, "4.3. Check breathing - begin CPR if needed."),
                (1, False, "4.4. Administer 100% oxygen."),
                (1, False, "4.5. Transport to medical facility."),
            ],
            checklist=[
                ("H₂S contingency plan on site", "Documentation"),
                ("All personnel H₂S certified", "Training"),
                ("H₂S safety briefing conducted", "Training"),
                ("H₂S drill/practice evacuation done", "Training"),
                ("Fixed detectors installed and calibrated", "Detection"),
                ("Personal monitors issued to ALL", "Detection"),
                ("SCBA at rig floor (min 2 sets)", "Equipment"),
                ("SCBA at every muster point", "Equipment"),
                ("All SCBA bottles full and tested", "Equipment"),
                ("Wind socks installed (4 corners)", "Monitoring"),
                ("Escape routes marked (min 2)", "Evacuation"),
                ("Muster points identified", "Evacuation"),
                ("Buddy system enforced", "Personnel"),
                ("All materials NACE compliant", "Materials"),
                ("ZnO scavenger in mud system", "Mud"),
                ("Emergency contacts posted", "Emergency"),
                ("Oxygen equipment available", "Medical"),
                ("Hospital notified", "Medical"),
            ]
        )

        self.conn.commit()

    def _seed_procedure(self, cat_id, key, name, desc, now,
                         steps=None, checklist=None):
        cur = self.conn.execute(
            "INSERT INTO procedures "
            "(proc_key, name, category_id, description, has_checklist, "
            "is_builtin, is_active, created_date, modified_date) "
            "VALUES (?, ?, ?, ?, 1, 1, 1, ?, ?)",
            (key, name, cat_id, desc, now, now)
        )
        proc_id = cur.lastrowid

        if steps:
            for idx, step_data in enumerate(steps):
                indent, is_hdr, text = step_data
                self.conn.execute(
                    "INSERT INTO procedure_steps "
                    "(procedure_id, step_number, indent_level, text, "
                    "is_header, is_note, is_warning) "
                    "VALUES (?, ?, ?, ?, ?, 0, 0)",
                    (proc_id, idx + 1, indent, text, int(is_hdr))
                )

        if checklist:
            for idx, (text, category) in enumerate(checklist):
                self.conn.execute(
                    "INSERT INTO checklist_items "
                    "(procedure_id, item_number, text, category) "
                    "VALUES (?, ?, ?, ?)",
                    (proc_id, idx + 1, text, category)
                )
                

    # ================================================================
    # INPUTS CRUD
    # ================================================================

    def add_input(self, proc_id: int, input_key: str,
                  input_label: str, input_type: str = 'text',
                  input_options: str = '', input_default: str = '',
                  input_unit: str = '', input_order: int = 0,
                  is_required: bool = False) -> int:
        cur = self.conn.execute(
            "INSERT INTO procedure_inputs "
            "(procedure_id, input_key, input_label, input_type, "
            "input_options, input_default, input_unit, "
            "input_order, is_required) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (proc_id, input_key, input_label, input_type,
             input_options, input_default, input_unit,
             input_order, int(is_required)))
        self.conn.commit()
        return cur.lastrowid

    def get_inputs(self, proc_id: int) -> List[Dict]:
        rows = self.conn.execute(
            "SELECT * FROM procedure_inputs "
            "WHERE procedure_id=? ORDER BY input_order",
            (proc_id,)).fetchall()
        return [dict(r) for r in rows]

    def replace_all_inputs(self, proc_id: int, inputs: List[Dict]):
        self.conn.execute(
            "DELETE FROM procedure_inputs WHERE procedure_id=?",
            (proc_id,))
        for idx, inp in enumerate(inputs):
            self.conn.execute(
                "INSERT INTO procedure_inputs "
                "(procedure_id, input_key, input_label, input_type, "
                "input_options, input_default, input_unit, "
                "input_order, is_required) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (proc_id, inp.get('key', ''),
                 inp.get('label', ''),
                 inp.get('type', 'text'),
                 inp.get('options', ''),
                 inp.get('default', ''),
                 inp.get('unit', ''),
                 idx, int(inp.get('required', False))))
        self.conn.commit()

    # ================================================================
    # TABLES CRUD
    # ================================================================

    def add_table(self, proc_id: int, table_name: str,
                  columns: str, data: str = '[]',
                  order: int = 0) -> int:
        cur = self.conn.execute(
            "INSERT INTO procedure_tables "
            "(procedure_id, table_name, table_columns, "
            "table_data, table_order) "
            "VALUES (?, ?, ?, ?, ?)",
            (proc_id, table_name, columns, data, order))
        self.conn.commit()
        return cur.lastrowid

    def get_tables(self, proc_id: int) -> List[Dict]:
        rows = self.conn.execute(
            "SELECT * FROM procedure_tables "
            "WHERE procedure_id=? ORDER BY table_order",
            (proc_id,)).fetchall()
        return [dict(r) for r in rows]

    def replace_all_tables(self, proc_id: int, tables: List[Dict]):
        self.conn.execute(
            "DELETE FROM procedure_tables WHERE procedure_id=?",
            (proc_id,))
        for idx, tbl in enumerate(tables):
            self.conn.execute(
                "INSERT INTO procedure_tables "
                "(procedure_id, table_name, table_columns, "
                "table_data, table_order) "
                "VALUES (?, ?, ?, ?, ?)",
                (proc_id, tbl.get('name', ''),
                 tbl.get('columns', ''),
                 tbl.get('data', '[]'), idx))
        self.conn.commit()
# ============================================================================
# PART 2 - GUI: PROCEDURE EDITOR & GENERATOR DIALOG
# ============================================================================

from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QListWidget, QListWidgetItem, QTextEdit, QGroupBox,
    QSplitter, QWidget, QScrollArea, QFrame, QMessageBox,
    QGridLayout, QTabWidget, QApplication, QCheckBox,
    QComboBox, QSpinBox, QDoubleSpinBox, QFormLayout,
    QTreeWidget, QTreeWidgetItem, QHeaderView, QLineEdit,
    QPlainTextEdit, QFileDialog, QProgressBar, QInputDialog,
    QMenu, QToolButton, QAbstractItemView
)
from PySide6.QtCore import Qt, Signal, QSize
from PySide6.QtGui import QFont, QColor, QAction


# ============================================================================
# STYLES
# ============================================================================

DB_DIALOG_STYLE = """
QDialog, QWidget { background: #0a0f1a; color: #e0e0e0;
    font-family: 'Segoe UI'; font-size: 11px; }
QLabel { color: #c0ccd8; }
QLabel#title { color: #e94560; font-size: 20px; font-weight: bold; }
QLabel#sub { color: #8899aa; font-size: 10px; }
QLabel#catH { color: #e94560; font-size: 12px; font-weight: bold; }
QLabel#stat { color: #27ae60; font-size: 10px; font-weight: bold; }

QTreeWidget, QListWidget {
    background: #0d1525; border: 2px solid #1a2744;
    border-radius: 6px; color: #d0d8e0; outline: none; }
QTreeWidget::item, QListWidget::item { padding: 3px; }
QTreeWidget::item:selected, QListWidget::item:selected {
    background: #1a2744; border: 1px solid #e94560; }
QTreeWidget::indicator:unchecked {
    border: 2px solid #0f3460; background: #0a0f1a;
    width: 14px; height: 14px; border-radius: 3px; }
QTreeWidget::indicator:checked {
    border: 2px solid #e94560; background: #e94560;
    width: 14px; height: 14px; border-radius: 3px; }

QGroupBox { border: 2px solid #1a2744; border-radius: 8px;
    margin-top: 12px; padding-top: 18px;
    font-weight: bold; color: #e94560; background: #0d1525; }
QGroupBox::title { subcontrol-origin: margin; left: 12px;
    padding: 0 8px; }

QTabWidget::pane { border: 1px solid #1a2744; background: #0d1525; }
QTabBar::tab { background: #0a0f1a; color: #8899aa;
    padding: 7px 14px; border-radius: 4px 4px 0 0; font-size: 10px; }
QTabBar::tab:selected { background: #1a2744; color: #e94560;
    border-bottom: 2px solid #e94560; }

QLineEdit, QComboBox, QSpinBox, QTextEdit, QPlainTextEdit {
    background: #0d1525; border: 1px solid #1a2744;
    border-radius: 4px; padding: 5px; color: #e0e0e0; min-height: 24px; }
QLineEdit:focus, QComboBox:focus { border: 2px solid #e94560; }
QComboBox QAbstractItemView { background: #0d1525; color: #e0e0e0;
    selection-background-color: #e94560; }

QCheckBox { spacing: 6px; }
QCheckBox::indicator { width: 14px; height: 14px;
    border: 2px solid #0f3460; border-radius: 3px; background: #0a0f1a; }
QCheckBox::indicator:checked { background: #e94560; border-color: #e94560; }

QPushButton { background: #0f3460; color: #fff; border: none;
    border-radius: 5px; padding: 7px 16px; font-weight: bold;
    font-size: 10px; }
QPushButton:hover { background: #1a5276; }
QPushButton:disabled { background: #2c3e50; color: #666; }
QPushButton#gen { background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
    stop:0 #e94560, stop:1 #c0392b);
    font-size: 13px; padding: 12px 30px; min-height: 40px; }
QPushButton#gen:hover { background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
    stop:0 #ff6b81, stop:1 #e74c3c); }
QPushButton#add { background: #27ae60; }
QPushButton#add:hover { background: #2ecc71; }
QPushButton#del { background: #c0392b; }
QPushButton#del:hover { background: #e74c3c; }
QPushButton#edit { background: #e67e22; }
QPushButton#edit:hover { background: #f39c12; }
QPushButton#dup { background: #8e44ad; }
QPushButton#cancel { background: #1a2744; color: #a0b0c0;
    border: 1px solid #2c3e50; }

QProgressBar { border: 1px solid #1a2744; border-radius: 4px;
    text-align: center; color: #fff; background: #0a0f1a; min-height: 20px; }
QProgressBar::chunk { background: qlineargradient(x1:0,y1:0,x2:1,y2:0,
    stop:0 #e94560, stop:1 #0f3460); border-radius: 3px; }

QScrollArea { border: none; background: transparent; }
"""


# ============================================================================
# STEP EDITOR DIALOG
# ============================================================================

class StepEditorDialog(QDialog):
    """ادیتور مرحله‌ای پروسیجر — structured step model (audit P1):
    پیش‌شرط، معیار پذیرش، Hold Point و Witness Point به‌همراه دکمه
    ساختاردهی خودکار از متن مرحله."""

    def __init__(self, step_text="", indent=0, is_header=False,
                 is_note=False, is_warning=False,
                 precondition="", acceptance="",
                 hold_point=False, witness_point=False, role="",
                 parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Step")
        self.setMinimumSize(560, 460)
        self.setStyleSheet(DB_DIALOG_STYLE)

        layout = QVBoxLayout(self)

        form = QFormLayout()
        self.txt = QPlainTextEdit()
        self.txt.setPlainText(step_text)
        self.txt.setMaximumHeight(80)
        form.addRow("Step Text:", self.txt)

        # Auto-structure from free text
        auto_row = QHBoxLayout()
        btn_auto = QPushButton("✨ Auto-structure from text")
        btn_auto.setMaximumWidth(220)
        btn_auto.clicked.connect(self._auto_structure)
        auto_row.addWidget(btn_auto)
        auto_row.addStretch(1)
        form.addRow("", auto_row)

        self.indent = QComboBox()
        self.indent.addItems(["0 - Main", "1 - Sub-step", "2 - Sub-sub"])
        self.indent.setCurrentIndex(min(indent, 2))
        form.addRow("Indent Level:", self.indent)

        self.chk_header = QCheckBox("Section Header (bold)")
        self.chk_header.setChecked(is_header)
        form.addRow("", self.chk_header)

        self.chk_note = QCheckBox("Note (italic)")
        self.chk_note.setChecked(is_note)
        form.addRow("", self.chk_note)

        self.chk_warning = QCheckBox("⚠️ Warning")
        self.chk_warning.setChecked(is_warning)
        form.addRow("", self.chk_warning)

        # ---- Structured execution fields (audit P1) ----
        self.txt_pre = QPlainTextEdit()
        self.txt_pre.setPlainText(precondition)
        self.txt_pre.setMaximumHeight(50)
        form.addRow("Precondition:", self.txt_pre)

        self.txt_acc = QPlainTextEdit()
        self.txt_acc.setPlainText(acceptance)
        self.txt_acc.setMaximumHeight(50)
        form.addRow("Acceptance criteria:", self.txt_acc)

        self.cmb_role = QComboBox()
        self.cmb_role.setEditable(True)
        self.cmb_role.addItems(["", "Driller", "Toolpusher", "Supervisor",
                                "Company Man", "Mud Engineer", "Cementer",
                                "Derrickman", "Electrician", "Mechanic",
                                "Safety Officer", "Mud Logger"])
        if role:
            self.cmb_role.setCurrentText(role)
        form.addRow("Responsible role:", self.cmb_role)

        hp_row = QHBoxLayout()
        self.chk_hold = QCheckBox("🚧 Hold Point (stop for approval)")
        self.chk_hold.setChecked(hold_point)
        self.chk_wit = QCheckBox("👁️ Witness Point (witness required)")
        self.chk_wit.setChecked(witness_point)
        hp_row.addWidget(self.chk_hold)
        hp_row.addWidget(self.chk_wit)
        hp_row.addStretch(1)
        form.addRow("Execution control:", hp_row)

        layout.addLayout(form)

        btns = QHBoxLayout()
        btns.addStretch()
        btn_ok = QPushButton("✅ OK")
        btn_ok.clicked.connect(self.accept)
        btn_cancel = QPushButton("Cancel")
        btn_cancel.setObjectName("cancel")
        btn_cancel.clicked.connect(self.reject)
        btns.addWidget(btn_cancel)
        btns.addWidget(btn_ok)
        layout.addLayout(btns)

    def _auto_structure(self):
        """Parse the step text and fill structured fields (structured_steps)."""
        text = self.txt.toPlainText().strip()
        if not text:
            return
        try:
            from structured_steps import structure_step
            s = structure_step(text)
            if s.precondition and not self.txt_pre.toPlainText().strip():
                self.txt_pre.setPlainText(s.precondition)
            if s.acceptance and not self.txt_acc.toPlainText().strip():
                self.txt_acc.setPlainText(s.acceptance)
            if s.hold_point:
                self.chk_hold.setChecked(True)
            if s.witness_point:
                self.chk_wit.setChecked(True)
        except Exception:
            import traceback
            traceback.print_exc()

    def get_data(self) -> dict:
        return {
            'text': self.txt.toPlainText().strip(),
            'indent': self.indent.currentIndex(),
            'is_header': self.chk_header.isChecked(),
            'is_note': self.chk_note.isChecked(),
            'is_warning': self.chk_warning.isChecked(),
            'precondition': self.txt_pre.toPlainText().strip(),
            'acceptance': self.txt_acc.toPlainText().strip(),
            'hold_point': self.chk_hold.isChecked(),
            'witness_point': self.chk_wit.isChecked(),
            'role': self.cmb_role.currentText().strip(),
        }


# ============================================================================
# PROCEDURE EDITOR DIALOG
# ============================================================================

class ProcedureEditorDialog(QDialog):
    """ادیتور کامل پروسیجر"""

    def __init__(self, db: ProcedureDatabase, proc_id: int = 0,
                 parent=None):
        super().__init__(parent)
        self.db = db
        self.proc_id = proc_id
        self.is_new = (proc_id == 0)
        self.setWindowTitle(
            "New Procedure" if self.is_new else "Edit Procedure")
        self.setMinimumSize(900, 700)
        self.setStyleSheet(DB_DIALOG_STYLE)

        self._build_ui()
        if not self.is_new:
            self._load_data()

    def _build_ui(self):
        layout = QVBoxLayout(self)

        # Header info
        grp_info = QGroupBox("📋 Procedure Information")
        form = QFormLayout()

        self.txt_name = QLineEdit()
        self.txt_name.setPlaceholderText("Procedure name...")
        form.addRow("Name:", self.txt_name)

        self.cmb_category = QComboBox()
        cats = self.db.get_all_categories()
        for c in cats:
            self.cmb_category.addItem(f"{c.icon} {c.name}", c.id)
        form.addRow("Category:", self.cmb_category)

        self.txt_desc = QLineEdit()
        self.txt_desc.setPlaceholderText("Brief description...")
        form.addRow("Description:", self.txt_desc)

        self.txt_tags = QLineEdit()
        self.txt_tags.setPlaceholderText("Tags (comma separated)...")
        form.addRow("Tags:", self.txt_tags)

        self.chk_checklist = QCheckBox("Has Checklist")
        self.chk_checklist.setChecked(True)
        form.addRow("", self.chk_checklist)

        # ---- Procedure <-> Well / Risk linking (audit P1) ----
        link_row = QHBoxLayout()
        self.cmb_well = QComboBox()
        self.cmb_well.setEditable(True)
        self.cmb_well.setPlaceholderText("Select / type well id or name...")
        self._load_wells()
        link_row.addWidget(self.cmb_well, 1)
        btn_clear = QPushButton("✖")
        btn_clear.setMaximumWidth(34)
        btn_clear.setToolTip("Clear well link")
        btn_clear.clicked.connect(
            lambda: self.cmb_well.setCurrentIndex(0))
        link_row.addWidget(btn_clear)
        form.addRow("Linked Well:", link_row)

        self.txt_section = QLineEdit()
        self.txt_section.setPlaceholderText(
            "e.g. 12.25 in hole section / 9.625 in casing run")
        form.addRow("Linked Section:", self.txt_section)

        risk_row = QHBoxLayout()
        self.cmb_risk = QComboBox()
        self.cmb_risk.setPlaceholderText("Select drilling risk...")
        self._load_risks()
        risk_row.addWidget(self.cmb_risk, 1)
        btn_addrisk = QPushButton("➕")
        btn_addrisk.setMaximumWidth(40)
        btn_addrisk.setToolTip("Add selected risk to the link list")
        btn_addrisk.clicked.connect(self._add_linked_risk)
        risk_row.addWidget(btn_addrisk)
        form.addRow("Linked Risks:", risk_row)

        self.lbl_risks = QLabel("No linked risks.")
        self.lbl_risks.setWordWrap(True)
        self.lbl_risks.setStyleSheet("color:#8a8a9a;font-size:9px;")
        form.addRow("", self.lbl_risks)

        grp_info.setLayout(form)
        layout.addWidget(grp_info)

        # Tabs for steps and checklist
        self.tabs = QTabWidget()

        # ---- Steps Tab ----
        steps_widget = QWidget()
        sl = QVBoxLayout(steps_widget)

        steps_btns = QHBoxLayout()
        btn_add_step = QPushButton("➕ Add Step")
        btn_add_step.setObjectName("add")
        btn_add_step.clicked.connect(self._add_step)
        steps_btns.addWidget(btn_add_step)

        btn_edit_step = QPushButton("✏️ Edit")
        btn_edit_step.setObjectName("edit")
        btn_edit_step.clicked.connect(self._edit_step)
        steps_btns.addWidget(btn_edit_step)

        btn_del_step = QPushButton("🗑️ Delete")
        btn_del_step.setObjectName("del")
        btn_del_step.clicked.connect(self._del_step)
        steps_btns.addWidget(btn_del_step)

        btn_up = QPushButton("⬆️ Up")
        btn_up.clicked.connect(self._move_step_up)
        steps_btns.addWidget(btn_up)

        btn_down = QPushButton("⬇️ Down")
        btn_down.clicked.connect(self._move_step_down)
        steps_btns.addWidget(btn_down)

        steps_btns.addStretch()
        sl.addLayout(steps_btns)

        self.steps_list = QListWidget()
        self.steps_list.setSelectionMode(QAbstractItemView.SingleSelection)
        self.steps_list.itemDoubleClicked.connect(self._edit_step)
        sl.addWidget(self.steps_list)

        self.tabs.addTab(steps_widget, "📝 Steps")

    def _load_wells(self):
        """Populate the linked-well combo from the canonical wells DB."""
        try:
            from well_model import WellDatabase
            wdb = WellDatabase()
            wells = wdb.list_wells()
            wdb.close()
        except Exception:
            wells = []
        self.cmb_well.clear()
        self.cmb_well.addItem("(none)", "")
        for w in wells:
            wid = w.get("id", "")
            name = w.get("well_name", "") or w.get("name", "")
            label = f"{name} [{wid}]" if wid else str(name)
            self.cmb_well.addItem(label, wid)

    def _load_risks(self):
        """Populate the linked-risk combo from the drilling problems DB."""
        try:
            from drilling_problems_db import DrillingProblemsDatabase
            pdb = DrillingProblemsDatabase()
            probs = pdb.get_all()
            pdb.close()
        except Exception:
            probs = []
        self.cmb_risk.clear()
        self.cmb_risk.addItem("(none)", None)
        for p in probs or []:
            pid = getattr(p, "id", None) or p.get("id") if isinstance(
                p, dict) else getattr(p, "id", None)
            name = (p.get("name") if isinstance(p, dict)
                    else getattr(p, "name", ""))
            self.cmb_risk.addItem(str(name), pid)

    def _add_linked_risk(self):
        rid = self.cmb_risk.currentData()
        if not rid:
            return
        cur = json.loads(self.lbl_risks.property("risk_ids") or "[]")
        if rid not in cur:
            cur.append(rid)
        self.lbl_risks.setProperty("risk_ids", json.dumps(cur))
        names = self._risk_names(cur)
        self.lbl_risks.setText("Linked risks: " + (", ".join(names)
                               if names else "(none)"))

    def _risk_names(self, ids):
        try:
            from drilling_problems_db import DrillingProblemsDatabase
            pdb = DrillingProblemsDatabase()
            by_id = {p.id: p.name for p in pdb.get_all()}
            pdb.close()
        except Exception:
            by_id = {}
        return [str(by_id.get(i, i)) for i in ids]

    def _linked_risk_ids(self):
        try:
            return json.loads(self.lbl_risks.property("risk_ids") or "[]")
        except Exception:
            return []

        # ---- Checklist Tab ----
        cl_widget = QWidget()
        cll = QVBoxLayout(cl_widget)

        cl_btns = QHBoxLayout()
        btn_add_cl = QPushButton("➕ Add Item")
        btn_add_cl.setObjectName("add")
        btn_add_cl.clicked.connect(self._add_checklist)
        cl_btns.addWidget(btn_add_cl)

        btn_edit_cl = QPushButton("✏️ Edit")
        btn_edit_cl.setObjectName("edit")
        btn_edit_cl.clicked.connect(self._edit_checklist)
        cl_btns.addWidget(btn_edit_cl)

        btn_del_cl = QPushButton("🗑️ Delete")
        btn_del_cl.setObjectName("del")
        btn_del_cl.clicked.connect(self._del_checklist)
        cl_btns.addWidget(btn_del_cl)

        cl_btns.addStretch()
        cll.addLayout(cl_btns)

        self.cl_list = QListWidget()
        self.cl_list.itemDoubleClicked.connect(self._edit_checklist)
        cll.addWidget(self.cl_list)

        self.tabs.addTab(cl_widget, "✅ Checklist")

        layout.addWidget(self.tabs)

        # ---- Inputs Tab ----
        inp_widget = QWidget()
        inpl = QVBoxLayout(inp_widget)
        inpl.setContentsMargins(0, 0, 0, 0)

        inp_btns = QHBoxLayout()
        btn_add_inp = QPushButton("➕ Add Input")
        btn_add_inp.setObjectName("add")
        btn_add_inp.clicked.connect(self._add_input)
        inp_btns.addWidget(btn_add_inp)

        btn_edit_inp = QPushButton("✏️ Edit")
        btn_edit_inp.setObjectName("edit")
        btn_edit_inp.clicked.connect(self._edit_input)
        inp_btns.addWidget(btn_edit_inp)

        btn_del_inp = QPushButton("🗑️ Delete")
        btn_del_inp.setObjectName("del")
        btn_del_inp.clicked.connect(self._del_input)
        inp_btns.addWidget(btn_del_inp)

        inp_btns.addStretch()
        inpl.addLayout(inp_btns)

        self.inp_list = QListWidget()
        inpl.addWidget(self.inp_list)

        self.tabs.addTab(inp_widget, "⚙️ Variable Inputs")
        
        # Bottom buttons
        bottom = QHBoxLayout()
        bottom.addStretch()

        btn_cancel = QPushButton("Cancel")
        btn_cancel.setObjectName("cancel")
        btn_cancel.clicked.connect(self.reject)
        bottom.addWidget(btn_cancel)

        btn_save = QPushButton("💾 Save Procedure")
        btn_save.setObjectName("add")
        btn_save.clicked.connect(self._save)
        bottom.addWidget(btn_save)

        layout.addLayout(bottom)

    def _load_data(self):
        rec = self.db.get_procedure(self.proc_id)
        if not rec:
            return

        self.txt_name.setText(rec.name)
        self.txt_desc.setText(rec.description)
        self.txt_tags.setText(rec.tags)
        self.chk_checklist.setChecked(rec.has_checklist)

        # Set category
        for i in range(self.cmb_category.count()):
            if self.cmb_category.itemData(i) == rec.category_id:
                self.cmb_category.setCurrentIndex(i)
                break

        # Linked well / section / risks
        if rec.linked_well_id:
            for i in range(self.cmb_well.count()):
                if str(self.cmb_well.itemData(i) or "") == \
                        str(rec.linked_well_id):
                    self.cmb_well.setCurrentIndex(i)
                    break
            else:
                self.cmb_well.setEditText(rec.linked_well_id)
        self.txt_section.setText(rec.linked_section)
        try:
            risk_ids = json.loads(rec.linked_risk_ids or "[]")
        except Exception:
            risk_ids = []
        self.lbl_risks.setProperty("risk_ids", json.dumps(risk_ids))
        names = self._risk_names(risk_ids)
        self.lbl_risks.setText("Linked risks: " + (", ".join(names)
                               if names else "(none)"))

        # Load steps
        for s in rec.steps:
            prefix = "  " * s.indent_level
            icon = ""
            if s.is_header:
                icon = "🔵 "
            elif s.is_warning:
                icon = "⚠️ "
            elif s.is_note:
                icon = "📌 "

            item = QListWidgetItem(f"{icon}{prefix}{s.text}")
            item.setData(Qt.UserRole, {
                'id': s.id, 'text': s.text,
                'indent': s.indent_level,
                'is_header': s.is_header,
                'is_note': s.is_note,
                'is_warning': s.is_warning,
                'precondition': s.precondition,
                'acceptance': s.acceptance,
                'hold_point': s.hold_point,
                'witness_point': s.witness_point,
                'role': s.role,
            })
            self.steps_list.addItem(item)

        # Load checklist
        for c in rec.checklist:
            item = QListWidgetItem(f"[{c.category}] {c.text}")
            item.setData(Qt.UserRole, {
                'id': c.id, 'text': c.text,
                'category': c.category
            })
            self.cl_list.addItem(item)

        # Load inputs
        inputs = self.db.get_inputs(self.proc_id)
        for inp in inputs:
            item = QListWidgetItem(
                f"[{inp['input_type']}] {inp['input_key']}: "
                f"{inp['input_label']}"
                f"{' (' + inp['input_options'] + ')' if inp['input_options'] else ''}"
                f"{' [' + inp['input_unit'] + ']' if inp['input_unit'] else ''}")
            item.setData(Qt.UserRole, {
                'key': inp['input_key'],
                'label': inp['input_label'],
                'type': inp['input_type'],
                'options': inp['input_options'],
                'default': inp['input_default'],
                'unit': inp['input_unit'],
                'required': bool(inp['is_required'])
            })
            self.inp_list.addItem(item)
            
    def _add_step(self):
        dlg = StepEditorDialog(parent=self)
        if dlg.exec() == QDialog.Accepted:
            data = dlg.get_data()
            if data['text']:
                item = self._steps_item(data)
                self.steps_list.addItem(item)

    def _steps_item(self, data):
        prefix = "  " * data['indent']
        icon = "🔵 " if data['is_header'] else \
               "⚠️ " if data['is_warning'] else \
               "📌 " if data['is_note'] else ""
        badges = ""
        if data.get('hold_point'):
            badges += " 🚧HP"
        if data.get('witness_point'):
            badges += " 👁️WP"
        if data.get('role'):
            badges += f" [{data['role']}]"
        item = QListWidgetItem(f"{icon}{prefix}{data['text']}{badges}")
        item.setData(Qt.UserRole, data)
        return item

    def _edit_step(self):
        row = self.steps_list.currentRow()
        if row < 0:
            return
        item = self.steps_list.item(row)
        data = item.data(Qt.UserRole) or {}

        dlg = StepEditorDialog(
            step_text=data.get('text', ''),
            indent=data.get('indent', 0),
            is_header=data.get('is_header', False),
            is_note=data.get('is_note', False),
            is_warning=data.get('is_warning', False),
            precondition=data.get('precondition', ''),
            acceptance=data.get('acceptance', ''),
            hold_point=data.get('hold_point', False),
            witness_point=data.get('witness_point', False),
            role=data.get('role', ''),
            parent=self
        )
        if dlg.exec() == QDialog.Accepted:
            new_data = dlg.get_data()
            new_item = self._steps_item(new_data)
            item.setText(new_item.text())
            item.setData(Qt.UserRole, new_data)

    def _del_step(self):
        row = self.steps_list.currentRow()
        if row >= 0:
            self.steps_list.takeItem(row)

    def _move_step_up(self):
        row = self.steps_list.currentRow()
        if row > 0:
            item = self.steps_list.takeItem(row)
            self.steps_list.insertItem(row - 1, item)
            self.steps_list.setCurrentRow(row - 1)

    def _move_step_down(self):
        row = self.steps_list.currentRow()
        if row < self.steps_list.count() - 1:
            item = self.steps_list.takeItem(row)
            self.steps_list.insertItem(row + 1, item)
            self.steps_list.setCurrentRow(row + 1)

    def _add_checklist(self):
        text, ok = QInputDialog.getText(
            self, "Add Checklist Item", "Item text:")
        if ok and text:
            cat, ok2 = QInputDialog.getText(
                self, "Category", "Category:",
                text="General")
            if ok2:
                item = QListWidgetItem(f"[{cat}] {text}")
                item.setData(Qt.UserRole, {
                    'text': text, 'category': cat or 'General'})
                self.cl_list.addItem(item)

    def _edit_checklist(self):
        row = self.cl_list.currentRow()
        if row < 0:
            return
        item = self.cl_list.item(row)
        data = item.data(Qt.UserRole) or {}

        text, ok = QInputDialog.getText(
            self, "Edit Item", "Text:",
            text=data.get('text', ''))
        if ok and text:
            cat, ok2 = QInputDialog.getText(
                self, "Category", "Category:",
                text=data.get('category', 'General'))
            if ok2:
                item.setText(f"[{cat}] {text}")
                item.setData(Qt.UserRole, {
                    'text': text, 'category': cat or 'General'})

    def _del_checklist(self):
        row = self.cl_list.currentRow()
        if row >= 0:
            self.cl_list.takeItem(row)

    def _save(self):
        name = self.txt_name.text().strip()
        if not name:
            QMessageBox.warning(self, "Error", "Name is required.")
            return

        cat_id = self.cmb_category.currentData()
        desc = self.txt_desc.text().strip()
        tags = self.txt_tags.text().strip()
        has_cl = self.chk_checklist.isChecked()

        # Collect steps
        steps = []
        for i in range(self.steps_list.count()):
            data = self.steps_list.item(i).data(Qt.UserRole)
            if data and data.get('text'):
                steps.append(data)

        # Collect checklist
        checklist = []
        for i in range(self.cl_list.count()):
            data = self.cl_list.item(i).data(Qt.UserRole)
            if data and data.get('text'):
                checklist.append(data)

        # Linked well / section / risks (audit P1)
        linked_well = ""
        if self.cmb_well.currentIndex() > 0:
            linked_well = str(self.cmb_well.currentData() or "")
        elif self.cmb_well.currentText().strip() and \
                self.cmb_well.currentText().strip() != "(none)":
            linked_well = self.cmb_well.currentText().strip()
        linked_section = self.txt_section.text().strip()
        linked_risk_ids = json.dumps(self._linked_risk_ids())

        try:
            if self.is_new:
                self.proc_id = self.db.add_procedure(
                    name=name, category_id=cat_id,
                    description=desc, has_checklist=has_cl, tags=tags,
                    linked_well_id=linked_well,
                    linked_section=linked_section,
                    linked_risk_ids=linked_risk_ids)
            else:
                self.db.update_procedure(
                    self.proc_id, name=name, category_id=cat_id,
                    description=desc, has_checklist=has_cl, tags=tags,
                    linked_well_id=linked_well,
                    linked_section=linked_section,
                    linked_risk_ids=linked_risk_ids)

            # Save steps
            self.db.replace_all_steps(self.proc_id, steps)

            # Save checklist
            self.db.replace_all_checklist(self.proc_id, checklist)

            # Save inputs
            inputs = []
            for i in range(self.inp_list.count()):
                data = self.inp_list.item(i).data(Qt.UserRole)
                if data:
                    inputs.append(data)
            self.db.replace_all_inputs(self.proc_id, inputs)
            
            QMessageBox.information(self, "Saved",
                                     f"Procedure '{name}' saved!")
            self.accept()

        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))

    def _add_input(self):
        key, ok = QInputDialog.getText(
            self, "Input Key", "Variable key (e.g. well_name):")
        if not ok or not key:
            return
        label, ok2 = QInputDialog.getText(
            self, "Label", "Display label:")
        if not ok2:
            return

        inp_type, ok3 = QInputDialog.getItem(
            self, "Type", "Input type:",
            ["text", "combo", "spin", "dspin", "check"],
            0, False)
        if not ok3:
            return

        options = ""
        default = ""
        unit = ""

        if inp_type == "combo":
            options, _ = QInputDialog.getText(
                self, "Options",
                "Options (comma separated):")
        if inp_type in ("spin", "dspin"):
            unit, _ = QInputDialog.getText(
                self, "Unit", "Unit (e.g. ft, psi, ppg):")
            default, _ = QInputDialog.getText(
                self, "Default", "Default value:")
        if inp_type == "text":
            default, _ = QInputDialog.getText(
                self, "Default", "Default value:")
        required, okr = QMessageBox.question(
            self, "Required?",
            "Is this input REQUIRED (must be filled before the procedure "
            "can be released)?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        required = (required == QMessageBox.Yes)

        item = QListWidgetItem(
            f"[{inp_type}] {key}: {label}"
            f"{' (required)' if required else ''}"
            f"{' (' + options + ')' if options else ''}"
            f"{' [' + unit + ']' if unit else ''}")
        item.setData(Qt.UserRole, {
            'key': key, 'label': label, 'type': inp_type,
            'options': options, 'default': default, 'unit': unit,
            'required': required
        })
        self.inp_list.addItem(item)

    def _del_input(self):
        row = self.inp_list.currentRow()
        if row >= 0:
            self.inp_list.takeItem(row)

    def _edit_input(self):
        """Batch Y — full edit of a variable input: key/label/type/unit/
        default/options/required."""
        row = self.inp_list.currentRow()
        if row < 0:
            QMessageBox.information(self, "Edit Input",
                                    "Select an input first.")
            return
        item = self.inp_list.item(row)
        data = item.data(Qt.UserRole) or {}

        key, ok = QInputDialog.getText(self, "Input Key", "Variable key:",
                                       text=data.get('key', ''))
        if not ok or not key:
            return
        label, ok2 = QInputDialog.getText(self, "Label", "Display label:",
                                          text=data.get('label', ''))
        if not ok2:
            return
        inp_type, ok3 = QInputDialog.getItem(
            self, "Type", "Input type:",
            ["text", "combo", "spin", "dspin", "check"],
            ["text", "combo", "spin", "dspin", "check"].index(
                data.get('type', 'text')), False)
        if not ok3:
            return
        options = data.get('options', '')
        default = data.get('default', '')
        unit = data.get('unit', '')
        if inp_type == "combo":
            options, _ = QInputDialog.getText(
                self, "Options", "Options (comma separated):",
                text=options)
        if inp_type in ("spin", "dspin"):
            unit, _ = QInputDialog.getText(self, "Unit",
                                           "Unit (e.g. ft, psi, ppg):",
                                           text=unit)
            default, _ = QInputDialog.getText(self, "Default",
                                              "Default value:",
                                              text=default)
        if inp_type == "text":
            default, _ = QInputDialog.getText(self, "Default",
                                              "Default value:",
                                              text=default)
        required, okr = QMessageBox.question(
            self, "Required?",
            "Is this input REQUIRED (must be filled before release)?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes if data.get('required') else QMessageBox.No)
        required = (required == QMessageBox.Yes)

        new_data = {
            'key': key, 'label': label, 'type': inp_type,
            'options': options, 'default': default, 'unit': unit,
            'required': required,
        }
        item.setText(
            f"[{inp_type}] {key}: {label}"
            f"{' (required)' if required else ''}"
            f"{' (' + options + ')' if options else ''}"
            f"{' [' + unit + ']' if unit else ''}")
        item.setData(Qt.UserRole, new_data)
            
# ============================================================================
# MAIN PROCEDURE MANAGER DIALOG
# ============================================================================

def _subst(t, vals):
    """Replace {{key}} / {key} with the resolved value; leave keys
    without a value untouched so they can be reported."""
    if not t or not vals:
        return t, []
    try:
        from wizard_engine import PLACEHOLDER_RE
    except Exception:
        PLACEHOLDER_RE = None
    unresolved = []
    if PLACEHOLDER_RE is None:
        return t, []
    def _repl(m):
        key = m.group(1) or m.group(2)
        if key in vals and str(vals[key]).strip():
            return str(vals[key])
        unresolved.append(key)
        return m.group(0)
    out = PLACEHOLDER_RE.sub(_repl, t)
    return out, sorted(set(unresolved))


def generate_procedures_docx(db, proc_ids, out_path,
                                  progress=None, operator_name="",
                                  contractor_name="",
                                  input_values=None):
    """Headless Word export of procedures (+ checklists).

    input_values (Batch X — unified procedure parameter engine):
      {proc_id: {input_key: value}} or flat {input_key: value} (applied to
      every procedure).  Values are substituted into step texts (both
      {{key}} and {key} syntax); inputs without a user value fall back to
      the DB default (marked "(default)" in the parameter table); inputs
      without any value are reported in the result as unresolved so the
      caller can warn/block before release.

    Scrub is applied to every text field (Batch T defense in depth).
    """
    ids = [i for i in proc_ids if i]
    if not ids:
        return {"ok": False, "error": "no procedures selected"}

    def _scrub(t):
        if not t:
            return ""
        try:
            from wizard_engine import neutralize_text
            return neutralize_text(t, operator_name, contractor_name)
        except Exception:
            return t

    def _prog(v):
        if progress:
            progress(v)

    unresolved_report = {}

    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        doc = Document()
        sec = doc.sections[0]
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.0)

        # ---- COVER PAGE ----
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("DRILLING OPERATIONS\nPROCEDURES && CHECKLISTS")
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(
            f"Generated: {datetime.now().strftime('%d-%B-%Y')}")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)

        doc.add_page_break()
        _prog(5)

        # ---- TABLE OF CONTENTS ----
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'))
        r = p.add_run("  TABLE OF CONTENTS")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph("")
        for idx, pid in enumerate(ids):
            rec = db.get_procedure(pid)
            if rec:
                tp = doc.add_paragraph()
                tr1 = tp.add_run(f"{idx+1}. ")
                tr1.bold = True
                tr1.font.size = Pt(10)
                tr2 = tp.add_run(rec.name)
                tr2.font.size = Pt(10)

        doc.add_page_break()
        _prog(10)

        # ---- EACH PROCEDURE ----
        total = len(ids)
        for idx, pid in enumerate(ids):
            rec = db.get_procedure(pid)
            if not rec:
                continue

            # Heading
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="0C2D48" w:val="clear"/>'))
            r = p.add_run(f"  {idx+1}. {rec.name}")
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            doc.add_paragraph("")

            # ---- VARIABLE INPUTS + PARAMETER SUBSTITUTION (Batch X) ----
            inputs = db.get_inputs(pid)
            # resolve values: user-provided > DB default > unresolved
            pv = {}
            if isinstance(input_values, dict):
                pv = dict(input_values.get(pid, {}) or {})
                # flat map applies to every procedure
                if not any(k in input_values for k in (pid,)):
                    flat = {k: v for k, v in input_values.items()
                            if not isinstance(v, (dict, list))}
                    for k, v in flat.items():
                        pv.setdefault(k, v)
            vals = {}
            used_default = set()
            for inp in inputs:
                key = inp["input_key"]
                user_val = pv.get(key)
                if user_val not in (None, ""):
                    vals[key] = user_val
                elif inp.get("input_default"):
                    vals[key] = inp["input_default"]
                    used_default.add(key)
            proc_unresolved = []
            if inputs:
                doc.add_paragraph("")
                ip = doc.add_paragraph()
                ipPr = ip._p.get_or_add_pPr()
                ipPr.append(parse_xml(
                    f'<w:shd {nsdecls("w")} '
                    f'w:fill="E67E22" w:val="clear"/>'))
                ir = ip.add_run("  ⚙️ VARIABLE PARAMETERS")
                ir.bold = True
                ir.font.size = Pt(11)
                ir.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                inp_table = doc.add_table(
                    rows=len(inputs) + 1, cols=3)
                inp_table.style = 'Table Grid'
                inp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                for i, h in enumerate(
                        ["Parameter", "Value", "Unit"]):
                    cell = inp_table.rows[0].cells[i]
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(h)
                    cr.bold = True
                    cr.font.size = Pt(9)
                    cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shd = parse_xml(
                        f'<w:shd {nsdecls("w")} '
                        f'w:fill="E67E22"/>')
                    cell._tc.get_or_add_tcPr().append(shd)

                # Data rows — show the RESOLVED value; defaults are
                # explicitly marked so they are never mistaken for
                # user-entered values (Batch X).
                for i, inp in enumerate(inputs):
                    key = inp["input_key"]
                    row = inp_table.rows[i + 1]
                    row.cells[0].text = ""
                    row.cells[0].paragraphs[0].add_run(
                        inp['input_label']
                    ).font.size = Pt(9)
                    row.cells[1].text = ""
                    if key in vals:
                        mark = " (default)" if key in used_default else ""
                        row.cells[1].paragraphs[0].add_run(
                            str(vals[key]) + mark
                        ).font.size = Pt(9)
                    else:
                        row.cells[1].paragraphs[0].add_run(
                            "[Not Entered]"
                        ).font.size = Pt(9)
                        proc_unresolved.append(key)
                    row.cells[2].text = ""
                    row.cells[2].paragraphs[0].add_run(
                        inp['input_unit'] or ""
                    ).font.size = Pt(9)
                    if i % 2 == 0:
                        for c in range(3):
                            shd = parse_xml(
                                f'<w:shd {nsdecls("w")} '
                                f'w:fill="FEF9E7"/>')
                            row.cells[c]._tc.get_or_add_tcPr(
                            ).append(shd)

                doc.add_paragraph("")

            # ---- STEPS (parameter substitution applied) ----
            for s in rec.steps:
                raw_text, un = _subst(s.text, vals)
                proc_unresolved.extend(un)
                step_text = _scrub(raw_text)
                if not step_text.strip():
                    doc.add_paragraph("")
                    continue

                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Cm(
                    0.2 + s.indent_level * 0.5)
                sp.paragraph_format.space_after = Pt(2)

                if s.is_header:
                    sr = sp.add_run(step_text)
                    sr.bold = True
                    sr.font.size = Pt(11)
                    sr.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
                    sp.paragraph_format.space_before = Pt(8)
                elif s.is_warning:
                    sr = sp.add_run(f"⚠️ {step_text}")
                    sr.bold = True
                    sr.font.size = Pt(10)
                    sr.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
                elif s.is_note:
                    sr = sp.add_run(f"📌 {step_text}")
                    sr.font.size = Pt(9)
                    sr.font.italic = True
                    sr.font.color.rgb = RGBColor(0x85, 0x92, 0x9E)
                else:
                    parts = step_text.split(' ', 1)
                    if (len(parts) == 2 and '.' in parts[0] and
                            any(c.isdigit() for c in parts[0])):
                        sr1 = sp.add_run(parts[0] + " ")
                        sr1.bold = True
                        sr1.font.size = Pt(10)
                        sr1.font.color.rgb = RGBColor(0x21, 0x61, 0x8C)
                        sr2 = sp.add_run(parts[1])
                        sr2.font.size = Pt(10)
                    else:
                        sr = sp.add_run(step_text)
                        sr.font.size = Pt(10)

                # ---- Structured execution info (audit P1) ----
                if s.precondition:
                    pp_ = doc.add_paragraph()
                    pp_.paragraph_format.left_indent = Cm(
                        0.5 + s.indent_level * 0.5)
                    pp_.paragraph_format.space_after = Pt(1)
                    pr1 = pp_.add_run("▸ Precondition: ")
                    pr1.bold = True
                    pr1.font.size = Pt(9)
                    pr1.font.color.rgb = RGBColor(0xB9, 0x77, 0x0E)
                    pr2 = pp_.add_run(_scrub(_subst(s.precondition, vals)[0] if vals else s.precondition))
                    pr2.font.size = Pt(9)
                if s.acceptance:
                    pa_ = doc.add_paragraph()
                    pa_.paragraph_format.left_indent = Cm(
                        0.5 + s.indent_level * 0.5)
                    pa_.paragraph_format.space_after = Pt(1)
                    ar1 = pa_.add_run("✓ Acceptance: ")
                    ar1.bold = True
                    ar1.font.size = Pt(9)
                    ar1.font.color.rgb = RGBColor(0x1E, 0x84, 0x49)
                    ar2 = pa_.add_run(_scrub(_subst(s.acceptance, vals)[0] if vals else s.acceptance))
                    ar2.font.size = Pt(9)
                if s.hold_point or s.witness_point:
                    ph_ = doc.add_paragraph()
                    ph_.paragraph_format.left_indent = Cm(
                        0.5 + s.indent_level * 0.5)
                    ph_.paragraph_format.space_after = Pt(4)
                    tags = []
                    if s.hold_point:
                        tags.append("🚧 HOLD POINT — stop and obtain "
                                    "approval before continuing")
                    if s.witness_point:
                        tags.append("👁️ WITNESS POINT — witness "
                                    "required at this step")
                    hr1 = ph_.add_run("   |   ".join(tags))
                    hr1.bold = True
                    hr1.font.size = Pt(9)
                    hr1.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                if s.role:
                    pr_ = doc.add_paragraph()
                    pr_.paragraph_format.left_indent = Cm(
                        0.5 + s.indent_level * 0.5)
                    pr_.paragraph_format.space_after = Pt(4)
                    rr1 = pr_.add_run(f"👤 Responsible: {s.role}")
                    rr1.font.size = Pt(9)
                    rr1.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

            # ---- CHECKLIST ----
            if rec.checklist:
                doc.add_paragraph("")
                cp = doc.add_paragraph()
                cpPr = cp._p.get_or_add_pPr()
                cpPr.append(parse_xml(
                    f'<w:shd {nsdecls("w")} '
                    f'w:fill="1B4F72" w:val="clear"/>'))
                cr = cp.add_run(f"  ✅ CHECKLIST - {rec.name}")
                cr.bold = True
                cr.font.size = Pt(12)
                cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                # Group by category
                cats = {}
                for c in rec.checklist:
                    cats.setdefault(_scrub(c.category),
                                    []).append(_scrub(c.text))

                for cat, items in cats.items():
                    cp2 = doc.add_paragraph()
                    cp2.paragraph_format.space_before = Pt(6)
                    cr2 = cp2.add_run(f"▸ {cat}")
                    cr2.bold = True
                    cr2.font.size = Pt(10)
                    cr2.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

                    table = doc.add_table(
                        rows=len(items) + 1, cols=4)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Checklist header
                    for i, h in enumerate(
                            ["☐", "Item", "Init.", "Date"]):
                        cell = table.rows[0].cells[i]
                        cell.text = ""
                        cp3 = cell.paragraphs[0]
                        cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cr3 = cp3.add_run(h)
                        cr3.bold = True
                        cr3.font.size = Pt(8)
                        cr3.font.color.rgb = RGBColor(
                            0xFF, 0xFF, 0xFF)
                        shd = parse_xml(
                            f'<w:shd {nsdecls("w")} '
                            f'w:fill="0C2D48"/>')
                        cell._tc.get_or_add_tcPr().append(shd)

                    # Set column widths
                    for row in table.rows:
                        row.cells[0].width = Inches(0.4)
                        row.cells[1].width = Inches(4.2)
                        row.cells[2].width = Inches(0.6)
                        row.cells[3].width = Inches(1.0)

                    # Checklist items
                    for i, item_text in enumerate(items):
                        row = table.rows[i + 1]
                        row.cells[0].text = ""
                        cp4 = row.cells[0].paragraphs[0]
                        cp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cp4.add_run("☐").font.size = Pt(11)

                        row.cells[1].text = ""
                        row.cells[1].paragraphs[0].add_run(
                            item_text).font.size = Pt(9)

                        row.cells[2].text = ""
                        row.cells[3].text = ""

                        if i % 2 == 0:
                            for c in range(4):
                                shd = parse_xml(
                                    f'<w:shd {nsdecls("w")} '
                                    f'w:fill="EBF5FB"/>')
                                row.cells[c]._tc.get_or_add_tcPr(
                                ).append(shd)

                # Sign-off
                doc.add_paragraph("")
                sp2 = doc.add_paragraph()
                sp2.add_run("SIGN-OFF:").bold = True

                sig = doc.add_table(rows=2, cols=4)
                sig.style = 'Table Grid'
                for i, h in enumerate(
                        ["Driller", "Toolpusher",
                        "Company Man", "Date"]):
                    cell = sig.rows[0].cells[i]
                    cell.text = ""
                    cp5 = cell.paragraphs[0]
                    cp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr5 = cp5.add_run(h)
                    cr5.bold = True
                    cr5.font.size = Pt(9)
                    cr5.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shd = parse_xml(
                        f'<w:shd {nsdecls("w")} '
                        f'w:fill="1B4F72"/>')
                    cell._tc.get_or_add_tcPr().append(shd)
                    sig.rows[1].cells[i].text = ""
                    sig.rows[1].cells[i].paragraphs[
                        0].paragraph_format.space_after = Pt(25)

            # Record unresolved parameters for this procedure
            if proc_unresolved:
                unresolved_report[rec.name] = sorted(set(proc_unresolved))

            # Page break after each procedure
            doc.add_page_break()

            pct = 10 + int(85 * (idx + 1) / total)
            _prog(pct)

        doc.save(out_path)
        _prog(100)
        return {"ok": True, "path": out_path, "procedures": len(ids),
                "unresolved": unresolved_report}
    except Exception as e:
        import traceback
        return {"ok": False,
                "error": f"{e}\n{traceback.format_exc()[-400:]}"}


class ProcedureInputsDialog(QDialog):
    """Batch X — collect parameter values for the selected procedures.

    Each input of each procedure is shown with its DB default prefilled;
    the user edits values which are then substituted into the step texts
    during Word export."""

    def __init__(self, db, proc_ids, parent=None):
        super().__init__(parent)
        self.setWindowTitle("⚙️ Procedure Parameters — enter values")
        self.setMinimumSize(760, 520)
        self.setStyleSheet(DB_DIALOG_STYLE)
        self.db = db
        self.proc_ids = [i for i in proc_ids if i]

        lay = QVBoxLayout(self)
        head = QLabel(
            "Enter the parameter values for the selected procedures.\n"
            "Prefilled values are the database defaults — values you "
            "enter here are substituted into the procedure steps and "
            "marked as user-entered in the Word document.")
        head.setWordWrap(True)
        lay.addWidget(head)

        tabs = QTabWidget()
        self._editors = {}
        for pid in self.proc_ids:
            rec = db.get_procedure(pid)
            if not rec:
                continue
            inputs = db.get_inputs(pid)
            w = QWidget()
            fl = QFormLayout(w)
            edits = {}
            if not inputs:
                fl.addRow(QLabel("(no variable parameters)"))
            for inp in inputs:
                key = inp["input_key"]
                le = QLineEdit(inp["input_default"] or "")
                le.setPlaceholderText(inp["input_label"])
                label = inp["input_label"] + (
                    " *" if inp.get("is_required") else "")
                if inp.get("input_unit"):
                    label += f" [{inp['input_unit']}]"
                fl.addRow(label, le)
                edits[key] = le
            tabs.addTab(w, rec.name[:30])
            self._editors[pid] = edits
        lay.addWidget(tabs)

        btns = QHBoxLayout()
        btns.addStretch()
        b_ok = QPushButton("✅ Generate with these values")
        b_ok.clicked.connect(self.accept)
        b_cancel = QPushButton("Cancel")
        b_cancel.setObjectName("cancel")
        b_cancel.clicked.connect(self.reject)
        btns.addWidget(b_cancel)
        btns.addWidget(b_ok)
        lay.addLayout(btns)

    def values(self) -> Dict:
        out = {}
        for pid, edits in self._editors.items():
            pv = {}
            for key, le in edits.items():
                v = le.text().strip()
                if v:
                    pv[key] = v
            if pv:
                out[pid] = pv
        return out


class ProcedureManagerDialog(QDialog):
    """دیالوگ اصلی مدیریت و تولید پروسیجرها"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle(
            "🛢️ Drilling Procedures & Checklists Manager")
        self.setMinimumSize(1300, 850)
        self.setStyleSheet(DB_DIALOG_STYLE)

        self.db = ProcedureDatabase()
        self._build_ui()
        self._refresh_tree()
        self._update_stats()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)
        layout.setContentsMargins(10, 10, 10, 10)

        # Title
        hdr = QHBoxLayout()
        t = QLabel("🛢️ Drilling Procedures & Checklists Manager")
        t.setObjectName("title")
        hdr.addWidget(t)
        hdr.addStretch()
        self.stat_label = QLabel("")
        self.stat_label.setObjectName("stat")
        hdr.addWidget(self.stat_label)
        layout.addLayout(hdr)

        # Toolbar
        tb = QHBoxLayout()

        btn_new = QPushButton("➕ New Procedure")
        btn_new.setObjectName("add")
        btn_new.clicked.connect(self._new_procedure)
        tb.addWidget(btn_new)

        btn_new_cat = QPushButton("📁 New Category")
        btn_new_cat.clicked.connect(self._new_category)
        tb.addWidget(btn_new_cat)

        tb.addWidget(self._separator())

        btn_import = QPushButton("📥 Import")
        btn_import.clicked.connect(self._import)
        tb.addWidget(btn_import)

        btn_export = QPushButton("📤 Export Selected")
        btn_export.clicked.connect(self._export)
        tb.addWidget(btn_export)

        tb.addWidget(self._separator())

        # lifecycle toolbar (audit P1)
        self.status_combo = QComboBox()
        self.status_combo.addItems(ProcedureDatabase.LIFECYCLE_STATES)
        self.status_combo.setMaximumWidth(170)
        tb.addWidget(self.status_combo)
        btn_set_status = QPushButton("↩ Set Status")
        btn_set_status.clicked.connect(self._set_status)
        tb.addWidget(btn_set_status)
        btn_approve = QPushButton("✅ Approve")
        btn_approve.clicked.connect(self._approve)
        tb.addWidget(btn_approve)

        tb.addWidget(self._separator())

        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("🔍 Search procedures...")
        self.search_box.setMaximumWidth(250)
        self.search_box.textChanged.connect(self._search)
        tb.addWidget(self.search_box)

        tb.addStretch()

        btn_sel_all = QPushButton("✅ All")
        btn_sel_all.clicked.connect(self._select_all)
        tb.addWidget(btn_sel_all)

        btn_sel_none = QPushButton("❌ None")
        btn_sel_none.clicked.connect(self._deselect_all)
        tb.addWidget(btn_sel_none)

        layout.addLayout(tb)

        # Main splitter
        splitter = QSplitter(Qt.Horizontal)

        # ---- LEFT: Tree ----
        left = QWidget()
        ll = QVBoxLayout(left)
        ll.setContentsMargins(0, 0, 0, 0)

        self.sel_count = QLabel("Selected: 0")
        self.sel_count.setObjectName("sub")
        ll.addWidget(self.sel_count)

        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Procedure", "Type", ""])
        self.tree.setMinimumWidth(450)
        self.tree.header().setSectionResizeMode(0, QHeaderView.Stretch)
        self.tree.header().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self._context_menu)
        self.tree.itemChanged.connect(self._on_check_changed)
        self.tree.currentItemChanged.connect(self._on_select)
        ll.addWidget(self.tree)

        splitter.addWidget(left)

        # ---- RIGHT: Preview ----
        right = QWidget()
        rl = QVBoxLayout(right)
        rl.setContentsMargins(0, 0, 0, 0)

        self.preview_tabs = QTabWidget()

        # Procedure preview
        self.preview_proc = QTextEdit()
        self.preview_proc.setReadOnly(True)
        self.preview_tabs.addTab(self.preview_proc, "📄 Procedure")

        # Checklist preview
        self.preview_cl = QTextEdit()
        self.preview_cl.setReadOnly(True)
        self.preview_tabs.addTab(self.preview_cl, "✅ Checklist")

        # Info
        self.preview_info = QTextEdit()
        self.preview_info.setReadOnly(True)
        self.preview_tabs.addTab(self.preview_info, "ℹ️ Info")

        rl.addWidget(self.preview_tabs)

        # Edit buttons
        edit_row = QHBoxLayout()
        btn_edit = QPushButton("✏️ Edit Procedure")
        btn_edit.setObjectName("edit")
        btn_edit.clicked.connect(self._edit_procedure)
        edit_row.addWidget(btn_edit)

        btn_dup = QPushButton("📋 Duplicate")
        btn_dup.setObjectName("dup")
        btn_dup.clicked.connect(self._duplicate)
        edit_row.addWidget(btn_dup)

        btn_delete = QPushButton("🗑️ Delete")
        btn_delete.setObjectName("del")
        btn_delete.clicked.connect(self._delete_procedure)
        edit_row.addWidget(btn_delete)

        edit_row.addStretch()
        rl.addLayout(edit_row)

        splitter.addWidget(right)
        splitter.setSizes([500, 800])
        layout.addWidget(splitter)

        # Bottom
        bottom = QHBoxLayout()
        bottom.addStretch()

        self.progress = QProgressBar()
        self.progress.setMinimumWidth(200)
        self.progress.setVisible(False)
        bottom.addWidget(self.progress)

        btn_cancel = QPushButton("✖ Close")
        btn_cancel.setObjectName("cancel")
        btn_cancel.clicked.connect(self.accept)
        bottom.addWidget(btn_cancel)

        btn_gen = QPushButton("📄 GENERATE WORD DOCUMENT")
        btn_gen.setObjectName("gen")
        btn_gen.clicked.connect(self._generate_word)
        bottom.addWidget(btn_gen)

        bottom.addStretch()
        layout.addLayout(bottom)

    def _separator(self):
        f = QFrame()
        f.setFrameShape(QFrame.VLine)
        f.setStyleSheet("color: #1a2744;")
        return f

    # ================================================================
    # TREE MANAGEMENT
    # ================================================================

    def _current_proc_id(self):
        item = self.tree.currentItem()
        if item is None:
            return None
        data = item.data(0, Qt.UserRole)
        if isinstance(data, dict) and data.get("type") == "procedure":
            return data.get("id")
        # maybe parent
        p = item.parent()
        if p is not None:
            d = p.data(0, Qt.UserRole)
            if isinstance(d, dict) and d.get("type") == "procedure":
                return d.get("id")
        return None

    def _set_status(self):
        pid = self._current_proc_id()
        if not pid:
            QMessageBox.information(self, "Select", "Select a procedure first.")
            return
        status = self.status_combo.currentText()
        self.db.set_status(pid, status)
        self._refresh_tree()
        self._update_stats()

    def _approve(self):
        pid = self._current_proc_id()
        if not pid:
            QMessageBox.information(self, "Select", "Select a procedure first.")
            return
        self.db.approve(pid, approver="admin")
        QMessageBox.information(self, "Approved",
                                "Procedure approved and released for use.")
        self._refresh_tree()
        self._update_stats()

    def _refresh_tree(self):
        self.tree.blockSignals(True)
        self.tree.clear()
        self.cat_items = {}

        categories = self.db.get_all_categories()
        procedures = self.db.get_all_procedures(active_only=True)

        # Group
        grouped = {}
        for p in procedures:
            cat = p.category_name or "Uncategorized"
            if cat not in grouped:
                grouped[cat] = []
            grouped[cat].append(p)

        for cat in categories:
            procs = grouped.get(cat.name, [])
            cat_item = QTreeWidgetItem(self.tree)
            cat_item.setText(0, f"{cat.icon} {cat.name} ({len(procs)})")
            cat_item.setText(1, "")
            cat_item.setFlags(
                cat_item.flags() | Qt.ItemIsAutoTristate |
                Qt.ItemIsUserCheckable)
            cat_item.setCheckState(0, Qt.Unchecked)
            cat_item.setExpanded(False)
            cat_item.setData(0, Qt.UserRole, {'type': 'category',
                                               'id': cat.id})
            self.cat_items[cat.name] = cat_item

            for p in procs:
                child = QTreeWidgetItem(cat_item)
                icon = "⭐" if p.is_builtin else "📝"
                child.setText(0, f"{icon} {p.name}")
                child.setText(1, "✓CL" if p.has_checklist else "")
                child.setFlags(child.flags() | Qt.ItemIsUserCheckable)
                child.setCheckState(0, Qt.Unchecked)
                child.setData(0, Qt.UserRole, {'type': 'procedure',
                                                'id': p.id})

        self.tree.blockSignals(False)
        self._update_count()

    def _on_check_changed(self, item, col):
        self._update_count()

    def _on_select(self, current, previous):
        if not current:
            return
        data = current.data(0, Qt.UserRole)
        if not data or data.get('type') != 'procedure':
            self.preview_proc.clear()
            self.preview_cl.clear()
            self.preview_info.clear()
            return

        proc_id = data['id']
        rec = self.db.get_procedure(proc_id)
        if not rec:
            return

        # Procedure preview — parameters substituted with defaults so raw
        # {key} placeholders never appear (Batch X)
        try:
            _inputs = self.db.get_inputs(proc_id)
            _dflt = {i["input_key"]: i.get("input_default") or ""
                     for i in _inputs if i.get("input_default")}
        except Exception:
            _dflt = {}
        html = f"<h3 style='color:#e94560'>{rec.name}</h3>"
        for s in rec.steps:
            _s_text, _ = _subst(s.text, _dflt)
            _s_pre, _ = _subst(s.precondition, _dflt)
            _s_acc, _ = _subst(s.acceptance, _dflt)
            prefix = "&nbsp;" * (s.indent_level * 6)
            badges = ""
            if s.hold_point:
                badges += " <span style='color:#c0392b'>🚧HOLD POINT</span>"
            if s.witness_point:
                badges += " <span style='color:#2471a3'>👁️WITNESS</span>"
            if s.is_header:
                html += f"<p style='color:#1a5276;font-weight:bold;margin:8px 0 2px'>{prefix}{_s_text}{badges}</p>"
            elif s.is_warning:
                html += f"<p style='color:#e74c3c;margin:1px 0'>{prefix}⚠️ {_s_text}{badges}</p>"
            elif s.is_note:
                html += f"<p style='color:#8899aa;font-style:italic;margin:1px 0'>{prefix}📌 {_s_text}{badges}</p>"
            elif not _s_text.strip():
                html += "<br>"
            else:
                html += f"<p style='margin:1px 0;font-size:10px'>{prefix}{_s_text}{badges}</p>"
            if _s_pre:
                html += (f"<p style='margin:0 0 1px 12px;font-size:9px;"
                         f"color:#b9770e'>▸ Precondition: {_s_pre}</p>")
            if _s_acc:
                html += (f"<p style='margin:0 0 1px 12px;font-size:9px;"
                         f"color:#1e8449'>✓ Acceptance: {_s_acc}</p>")
            if s.role:
                html += (f"<p style='margin:0 0 1px 12px;font-size:9px;"
                         f"color:#2e86c1'>👤 Role: {s.role}</p>")
        self.preview_proc.setHtml(html)

        # Checklist preview
        cl_html = f"<h3 style='color:#27ae60'>Checklist - {rec.name}</h3>"
        if rec.checklist:
            cats = {}
            for c in rec.checklist:
                cats.setdefault(c.category, []).append(c.text)
            for cat, items in cats.items():
                cl_html += f"<p style='color:#e94560;font-weight:bold'>▸ {cat}</p>"
                for item in items:
                    cl_html += f"<p style='margin:1px 0;font-size:10px'>☐ {item}</p>"
        else:
            cl_html += "<p style='color:#666'>No checklist items.</p>"
        self.preview_cl.setHtml(cl_html)

        # Info
        info = (
            f"<b>ID:</b> {rec.id}<br>"
            f"<b>Key:</b> {rec.proc_key}<br>"
            f"<b>Category:</b> {rec.category_name}<br>"
            f"<b>Description:</b> {rec.description}<br>"
            f"<b>Built-in:</b> {'Yes' if rec.is_builtin else 'No'}<br>"
            f"<b>Steps:</b> {len(rec.steps)}<br>"
            f"<b>Checklist Items:</b> {len(rec.checklist)}<br>"
            f"<b>Tags:</b> {rec.tags}<br>"
            f"<b>Created:</b> {rec.created_date}<br>"
            f"<b>Modified:</b> {rec.modified_date}<br>"
            f"<b>Version:</b> {rec.version}"
        )
        if rec.linked_well_id or rec.linked_section:
            info += (f"<br><b>Linked Well:</b> {rec.linked_well_id or '—'}"
                     f"{' / ' + rec.linked_section if rec.linked_section else ''}<br>")
        try:
            risk_ids = json.loads(rec.linked_risk_ids or "[]")
        except Exception:
            risk_ids = []
        if risk_ids:
            names = self._risk_names(risk_ids)
            info += (f"<b>Linked Risks:</b> {', '.join(names)}<br>")
        self.preview_info.setHtml(info)

        # Inputs preview
        inputs = self.db.get_inputs(proc_id)
        if inputs:
            info += "<br><b>Variable Inputs:</b><br>"
            for inp in inputs:
                info += (f"• <b>{inp['input_label']}</b> "
                         f"[{inp['input_type']}]"
                         f"{' (' + inp['input_options'] + ')' if inp['input_options'] else ''}"
                         f"{' ' + inp['input_unit'] if inp['input_unit'] else ''}<br>")
                         
    def _update_count(self):
        ids = self._get_selected_ids()
        self.sel_count.setText(f"Selected: {len(ids)} procedures")

    def _update_stats(self):
        stats = self.db.get_stats()
        self.stat_label.setText(
            f"📊 {stats['total_procedures']} procedures "
            f"({stats['builtin']} built-in, {stats['custom']} custom) | "
            f"{stats['categories']} categories | "
            f"{stats['total_steps']} steps | "
            f"{stats['total_checklist_items']} checklist items"
        )

    def _get_selected_ids(self) -> List[int]:
        ids = []
        for cat_name, cat_item in self.cat_items.items():
            for i in range(cat_item.childCount()):
                child = cat_item.child(i)
                if child.checkState(0) == Qt.Checked:
                    data = child.data(0, Qt.UserRole)
                    if data and data.get('id'):
                        ids.append(data['id'])
        return ids

    # ================================================================
    # ACTIONS
    # ================================================================

    def _new_procedure(self):
        dlg = ProcedureEditorDialog(self.db, proc_id=0, parent=self)
        if dlg.exec() == QDialog.Accepted:
            self._refresh_tree()
            self._update_stats()

    def _edit_procedure(self):
        current = self.tree.currentItem()
        if not current:
            return
        data = current.data(0, Qt.UserRole)
        if not data or data.get('type') != 'procedure':
            return

        dlg = ProcedureEditorDialog(
            self.db, proc_id=data['id'], parent=self)
        if dlg.exec() == QDialog.Accepted:
            self._refresh_tree()
            self._update_stats()
            # Re-select to refresh preview
            self._on_select(self.tree.currentItem(), None)

    def _duplicate(self):
        current = self.tree.currentItem()
        if not current:
            return
        data = current.data(0, Qt.UserRole)
        if not data or data.get('type') != 'procedure':
            return

        new_id = self.db.duplicate_procedure(data['id'])
        if new_id:
            self._refresh_tree()
            self._update_stats()
            QMessageBox.information(self, "Duplicated",
                                     "Procedure duplicated!")

    def _delete_procedure(self):
        current = self.tree.currentItem()
        if not current:
            return
        data = current.data(0, Qt.UserRole)
        if not data or data.get('type') != 'procedure':
            return

        reply = QMessageBox.question(
            self, "Delete",
            "Delete this procedure?\n"
            "(Built-in procedures will be deactivated, "
            "custom will be permanently deleted.)")
        if reply == QMessageBox.Yes:
            self.db.delete_procedure(data['id'])
            self._refresh_tree()
            self._update_stats()

    def _new_category(self):
        name, ok = QInputDialog.getText(
            self, "New Category", "Category name:")
        if ok and name:
            icon, ok2 = QInputDialog.getText(
                self, "Icon", "Icon (emoji):", text="📁")
            self.db.add_category(name, icon or "📁")
            self._refresh_tree()
            self._update_stats()

    def _select_all(self):
        self.tree.blockSignals(True)
        for cat_item in self.cat_items.values():
            cat_item.setCheckState(0, Qt.Checked)
        self.tree.blockSignals(False)
        self._update_count()

    def _deselect_all(self):
        self.tree.blockSignals(True)
        for cat_item in self.cat_items.values():
            cat_item.setCheckState(0, Qt.Unchecked)
        self.tree.blockSignals(False)
        self._update_count()

    def _search(self, text):
        if not text.strip():
            self._refresh_tree()
            return

        results = self.db.search_procedures(text)
        result_ids = {r.id for r in results}

        for cat_item in self.cat_items.values():
            any_visible = False
            for i in range(cat_item.childCount()):
                child = cat_item.child(i)
                data = child.data(0, Qt.UserRole)
                visible = data and data.get('id') in result_ids
                child.setHidden(not visible)
                if visible:
                    any_visible = True
            cat_item.setHidden(not any_visible)
            if any_visible:
                cat_item.setExpanded(True)

    def _import(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Import Procedures", "",
            "JSON Files (*.json)")
        if path:
            try:
                count = self.db.import_procedures(path)
                self._refresh_tree()
                self._update_stats()
                QMessageBox.information(
                    self, "Imported",
                    f"{count} procedures imported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))

    def _export(self):
        ids = self._get_selected_ids()
        if not ids:
            QMessageBox.warning(self, "No Selection",
                                "Select procedures to export.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Export Procedures",
            f"procedures_export_{datetime.now().strftime('%Y%m%d')}.json",
            "JSON Files (*.json)")
        if path:
            try:
                self.db.export_procedures(ids, path)
                QMessageBox.information(
                    self, "Exported",
                    f"{len(ids)} procedures exported to:\n{path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))

    def _context_menu(self, pos):
        item = self.tree.itemAt(pos)
        if not item:
            return
        data = item.data(0, Qt.UserRole)
        if not data:
            return

        menu = QMenu(self)
        menu.setStyleSheet(
            "QMenu { background: #0d1525; color: #e0e0e0; "
            "border: 1px solid #1a2744; } "
            "QMenu::item:selected { background: #e94560; }")

        if data.get('type') == 'procedure':
            menu.addAction("✏️ Edit", self._edit_procedure)
            menu.addAction("📋 Duplicate", self._duplicate)
            menu.addAction("📤 Export", lambda: self._export_single(data['id']))
            menu.addSeparator()
            menu.addAction("🗑️ Delete", self._delete_procedure)
        elif data.get('type') == 'category':
            menu.addAction("➕ New Procedure Here", self._new_procedure)

        menu.exec(self.tree.mapToGlobal(pos))

    def _export_single(self, proc_id):
        path, _ = QFileDialog.getSaveFileName(
            self, "Export Procedure", "procedure.json",
            "JSON Files (*.json)")
        if path:
            self.db.export_procedures([proc_id], path)
            QMessageBox.information(self, "Done", "Exported!")

    # ================================================================
    # GENERATE WORD
    # ================================================================
    def _scrub(self, text: str) -> str:
        """Defense in depth: neutralize company/well/reservoir names before
        any document export (Batch T — leak-free outputs)."""
        if not text:
            return ""
        try:
            from wizard_engine import neutralize_text
            return neutralize_text(text, self._op_name, self._con_name)
        except Exception:
            return text

    def _generate_word(self):
        ids = self._get_selected_ids()
        if not ids:
            QMessageBox.warning(self, "No Selection",
                                "Select procedures to generate.")
            return
        self._op_name = ""
        self._con_name = ""

        path, _ = QFileDialog.getSaveFileName(
            self, "Save Procedures Document",
            f"Drilling_Procedures_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)")
        if not path:
            return

        # Batch X — collect parameter values from the user before export
        input_values = {}
        try:
            dlg = ProcedureInputsDialog(self.db, ids, self)
            if dlg.exec() == QDialog.Accepted:
                input_values = dlg.values()
            else:
                return
        except Exception:
            input_values = {}

        self.progress.setVisible(True)
        self.progress.setValue(0)
        QApplication.processEvents()

        res = generate_procedures_docx(
            self.db, ids, path,
            progress=self.progress.setValue,
            operator_name=self._op_name, contractor_name=self._con_name,
            input_values=input_values)
        self.progress.setValue(100 if res.get("ok") else 0)

        if res.get("ok"):
            unres = res.get("unresolved") or {}
            total_unres = sum(len(v) for v in unres.values())
            msg = (f"Document generated!\n\n"
                   f"File: {path}\n"
                   f"Procedures: {len(ids)}")
            if total_unres:
                msg += ("\n\n⚠️ UNRESOLVED PARAMETERS:\n"
                        f"{total_unres} parameter(s) still unresolved in "
                        f"{len(unres)} procedure(s):\n")
                for name, keys in list(unres.items())[:6]:
                    msg += f"\n• {name}: {', '.join(keys[:8])}"
                QMessageBox.warning(self, "Generated with Warnings", msg)
            else:
                QMessageBox.information(self, "✅ Success", msg)
            if os.name == 'nt':
                os.startfile(path)
        else:
            QMessageBox.critical(self, "Export Failed",
                                 res.get("error", "unknown error"))

